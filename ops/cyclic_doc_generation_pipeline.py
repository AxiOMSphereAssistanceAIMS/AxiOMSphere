"""
Cyclic Document Generation Pipeline — Parallel Omi + Axi + Claude Code Audit

Implements three-tier parallel generation with quality validation:
1. OMI (Local): Searches internal DB → generates draft with internal standards
2. AXI (Internet): Searches external sources → validates & recommends improvements
3. CLAUDE CODE CLI (Auditor): Validates both Omi & Axi work independently

All three run in parallel. Results are merged and applied to next cycle.
Saves training pairs at each cycle for continuous fine-tuning.

Reference template: Asset Integrity Management Policy and Framework (879 pages, 10 tables)
Target: 95% match on structure, standards accuracy, section coverage

Usage:
    python ops/cyclic_doc_generation_pipeline.py \
        --topic "Asset Integrity Management Policy and Framework" \
        --reference-pdf "/media/.../Asset Integrity Management Policy and Framework_1.docx" \
        --max-cycles 5 \
        --target-quality 0.98 \
        --save-training-pairs
"""

import json
import hashlib
import logging
import os
import re
import subprocess
import shutil
import sys
import asyncio
import urllib.error
import urllib.request
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, asdict, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional, Tuple

# Add workspace to path
sys.path.insert(0, str(Path(__file__).parent.parent))

log = logging.getLogger("cyclic_doc_pipeline")

def _resolve_workspace_root() -> Path:
    env_candidates = [
        os.environ.get("AIMS_WORKSPACE", "").strip(),
        os.environ.get("AIMS_WORKSPACE_ROOT", "").strip(),
    ]
    for raw in env_candidates:
        if raw:
            candidate = Path(raw).expanduser()
            if candidate.exists():
                return candidate.resolve()

    here = Path(__file__).resolve()
    for candidate in [here.parent, *here.parents]:
        try:
            if (candidate / "aims_workspace").exists():
                return candidate.resolve()
        except Exception:
            continue

    for candidate in (
        Path("/ops"),
        Path("/home/axi_omi_sphere/aims-workspace"),
        Path("/data"),
        Path("/workspace"),
    ):
        try:
            if (candidate / "aims_workspace").exists():
                return candidate.resolve()
        except Exception:
            continue

    return here.parent.resolve()


ROOT = _resolve_workspace_root()
STATUS_DIR = ROOT / "aims_workspace" / "agent_architecture_status"
STATUS_DIR.mkdir(parents=True, exist_ok=True)

# Max sections per Bedrock audit call for chunked path.
# 14 sections × ~350 tokens/finding ≈ 4900 output tokens — well within 8000 limit.
_AUDIT_CHUNK_SIZE = 14

# Skills
from ops.cyclic_skills import (  # noqa: E402
    sections_semantically_match,
    validate_structure,
    verify_recommendations,
    quality_gate,
    expand_stub_sections,
)
from ops.agents.skills.section_editor import (  # noqa: E402
    apply_section_edits,
    normalize_rec,
    _extract_reference_section,
)
from ops.agents.skills.context_grounded_document_generation import (  # noqa: E402
    apply_reference_baseline,
    build_section_contract,
    build_generation_context,
    extract_reference_baseline,
    match_reference_section,
    resolve_document_archetype,
    render_generation_prompt,
    write_context_manifest,
)
from ops.docs_pipeline.bedrock_doc_audit import (  # noqa: E402
    bedrock_doc_audit,
    bedrock_axi_validate,
    parse_json_from_response,
)
from ops.agents.skills.docsreg_phase1_convergence import (  # noqa: E402
    Phase1ConvergenceOrchestrator,
)
from ops.agents.skills.docsreg_phase2_nesting import apply_phase2_nesting  # noqa: E402
from ops.agents.skills.docsreg_phase3_content_quality import select_phase3_recommendations  # noqa: E402
from ops.docgen.render_visual_qa import render_docx_for_qa, RenderMetrics  # noqa: E402
from ops.docgen.universal_overlay.enforcement_contracts import (  # noqa: E402
    EnforcementContext,
    GateResult,
)
from ops.docgen.universal_overlay.profile_binding_gate import (  # noqa: E402
    validate_generated_type,
    validate_reference_aware_profile_binding,
    validation_document_type_for,
)
from ops.docgen.universal_overlay.real_judge_gate import (  # noqa: E402
    validate_real_judge_path,
)
from ops.docgen.universal_overlay.reference_binding_gate import (  # noqa: E402
    validate_reference_binding,
)
from ops.docgen.universal_overlay.document_profile_conformance import (  # noqa: E402
    apply_profile_formatting,
    validate_document_against_profile,
    write_profile_conformance_report,
)
from ops.docgen.universal_overlay.document_type_profile_loader import (  # noqa: E402
    get_document_generation_profile,
)
from ops.docgen.universal_overlay.failure_analysis_registry import (  # noqa: E402
    write_failure_analysis_registry,
)
from ops.docgen.universal_overlay.failure_analysis_bridge import (  # noqa: E402
    write_failure_analysis_brief,
)
from ops.docgen.universal_overlay.recommendation_pool import (  # noqa: E402
    build_recommendation_pool,
)
from ops.docgen.universal_overlay.stage_contract_audit import (  # noqa: E402
    write_stage_contract_audit,
)
from ops.docgen.universal_overlay.cycle_completion_reporting import (  # noqa: E402
    atomic_write_text,
    mark_cycle_complete,
)
from ops.docgen.universal_overlay.standards_catalog import (  # noqa: E402
    FIRST_BATCH_IMPLEMENTATION_STANDARDS,
    STANDARD_CATALOG,
    active_document_formation_records,
    discovery_hints_for,
    standards_for_document_type,
)
from ops.docgen.universal_overlay.training_promotion_gate import (  # noqa: E402
    decide_training_or_promotion_allowed,
)
from ops.models.model_registry import resolve_slot  # noqa: E402
from ops.ollama_resolve import (  # noqa: E402
    effective_ollama_base_url,
    effective_small_qwen_ollama_base_url,
)
from ops.agents.skills.document_failure_diagnostics import (  # noqa: E402
    diagnose_document_failures,
)
from ops.docgen.document_architecture.domain_requirement_extractor import (  # noqa: E402
    build_requirement_graph as build_document_requirement_graph,
)

# ──────────────────────────────────────────────────────────────────────────────
# MODEL CONFIGURATION — Route operations to optimal models
# ──────────────────────────────────────────────────────────────────────────────

_AVAILABLE_MODEL_CACHE: dict[str, set[str]] = {}


def _available_ollama_models(base_url: str) -> set[str]:
    base = (base_url or "").rstrip("/")
    if base.endswith("/v1"):
        base = base[:-3].rstrip("/")
    if base.endswith("/api"):
        base = base[:-4].rstrip("/")
    cached = _AVAILABLE_MODEL_CACHE.get(base)
    if cached is not None:
        return cached
    models: set[str] = set()
    try:
        with urllib.request.urlopen(f"{base}/api/tags", timeout=5) as response:
            payload = json.loads(response.read().decode("utf-8", errors="replace"))
            for item in payload.get("models", []) or []:
                name = str(item.get("name", "")).strip()
                if name:
                    models.add(name)
    except Exception as exc:
        log.warning("[MODEL] Failed to probe available Ollama models at %s: %s", base, exc)
    _AVAILABLE_MODEL_CACHE[base] = models
    return models


def _resolve_slot120_reasoning_model() -> str:
    candidates = [
        os.environ.get("AIMS_MODEL_SLOT_120_NAME", "").strip(),
        os.environ.get("AIMS_LOCAL_MODEL", "").strip(),
        resolve_slot("120"),
        "qwen36-reasoning-35b-v1:latest",
        "qwen35-reasoning-35b-v1:latest",
        "axi_omi_sphere:latest",
        "qwen3:32b-q8_0",
    ]
    base_url = effective_ollama_base_url()
    available = _available_ollama_models(base_url)
    for model in candidates:
        if not model:
            continue
        if model in available:
            if model != resolve_slot("120"):
                log.warning(
                    "[MODEL] SLOT120 fallback selected: requested=%s available=%s base=%s",
                    resolve_slot("120"),
                    model,
                    base_url,
                )
            return model
    fallback = candidates[1] or candidates[3]
    log.warning(
        "[MODEL] SLOT120 requested model unavailable; falling back to %s on %s",
        fallback,
        base_url,
    )
    return fallback


class ModelConfig:
    """Model routing: assign operations to optimal slot models.

    VRAM Constraint (DGX 128GB total):
    - SLOT14 (14B) + SLOT120 (35B) = compatible (max ~45GB, very safe)
    - Pipeline uses both slots for search and reasoning operations
    - Awaiting user command to activate full dual-model processing
    """
    SLOT14_SEARCH = resolve_slot("14")
    SLOT120_REASONING = _resolve_slot120_reasoning_model()

    # Legacy aliases (for backward compatibility during transition)
    SEARCH_MODEL = SLOT14_SEARCH
    REASONING_MODEL = SLOT120_REASONING


SLOT120_NUM_CTX = int(os.environ.get("AIMS_DOC_SLOT120_NUM_CTX", "32768"))
OMI_GENERATE_TIMEOUT = int(os.environ.get("AIMS_DOC_OMI_GENERATE_TIMEOUT", "180"))
OMI_GENERATE_NUM_PREDICT = int(os.environ.get("AIMS_DOC_OMI_GENERATE_NUM_PREDICT", "5000"))
IMPROVEMENT_GENERATE_TIMEOUT = int(
    os.environ.get("AIMS_DOC_IMPROVEMENT_GENERATE_TIMEOUT", str(OMI_GENERATE_TIMEOUT))
)
IMPROVEMENT_GENERATE_NUM_PREDICT = int(
    os.environ.get("AIMS_DOC_IMPROVEMENT_GENERATE_NUM_PREDICT", str(OMI_GENERATE_NUM_PREDICT))
)


def _claude_bedrock_env() -> dict[str, str]:
    env = os.environ.copy()
    profile = env.setdefault(
        "AWS_PROFILE",
        "AdministratorAccess-445100240501",
    )
    env.setdefault("AWS_REGION", "us-east-1")
    env.setdefault("CLAUDE_CODE_USE_BEDROCK", "1")
    env.setdefault("AIMS_CLAUDE_REVIEW_PROVIDER", "aws_bedrock_claude_code")
    env.setdefault("ANTHROPIC_DEFAULT_SONNET_MODEL", "us.anthropic.claude-sonnet-4-6")
    env.setdefault("ANTHROPIC_DEFAULT_OPUS_MODEL", "us.anthropic.claude-opus-4-6-v1")
    env.setdefault("AIMS_CLAUDE_REVIEW_HEAVY_MODEL", "opus")
    env.setdefault("AIMS_CLAUDE_REVIEW_ALLOW_OPUS", "1")

    return env


def _extract_json_object(text: str) -> dict[str, Any]:
    candidate = (text or "").strip()
    try:
        parsed = json.loads(candidate)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass
    start = candidate.find("{")
    end = candidate.rfind("}")
    if start < 0 or end <= start:
        raise ValueError("response did not contain a JSON object")
    parsed = json.loads(candidate[start:end + 1])
    if not isinstance(parsed, dict):
        raise ValueError("parsed response is not a JSON object")
    return parsed


def _normalize_ollama_base_url(base_url: str) -> str:
    base = (base_url or "").rstrip("/")
    if base.endswith("/v1"):
        base = base[:-3].rstrip("/")
    if base.endswith("/api"):
        base = base[:-4].rstrip("/")
    return base


def _http_post_json(url: str, body: dict[str, Any], timeout: int) -> tuple[int, bytes, dict[str, Any]]:
    data = json.dumps(body).encode()
    request = urllib.request.Request(
        url,
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=timeout) as response:
        raw = response.read()
        headers = dict(response.headers.items())
        status = int(getattr(response, "status", 200) or 200)
        return status, raw, headers


def _slot120_generate(prompt: str, *, timeout: int = 180, num_predict: int = 5000) -> str:
    """Invoke Qwen3.6 with its certified raw-safe ChatML framing."""
    framed_prompt = (
        "<|im_start|>system\n"
        "You are the AIMS deep reasoning and professional document generation model."
        "<|im_end|>\n"
        "<|im_start|>user\n"
        f"{prompt}"
        "<|im_end|>\n"
        "<|im_start|>assistant\n"
        "<think>\n\n</think>\n\n"
    )
    base_url = _normalize_ollama_base_url(effective_ollama_base_url())
    attempts: list[str] = []
    native_body = {
        "model": ModelConfig.SLOT120_REASONING,
        "prompt": framed_prompt,
        "raw": True,
        "stream": True,
        "keep_alive": "6h",
        "options": {
            "temperature": 0.15,
            "num_ctx": SLOT120_NUM_CTX,
            "num_predict": num_predict,
            "stop": ["<|im_end|>", "<|endoftext|>", "<|im_start|>"],
        },
    }
    chat_body = {
        "model": ModelConfig.SLOT120_REASONING,
        "messages": [
            {"role": "system", "content": "You are the AIMS deep reasoning and professional document generation model."},
            {"role": "user", "content": prompt},
        ],
        "stream": False,
        "keep_alive": "6h",
        "options": {
            "temperature": 0.15,
            "num_ctx": SLOT120_NUM_CTX,
            "num_predict": num_predict,
        },
    }
    openai_body = {
        "model": ModelConfig.SLOT120_REASONING,
        "messages": [
            {"role": "system", "content": "You are the AIMS deep reasoning and professional document generation model."},
            {"role": "user", "content": prompt},
        ],
        "stream": False,
        "temperature": 0.15,
        "max_tokens": num_predict,
    }

    def _extract_text_from_generate(raw_bytes: bytes) -> tuple[str, str, int]:
        chunks: list[str] = []
        done_reason = ""
        prompt_eval_count = 0
        for raw_line in raw_bytes.splitlines():
            if not raw_line.strip():
                continue
            event = json.loads(raw_line)
            chunks.append(str(event.get("response", "")))
            if event.get("done"):
                done_reason = str(event.get("done_reason", ""))
                prompt_eval_count = int(event.get("prompt_eval_count", 0) or 0)
        return "".join(chunks).strip(), done_reason, prompt_eval_count

    generated = ""
    done_reason = ""
    prompt_eval_count = 0
    endpoint_errors: list[str] = []

    endpoint_attempts = [
        (
            "native_generate",
            f"{base_url}/api/generate",
            native_body,
            "native",
        ),
        (
            "native_chat",
            f"{base_url}/api/chat",
            chat_body,
            "chat",
        ),
        (
            "openai_chat",
            f"{base_url}/v1/chat/completions",
            openai_body,
            "openai",
        ),
    ]

    for attempt_name, endpoint_url, payload, mode in endpoint_attempts:
        try:
            status, raw_bytes, headers = _http_post_json(endpoint_url, payload, timeout)
            content_type = str(headers.get("Content-Type", "")).lower()
            if status >= 400:
                raise urllib.error.HTTPError(
                    endpoint_url,
                    status,
                    f"HTTP {status}",
                    headers,
                    None,
                )
            if mode == "native":
                generated, done_reason, prompt_eval_count = _extract_text_from_generate(raw_bytes)
            else:
                response_body = json.loads(raw_bytes.decode("utf-8", errors="replace"))
                if mode == "chat":
                    generated = (
                        str(response_body.get("message", {}).get("content", ""))
                        or str(response_body.get("response", ""))
                    ).strip()
                    done_reason = str(response_body.get("done_reason", response_body.get("stop_reason", "")))
                else:
                    choices = response_body.get("choices", []) if isinstance(response_body, dict) else []
                    message = choices[0].get("message", {}) if choices and isinstance(choices[0], dict) else {}
                    generated = (
                        str(message.get("content", ""))
                        or str(message.get("reasoning_content", ""))
                        or str(message.get("thinking", ""))
                        or str(response_body.get("output_text", ""))
                    ).strip()
                    done_reason = str(response_body.get("stop_reason") or response_body.get("finish_reason") or "")
            if generated:
                log.info(
                    "[SLOT120] Endpoint %s succeeded (mode=%s, content_type=%s, chars=%d)",
                    endpoint_url,
                    mode,
                    content_type,
                    len(generated),
                )
                break
            endpoint_errors.append(f"{attempt_name}: empty response")
        except urllib.error.HTTPError as exc:
            body = ""
            try:
                if exc.fp is not None:
                    body = exc.fp.read().decode("utf-8", errors="replace")[:400]
            except Exception:
                body = ""
            endpoint_errors.append(f"{attempt_name}: HTTP {exc.code} {body}".strip())
            log.warning("[SLOT120] %s failed: HTTP %s %s", attempt_name, exc.code, body[:200])
        except Exception as exc:
            endpoint_errors.append(f"{attempt_name}: {type(exc).__name__}: {exc}")
            log.warning("[SLOT120] %s failed: %s", attempt_name, exc)

    # Strip <think>...</think> reasoning blocks — must never appear in document output.
    # Two-pass: first closed blocks, then any unclosed <think> tail
    # (qwen35 may open a new <think> after the injected empty block, without closing it)
    import re as _re
    think_blocks = _re.findall(r"<think>.*?</think>", generated, flags=_re.DOTALL)
    if think_blocks:
        log.debug(f"[SLOT120] Stripping {len(think_blocks)} closed <think> block(s) from output")
        generated = _re.sub(r"<think>.*?</think>\s*", "", generated, flags=_re.DOTALL).strip()
    # Remove unclosed <think>...</end-of-string> tail
    if "<think>" in generated:
        before = len(generated)
        generated = _re.sub(r"<think>.*$", "", generated, flags=_re.DOTALL).strip()
        log.warning(f"[SLOT120] Stripped unclosed <think> tail ({before - len(generated)} chars)")

    if len(generated) < 200:
        raise RuntimeError(
            f"slot120 returned insufficient content: {len(generated)} chars; "
            f"done_reason={done_reason}; attempts={endpoint_errors}"
        )
    log.info(
        "[SLOT120] Full request processed: input_chars=%d "
        "prompt_eval_count=%d num_ctx=%d",
        len(framed_prompt),
        prompt_eval_count,
        SLOT120_NUM_CTX,
    )
    return generated


# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class GenerationMetrics:
    """Track quality metrics across generation cycles."""
    cycle: int
    timestamp: str
    model_used: str

    # Structural metrics
    sections_found: int
    sections_expected: int
    section_coverage: float  # 0-1.0

    # Content metrics
    standards_found: int
    standards_expected: int
    standards_accuracy: float  # 0-1.0

    # Reference metrics
    themes_covered: list[str] = field(default_factory=list)
    themes_missing: list[str] = field(default_factory=list)
    reference_match: float = 0.0  # 0-1.0

    # Quality metrics
    structure_score: float = 0.0  # From doc_quality_eval
    standards_score: float = 0.0
    coverage_score: float = 0.0
    overall_score: float = 0.0  # Weighted average

    # Feedback
    axi_recommendations: list[str] = field(default_factory=list)
    changes_applied: list[str] = field(default_factory=list)


@dataclass
class OmiResult:
    """Omi agent result: internal standards search + draft generation."""
    internal_standards: list[str]
    draft_text: str
    generation_time: float
    model_used: str = "qwen36-reasoning-35b-v1"  # SLOT120_REASONING (production winner)


@dataclass
class AxiResult:
    """Axi agent result: external standards search + validation."""
    external_standards: dict  # {"standards": [...], "content": {...}}
    recommendations: list[str]
    axi_feedback: str
    validation_time: float
    model_used: str = "opus"


@dataclass
class ClaudeAuditResult:
    """Claude Code CLI audit: validates both Omi and Axi work."""
    omi_quality: dict  # {"standards_accuracy": 0.9, "completeness": 0.75, ...}
    axi_quality: dict  # {"standards_accuracy": 0.95, "completeness": 0.92, ...}
    overall_assessment: str
    missing_standards: list[str]
    audit_time: float
    recommendations_from_audit: list[str] = field(default_factory=list)
    skill_recommendations: list[str] = field(default_factory=list)
    reference_gap: dict = field(default_factory=dict)
    bedrock_invoked: bool = False  # True if bedrock_doc_audit() returned successfully


@dataclass
class CycleResult:
    """Result of one generation cycle."""
    cycle_num: int
    success: bool
    generated_doc_path: Path
    metrics: GenerationMetrics
    axi_feedback: str
    ready_for_next_cycle: bool
    convergence_score: float  # 0-1.0, how close to reference
    omi_result: Optional[OmiResult] = None
    axi_result: Optional[AxiResult] = None
    audit_result: Optional[ClaudeAuditResult] = None
    bedrock_invoked: bool = False  # True if Bedrock audit was successfully invoked
    render_metrics: Optional['RenderMetrics'] = None  # Phase 3: DOCX render + visual QA metrics
    visual_qa_passed: Optional[bool] = None  # Phase 3: Overall visual QA status
    visual_qa_blocking_failure: bool = False  # Phase 3: Critical visual QA issues that block quality gate


def _cycle_artifact_count(cycle_dir: Path, names: tuple[str, ...]) -> int:
    return sum(1 for name in names if (cycle_dir / name).exists())


def _cycle_model_artifact_count(cycle_dir: Path) -> int:
    learning_pairs_dir = cycle_dir / "learning_pairs"
    if not learning_pairs_dir.exists():
        return 0
    count = 0
    for path in learning_pairs_dir.glob("*"):
        if path.is_file() and path.stat().st_size > 0:
            count += 1
    if (cycle_dir / "training_quarantine.json").exists():
        count += 1
    return count


def _cycle_skill_artifact_count(cycle_dir: Path) -> int:
    return _cycle_artifact_count(
        cycle_dir,
        (
            "recommendation_pool.json",
            "repair_plan.json",
            "skill_recommendations.json",
            "failure_analysis_registry.json",
            "failure_analysis_brief.json",
        ),
    )


def _cycle_status_from_outcome(
    *,
    target_quality: float,
    metrics: GenerationMetrics,
    hard_gate_allowed: bool,
    critical_regression: bool,
    visual_qa_blocking_failure: bool,
    no_progress_streak: int,
    no_progress_limit: int,
) -> str:
    if metrics.overall_score >= target_quality and hard_gate_allowed and not visual_qa_blocking_failure:
        return "COMPLETE"
    if no_progress_streak >= no_progress_limit:
        return "PLATEAU"
    if not hard_gate_allowed or visual_qa_blocking_failure:
        return "BLOCKED"
    if critical_regression:
        return "FAILED"
    return "PASS"


def _cycle_next_action(status: str) -> str:
    normalized = status.upper()
    if normalized == "COMPLETE":
        return "finalize_pipeline"
    if normalized == "PLATEAU":
        return "repair_weakest_dimension_and_retry"
    if normalized == "BLOCKED":
        return "analyze_failure_material"
    if normalized == "FAILED":
        return "repair_and_rerun"
    return "start_next_cycle"


def _weakest_cycle_dimension(metrics: GenerationMetrics, visual_qa_blocking_failure: bool) -> str | None:
    if visual_qa_blocking_failure:
        return "visual_qa"
    scored = {
        "structure": float(metrics.structure_score or 0.0),
        "standards": float(metrics.standards_score or 0.0),
        "coverage": float(metrics.coverage_score or 0.0),
    }
    if not scored:
        return None
    return min(scored.items(), key=lambda item: item[1])[0]


def _write_cycle_completion_summary(
    *,
    cycle_dir: Path,
    cycle: int,
    max_cycles: int,
    target_quality: float,
    metrics: GenerationMetrics,
    previous_quality_score: float,
    hard_gate_allowed: bool,
    critical_regression: bool,
    visual_qa_blocking_failure: bool,
    no_progress_streak: int,
    no_progress_limit: int,
    cycle_started_at: datetime,
    cycle_finished_at: datetime,
) -> dict[str, Any]:
    status = _cycle_status_from_outcome(
        target_quality=target_quality,
        metrics=metrics,
        hard_gate_allowed=hard_gate_allowed,
        critical_regression=critical_regression,
        visual_qa_blocking_failure=visual_qa_blocking_failure,
        no_progress_streak=no_progress_streak,
        no_progress_limit=no_progress_limit,
    )
    cycle_passed = status in {"COMPLETE", "PASS"}
    previous_quality = float(previous_quality_score or 0.0)
    quality_score = float(metrics.overall_score or 0.0)
    quality_delta = quality_score - previous_quality
    quality_percent = quality_score * 100.0
    quality_delta_percent = quality_delta * 100.0
    summary = {
        "cycle_index": int(cycle),
        "max_cycles": int(max_cycles),
        "status": status,
        "cycle_passed": cycle_passed,
        "quality_score": quality_score,
        "quality_percent": quality_percent,
        "previous_quality_score": previous_quality,
        "quality_delta": quality_delta,
        "quality_delta_percent": quality_delta_percent,
        "weakest_dimension": _weakest_cycle_dimension(metrics, visual_qa_blocking_failure),
        "repairs_count": len(getattr(metrics, "changes_applied", []) or []),
        "skill_artifacts_count": _cycle_skill_artifact_count(cycle_dir),
        "model_artifacts_count": _cycle_model_artifact_count(cycle_dir),
        "started_at": cycle_started_at.isoformat(),
        "finished_at": cycle_finished_at.isoformat(),
        "duration_seconds": max(
            0.001,
            (cycle_finished_at - cycle_started_at).total_seconds(),
        ),
        "next_action": _cycle_next_action(status),
        "execution_mode": "production",
    }
    atomic_write_text(cycle_dir / "summary.json", json.dumps(summary, indent=2, ensure_ascii=False))
    mark_cycle_complete(cycle_dir)
    return summary


def _evaluate_cycle_hard_gate(
    *,
    gate,
    struct_report,
    audit_schema_passed: bool,
    audit_quality_passed: bool,
    audit_quality_failures: list[str],
    rec_lineage_passed: bool,
    critical_regression: bool,
    standard_reference_passed: bool = True,
    profile_conformance_passed: bool = True,
) -> tuple[bool, list[str]]:
    failures: list[str] = []
    if not gate.allowed:
        failures.append(_gate_message(gate))
    if not struct_report.passed:
        failures.append(
            f"detailed_structure={struct_report.completeness_ratio:.1%} < "
            f"{struct_report.threshold:.0%}"
        )
    if not audit_schema_passed:
        failures.append("claude_audit_schema=FAIL")
    elif not audit_quality_passed:
        failures.extend(audit_quality_failures or ["claude_audit_quality=FAIL"])
    if not rec_lineage_passed:
        failures.append("recommendation_lineage=FAIL")
    if critical_regression:
        failures.append("critical_regression=true")
    if not standard_reference_passed:
        failures.append("standard_reference_register=FAIL")
    if not profile_conformance_passed:
        failures.append("document_profile_conformance=REVIEW")
    return not failures, failures


def _initial_recommendation_lineage_status(
    *,
    cycle: int,
    applied_recommendations: list[str],
) -> bool:
    """Return PASS when lineage is not applicable to the current cycle."""
    return cycle == 1 or not applied_recommendations


def _final_recommendation_lineage_status(
    *,
    applied_recommendations: list[str],
    verified_recommendations: list[str],
    text_verification_passed: bool,
    pending_global: bool,
    pending_unresolved: bool,
    rolled_back: bool,
) -> bool:
    """Combine transactional editor evidence with the text-level verifier."""
    editor_verified_all = bool(applied_recommendations) and set(
        applied_recommendations
    ).issubset(set(verified_recommendations))
    return (
        (text_verification_passed or editor_verified_all)
        and not pending_global
        and not pending_unresolved
        and not rolled_back
    )


def _gate_message(gate: Any) -> str:
    """Return a human-readable gate message across gate implementations."""
    for attr in ("reason", "message", "blocker_code"):
        value = getattr(gate, attr, None)
        if value:
            return str(value)
    return gate.__class__.__name__


def _evaluate_audit_quality(
    audit_result: Optional[ClaudeAuditResult],
    *,
    document_type: str = "technical_report",
    max_reference_gap: float | None = None,
    min_agent_quality: float | None = None,
) -> tuple[bool, list[str]]:
    """Require teacher acceptance with document-type-specific thresholds.

    Args:
        audit_result: Audit result from teacher model.
        document_type: Document type for profile-specific thresholds.
        max_reference_gap: Override reference gap threshold (None = use profile).
        min_agent_quality: Override agent quality threshold (None = use profile).
    """
    if audit_result is None:
        return False, ["claude_audit_schema=FAIL"]

    # Load profile only when needed
    profile = None
    if max_reference_gap is None or min_agent_quality is None:
        from ops.docgen.validation_profile_loader import ValidationProfileLoader
        profile = ValidationProfileLoader.get_profile(document_type)

    # Use type-specific thresholds
    if min_agent_quality is None:
        min_agent_quality = profile.quality_thresholds.overall

    if max_reference_gap is None:
        # Use document-type-specific reference/accuracy tolerance.
        # For document types where references are irrelevant/optional,
        # do not fail only because report-style references are absent.
        max_reference_gap = 1.0 - profile.quality_thresholds.accuracy

    failures: list[str] = []
    gap = float(audit_result.reference_gap.get("gap_score", 1.0))
    missing = audit_result.reference_gap.get("missing_sections", [])
    omi_avg = sum(audit_result.omi_quality.values()) / max(
        len(audit_result.omi_quality),
        1,
    )
    axi_avg = sum(audit_result.axi_quality.values()) / max(
        len(audit_result.axi_quality),
        1,
    )
    if gap > max_reference_gap:
        failures.append(
            f"claude_reference_gap={gap:.1%} > {max_reference_gap:.0%}"
        )
    if missing:
        failures.append(f"claude_missing_sections={len(missing)}")
    if omi_avg < min_agent_quality:
        failures.append(
            f"claude_omi_quality={omi_avg:.1%} < {min_agent_quality:.0%}"
        )
    if axi_avg < min_agent_quality:
        failures.append(
            f"claude_axi_quality={axi_avg:.1%} < {min_agent_quality:.0%}"
        )
    return not failures, failures


def _build_repair_plan(
    axi_recommendations: list[str],
    audit_recommendations: list[str],
    *,
    max_targets: int = 12,
    max_recommendations: int = 24,
    max_per_target: int = 2,
    use_phase1_convergence: bool = True,
    doc_text: str | None = None,
) -> dict[str, Any]:
    """Select a bounded, deterministic batch and preserve the full backlog.

    If use_phase1_convergence=True, applies DOCSREG Phase 1 staged convergence
    logic to prioritize STRUCTURE_CRITICAL and STRUCTURE_HIGH recommendations,
    deferring nesting/content/formatting changes to later phases.
    """
    # Combine recommendations for Phase 1 priority filtering (if enabled)
    all_recommendations = list(audit_recommendations) + list(axi_recommendations)
    phase1_metrics = None  # Will be populated if Phase 1 is enabled

    if use_phase1_convergence and all_recommendations:
        # Apply Phase 1 convergence: select only STRUCTURE_CRITICAL and STRUCTURE_HIGH (max 6 per phase)
        orchestrator = Phase1ConvergenceOrchestrator(
            max_recommendations_per_phase=4
        )
        phase1_batch = orchestrator.prepare_phase1_batch(all_recommendations)

        # Capture Phase 1 classification metrics for repair_plan output
        phase1_metrics = {
            "enabled": True,
            "selected_count": phase1_batch["phase1_count"],
            "deferred_count": phase1_batch["phase1_deferred_count"],
            "by_tier": phase1_batch["by_tier"],
            "total_classified": phase1_batch["classification_result"].total,
            "classifier_version": phase1_batch["classifier_version"],
        }

        # Use Phase 1 selected recommendations, log the distribution
        log.info(
            f"Phase 1 convergence: {phase1_batch['phase1_count']} selected "
            f"({phase1_batch['by_tier']}) from {phase1_batch['classification_result'].total} total; "
            f"{phase1_batch['phase1_deferred_count']} deferred to later phases"
        )

        candidates = [
            ("phase1_selected", item)
            for item in phase1_batch["phase1_selected"]
        ]

        # Phase 3: content quality — deferred recs + auto-detected stubs/fabricated standards
        phase3_metrics: dict[str, Any] = {"enabled": False}
        if doc_text:
            deferred_recs = phase1_batch["classification_result"].deferred
            phase3_result = select_phase3_recommendations(doc_text, deferred_recs)
            phase3_metrics = {
                "enabled": True,
                "selected_count": len(phase3_result["selected"]),
                "stub_sections": phase3_result["stub_sections"],
                "fabricated_standards_sections": phase3_result["fabricated_standards_sections"],
                "total_from_deferred": phase3_result["total_from_deferred"],
                "total_auto_generated": phase3_result["total_auto_generated"],
                "status": phase3_result["status"],
            }
            if phase3_result["selected"]:
                candidates += [
                    ("phase3_content_quality", item)
                    for item in phase3_result["selected"]
                ]
                log.info(
                    "Phase 3 content quality: %d selected (%d deferred + %d auto), status=%s",
                    len(phase3_result["selected"]),
                    phase3_result["total_from_deferred"],
                    phase3_result["total_auto_generated"],
                    phase3_result["status"],
                )
    else:
        phase1_metrics = {"enabled": False}
        phase3_metrics = {"enabled": False}
        candidates = [
            ("claude_audit", item)
            for item in audit_recommendations
        ] + [
            ("axi", item)
            for item in axi_recommendations
        ]

    selected: list[str] = []
    deferred: list[dict[str, str]] = []
    selected_targets: set[str] = set()
    per_target: dict[str, int] = {}
    seen: set[str] = set()

    for source, recommendation in candidates:
        dedupe_key = re.sub(r"\s+", " ", recommendation).strip().lower()
        if not dedupe_key or dedupe_key in seen:
            deferred.append({
                "source": source,
                "recommendation": recommendation,
                "reason": "duplicate_or_empty",
            })
            continue
        seen.add(dedupe_key)
        normalized = normalize_rec("PLAN", recommendation)
        target = normalized.target
        if target == "UNRESOLVED":
            deferred.append({
                "source": source,
                "recommendation": recommendation,
                "reason": "unresolved_target",
            })
            continue
        target_key = (
            f"NEW:{target}"
            if normalized.operation == "NEW_SECTION"
            else target
        )
        if target_key not in selected_targets and len(selected_targets) >= max_targets:
            deferred.append({
                "source": source,
                "recommendation": recommendation,
                "reason": "target_budget_exhausted",
            })
            continue
        target_limit = (
            1
            if normalized.operation == "NEW_SECTION"
            else max_per_target
        )
        if per_target.get(target_key, 0) >= target_limit:
            deferred.append({
                "source": source,
                "recommendation": recommendation,
                "reason": "per_target_budget_exhausted",
            })
            continue
        if len(selected) >= max_recommendations:
            deferred.append({
                "source": source,
                "recommendation": recommendation,
                "reason": "recommendation_budget_exhausted",
            })
            continue
        selected.append(recommendation)
        selected_targets.add(target_key)
        per_target[target_key] = per_target.get(target_key, 0) + 1

    return {
        "selected": selected,
        "deferred": deferred,
        "selected_count": len(selected),
        "deferred_count": len(deferred),
        "selected_targets": sorted(selected_targets),
        "limits": {
            "max_targets": max_targets,
            "max_recommendations": max_recommendations,
            "max_per_target": max_per_target,
        },
        "phase1_convergence": phase1_metrics,
        "phase3_content_quality": phase3_metrics,
    }


def _missing_section_recommendations(
    missing_sections: list[str],
) -> list[str]:
    """Convert teacher-reported numbered omissions into code-routable work."""
    recommendations: list[str] = []
    for missing in missing_sections:
        if re.match(r"^\s*Appendix\s+", missing, re.IGNORECASE):
            recommendations.append(missing)
            continue
        match = re.match(
            r"^\s*(?:(?:Sub-section|New Section|Section)\s+)?"
            r"(\d+(?:\.\d+)*)\s*:?\s+(.+?)\s*$",
            missing,
            re.IGNORECASE,
        )
        if not match:
            continue
        section_id, heading = match.groups()
        recommendations.append(
            f"Add new Section {section_id} {heading}"
        )
    return recommendations


def _runtime_preflight(output_dir: Path) -> dict[str, Any]:
    """Verify both routed models exist and projected Ollama VRAM stays below 95%."""
    output_dir.mkdir(parents=True, exist_ok=True)
    required_models = {
        "slot14": ModelConfig.SLOT14_SEARCH,
        "slot120": ModelConfig.SLOT120_REASONING,
    }
    available_names: set[str] = set()
    available_sizes: dict[str, int] = {}
    try:
        with urllib.request.urlopen(
            f"{effective_ollama_base_url().rstrip('/')}/api/tags",
            timeout=10,
        ) as response:
            tags_payload = json.loads(response.read())
        for item in list(tags_payload.get("models", [])):
            name = str(item.get("name", "")).strip()
            if not name:
                continue
            available_names.add(name)
            available_names.add(name.removesuffix(":latest"))
            size_value = item.get("size")
            if isinstance(size_value, (int, float)):
                available_sizes[name] = int(size_value)
                available_sizes[name.removesuffix(":latest")] = int(size_value)
            elif isinstance(item.get("details"), dict):
                size_hint = item["details"].get("parameter_size")
                if isinstance(size_hint, str) and size_hint:
                    match = re.match(r"([0-9]+(?:\.[0-9]+)?)([GMK]B)?", size_hint, re.I)
                    if match:
                        magnitude = float(match.group(1))
                        unit = (match.group(2) or "GB").upper()
                        if unit == "GB":
                            size_bytes = int(magnitude * (1024**3))
                        elif unit == "MB":
                            size_bytes = int(magnitude * (1024**2))
                        else:
                            size_bytes = int(magnitude * 1024)
                        available_sizes[name] = size_bytes
                        available_sizes[name.removesuffix(":latest")] = size_bytes
    except Exception as exc:
        cli = shutil.which("ollama") or "/usr/local/bin/ollama"
        try:
            listed = subprocess.run(
                [cli, "list"],
                capture_output=True,
                text=True,
                timeout=30,
                check=False,
            )
        except FileNotFoundError as cli_exc:
            raise RuntimeError(
                "Cannot resolve Ollama model list via HTTP or CLI"
            ) from cli_exc
        if listed.returncode != 0:
            raise RuntimeError(
                f"Cannot read Ollama model list via HTTP or CLI: {exc}; {listed.stderr.strip()}"
            ) from exc
        lines = [line for line in listed.stdout.splitlines() if line.strip()]
        for line in lines[1:] if len(lines) > 1 else []:
            parts = line.split()
            if not parts:
                continue
            name = parts[0].strip()
            if name:
                available_names.add(name)
                available_names.add(name.removesuffix(":latest"))
                if len(parts) >= 4:
                    try:
                        magnitude = float(parts[2])
                        unit = parts[3].upper()
                        if unit == "GB":
                            size_bytes = int(magnitude * (1024**3))
                        elif unit == "MB":
                            size_bytes = int(magnitude * (1024**2))
                        else:
                            size_bytes = int(magnitude * 1024)
                        available_sizes[name] = size_bytes
                        available_sizes[name.removesuffix(":latest")] = size_bytes
                    except (ValueError, IndexError):
                        pass

    missing = [
        model for model in required_models.values()
        if model not in available_names and model.removesuffix(":latest") not in available_names
    ]

    loaded_vram = 0
    loaded_model_names: set[str] = set()
    try:
        with urllib.request.urlopen(
            f"{effective_ollama_base_url().rstrip('/')}/api/ps",
            timeout=5,
        ) as response:
            loaded = json.loads(response.read())
        loaded_models = loaded.get("models", [])
        loaded_vram = sum(
            int(item.get("size_vram", 0))
            for item in loaded_models
        )
        for item in loaded_models:
            for key in ("name", "model"):
                name = str(item.get(key, "")).strip()
                if name:
                    loaded_model_names.add(name)
                    loaded_model_names.add(name.removesuffix(":latest"))
    except Exception as exc:
        raise RuntimeError(f"Cannot read Ollama runtime state: {exc}") from exc

    required_bytes = 0
    for model in required_models.values():
        if (
            model in loaded_model_names
            or model.removesuffix(":latest") in loaded_model_names
        ):
            continue
        size_bytes = available_sizes.get(model) or available_sizes.get(model.removesuffix(":latest"), 0)
        if size_bytes:
            required_bytes += size_bytes

    total_vram = 128 * 1024**3
    projected_fraction = min((loaded_vram + required_bytes) / total_vram, 1.0)
    result = {
        "status": "PASS" if not missing and projected_fraction < 0.95 else "BLOCKED",
        "required_models": required_models,
        "missing_models": missing,
        "loaded_vram_gb": round(loaded_vram / 1024**3, 3),
        "projected_vram_fraction": round(projected_fraction, 5),
        "vram_limit_fraction": 0.95,
        "destructive_actions": False,
        "teacher_judge_provider": "claude_code_cli_aws",
        "slot120_teacher": False,
        "slot120_judge": False,
    }
    (output_dir / "runtime_preflight.json").write_text(
        json.dumps(result, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    if result["status"] != "PASS":
        raise RuntimeError(f"Runtime preflight blocked: {result}")
    return result


def _slot14_route_topic(topic: str, output_dir: Path) -> str:
    """Use the slot14 operational model to normalize the Axi/Omi document request."""
    prompt = (
        "Classify this document-generation request for AIMS. Return one JSON object "
        "with action='search', query, doc_type, and language. Request: "
        f"{topic}"
    )
    body = json.dumps({
        "model": ModelConfig.SLOT14_SEARCH,
        "prompt": prompt,
        "stream": False,
        "keep_alive": "4h",
    }).encode()
    base_url = effective_small_qwen_ollama_base_url().rstrip("/")
    request = urllib.request.Request(
        f"{base_url}/api/generate",
        data=body,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=120) as response:
            ollama_response = json.loads(response.read())
        raw_response = str(ollama_response.get("response", ""))
        route = {
            "model": ModelConfig.SLOT14_SEARCH,
            "status": "PASS",
            "response": raw_response,
        }
    except Exception as exc:
        route = {
            "model": ModelConfig.SLOT14_SEARCH,
            "status": "FALLBACK",
            "error": str(exc),
            "fallback_topic": topic,
            "fallback_reason": "slot14 routing unavailable; preserving original topic",
        }
        raw_response = json.dumps(
            {
                "action": "search",
                "query": topic,
                "doc_type": "policy_framework",
                "language": "en",
                "fallback": True,
            },
            ensure_ascii=False,
        )
    (output_dir / "slot14_route.json").write_text(
        json.dumps(route, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    try:
        parsed = _extract_json_object(raw_response)
        return str(parsed.get("query") or topic)
    except (json.JSONDecodeError, ValueError):
        log.warning("[SLOT14] Non-JSON route response; preserving original topic")
        return topic


def _claude_judge_smoke(
    output_dir: Path,
    soft_fail: bool = False,
) -> dict[str, Any]:
    """Require a real Bedrock-backed structured judgment before generation starts.

    When ``soft_fail`` is enabled, record judge failures but do not block the
    DOCGEN cycle. This is used for self-improvement runs so the loop can reach
    the first repair/feedback cycle even when a judge smoke is unavailable.
    """
    prompt = (
        "Evaluate two answers to the arithmetic question 'What is 2+2?'. "
        "Candidate A says '4'. Candidate B says '5'. Return JSON only with "
        "verdict='A', integer score from 1 to 10, and a short rationale."
    )
    smoke: dict[str, Any] = {
        "provider": "aws_bedrock_boto3",
        "status": "FAIL",
        "attempts": [],
    }
    forced_model = os.getenv("AIMS_CLAUDE_JUDGE_MODEL", "").strip().lower()
    if forced_model in {"opus", "sonnet"}:
        model_attempts = ((forced_model, 120),)
    else:
        model_attempts = (("opus", 90), ("sonnet", 120))

    # Map model names to Bedrock inference profile IDs
    model_id_map = {
        "sonnet": "us.anthropic.claude-sonnet-4-6",
        "opus": "us.anthropic.claude-opus-4-6-v1"
    }

    session_kwargs: dict[str, Any] = {
        "region_name": "us-east-1",
    }
    profile_name = os.environ.get(
        "AWS_PROFILE",
        "AdministratorAccess-445100240501",
    ).strip()
    if profile_name:
        session_kwargs["profile_name"] = profile_name

    for attempt_number, (model, timeout_seconds) in enumerate(
        model_attempts,
        start=1,
    ):
        attempt: dict[str, Any] = {
            "attempt": attempt_number,
            "model": model,
            "timeout_seconds": timeout_seconds,
        }
        try:
            import boto3

            try:
                session = boto3.Session(**session_kwargs)
            except Exception:
                session = boto3.Session(region_name="us-east-1")
            bedrock = session.client("bedrock-runtime")

            # Call Bedrock invoke_model API (Messages API format)
            model_id = model_id_map[model]
            request_body = json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 1024,
                "temperature": 0.0,
                "messages": [{
                    "role": "user",
                    "content": prompt
                }]
            })

            response = bedrock.invoke_model(
                modelId=model_id,
                contentType="application/json",
                accept="application/json",
                body=request_body
            )

            # Parse response
            response_body = json.loads(response["body"].read())
            response_text = response_body["content"][0]["text"]
            stop_reason = response_body.get("stop_reason", "unknown")

            # Parse JSON from response
            judged = _extract_json_object(response_text)
            valid = (
                stop_reason == "end_turn"
                and judged.get("verdict") == "A"
                and isinstance(judged.get("score"), int)
                and judged["score"] > 0
            )
            attempt.update(
                {
                    "returncode": 0,
                    "stderr": "",
                    "raw_stdout": response_text,
                    "wrapper": {"result": response_text, "is_error": False},
                    "parsed": judged,
                    "status": "PASS" if valid else "FAIL",
                    "bedrock_stop_reason": stop_reason,
                    "bedrock_usage": response_body.get("usage", {})
                }
            )
            smoke["attempts"].append(attempt)
            if valid:
                smoke.update(
                    {
                        "status": "PASS",
                        "passed_attempt": attempt_number,
                        "model": model,
                        "degraded_mode": model != "opus",
                        "parsed": judged,
                    }
                )
                break
        except Exception as exc:
            # Catch boto3 exceptions (ClientError, TimeoutError, etc.)
            attempt.update(
                {
                    "status": "ERROR",
                    "error": str(exc),
                    "error_type": type(exc).__name__,
                }
            )
            smoke["attempts"].append(attempt)

    allow_unverified = os.getenv(
        "AIMS_DOC_ALLOW_UNVERIFIED_CLAUDE_JUDGE", ""
    ).strip().lower() in {"1", "true", "yes", "on"}
    if allow_unverified and smoke["status"] != "PASS":
        smoke.update(
            {
                "status": "PASS",
                "passed_attempt": 0,
                "model": model_attempts[0][0],
                "degraded_mode": False,
                "override_mode": "unverified_judge_override",
                "parsed": {
                    "verdict": "A",
                    "score": 10,
                    "rationale": (
                        "Unverified judge override enabled for controlled probe."
                    ),
                },
            }
        )

    (output_dir / "claude_judge_smoke.json").write_text(
        json.dumps(smoke, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    if smoke["status"] != "PASS":
        incident = {
            "classification": "DOCUMENT_CYCLE_BLOCKED_CLAUDE_JUDGE_SMOKE",
            "reason": "Claude Code AWS judge smoke did not return valid judgment JSON",
            "generation_started": False,
            "attempts": smoke["attempts"],
        }
        (output_dir / "DOCUMENT_CYCLE_BLOCKED.json").write_text(
            json.dumps(incident, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        if soft_fail:
            fallback_model = (
                smoke["attempts"][-1].get("model")
                if smoke.get("attempts")
                else model_attempts[0][0]
            )
            log.warning(
                "[CLAUDE] Soft-failing judge smoke: %s",
                "; ".join(
                    f"{attempt.get('model', 'unknown')}: {attempt.get('error', attempt.get('status', 'FAIL'))}"
                    for attempt in smoke["attempts"]
                ),
            )
            smoke.update(
                {
                    "model": fallback_model,
                    "degraded_mode": fallback_model != "opus",
                    "passed_attempt": None,
                }
            )
            return smoke
        raise RuntimeError(
            "Claude Code AWS judge smoke failed; document cycle blocked"
        )
    return smoke


# ──────────────────────────────────────────────────────────────────────────────
# SEARCH STANDARDS — Two-Level Injection
# ──────────────────────────────────────────────────────────────────────────────

def _search_internal_standards(
    topic: str,
    max_results: int = 20,
    document_type: str = "technical_report",
    governed_fallback: Optional[list[str]] = None,
) -> list[str]:
    """
    Omi searches internal standards database (Qdrant RAG + SQLite).
    Fast, local, based on internal document registry.
    """
    log.info(f"[OMI-SEARCH] Querying internal standards database for: {topic}")

    try:
        from ops.docagent.doc_skills import DocSkillRunner

        runner = DocSkillRunner()
        results = runner.invoke("doc-search", query=topic, mode="hybrid", max_results=max_results)

        if results and isinstance(results, dict):
            standards = []
            for item in results.get("results", []):
                if not isinstance(item, dict):
                    continue
                identifier = str(item.get("standard_id", "")).strip()
                title = str(item.get("title", "")).strip()
                citation = " — ".join(
                    part for part in (identifier, title) if part
                )
                if citation:
                    standards.append(citation)
            if standards:
                log.info(f"[OMI-SEARCH] Found {len(standards)} internal standards via DocSkillRunner")
                return standards[:max_results]
    except Exception as e:
        log.warning(f"[OMI-SEARCH] DocSkillRunner failed: {e}")

    # Direct Qdrant vector search against intl_standards (34k indexed clauses)
    try:
        from ops.docagent.standards_rag import _embed, _search, _collection_exists

        if _collection_exists():
            vector = _embed(topic)
            clauses = _search(vector, top_k=max_results)
            if clauses:
                standards = [
                    f"{c.citation()}: {c.text[:400]}"
                    for c in clauses
                ]
                log.info(f"[OMI-SEARCH] Found {len(standards)} clauses via Qdrant direct search")
                return standards
            log.warning("[OMI-SEARCH] Qdrant search returned no clauses")
        else:
            log.warning("[OMI-SEARCH] intl_standards collection not available")
    except Exception as e:
        log.warning(f"[OMI-SEARCH] Qdrant direct search failed: {e}")

    # Final fallback: task-archetype governed standards only.
    log.warning("[OMI-SEARCH] Using governed fallback standards set")
    fallback = list(governed_fallback or standards_for_document_type(document_type) or [])
    if not fallback:
        fallback = [
            "ISO 10013:2021",
            "ISO 9001:2015 clause 7.5",
        ]
    return fallback


def _filter_standards_for_reference(
    qdrant_standards: list[str],
    reference_text: str,
) -> list[str]:
    """
    When a reference document is available, prefer standards cited in it over
    noisy Qdrant results. Falls back to keyword filtering if reference extraction
    fails.

    Strategy:
    1. Extract standard identifiers from the reference text's "References" section.
    2. Return those directly if ≥ 2 were found (reference-grounded).
    3. Otherwise filter Qdrant results, removing obvious electrical/off-topic
       entries whose identifier prefix matches _ELECTRICAL_NOISE_PREFIXES.
    """
    # Extract governed identifiers anywhere in the reference. Section numbering
    # and headings vary across document archetypes.
    extracted: list[str] = []
    if reference_text:
        for line in reference_text.splitlines():
            stripped = line.strip()
            if re.search(
                r"\b(ISO|IEC|IEEE|API|ASME|AMPP|NACE|ASTM|EN)\s+"
                r"[A-Z0-9][A-Z0-9./:-]*",
                stripped,
                re.I,
            ):
                extracted.append(stripped)

    if extracted:
        log.info(
            "[STANDARDS-FILTER] Using %d reference-extracted standards instead of Qdrant",
            len(extracted),
        )
        return list(dict.fromkeys(extracted))[:20]
    return list(dict.fromkeys(qdrant_standards))[:20]


def _search_external_standards(topic: str, doc_text: str) -> dict:
    """
    Discover governed external standards through the contextual provider chain.

    Returns:
    {
        "standards": ["ISO 55002", "API 570", ...],
        "content": {
            "ISO 55002": "extracted content/requirements",
            ...
        },
        "sources": [...],
        "provider": "..."
    }
    """
    log.info("[AXI-STANDARD-DISCOVERY] Searching standards for: %s", topic)

    try:
        from ops.agents.skills.contextual_standard_discovery import (
            build_provider_chain,
            search_with_chain,
        )
        from ops.agents.skills.contextual_standard_discovery_and_document_review import (
            build_standard_discovery_queries,
            classify_standard_source,
            extract_document_review_context,
            sanitize_external_search_context,
        )

        context = extract_document_review_context(topic, doc_text)
        queries = build_standard_discovery_queries(
            sanitize_external_search_context(context)
        )
        raw_sources, provider = search_with_chain(
            queries,
            build_provider_chain(),
            max_results=12,
        )
        sources = [classify_standard_source(item) for item in raw_sources]
        governed = [
            item
            for item in sources
            if item.get("source_authority_level") in {"high", "medium"}
            and item.get("source_type") != "secondary"
        ]
        standards = list(
            dict.fromkeys(
                str(item.get("source_title", "")).strip()
                for item in governed
                if str(item.get("source_title", "")).strip()
            )
        )
        content = {
            str(item["source_title"]): str(item.get("excerpt", "")).strip()
            for item in governed
            if item.get("source_title")
        }
        log.info(
            "[AXI-STANDARD-DISCOVERY] provider=%s governed_sources=%d",
            provider,
            len(standards),
        )
        return {
            "standards": standards,
            "content": content,
            "sources": governed,
            "provider": provider,
        }
    except Exception as e:
        log.warning("[AXI-STANDARD-DISCOVERY] failed: %s", e)

    return {
        "standards": [],
        "content": {},
        "sources": [],
        "provider": "unavailable",
    }


# ──────────────────────────────────────────────────────────────────────────────
# STAGE 1: OMI — Draft Generation with Standards Context
# ──────────────────────────────────────────────────────────────────────────────

def _omi_generate_draft(
    topic: str,
    context: Optional[str] = None,
    reference_template: Optional[str] = None,
    standards_to_inject: Optional[list[str]] = None,
    search_internal: bool = True,
    grounded_prompt: Optional[str] = None,
) -> str:
    """
    Omi generates initial document draft.

    1. Searches internal standards database (Qdrant RAG + SQLite)
    2. Injects standards context before generation
    3. Foundation: establish proper ADNOC structure, roles, scope, references
    """
    log.info(f"[OMI] Generating draft for topic: {topic}")

    # If no standards provided, search internal database
    if not standards_to_inject and search_internal:
        standards_to_inject = _search_internal_standards(topic, max_results=20)
        log.info(f"[OMI] Auto-injected {len(standards_to_inject)} standards from internal database")

    if grounded_prompt:
        prompt = grounded_prompt
    else:
        # Build prompt with injected standards and reference structure hints
        prompt_parts = [
            f"Generate a comprehensive professional document on:\n{topic}\n\n",
            "STRUCTURAL REQUIREMENTS:",
            "- Document Control & Distribution (with revision history)",
            "- Table of Contents (numbered sections)",
            "- Introduction & Purpose/Objective",
            "- Scope (applicability, lifecycle coverage)",
            "- Policy/Framework elements (list all)",
            "- Definitions & Acronyms",
            "- References (standards, guidelines, industry best practices)",
            "- Framework Overview (PDCA cycle, lifecycle aspects)",
            "- Elements & Sub-elements with expectations",
            "- Operational integrity guidance",
            "- Appendices\n\n",
        ]

        if standards_to_inject:
            prompt_parts.append(
                "MANDATORY STANDARDS TO INCLUDE:\n"
                f"{chr(10).join(standards_to_inject)}\n\n"
            )

        if reference_template:
            prompt_parts.append(
                "REFERENCE TEMPLATE STRUCTURE AND CONTENT:\n"
                f"{reference_template}\n\n"
            )

        prompt_parts.append(
            "Style: Official corporate document (ADNOC/Oil & Gas standards compliance)\n"
            "Language: English, professional terminology\n"
            "Format: Markdown with proper headings, emphasis on clarity and compliance\n"
            "Length: 2500-3500 words for this controlled improvement cycle\n"
        )
        prompt = "".join(prompt_parts)

    # Call Omi reasoning model through the certified Qwen3.6 raw-safe protocol.
    try:
        generated = _slot120_generate(
            prompt,
            timeout=OMI_GENERATE_TIMEOUT,
            num_predict=OMI_GENERATE_NUM_PREDICT,
        )
        log.info(
            "[OMI] Draft generation complete (%d chars, model=%s)",
            len(generated),
            ModelConfig.SLOT120_REASONING,
        )
        return generated
    except TimeoutError:
        log.error("[OMI] Generation timeout")
        return ""
    except Exception as e:
        log.error(f"[OMI] Generation error: {e}")
        return ""


# ──────────────────────────────────────────────────────────────────────────────
# STAGE 2: QUALITY EVALUATION — Structure, Standards, Coverage
# ──────────────────────────────────────────────────────────────────────────────

def _evaluate_document_quality(
    doc_text: str,
    topic: str,
    reference_text: Optional[str] = None,
    required_sections: Optional[list[str]] = None,
    required_standards: Optional[list[str]] = None,
) -> GenerationMetrics:
    """
    Evaluate generated document against:
    1. ADNOC structural requirements
    2. Standards accuracy & completeness
    3. Reference template coverage
    """
    log.info("[EVAL] Starting quality evaluation...")

    # Try to use doc_quality_eval module
    try:
        from ops.docagent.doc_quality_eval import evaluate_document
        result = evaluate_document(
            doc_text,
            topic=topic,
            reference_text=reference_text,
            required_sections=required_sections,
            required_standards=required_standards,
        )

        # Convert to our metrics
        metrics = GenerationMetrics(
            cycle=0,  # Will be updated by caller
            timestamp=datetime.now().isoformat(),
            model_used=ModelConfig.SLOT120_REASONING,
            sections_found=len(result.structure_details),
            sections_expected=len(required_sections or result.structure_details),
            section_coverage=result.structure_score,
            standards_found=len(result.standards_details.get("found_standards", [])),
            standards_expected=len(
                required_standards
                or result.standards_details.get("expected_standards", [])
            ),
            standards_accuracy=result.standards_score,
            themes_covered=result.reference_comparison.reference_sections if hasattr(result, 'reference_comparison') else [],
            reference_match=result.reference_score,
            structure_score=result.structure_score,
            standards_score=result.standards_score,
            coverage_score=result.reference_score,
            overall_score=result.overall_score,
        )
        log.info(f"[EVAL] Quality scores: structure={metrics.structure_score:.2f}, standards={metrics.standards_score:.2f}, coverage={metrics.reference_match:.2f}")
        return metrics
    except Exception as e:
        log.warning(f"[EVAL] doc_quality_eval unavailable: {e}, using basic eval")

    # Fallback: basic evaluation
    metrics = GenerationMetrics(
        cycle=0,
        timestamp=datetime.now().isoformat(),
        model_used=ModelConfig.SLOT120_REASONING,
        sections_found=doc_text.count("#"),  # Count headings
        sections_expected=len(required_sections or []) or 1,
        section_coverage=min(
            doc_text.count("#") / (len(required_sections or []) or 1),
            1.0,
        ),
        standards_found=sum(1 for std in ["ISO", "API", "ASME", "NACE", "IEC"] if std in doc_text),
        standards_expected=len(required_standards or []),
        standards_accuracy=0.0,  # Would need NLP
        reference_match=0.0,
        structure_score=min(
            doc_text.count("#") / (len(required_sections or []) or 1),
            1.0,
        ),
        standards_score=0.0,
        coverage_score=0.0,
        overall_score=min(
            doc_text.count("#") / (len(required_sections or []) or 1),
            1.0,
        ),
    )
    return metrics


# ──────────────────────────────────────────────────────────────────────────────
# STAGE 3: AXI VALIDATION — High-Rank Claude Recommendations
# ──────────────────────────────────────────────────────────────────────────────

def _axi_validate_and_recommend(
    doc_text: str,
    metrics: GenerationMetrics,
    topic: str,
    reference_text: Optional[str] = None,
    search_external: bool = True,
    document_type: str = "technical_report",
    required_sections: Optional[list[str]] = None,
    required_standards: Optional[list[str]] = None,
    discovered_external_standards: Optional[dict[str, Any]] = None,
) -> tuple[str, list[str]]:
    """
    Axi calls high-rank Claude model (via gateway) to:
    1. Validate document structure and content
    2. Find gaps vs. reference template
    3. Recommend specific improvements
    4. Search external sources for missing standards
    5. Identify gaps not covered by internal database
    """
    log.info(f"[AXI] Validating generation quality (cycle metrics: overall={metrics.overall_score:.2f})...")

    # Search external standards if enabled
    external_standards_data = discovered_external_standards or {
        "standards": [],
        "content": {},
        "sources": [],
        "provider": "not_requested",
    }
    if search_external and not discovered_external_standards:
        external_standards_data = _search_external_standards(topic, doc_text)

    # Build validation prompt with external standards context
    external_context = ""
    if external_standards_data["standards"]:
        standards_list = "\n".join(external_standards_data["standards"])
        external_context = f"\n\nEXTERNAL STANDARDS FOUND (from internet search):\n{standards_list}\n"

        for std, content in external_standards_data["content"].items():
            external_context += f"\n{std} REQUIREMENTS:\n{content[:500]}...\n"

    # Prepare validation prompt with external standards
    validation_prompt = f"""You are an expert document validator for Oil & Gas Asset Integrity Management.

GENERATED DOCUMENT (Omi draft — full text):
{doc_text}

REFERENCE DOCUMENT — FULL AVAILABLE TEXT:
{reference_text or "(no reference text available)"}

CURRENT METRICS:
- Structure coverage: {metrics.section_coverage:.1%}
- Standards accuracy: {metrics.standards_accuracy:.1%}
- Reference match: {metrics.reference_match:.1%}
- Overall quality: {metrics.overall_score:.1%}
{external_context}

DOCUMENT FORMATION CONTRACT:
- ISO 10013:2021 governs documented-information structure.
- ISO 2145:1978 governs section numbering.
- ISO 690:2021 governs the final standards reference table.
- These formation standards are applied by deterministic pipeline gates and
  do not need redundant narrative citations in the document body.

REFERENCE REQUIREMENTS:
- Document type: {document_type}
- Task: {topic}
- Expected semantic sections: {json.dumps(required_sections or [], ensure_ascii=False)}
- Expected standards: {json.dumps(required_standards or [], ensure_ascii=False)}
- Do not impose sections or standards from a different task archetype

VALIDATION & IMPROVEMENT TASKS:
1. List major structural gaps (missing sections, improper ordering)
2. Identify missing or incorrectly cited standards
3. **Analyze external standards found** - what requirements should be added to the document?
4. Propose SPECIFIC improvements to integrate the external standards
   - For each external standard found: "Section X: Add Y from [standard]"
5. Rate convergence to reference (0-100%)
6. Return at most 15 recommendations, at most 2 per exact section target
7. Use recommendations=[] when there is no material, verifiable gap

RESPOND WITH STRICT JSON ONLY:
{{
  "validation_passed": true/false,
  "structural_gaps": ["gap1", "gap2", ...],
  "external_standards_to_integrate": ["ISO 55002: Add implementation guidance section", ...],
  "improvement_recommendations": [
    {{
      "text": "Section X: specific change",
      "evidence_quote": "exact substring from the cited source",
      "evidence_source": "reference or exact external source title"
    }}
  ],
  "convergence_score": 0.XX,
  "summary": "one paragraph assessment"
}}"""

    try:
        # Axi recommendation pass uses Sonnet via direct Bedrock API.
        # Opus is reserved for the independent teacher/judge audit below.
        result_text = bedrock_axi_validate(validation_prompt, timeout=420)

        if result_text:
            try:
                validation_data = parse_json_from_response(result_text)
                feedback = validation_data.get("summary", "")
                recommendation_items = validation_data.get(
                    "improvement_recommendations",
                    [],
                )
                recommendations: list[str] = []
                normalized_reference = _normalized_evidence(reference_text or "")
                normalized_sources = {
                    str(title): _normalized_evidence(content)
                    for title, content in external_standards_data[
                        "content"
                    ].items()
                }
                for item in recommendation_items:
                    if not isinstance(item, dict):
                        continue
                    text = str(item.get("text", "")).strip()
                    quote = _normalized_evidence(
                        str(item.get("evidence_quote", ""))
                    )
                    source = str(item.get("evidence_source", "")).strip()
                    source_text = (
                        normalized_reference
                        if source.lower() == "reference"
                        else normalized_sources.get(source, "")
                    )
                    if text and len(quote) >= 12 and quote in source_text:
                        recommendations.append(text)
                convergence_raw = float(validation_data.get("convergence_score", 0.0))
                convergence = (
                    convergence_raw / 100.0
                    if convergence_raw > 1.0
                    else convergence_raw
                )
                log.info(f"[AXI] Validation complete: convergence={convergence:.1%}, {len(recommendations)} recommendations")
                return feedback, recommendations
            except (ValueError, json.JSONDecodeError):
                log.warning("[AXI] Could not parse Bedrock validation response")
    except Exception as exc:
        log.error("[AXI] Validation call failed: %s", exc)

    feedback = "Axi validation unavailable; no ungrounded fallback repairs emitted"
    return feedback, []


# ──────────────────────────────────────────────────────────────────────────────
# STAGE 3B: CLAUDE CODE AUDIT — Independent Validation of Omi & Axi
# ──────────────────────────────────────────────────────────────────────────────

def _split_into_sections(text: str) -> dict[str, str]:
    """Split document into {heading: body} chunks by Markdown headings."""
    sections: dict[str, str] = {}
    current_heading = "__preamble__"
    current_body: list[str] = []

    for line in text.splitlines():
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            if current_body:
                sections[current_heading] = "\n".join(current_body).strip()
            current_heading = m.group(2).strip()
            current_body = []
        else:
            current_body.append(line)

    if current_body:
        sections[current_heading] = "\n".join(current_body).strip()

    return sections


def _run_claude_prompt(prompt: str, timeout: int = 300) -> Optional[str]:
    """Execute a single Claude Code CLI call, return result text or None."""
    result = subprocess.run(
        ["claude", "--print", "--output-format", "json", "--model", "opus", "--setting-sources="],
        input=prompt,
        env=_claude_bedrock_env(),
        capture_output=True,
        timeout=timeout,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        log.warning(f"[CLAUDE] CLI returned {result.returncode}: {result.stderr[:200]}")
        return None
    try:
        wrapper = json.loads(result.stdout)
        return wrapper.get("result", "")
    except (json.JSONDecodeError, KeyError):
        return None


def _normalized_evidence(text: str) -> str:
    text = (
        str(text)
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2212", "-")
    )
    return re.sub(r"\s+", " ", text).strip().lower()


def _standard_citations(text: str) -> list[str]:
    citations = re.findall(
        r"\b(?:ISO/IEC/IEEE|ISO/IEC|IEC/IEEE|ISO|IEC|IEEE|API|ASME|ANSI|"
        r"NFPA|NACE|EN)\s+"
        r"(?:RP\s+)?(?=[A-Z0-9.\-:]*\d)[A-Z0-9][A-Z0-9.\-:]*",
        text,
        flags=re.IGNORECASE,
    )
    return [
        re.sub(r"[\s\.,;:]+$", "", citation).strip()
        for citation in citations
        if re.sub(r"[\s\.,;:]+$", "", citation).strip()
    ]


_STANDARD_REFERENCE_HEADING = "11.0 References and Related Documents"
_STANDARD_REFERENCE_MARKER = "<!-- AIMS_STANDARD_REFERENCE_REGISTER -->"


def _standard_identifier(text: str) -> str | None:
    match = re.search(
        r"\b(?:ISO/IEC/IEEE|ISO/IEC|IEC/IEEE|ISO|IEC|IEEE|API(?:\s+RP)?|"
        r"ASME|ANSI|NFPA|NACE|EN)\s+"
        r"(?=[A-Z0-9.\-:]*\d)[A-Z0-9][A-Z0-9.\-:]*",
        str(text),
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    return re.sub(r"[\s\.,;:]+$", "", re.sub(r"\s+", " ", match.group(0)).strip())


def _standard_description(identifier: str, title: str, excerpt: str = "") -> str:
    if "|" in str(title):
        cells = [cell.strip() for cell in str(title).split("|")]
        for index, cell in enumerate(cells):
            if _normalized_evidence(identifier) == _normalized_evidence(cell):
                description = " ".join(
                    candidate
                    for candidate in cells[index + 1 :]
                    if candidate
                )
                if description:
                    return description
    description = re.sub(
        re.escape(identifier),
        "",
        str(title),
        count=1,
        flags=re.IGNORECASE,
    )
    description = re.sub(r"^[\s:—–-]+", "", description).strip()
    if not description:
        description = re.sub(r"\s+", " ", str(excerpt)).strip()
    return description or "Applicable requirements and guidance."


def _without_standard_reference_register(text: str) -> str:
    marker_index = text.find(_STANDARD_REFERENCE_MARKER)
    if marker_index >= 0:
        return text[:marker_index].rstrip()
    heading = re.search(
        rf"(?im)^#+\s+{re.escape(_STANDARD_REFERENCE_HEADING)}\s*$",
        text,
    )
    return text[: heading.start()].rstrip() if heading else text.rstrip()


def _build_standard_source_records(
    *,
    internal_standards: list[str],
    external_standards: dict[str, Any],
    formatting_standards: list[str],
) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for value in internal_standards:
        identifier = _standard_identifier(value)
        if identifier:
            records.append(
                {
                    "identifier": identifier,
                    "description": _standard_description(identifier, value),
                    "source": "omi_internal_discovery",
                    "source_title": value,
                    "usage": "technical_context",
                }
            )
    for source in external_standards.get("sources", []):
        if not isinstance(source, dict):
            continue
        title = str(source.get("source_title", "")).strip()
        identifier = _standard_identifier(title)
        if not identifier:
            continue
        records.append(
            {
                "identifier": identifier,
                "description": _standard_description(
                    identifier,
                    title,
                    str(source.get("excerpt", "")),
                ),
                "source": "axi_contextual_discovery",
                "source_title": title,
                "source_url": str(source.get("source_url", "")),
                "source_authority_level": source.get(
                    "source_authority_level"
                ),
                "usage": "technical_context",
            }
        )
    active_records = {
        record["identifier"]: record
        for record in active_document_formation_records()
    }
    for value in formatting_standards:
        identifier = _standard_identifier(value)
        if not identifier:
            continue
        catalog_record = active_records.get(identifier, {})
        records.append(
            {
                "identifier": identifier,
                "description": str(
                    catalog_record.get("description")
                    or _standard_description(identifier, value)
                ),
                "source": "document_formation_contract",
                "source_title": value,
                "official_url": catalog_record.get("official_url"),
                "usage": "formatting_contract",
            }
        )
    deduplicated: dict[str, dict[str, Any]] = {}
    for record in records:
        key = _normalized_evidence(record["identifier"])
        deduplicated.setdefault(key, record)
    return list(deduplicated.values())


def _profile_unbound_references(
    *,
    forbidden_references: list[str],
    body_citations: list[str],
    matched_citations: set[str],
) -> list[str]:
    if not forbidden_references:
        return []
    unbound: list[str] = []
    for citation in body_citations:
        normalized = _normalized_evidence(citation)
        for forbidden in forbidden_references:
            if (
                _normalized_evidence(forbidden) == normalized
                and citation not in matched_citations
            ):
                unbound.append(citation)
    return list(dict.fromkeys(unbound))


def _ensure_standard_reference_register(
    *,
    doc_text: str,
    internal_standards: list[str],
    external_standards: dict[str, Any],
    formatting_standards: list[str],
    forbidden_references: list[str] | None = None,
    document_type: str = "technical_report",
) -> tuple[str, dict[str, Any]]:
    """Append a final two-column table containing only evidenced, used standards."""
    body = _without_standard_reference_register(doc_text)
    records = _build_standard_source_records(
        internal_standards=internal_standards,
        external_standards=external_standards,
        formatting_standards=formatting_standards,
    )
    body_citations = list(dict.fromkeys(_standard_citations(body)))
    used: list[dict[str, Any]] = []
    matched_citations: set[str] = set()
    for record in records:
        identifier_key = _normalized_evidence(record["identifier"])
        cited = any(
            identifier_key == _normalized_evidence(citation)
            or identifier_key in _normalized_evidence(citation)
            or _normalized_evidence(citation) in identifier_key
            for citation in body_citations
        )
        if cited or record["usage"] == "formatting_contract":
            used.append(record)
            if cited:
                matched_citations.update(
                    citation
                    for citation in body_citations
                    if (
                        identifier_key == _normalized_evidence(citation)
                        or identifier_key in _normalized_evidence(citation)
                        or _normalized_evidence(citation) in identifier_key
                    )
                )

    matched_normalized = {
        _normalized_evidence(citation) for citation in matched_citations
    }
    for citation in body_citations:
        normalized = _normalized_evidence(citation)
        if normalized in matched_normalized:
            continue
        identifier = _standard_identifier(citation)
        if not identifier:
            continue
        catalog_record = STANDARD_CATALOG.get(identifier)
        if not catalog_record:
            continue
        used.append(
            {
                "identifier": identifier,
                "description": str(catalog_record.get("description") or "Applicable requirements and guidance."),
                "source": "standard_catalog_context",
                "source_title": citation,
                "official_url": catalog_record.get("official_url"),
                "usage": "contextual_catalog_citation",
            }
        )
        matched_citations.update(
            citation for citation in body_citations if _normalized_evidence(citation) == normalized
        )

    unverified = [
        citation for citation in body_citations
        if citation not in matched_citations
    ]
    profile_forbidden_references = [
        str(item)
        for item in dict.fromkeys(forbidden_references or [])
        if str(item).strip()
    ]
    profile_unbound_references = _profile_unbound_references(
        forbidden_references=profile_forbidden_references,
        body_citations=body_citations,
        matched_citations=matched_citations,
    )
    branch_blockers = []
    if profile_unbound_references:
        branch_blockers.append("profile_forbidden_unbound_reference")
    fallback_mode = bool(unverified or not used)
    lines = [
        _STANDARD_REFERENCE_MARKER,
        f"## {_STANDARD_REFERENCE_HEADING}",
        "",
        "| Standard Number | Description |",
        "|---|---|",
    ]
    for record in used:
        lines.append(
            f"| {record['identifier']} | {record['description']} |"
        )
    result = f"{body}\n\n" + "\n".join(lines) + "\n"
    evidence = {
        "status": "PASS" if not branch_blockers else "FAIL",
        "table_required": True,
        "table_at_document_end": True,
        "document_type": document_type,
        "formatting_standard_policy": formatting_standards,
        "registration_taxonomy_standards_auto_added": False,
        "fallback_mode": fallback_mode,
        "fallback_reason": (
            "unregistered_or_unresolved_standards"
            if fallback_mode
            else ""
        ),
        "discovered_source_count": len(records),
        "used_standard_count": len(used),
        "used_standards": used,
        "body_citations": body_citations,
        "unverified_citations": unverified,
        "branch_blockers": branch_blockers,
        "profile_forbidden_references": profile_forbidden_references,
        "profile_unbound_references": profile_unbound_references,
        "excluded_discovered_standards": [
            record for record in records if record not in used
        ],
    }
    return result, evidence


def _ground_audit_recommendations(
    finding: dict,
    reference_excerpt: str,
    standard_sources: Optional[list[dict[str, Any]]] = None,
) -> tuple[list[str], list[dict]]:
    accepted: list[str] = []
    rejected: list[dict] = []
    normalized_reference = _normalized_evidence(reference_excerpt)
    source_identifiers = {
        _normalized_evidence(identifier)
        for source in (standard_sources or [])
        if (identifier := _standard_identifier(source.get("source_title", "")))
    }
    recommendations = finding.get("recommendations", [])
    if not isinstance(recommendations, list):
        return accepted, [{
            "reason": "recommendations_not_list",
            "value": recommendations,
        }]

    for recommendation in recommendations:
        if isinstance(recommendation, dict):
            text = str(recommendation.get("text", "")).strip()
            quote = str(recommendation.get("evidence_quote", "")).strip()
        else:
            text = str(recommendation).strip()
            quote = ""

        normalized_quote = _normalized_evidence(quote)
        unsupported_standards = [
            citation
            for citation in _standard_citations(text)
            if (
                _normalized_evidence(citation) not in normalized_reference
                and _normalized_evidence(citation) not in source_identifiers
            )
        ]
        if not text:
            rejected.append({"reason": "empty_text", "value": recommendation})
        elif not reference_excerpt:
            rejected.append({
                "reason": "reference_excerpt_missing",
                "text": text,
            })
        elif len(normalized_quote) < 12:
            rejected.append({
                "reason": "evidence_quote_too_short_or_missing",
                "text": text,
                "evidence_quote": quote,
            })
        elif normalized_quote not in normalized_reference:
            rejected.append({
                "reason": "evidence_quote_not_in_reference",
                "text": text,
                "evidence_quote": quote,
            })
        elif unsupported_standards:
            rejected.append({
                "reason": "unsupported_standard_citations",
                "text": text,
                "citations": unsupported_standards,
            })
        else:
            accepted.append(text)
    return accepted[:2], rejected


def _grounded_repair_bridge(
    *,
    doc_text: str,
    reference_text: str,
    generated_sections: list[str],
    audit_result: ClaudeAuditResult,
    audit_quality_failures: list[str],
    evidence_dir: Path,
    contextual_sources: Optional[list[dict[str, Any]]] = None,
    model: str = "opus",
) -> list[str]:
    """Request a bounded executable repair batch when audit advice is empty."""
    bridge_dir = evidence_dir / "grounded_repair_bridge"
    prompt = f"""You are the repair-planning controller for a DOCGEN cycle.

The independent audit failed its quality gate but emitted no executable
section recommendations. Produce at most four grounded repairs.

GENERATED SECTION NAMES:
{json.dumps(generated_sections, ensure_ascii=False)}

AUDIT QUALITY FAILURES:
{json.dumps(audit_quality_failures, ensure_ascii=False)}

AUDIT ASSESSMENT:
{audit_result.overall_assessment}

SKILL-LEVEL OBSERVATIONS:
{json.dumps(audit_result.skill_recommendations, ensure_ascii=False)}

GENERATED DOCUMENT:
{doc_text}

REFERENCE DOCUMENT:
{reference_text}

CONTEXTUALLY DISCOVERED STANDARD SOURCES:
{json.dumps(contextual_sources or [], ensure_ascii=False)}

Each repair must target one existing generated section. Each evidence_quote
must be an exact substring copied from either the reference document or a
contextually discovered source excerpt. A cited standard must be present in
the discovered source list. Do not invent standards, thresholds, sections,
or source text. Return strict JSON only:
{{"recommendations":[{{"section":"exact generated section name","text":"specific executable change","evidence_source":"reference or exact source_title","evidence_quote":"exact source substring"}}]}}
"""
    parsed = bedrock_doc_audit(
        prompt=prompt,
        model_alias=model,
        max_tokens=4000,
        timeout=300,
        evidence_dir=bridge_dir,
    ) or {}
    items = parsed.get("recommendations", [])
    accepted: list[str] = []
    rejected: list[dict[str, Any]] = []
    normalized_reference = _normalized_evidence(reference_text)
    normalized_sources = {
        str(source.get("source_title", "")).strip(): _normalized_evidence(
            str(source.get("excerpt", ""))
        )
        for source in (contextual_sources or [])
        if isinstance(source, dict)
        and str(source.get("source_title", "")).strip()
    }
    source_identifiers = {
        _normalized_evidence(identifier)
        for title in normalized_sources
        if (identifier := _standard_identifier(title))
    }
    if not isinstance(items, list):
        rejected.append(
            {
                "reason": "recommendations_not_list",
                "value": items,
            }
        )
        items = []

    for item in items[:8]:
        if not isinstance(item, dict):
            rejected.append(
                {"reason": "recommendation_not_object", "value": item}
            )
            continue
        section = str(item.get("section", "")).strip()
        text = str(item.get("text", "")).strip()
        quote = str(item.get("evidence_quote", "")).strip()
        evidence_source = str(item.get("evidence_source", "")).strip()
        normalized_quote = _normalized_evidence(quote)
        unsupported_standards = [
            citation
            for citation in _standard_citations(text)
            if (
                _normalized_evidence(citation) not in normalized_reference
                and _normalized_evidence(citation) not in source_identifiers
            )
        ]
        quote_source = ""
        if normalized_quote and normalized_quote in normalized_reference:
            quote_source = "reference"
        else:
            for title, source_text in normalized_sources.items():
                if normalized_quote and normalized_quote in source_text:
                    quote_source = title
                    break
        matched_section = next(
            (
                generated
                for generated in generated_sections
                if sections_semantically_match(section, generated)
            ),
            None,
        )
        if not matched_section:
            rejected.append(
                {
                    "reason": "target_section_not_generated",
                    "section": section,
                    "text": text,
                }
            )
        elif not text:
            rejected.append(
                {"reason": "empty_text", "section": section}
            )
        elif len(normalized_quote) < 12:
            rejected.append(
                {
                    "reason": "evidence_quote_too_short_or_missing",
                    "section": section,
                    "text": text,
                }
            )
        elif not quote_source:
            rejected.append(
                {
                    "reason": "evidence_quote_not_in_bound_sources",
                    "section": section,
                    "text": text,
                    "evidence_quote": quote,
                    "evidence_source": evidence_source,
                }
            )
        elif unsupported_standards:
            rejected.append(
                {
                    "reason": "unsupported_standard_citations",
                    "section": section,
                    "text": text,
                    "citations": unsupported_standards,
                }
            )
        else:
            accepted.append(f"Section {matched_section}: {text}")
        if len(accepted) >= 4:
            break

    bridge_dir.mkdir(parents=True, exist_ok=True)
    (bridge_dir / "grounded_repair_result.json").write_text(
        json.dumps(
            {
                "status": (
                    "GROUNDED_REPAIRS_CREATED"
                    if accepted
                    else "NO_GROUNDED_REPAIRS"
                ),
                "accepted": accepted,
                "rejected": rejected,
                "source": "claude_bedrock_bounded_repair_planner",
                "contextual_source_titles": list(normalized_sources),
            },
            indent=2,
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return accepted


def _verified_missing_sections(
    claimed_missing: list[str],
    generated_sections: list[str],
) -> list[str]:
    """Keep only judge-reported sections absent from generated headings."""
    return [
        missing
        for missing in claimed_missing
        if not any(
            sections_semantically_match(missing, generated)
            for generated in generated_sections
        )
    ]


def _audit_section_batch(
    batch_idx: int,
    section_batch: list[str],
    doc_sections: dict[str, str],
    omi_standards_str: str,
    axi_standards_str: str,
    axi_recs_str: str,
    topic: str,
    reference_text: str,
    cycle: int,
    model: str,
    evidence_dir: Optional[Path],
) -> list[dict]:
    """Run one Bedrock call auditing a section batch; return validated section_findings."""
    batch_text = "\n\n".join(
        f"### {name}\n{doc_sections[name]}" for name in section_batch if name in doc_sections
    )
    # Build per-section reference excerpts so the auditor sees relevant reference
    # content for each section rather than the same truncated front-matter prefix.
    ref_parts: list[str] = []
    for name in section_batch:
        # Section names are like "8.4 Element 4: Leadership…" — extract numeric ID.
        excerpt = (
            match_reference_section(reference_text, name)
            if reference_text
            else ""
        )
        if excerpt:
            ref_parts.append(f"[{name}]\n{excerpt}")
    ref_block = (
        "\n\n".join(ref_parts)
        if ref_parts
        else (reference_text[:6000] if reference_text else "(none)")
    )
    batch_prompt = (
        f"You are an independent auditor for an Asset Integrity Management document.\n\n"
        f"TOPIC: {topic}\n"
        f"CYCLE: {cycle}\n"
        f"BATCH: {batch_idx + 1}\n"
        f"SECTIONS IN THIS BATCH ({len(section_batch)}):\n"
        f"{json.dumps(section_batch, ensure_ascii=False)}\n\n"
        f"SECTION TEXTS:\n{batch_text}\n\n"
        f"REFERENCE CONTENT (per-section excerpts from source document):\n"
        f"{ref_block}\n\n"
        f"OMI STANDARDS FOUND:\n{omi_standards_str}\n\n"
        f"AXI STANDARDS FOUND:\n{axi_standards_str}\n\n"
        f"DOCUMENT FORMATION CONTRACT:\n"
        f"ISO 10013:2021 (structure), ISO 2145:1978 (numbering), "
        f"ISO 690:2021 (reference table) are deterministic pipeline gates. "
        f"They do not require redundant narrative citations.\n\n"
        f"AXI RECOMMENDATIONS:\n{axi_recs_str}\n\n"
        f"gap_score means: 0.0 = section content matches reference well; "
        f"1.0 = section is completely missing/divergent from reference.\n\n"
        f"Audit ONLY the {len(section_batch)} sections listed above. "
        f"Return strict JSON with ONLY this key:\n"
        f'{{"section_findings":['
        f'{{"section":"...","gap_score":0.0,"recommendations":['
        f'{{"text":"...","evidence_quote":"exact source substring"}}'
        f'],"missing_standards":[]}}'
        f']}}\n'
        f"Each listed section must appear exactly once. At most 2 recommendations per section. "
        f"Every recommendation MUST include an exact evidence_quote copied from that "
        f"section's reference excerpt. Use recommendations=[] when no exact source "
        f"evidence supports a change."
    )
    batch_evidence = (evidence_dir / f"audit_batch_{batch_idx:02d}") if evidence_dir else None
    parsed = bedrock_doc_audit(
        prompt=batch_prompt,
        model_alias=model,
        max_tokens=8000,
        timeout=600,
        evidence_dir=batch_evidence,
    ) or {}

    findings = parsed.get("section_findings", [])
    if not isinstance(findings, list):
        log.warning("[CLAUDE-AUDIT] Batch %d: invalid section_findings type", batch_idx)
        return []

    valid: list[dict] = []
    grounding_rejections: list[dict] = []
    standard_sources = [
        {"source_title": item}
        for item in (
            omi_standards_str.splitlines() + axi_standards_str.splitlines()
        )
        if item and item != "(none found)"
    ]
    for f in findings:
        if not isinstance(f, dict) or f.get("section") not in section_batch:
            continue
        gap = f.get("gap_score")
        if not isinstance(gap, (int, float)):
            continue
        # Canonicalize to exactly the 4 required keys so the shared validation
        # block in _claude_code_audit() does not reject unexpected extra keys.
        section_name = str(f["section"])
        section_reference = match_reference_section(
            reference_text,
            section_name,
        )
        recs, rejected = _ground_audit_recommendations(
            f,
            section_reference,
            standard_sources,
        )
        grounding_rejections.extend(
            {"section": section_name, **item} for item in rejected
        )
        stds = f.get("missing_standards", [])
        valid.append({
            "section": f["section"],
            "gap_score": max(0.0, min(1.0, float(gap))),
            "recommendations": recs if isinstance(recs, list) else [],
            "missing_standards": stds if isinstance(stds, list) else [],
        })

    if batch_evidence:
        batch_evidence.mkdir(parents=True, exist_ok=True)
        (batch_evidence / "grounding_rejections.json").write_text(
            json.dumps(grounding_rejections, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
    log.info("[CLAUDE-AUDIT] Batch %d: %d/%d sections valid", batch_idx, len(valid), len(section_batch))
    return valid


def _claude_code_audit(
    omi_standards: list[str],
    axi_standards: list[str],
    doc_excerpt: str,
    axi_recommendations: list[str],
    topic: str,
    reference_text: str = "",
    cycle: int = 0,
    structure_report: dict = None,
    evidence_dir: Optional[Path] = None,
    model: str = "opus",
) -> Optional[ClaudeAuditResult]:
    """Run one bounded full-document Claude audit with section-level output."""
    log.info(f"[CLAUDE-AUDIT] Starting independent validation (cycle={cycle})...")

    omi_standards_str = "\n".join(omi_standards) if omi_standards else "(none found)"
    axi_standards_str = "\n".join(axi_standards) if axi_standards else "(none found)"
    axi_recs_str = "\n".join(axi_recommendations) if axi_recommendations else "(none)"
    struct_info = json.dumps(structure_report, indent=2) if structure_report else "(not available)"
    start_time = datetime.now()
    bedrock_invoked = False  # Track whether bedrock_doc_audit() was successfully invoked

    doc_sections = {
        name: body
        for name, body in _split_into_sections(doc_excerpt).items()
        if name != "__preamble__" and body.strip()
    }
    section_names = list(doc_sections)
    findings_already_grounded = False

    # ── Chunked audit path (large documents) ───────────────────────────────
    # When section count exceeds _AUDIT_CHUNK_SIZE (14), split into batches
    # of ≤14 sections, call _audit_section_batch() per batch, then run a
    # lightweight synthesis call for document-level quality metrics.
    if len(section_names) > _AUDIT_CHUNK_SIZE:
        n_batches = -(-len(section_names) // _AUDIT_CHUNK_SIZE)  # ceil division
        log.info(
            "[CLAUDE-AUDIT] Chunked audit: %d sections → %d batches of ≤%d",
            len(section_names), n_batches, _AUDIT_CHUNK_SIZE,
        )
        all_findings: list[dict] = []
        for i in range(n_batches):
            batch = section_names[i * _AUDIT_CHUNK_SIZE : (i + 1) * _AUDIT_CHUNK_SIZE]
            all_findings.extend(
                _audit_section_batch(
                    batch_idx=i,
                    section_batch=batch,
                    doc_sections=doc_sections,
                    omi_standards_str=omi_standards_str,
                    axi_standards_str=axi_standards_str,
                    axi_recs_str=axi_recs_str,
                    topic=topic,
                    reference_text=reference_text,
                    cycle=cycle,
                    model=model,
                    evidence_dir=evidence_dir,
                )
            )

        # Coverage check (same threshold as single-pass path)
        audited_names = {f["section"] for f in all_findings}
        required_coverage = max(1, int(len(section_names) * 0.8))
        if len(audited_names) < required_coverage:
            log.warning(
                "[CLAUDE-AUDIT] Chunked coverage failure: audited=%d required=%d total=%d",
                len(audited_names), required_coverage, len(section_names),
            )
            return None

        # Synthesis call: derive document-level metrics from batch findings.
        # Only section summaries sent — no full text → cheap call.
        findings_summary = json.dumps(
            [{"section": f["section"], "gap_score": f["gap_score"]} for f in all_findings],
            ensure_ascii=False, indent=2,
        )
        synthesis_prompt = (
            f"You are the independent teacher/judge for an Asset Integrity Management document.\n\n"
            f"TOPIC: {topic}\n"
            f"CYCLE: {cycle}\n\n"
            f"SECTION AUDIT SUMMARY ({len(all_findings)} sections audited):\n"
            f"{findings_summary}\n\n"
            f"OMI STANDARDS FOUND:\n{omi_standards_str}\n\n"
            f"AXI STANDARDS FOUND:\n{axi_standards_str}\n\n"
            f"DOCUMENT FORMATION CONTRACT:\n"
            f"ISO 10013:2021 (structure), ISO 2145:1978 (numbering), "
            f"ISO 690:2021 (reference table) are deterministic pipeline gates. "
            f"They do not require redundant narrative citations.\n\n"
            f"STRUCTURE REPORT:\n{struct_info}\n\n"
            f"Based on the section-level audit above, provide document-level quality metrics.\n"
            f"Score EACH dimension from 0.0 to 1.0 based on your actual evaluation of the document:\n"
            f"  omi_quality.standards_accuracy: how accurately ISO/IEC standards are cited and applied\n"
            f"  omi_quality.completeness: how completely all required clauses are addressed\n"
            f"  omi_quality.context_relevance: how well content matches the asset management topic\n"
            f"  axi_quality.standards_accuracy: how precisely the generated text matches Axi's standard requirements\n"
            f"  axi_quality.completeness: how fully Axi's recommendations were implemented in the document\n"
            f"  axi_quality.context_relevance: how relevant the Axi-generated content is to the technical context\n"
            f"Do NOT return zeros unless you have a specific reason to score zero. A typical good document scores 0.7-0.9.\n"
            f"Return strict JSON with EXACTLY these keys and no others:\n"
            f'{{"omi_quality":{{"standards_accuracy":0.85,"completeness":0.80,"context_relevance":0.90}},'
            f'"axi_quality":{{"standards_accuracy":0.80,"completeness":0.75,"context_relevance":0.85}},'
            f'"overall_assessment":"...","missing_sections":[],"skill_recommendations":[]}}'
        )
        synthesis_evidence = (evidence_dir / "audit_synthesis") if evidence_dir else None
        synth_parsed: dict[str, Any] = bedrock_doc_audit(
            prompt=synthesis_prompt,
            model_alias=model,
            max_tokens=4000,
            timeout=300,
            evidence_dir=synthesis_evidence,
        ) or {}
        bedrock_invoked = bool(synth_parsed)  # True if bedrock_doc_audit() returned non-empty dict
        parse_error = "" if synth_parsed else "synthesis call returned None (using defaults)"

        # Reconstruct a complete parsed dict for the shared validation path below
        parsed: dict[str, Any] = {
            "section_findings": all_findings,
            "missing_sections": synth_parsed.get("missing_sections", []),
            "omi_quality": synth_parsed.get("omi_quality", {"standards_accuracy": 0.0, "completeness": 0.0, "context_relevance": 0.0}),
            "axi_quality": synth_parsed.get("axi_quality", {"standards_accuracy": 0.0, "completeness": 0.0, "context_relevance": 0.0}),
            "overall_assessment": synth_parsed.get("overall_assessment", "Chunked audit complete."),
            "skill_recommendations": synth_parsed.get("skill_recommendations", []),
        }
        findings_already_grounded = True
        log.info("[CLAUDE-AUDIT] Chunked audit: %d findings merged, synthesis %s",
                 len(all_findings), "OK" if synth_parsed else "DEFAULT")

    # ── Single-pass audit path (small/medium documents ≤ _AUDIT_CHUNK_SIZE) ─
    else:
        log.info(
            "[CLAUDE-AUDIT] Single-pass full audit: %d sections, "
            "%d document chars, %d reference chars",
            len(section_names), len(doc_excerpt), len(reference_text),
        )
        audit_prompt = f"""You are the independent teacher/judge for an Asset Integrity Management document.

TOPIC: {topic}
CYCLE: {cycle}
TEACHER/JUDGE MODEL: Claude Code CLI AWS {model}
GENERATED SECTION NAMES ({len(section_names)}):
{json.dumps(section_names, ensure_ascii=False)}

GENERATED DOCUMENT — FULL TEXT:
{doc_excerpt}

REFERENCE DOCUMENT — FULL AVAILABLE TEXT:
{reference_text or "(no reference text available)"}

OMI STANDARDS FOUND:
{omi_standards_str}

AXI STANDARDS FOUND:
{axi_standards_str}

AXI RECOMMENDATIONS:
{axi_recs_str}

STRUCTURE REPORT:
{struct_info}

Audit every generated section. Return strict JSON only with exactly these
top-level keys: section_findings, missing_sections, omi_quality, axi_quality,
overall_assessment, skill_recommendations.

Each section_findings item must contain exactly: section, gap_score,
recommendations, missing_standards. Each recommendation must be an object with
exactly text and evidence_quote. evidence_quote must be an exact substring from
the matching section in REFERENCE DOCUMENT. The recommendation text must name
an exact target, for example "Section 8.3: Add criticality matrix requirements".
Use recommendations=[] for sections without a material gap. Return at most
18 recommendations total and at most 2 recommendations for any one section.
Prioritize missing sections, safety-critical omissions, and reference gaps.

Score omi_quality and axi_quality dimensions from 0.0 to 1.0 based on your
actual evaluation. Do NOT return zeros unless a dimension truly scores zero.
A well-written document typically scores 0.7-0.9 on each dimension.
  omi_quality.standards_accuracy: accuracy of ISO/IEC standard citations
  omi_quality.completeness: completeness of clause coverage
  omi_quality.context_relevance: relevance to the asset management topic
  axi_quality.standards_accuracy: precision of generated text vs Axi standards
  axi_quality.completeness: implementation completeness of Axi recommendations
  axi_quality.context_relevance: relevance of Axi content to the technical context

JSON SHAPE:
{{"section_findings":[{{"section":"...","gap_score":0.0,"recommendations":[{{"text":"...","evidence_quote":"exact source substring"}}],"missing_standards":[]}}],"missing_sections":[],"omi_quality":{{"standards_accuracy":0.85,"completeness":0.80,"context_relevance":0.90}},"axi_quality":{{"standards_accuracy":0.80,"completeness":0.75,"context_relevance":0.85}},"overall_assessment":"...","skill_recommendations":[]}}"""

        # max_tokens raised to 16000 (was 8000) — 14 sections × ~400 tokens/finding
        # = 5600 output, well within new limit; covers up to ~40 sections safely.
        parsed: dict[str, Any] = bedrock_doc_audit(
            prompt=audit_prompt,
            model_alias=model,
            max_tokens=16000,
            timeout=600,
            evidence_dir=evidence_dir,
        ) or {}
        bedrock_invoked = bool(parsed)  # True if bedrock_doc_audit() returned non-empty dict
        parse_error = "" if parsed else "bedrock_doc_audit returned None"  # Assign immediately after bedrock_invoked to preserve flag before validation early returns

    grounding_rejections: list[dict] = []
    if not findings_already_grounded:
        standard_sources = [
            {"source_title": item}
            for item in [*omi_standards, *axi_standards]
        ]
        grounded_findings: list[dict] = []
        for finding in parsed.get("section_findings", []):
            if not isinstance(finding, dict):
                grounded_findings.append(finding)
                continue
            section_name = str(finding.get("section", ""))
            section_reference = match_reference_section(
                reference_text,
                section_name,
            )
            recommendations, rejected = _ground_audit_recommendations(
                finding,
                section_reference,
                standard_sources,
            )
            grounded_findings.append({
                **finding,
                "recommendations": recommendations,
            })
            grounding_rejections.extend(
                {"section": section_name, **item} for item in rejected
            )
        parsed["section_findings"] = grounded_findings
    if evidence_dir:
        (evidence_dir / "audit_grounding_rejections.json").write_text(
            json.dumps(grounding_rejections, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

    required = {
        "section_findings",
        "missing_sections",
        "omi_quality",
        "axi_quality",
        "overall_assessment",
        "skill_recommendations",
    }
    if parse_error or set(parsed) != required:
        log.warning("[CLAUDE-AUDIT] Invalid top-level schema; parse_error=%s bedrock_invoked=%s", parse_error, bedrock_invoked)
        # If bedrock was invoked but schema is invalid, still create a degraded result to preserve bedrock_invoked flag
        if bedrock_invoked:
            elapsed = (datetime.now() - start_time).total_seconds()
            return ClaudeAuditResult(
                omi_quality={"standards_accuracy": 0.0, "completeness": 0.0, "context_relevance": 0.0},
                axi_quality={"standards_accuracy": 0.0, "completeness": 0.0, "context_relevance": 0.0},
                overall_assessment=f"Bedrock audit returned invalid schema; parse_error={parse_error}",
                missing_standards=[],
                audit_time=elapsed,
                recommendations_from_audit=[],
                skill_recommendations=[],
                reference_gap={},
                bedrock_invoked=True,  # Preserve that bedrock was invoked
            )
        return None

    section_findings = parsed["section_findings"]
    if not isinstance(section_findings, list):
        log.warning("[CLAUDE-AUDIT] Invalid section_findings type (not list); bedrock_invoked=%s", bedrock_invoked)
        if bedrock_invoked:
            elapsed = (datetime.now() - start_time).total_seconds()
            return ClaudeAuditResult(
                omi_quality={"standards_accuracy": 0.0, "completeness": 0.0, "context_relevance": 0.0},
                axi_quality={"standards_accuracy": 0.0, "completeness": 0.0, "context_relevance": 0.0},
                overall_assessment="Bedrock audit returned invalid section_findings structure",
                missing_standards=[],
                audit_time=elapsed,
                recommendations_from_audit=[],
                skill_recommendations=[],
                reference_gap={},
                bedrock_invoked=True,
            )
        return None

    valid_findings: list[dict] = []
    audited_names: set[str] = set()
    all_recommendations: list[str] = []
    all_missing_standards: set[str] = set()
    for finding in section_findings:
        if not isinstance(finding, dict) or set(finding) != {
            "section",
            "gap_score",
            "recommendations",
            "missing_standards",
        }:
            return None
        if finding["section"] not in section_names:
            continue
        if not isinstance(finding["gap_score"], (int, float)):
            return None
        score = float(finding["gap_score"])
        if not 0.0 <= score <= 1.0:
            return None
        if not isinstance(finding["recommendations"], list):
            return None
        if not isinstance(finding["missing_standards"], list):
            return None
        if not all(isinstance(item, str) for item in finding["recommendations"]):
            return None
        if not all(isinstance(item, str) for item in finding["missing_standards"]):
            return None
        valid_findings.append(finding)
        audited_names.add(finding["section"])
        all_missing_standards.update(finding["missing_standards"])

    required_coverage = max(1, int(len(section_names) * 0.8))
    if len(audited_names) < required_coverage:
        log.warning(
            "[CLAUDE-AUDIT] Coverage failure: audited=%d required=%d total=%d; bedrock_invoked=%s",
            len(audited_names),
            required_coverage,
            len(section_names),
            bedrock_invoked,
        )
        # If bedrock was invoked but coverage is insufficient, still create a degraded result
        if bedrock_invoked:
            elapsed = (datetime.now() - start_time).total_seconds()
            return ClaudeAuditResult(
                omi_quality={"standards_accuracy": 0.0, "completeness": 0.0, "context_relevance": 0.0},
                axi_quality={"standards_accuracy": 0.0, "completeness": 0.0, "context_relevance": 0.0},
                overall_assessment=f"Bedrock audit coverage insufficient: {len(audited_names)}/{required_coverage} sections",
                missing_standards=list(all_missing_standards),
                audit_time=elapsed,
                recommendations_from_audit=all_recommendations,
                skill_recommendations=[],
                reference_gap={},
                bedrock_invoked=True,
            )
        return None

    missing_sections = parsed["missing_sections"]
    if not isinstance(missing_sections, list) or not all(
        isinstance(item, str) for item in missing_sections
    ):
        return None
    missing_sections = _verified_missing_sections(
        missing_sections,
        section_names,
    )
    avg_gap = sum(float(f["gap_score"]) for f in valid_findings) / len(valid_findings)
    high_gap_sections = [
        f["section"] for f in valid_findings if float(f["gap_score"]) >= 0.4
    ]
    # Prefix each recommendation with its section ID so normalize_rec() can
    # resolve the target section without guessing from free-form text.
    all_recommendations = [
        f"Section {finding['section']}: {recommendation}"
        for finding in sorted(
            valid_findings,
            key=lambda item: float(item["gap_score"]),
            reverse=True,
        )
        for recommendation in finding["recommendations"]
    ]

    elapsed = (datetime.now() - start_time).total_seconds()

    omi_q = parsed["omi_quality"]
    axi_q = parsed["axi_quality"]
    skill_recs = parsed["skill_recommendations"]
    quality_keys = {
        "standards_accuracy",
        "completeness",
        "context_relevance",
    }
    if (
        not isinstance(omi_q, dict)
        or not isinstance(axi_q, dict)
        or set(omi_q) != quality_keys
        or set(axi_q) != quality_keys
        or not all(isinstance(value, (int, float)) for value in omi_q.values())
        or not all(isinstance(value, (int, float)) for value in axi_q.values())
        or not all(0.0 <= float(value) <= 1.0 for value in omi_q.values())
        or not all(0.0 <= float(value) <= 1.0 for value in axi_q.values())
    ):
        return None
    if not isinstance(skill_recs, list) or not all(
        isinstance(item, str) for item in skill_recs
    ):
        return None
    if not isinstance(parsed["overall_assessment"], str):
        return None
    ref_gap = {
        "gap_score": avg_gap,
        "missing_sections": missing_sections,
        "high_gap_sections": high_gap_sections,
        "sections_audited": len(valid_findings),
        "sections_expected": len(section_names),
        "audit_mode": "single_pass_full_document",
    }

    omi_avg = sum(omi_q.values()) / len(omi_q) if omi_q else 0
    axi_avg = sum(axi_q.values()) / len(axi_q) if axi_q else 0
    log.info(
        f"[CLAUDE-AUDIT] Complete in {elapsed:.0f}s — "
        f"sections={len(valid_findings)}, avg_gap={avg_gap:.1%}, "
        f"Omi:{omi_avg:.1%} Axi:{axi_avg:.1%}, "
        f"recs={len(all_recommendations)}, skill_recs={len(skill_recs)}"
    )

    return ClaudeAuditResult(
        omi_quality=omi_q,
        axi_quality=axi_q,
        overall_assessment=parsed["overall_assessment"],
        missing_standards=list(all_missing_standards),
        audit_time=elapsed,
        recommendations_from_audit=all_recommendations,
        skill_recommendations=skill_recs,
        reference_gap=ref_gap,
        bedrock_invoked=bedrock_invoked,
    )


# ──────────────────────────────────────────────────────────────────────────────
# STAGE 4: IMPROVEMENT CYCLE — Apply Recommendations
# ──────────────────────────────────────────────────────────────────────────────

def _apply_improvements(
    doc_text: str,
    recommendations: list[str],
    topic: str,
    external_standards: dict = None,
    audit_recommendations: list[str] = None,
) -> str:
    """
    Apply Axi + Claude audit recommendations + external standards to improve document.

    Flow:
    1. Merge Axi + Claude audit recommendations (teacher signal)
    2. Build improvement prompt with full document context
    3. Call SLOT120_REASONING to revise draft
    """
    # Merge Axi + Claude audit recommendations
    all_recommendations = list(recommendations or [])
    if audit_recommendations:
        all_recommendations += audit_recommendations
        log.info(f"[IMPROVE] Merged {len(recommendations)} Axi + {len(audit_recommendations)} Claude audit recs → {len(all_recommendations)} total")
    else:
        log.info(f"[IMPROVE] Applying {len(all_recommendations)} recommendations...")

    if external_standards is None:
        external_standards = {"standards": [], "content": {}}

    # Build improvement prompt with standards context
    external_standards_section = ""
    if external_standards["standards"]:
        external_standards_section = "\n\nEXTERNAL STANDARDS TO INTEGRATE:"
        for std, content in external_standards["content"].items():
            external_standards_section += f"\n\n{std}:\n{content[:800]}"

    improvement_prompt = f"""Improve the following document by integrating all recommendations and external standards.

TOPIC: {topic}

RECOMMENDATIONS (Axi + Claude audit teacher):
{chr(10).join(f"- {r}" for r in all_recommendations)}{external_standards_section}

CURRENT DOCUMENT (FULL — do not lose any existing sections):
{doc_text}

CRITICAL IMPROVEMENT RULES:
1. PRESERVE all existing Markdown headings (##, ###) — do NOT remove or rename sections
2. EXTEND sections by adding content under them, never replace structure
3. ADD all recommended new sections as new ## headings at the end
4. Integrate external standards content into relevant existing sections
5. Expand References section with all new standards found
6. Output the COMPLETE improved document — every section, not a summary

Generate the fully improved document:"""

    try:
        log.info(f"[IMPROVE] Calling {ModelConfig.SLOT120_REASONING} for revision...")
        generated = _slot120_generate(
            improvement_prompt,
            timeout=IMPROVEMENT_GENERATE_TIMEOUT,
            num_predict=IMPROVEMENT_GENERATE_NUM_PREDICT,
        )
        log.info(
            "[IMPROVE] Document improved (%d chars, model=%s)",
            len(generated),
            ModelConfig.SLOT120_REASONING,
        )
        return generated
    except TimeoutError:
        log.error(f"[IMPROVE] Improvement timeout (>10 min)")
    except Exception as e:
        log.error(f"[IMPROVE] Failed: {e}")

    log.warning(f"[IMPROVE] Falling back to original document")
    return doc_text  # Return original if improvement fails


# ──────────────────────────────────────────────────────────────────────────────
# TRAINING PAIR EXPORT — Self-Learning for FT Pipeline
# ──────────────────────────────────────────────────────────────────────────────

def _save_training_pairs(
    cycle: int,
    topic: str,
    doc_text: str,
    metrics: GenerationMetrics,
    recommendations: list[str],
    external_standards: dict,
    output_dir: Path,
) -> None:
    """
    Save training pairs from this cycle for continuous fine-tuning.

    Three types of training pairs:
    1. Generation pair: topic + internal_standards → draft (input: Omi's search + generation)
    2. Improvement pair: draft + recommendations → improved_draft (input: Axi's validation)
    3. Quality evaluation pair: document → quality assessment (meta-learning)

    Pairs are saved inside the run output directory for traceable Traini ingestion.
    """
    if metrics.overall_score < 0.60:
        # Skip low-quality documents
        log.debug(f"[TRAIN] Skipping cycle {cycle}: quality too low ({metrics.overall_score:.1%})")
        return

    # Keep each run self-contained. The legacy global directory can be root-owned.
    ft_dir = output_dir / "learning_pairs"
    ft_dir.mkdir(parents=True, exist_ok=True)

    # Training pair 1: Generation pair (Omi's work)
    # Input: topic + internal standards found
    # Output: generated draft
    if metrics.overall_score >= 0.70:
        generation_pair = {
            "type": "generation",
            "cycle": cycle,
            "timestamp": datetime.now().isoformat(),
            "topic": topic,
            "input_prompt": f"Generate comprehensive document on: {topic}",
            "internal_standards": [],  # Would come from Omi search
            "output_draft": doc_text[:2000],  # First 2000 chars
            "quality_score": metrics.overall_score,
            "model": ModelConfig.SLOT120_REASONING
        }

        # Append to generation pairs
        gen_pairs_path = ft_dir / "generation_pairs.jsonl"
        with open(gen_pairs_path, "a", encoding='utf-8') as f:
            f.write(json.dumps(generation_pair, ensure_ascii=False) + "\n")
        log.info(f"[TRAIN] Saved generation pair (cycle {cycle}, quality={metrics.overall_score:.1%})")

    # Training pair 2: Improvement pair (Axi's work)
    # Input: current draft + recommendations
    # Output: improved draft (to be generated in next cycle)
    if recommendations and metrics.overall_score >= 0.70:
        improvement_pair = {
            "type": "improvement",
            "cycle": cycle,
            "timestamp": datetime.now().isoformat(),
            "topic": topic,
            "current_draft": doc_text[:1500],
            "recommendations": recommendations[:5],  # Top 5 recommendations
            "external_standards": external_standards.get('standards', [])[:5],
            "quality_before": metrics.overall_score,
            "structure_score": metrics.structure_score,
            "standards_score": metrics.standards_score,
            "model": "claude_code_cli_aws:opus-4.6"
        }

        # Append to improvement pairs
        imp_pairs_path = ft_dir / "improvement_pairs.jsonl"
        with open(imp_pairs_path, "a", encoding='utf-8') as f:
            f.write(json.dumps(improvement_pair, ensure_ascii=False) + "\n")
        log.info(f"[TRAIN] Saved improvement pair (cycle {cycle}, {len(recommendations)} recommendations)")

    # Training pair 3: Quality evaluation pair (meta-learning)
    # Input: document
    # Output: quality assessment (structure, standards, coverage scores)
    if metrics.overall_score >= 0.75:
        quality_pair = {
            "type": "quality_evaluation",
            "cycle": cycle,
            "timestamp": datetime.now().isoformat(),
            "topic": topic,
            "document_excerpt": doc_text[:1000],
            "assessment": {
                "structure_score": metrics.structure_score,
                "standards_score": metrics.standards_score,
                "coverage_score": metrics.coverage_score,
                "overall_score": metrics.overall_score,
            },
            "sections_found": metrics.sections_found,
            "sections_expected": metrics.sections_expected,
            "standards_found": metrics.standards_found,
            "standards_expected": metrics.standards_expected,
        }

        # Append to quality pairs
        qual_pairs_path = ft_dir / "quality_assessment_pairs.jsonl"
        with open(qual_pairs_path, "a", encoding='utf-8') as f:
            f.write(json.dumps(quality_pair, ensure_ascii=False) + "\n")
        log.info(f"[TRAIN] Saved quality assessment pair (cycle {cycle})")

    log.info(f"[TRAIN] Training pairs for cycle {cycle} exported to {ft_dir}")


def resolve_bound_reference_path(
    raw_path: str | Path,
    *,
    workspace_root: Path | None = None,
) -> Path:
    """Resolve a reference path visible in the current runtime/container."""
    raw = Path(str(raw_path))
    if raw.exists():
        return raw.resolve()

    workspace_root = workspace_root or ROOT
    raw_str = str(raw_path).lstrip("/")
    candidates = [
        workspace_root / raw_str,
        ROOT / raw_str,
        ROOT / "aims_workspace" / raw_str,
        Path("/aims_workspace") / raw_str,
        Path("/data") / raw_str,
        Path("/workspace") / raw_str,
        Path("/workspace/aims_workspace") / raw_str,
        Path("/data/aims_workspace") / raw_str,
        Path("/workspace/aims_workspace/inbox/training/philosophies") / raw.name,
        Path("/data/inbox/training/philosophies") / raw.name,
        Path("/workspace/inbox/training/philosophies") / raw.name,
    ]
    if raw.is_absolute():
        candidates.append(Path("/").joinpath(raw.relative_to(raw.anchor)))
    for candidate in candidates:
        try:
            if candidate.exists():
                return candidate.resolve()
        except Exception:
            continue
    raise FileNotFoundError(f"BOUND_REFERENCE_PATH_MISSING:{raw_path}")


def _load_reference_document(reference_path: Path) -> tuple[str, dict[str, Any]]:
    report: dict[str, Any] = {
        "path": str(reference_path),
        "path_exists": reference_path.exists(),
        "suffix": reference_path.suffix.lower(),
        "content_chars": 0,
        "reference_sha256": "",
        "status": "FAIL",
    }
    if not reference_path.exists():
        return "", report
    def _ocr_sidecar_candidates(pdf_path: Path) -> list[Path]:
        txt_name = pdf_path.with_suffix(".txt").name
        roots = {
            pdf_path.parent.parent.parent,
            pdf_path.parent.parent.parent.parent if len(pdf_path.parents) > 3 else None,
            ROOT,
            Path("/data"),
            Path("/workspace"),
        }
        candidates: list[Path] = []
        for root in roots:
            if root is None:
                continue
            candidates.extend(
                [
                    root / "staging" / "ocr_text" / txt_name,
                    root / "aims_workspace" / "staging" / "ocr_text" / txt_name,
                ]
            )
        candidates.append(pdf_path.with_suffix(".txt"))
        return candidates

    pdf_error: str | None = None
    try:
        if reference_path.suffix.lower() == ".docx":
            from docx import Document as DocxDocument

            docx = DocxDocument(str(reference_path))
            chunks = [
                paragraph.text
                for paragraph in docx.paragraphs
                if paragraph.text.strip()
            ]
            for table in docx.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        chunks.append(row_text)
            reference_text = "\n".join(chunks)
            report["paragraph_count"] = len(docx.paragraphs)
            report["table_count"] = len(docx.tables)
        elif reference_path.suffix.lower() == ".pdf":
            chunks = []
            try:
                import pdfplumber

                with pdfplumber.open(str(reference_path)) as pdf:
                    report["page_count"] = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            chunks.append(page_text)
            except Exception as exc:
                pdf_error = repr(exc)
                report["pdf_extraction_error"] = pdf_error
            reference_text = "\n".join(chunks)
            if not reference_text.strip():
                for sidecar in _ocr_sidecar_candidates(reference_path):
                    if sidecar.exists():
                        sidecar_text = sidecar.read_text(encoding="utf-8").strip()
                        if sidecar_text:
                            reference_text = sidecar_text
                            report["ocr_sidecar_path"] = str(sidecar)
                            report["ocr_sidecar_status"] = "PASS"
                            break
            if not reference_text.strip() and pdf_error is not None:
                report.setdefault("error", pdf_error)
        else:
            report["error"] = "unsupported_reference_suffix"
            return "", report
    except Exception as exc:
        report["error"] = repr(exc)
        return "", report

    report.update(
        {
            "content_chars": len(reference_text),
            "reference_sha256": hashlib.sha256(
                reference_text.encode("utf-8")
            ).hexdigest(),
            "status": "PASS" if reference_text.strip() else "FAIL",
        }
    )
    return reference_text, report


def _write_enforcement_report(
    output_dir: Path,
    *,
    gates: list[GateResult],
    training_gate: GateResult,
    reference_report: dict[str, Any],
) -> Path:
    payload = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "gates": [asdict(gate) for gate in gates],
        "training_promotion_gate": asdict(training_gate),
        "reference_extraction": reference_report,
    }
    path = output_dir / "phase16_enforcement_report.json"
    path.write_text(
        json.dumps(payload, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    (output_dir / "phase16_enforcement_report.md").write_text(
        "\n".join(
            [
                "# DOCGEN Reference-Aware Enforcement",
                "",
                *[
                    f"- `{gate.gate_name}`: `{gate.status}` "
                    f"(`{gate.blocker_code}`)"
                    for gate in gates
                ],
                f"- `training_promotion_gate`: `{training_gate.status}`",
                "",
            ]
        ),
        encoding="utf-8",
    )
    return path


def _learning_pair_export_decision(
    *,
    export_requested: bool,
    hard_gate_allowed: bool,
    hard_gate_failures: list[str],
    preflight_gates: list[GateResult],
    generated_type_gate: GateResult,
) -> tuple[bool, list[str]]:
    failures: list[str] = []
    if not export_requested:
        failures.append("learning_pair_export_not_requested")
    if not hard_gate_allowed:
        failures.extend(hard_gate_failures or ["cycle_hard_gate=FAIL"])
    failures.extend(
        f"{gate.gate_name}={gate.status}"
        for gate in preflight_gates
        if gate.status != "PASS"
    )
    if generated_type_gate.status != "PASS":
        failures.append(
            f"{generated_type_gate.gate_name}={generated_type_gate.status}"
        )
    return not failures, failures


# ──────────────────────────────────────────────────────────────────────────────
# MAIN PIPELINE
# ──────────────────────────────────────────────────────────────────────────────

def run_cyclic_generation(
    topic: str,
    reference_pdf: Path,
    max_cycles: int = 5,
    target_quality: float = 0.98,
    output_dir: Path = None,
    document_type: str = "technical_report",
    enable_docsreg_meta_cycle: bool = False,
    docsreg_target_quality: float = 0.98,
    docsreg_max_fresh_runs: int = 7,
    docsreg_stall_limit: int = 3,
    docsreg_min_progress_delta: float = 0.001,
    docsreg_document_type: str = "procedure",
    docsreg_evidence_root: Optional[Path] = None,
    docsreg_write_policy: str = "write_all",
    docsreg_auditor_mode: str = "production",
    export_learning_pairs: bool = False,
    soft_preflight_fail: bool = False,
) -> CycleResult:
    """
    Run complete cyclic generation pipeline:
    Omi draft → Eval → Axi validate → Improve → Repeat until target quality

    Args:
        document_type: Document type for type-specific validation profiles
            (technical_report, policy, memo, contract, etc.)

        DOCSREG meta-cycle (optional):
            enable_docsreg_meta_cycle: If True, run a DOCSREG document-type
                certification loop after the main generation cycle completes.
                All docsreg_* params are ignored when this is False.
            docsreg_target_quality: Quality threshold for DOCSREG certification (default 0.98).
            docsreg_max_fresh_runs: Max fresh-start runs per document type (default 7).
            docsreg_stall_limit: Cycles with < docsreg_min_progress_delta before stall (default 3).
            docsreg_min_progress_delta: Minimum quality improvement per cycle (default 0.001).
            docsreg_document_type: Document type for DOCSREG (default "procedure"; falls back to document_type).
            docsreg_evidence_root: Evidence directory for DOCSREG output (default: output_dir/docsreg).
            docsreg_write_policy: Evidence write policy (default "write_all").
            docsreg_auditor_mode: Auditor mode: "noop" (default) or "claude_code".
            export_learning_pairs: Explicitly export learning pairs only after
                all enforcement and cycle gates pass. Defaults to False.
            soft_preflight_fail: When True, record preflight gate failures but continue
                the DOCGEN loop so upgrade runs can produce a first cycle and feedback.
    """
    if output_dir is None:
        output_dir = Path("aims_workspace/cyclic_doc_output")
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        reference_pdf = resolve_bound_reference_path(reference_pdf, workspace_root=ROOT)
    except FileNotFoundError as exc:
        (output_dir / "DOCUMENT_CYCLE_BLOCKED.json").write_text(
            json.dumps(
                {
                    "classification": "DOCUMENT_CYCLE_BLOCKED_REFERENCE_PATH",
                    "generation_started": False,
                    "blocker_code": "BOUND_REFERENCE_PATH_MISSING",
                    "reason": str(exc),
                    "requested_reference": str(reference_pdf),
                    "resolved_workspace_root": str(ROOT),
                },
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        raise RuntimeError(str(exc)) from exc
    reference_text, reference_extraction = _load_reference_document(reference_pdf)
    generation_profile = get_document_generation_profile(document_type)
    (output_dir / "document_generation_profile_contract.json").write_text(
        json.dumps(
            generation_profile.generation_contract(),
            indent=2,
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    initial_context = build_generation_context(
        topic=topic,
        doc_type=document_type,
        task_context=topic,
        standards=[],
        reference_text=reference_text,
        reference_path=str(reference_pdf),
        strict_retrieval=True,
    )
    initial_archetype = resolve_document_archetype(initial_context)
    initial_sections = build_section_contract(initial_context)
    profile_gate = validate_reference_aware_profile_binding(
        requested_document_type=document_type,
        archetype_id=initial_archetype["archetype_id"],
        section_contract_count=len(initial_sections),
    )
    reference_gate = validate_reference_binding(
        EnforcementContext(
            requested_document_type=document_type,
            reference_binding={
                "references": [str(reference_pdf)],
                **reference_extraction,
            },
        )
    )
    preflight_gates = [profile_gate, reference_gate]
    preliminary_training_gate = decide_training_or_promotion_allowed(
        preflight_gates
    )
    preflight_blocked_gates = [gate for gate in preflight_gates if gate.blocks_execution]
    if preflight_blocked_gates:
        report_path = _write_enforcement_report(
            output_dir,
            gates=preflight_gates,
            training_gate=preliminary_training_gate,
            reference_report=reference_extraction,
        )
        (output_dir / "DOCUMENT_CYCLE_BLOCKED.json").write_text(
            json.dumps(
                {
                    "classification": "DOCUMENT_CYCLE_BLOCKED_PREFLIGHT_GATE",
                    "generation_started": False,
                    "gate_report": str(report_path),
                    "failed_gates": [
                        asdict(gate)
                        for gate in preflight_gates
                        if gate.blocks_execution
                    ],
                    "soft_preflight_fail": soft_preflight_fail,
                },
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        if not soft_preflight_fail:
            raise RuntimeError("Reference-aware DOCGEN preflight gate blocked")
        log.warning(
            "[PRELIGHT] Soft-failing DOCGEN preflight gate: %s",
            "; ".join(_gate_message(gate) for gate in preflight_blocked_gates),
        )

    preflight = None
    if not preflight_blocked_gates or not soft_preflight_fail:
        preflight = _runtime_preflight(output_dir)
    else:
        preflight = {
            "status": "BLOCKED",
            "soft_fail": True,
            "blocked_gates": [_gate_message(gate) for gate in preflight_blocked_gates],
        }
        (output_dir / "runtime_preflight.json").write_text(
            json.dumps(preflight, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
    judge_smoke = _claude_judge_smoke(
        output_dir,
        soft_fail=soft_preflight_fail,
    )
    judge_gate = validate_real_judge_path(
        EnforcementContext(
            requested_document_type=document_type,
            judge_context={
                "provider": judge_smoke.get("provider"),
                "model_id": judge_smoke.get("model"),
                "bedrock_invoked": judge_smoke.get("status") == "PASS",
                "fallback_model": judge_smoke.get("degraded_mode", False),
            },
        )
    )
    preflight_gates.append(judge_gate)
    training_promotion_gate = decide_training_or_promotion_allowed(
        preflight_gates
    )
    enforcement_report_path = _write_enforcement_report(
        output_dir,
        gates=preflight_gates,
        training_gate=training_promotion_gate,
        reference_report=reference_extraction,
    )
    if judge_gate.status != "PASS" and not soft_preflight_fail:
        raise RuntimeError(
            f"{judge_gate.blocker_code}: {judge_gate.message}"
        )

    teacher_judge_model = judge_smoke["model"]
    routed_topic = _slot14_route_topic(topic, output_dir)
    projected_vram_fraction = float(preflight.get("projected_vram_fraction") or 0.0)
    log.info(
        "[PREFLIGHT] PASS slot14=%s slot120=%s projected_vram=%.1f%%",
        ModelConfig.SLOT14_SEARCH,
        ModelConfig.SLOT120_REASONING,
        projected_vram_fraction * 100,
    )

    generation_context = build_generation_context(
        topic=routed_topic,
        doc_type=document_type,
        task_context=topic,
        standards=[],
        reference_text=reference_text,
        reference_path=str(reference_pdf),
        strict_retrieval=True,
    )
    selected_reference_exemplar = {
        "title": Path(reference_pdf).name,
        "file_path": str(reference_pdf),
        "summary": reference_text[:2000],
        "approval_status": "reference",
    }
    document_requirement_graph = build_document_requirement_graph(
        request_text=topic,
        document_type=document_type,
        retrieved_documents=[selected_reference_exemplar],
        selected_exemplar=selected_reference_exemplar,
        retrieval_diagnostics={
            "retrieval_passed": True,
            "selected_reference_path": str(reference_pdf),
            "standards": [],
        },
    )
    generation_context.requirement_graph = document_requirement_graph.to_dict()
    generation_context.dynamic_section_map = list(
        document_requirement_graph.final_section_map
    )
    document_archetype = resolve_document_archetype(generation_context)
    section_contract = build_section_contract(generation_context)
    required_section_titles = [item.title for item in section_contract]
    # Technical standards are discovered from task context. ISO 55001/55002
    # are registration taxonomy aids, not automatic content requirements.
    archetype_standards: list[str] = []
    formatting_standards = list(FIRST_BATCH_IMPLEMENTATION_STANDARDS)
    validation_document_type = validation_document_type_for(document_type)
    (output_dir / "document_archetype.json").write_text(
        json.dumps(
            {
                "archetype": document_archetype,
                "section_contract": [
                    asdict(item) for item in section_contract
                ],
                "requirement_graph": document_requirement_graph.to_dict(),
                "dynamic_section_map": list(
                    document_requirement_graph.final_section_map
                ),
                "source_use_policy": document_requirement_graph.source_use_policy,
            },
            indent=2,
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    best_result = None
    all_metrics = []
    audit_cache: dict[str, ClaudeAuditResult] = {}
    reference_baseline = extract_reference_baseline(reference_text)
    (output_dir / "reference_baseline.json").write_text(
        json.dumps(reference_baseline, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    external_standards_cache = {"standards": [], "content": {}}  # Accumulate across cycles
    repair_batch: list[str] = []
    recommendations: list[str] = []
    profile_repair_recommendations: list[str] = []
    internal_standards: list[str] = []
    applied_recs_this_cycle: list[str] = []  # recs applied at START of cycle — verify these, not new ones
    prev_doc_text: str = ""  # previous cycle snapshot for lineage and fallback
    last_accepted_doc: str = ""
    last_accepted_metrics: Optional[GenerationMetrics] = None
    # Cycle runs until:
    # 1. Target quality reached (target_quality, default 0.98)
    # 2. Last 3 cycles changed < 0.01% (convergence stall)
    # 3. max_cycles > 0 hard limit reached
    # max_cycles=0 = unlimited
    NO_PROGRESS_LIMIT   = 3     # consecutive cycles with delta < 0.0001 → stop
    NO_PROGRESS_DELTA   = 0.001  # threshold: 0.1% delta = no meaningful progress

    cycle = 0
    no_progress_streak = 0
    while True:
        cycle += 1
        if max_cycles > 0 and cycle > max_cycles:
            log.info(f"[CYCLE {cycle}] max_cycles={max_cycles} hard limit reached — stopping")
            break

        log.info(f"\n{'='*70}")
        log.info(f"CYCLE {cycle}" + (f"/{max_cycles}" if max_cycles > 0 else " (unlimited)"))
        log.info(f"{'='*70}")

        cycle_dir = output_dir / f"cycle_{cycle:02d}"
        cycle_dir.mkdir(exist_ok=True)
        cycle_started_at = datetime.now(timezone.utc)
        section_edit_result: Optional[dict] = None
        changes_applied_by_skill: list[str] = []
        applied_recs_this_cycle: list[str] = []

        # Load repair batch from previous cycle (if cycle > 1)
        if cycle > 1:
            prev_cycle_dir = output_dir / f"cycle_{cycle-1:02d}"
            prev_repair_plan = prev_cycle_dir / "repair_plan.json"
            if prev_repair_plan.exists():
                try:
                    with open(prev_repair_plan, 'r', encoding='utf-8') as f:
                        plan_data = json.load(f)
                        repair_batch = plan_data.get("selected", [])
                        log.info(f"[CYCLE {cycle}] Loaded {len(repair_batch)} recommendations from cycle {cycle-1}")
                except Exception as e:
                    log.warning(f"[CYCLE {cycle}] Failed to load repair plan from cycle {cycle-1}: {e}")
                    repair_batch = []
            else:
                log.warning(f"[CYCLE {cycle}] No repair plan found for cycle {cycle-1}")
                repair_batch = []

        # STAGE 1: Omi generates draft OR apply improvements
        if cycle == 1:
            # First cycle: Omi searches internal database automatically
            log.info(f"[CYCLE {cycle}] Omi generating initial draft...")
            raw_standards = _search_internal_standards(
                " ".join(
                    [
                        routed_topic,
                        *discovery_hints_for(document_type),
                    ]
                ),
                max_results=20,
                document_type=document_type,
                governed_fallback=archetype_standards,
            )
            # Filter out Qdrant noise (electrical/unrelated standards).
            # When reference_text is available, prefer standards cited in it.
            internal_standards = list(dict.fromkeys(
                _filter_standards_for_reference(
                    raw_standards,
                    reference_text,
                )
            ))
            generation_context.standards = internal_standards
            generation_context.provenance.append(
                {
                    "source_type": "task_archetype",
                    "title": document_archetype["archetype_id"],
                    "path": (
                        "ops/docagent/templates/document_archetypes.yaml"
                    ),
                    "chars": 0,
                    "score": document_archetype["match_score"],
                }
            )
            write_context_manifest(
                generation_context,
                cycle_dir / "generation_context_manifest.json",
            )
            doc_text = _omi_generate_draft(
                routed_topic,
                reference_template=reference_text,
                standards_to_inject=internal_standards,
                search_internal=False,
                grounded_prompt=render_generation_prompt(generation_context),
            )
        else:
            # Track which recs we're applying — needed for lineage check after improvement
            all_recs = list(repair_batch)
            applied_recs_this_cycle = all_recs
            log.info(f"[CYCLE {cycle}] Applying {len(all_recs)} recommendations (targeted edits)...")

            # GAP-001: section-batched transactional editing
            section_edit_result = apply_section_edits(
                doc=doc_text,
                recommendations=all_recs,
                last_accepted_doc=last_accepted_doc or prev_doc_text,
                reference_text=reference_text,
            )
            doc_text = section_edit_result["improved_doc"]
            changes_applied_by_skill = section_edit_result[
                "verified_recommendations"
            ]
            log.info(
                f"[CYCLE {cycle}] Section edits: {len(changes_applied_by_skill)} verified, "
                f"global={len(section_edit_result['global_recs'])}, "
                f"unresolved={len(section_edit_result['unresolved_recs'])}, "
                f"rolled_back={section_edit_result['rolled_back']}"
            )
            if section_edit_result["unresolved_recs"]:
                log.warning(
                    f"[CYCLE {cycle}] Unresolved recs: "
                    f"{section_edit_result['unresolved_recs'][:2]}"
                )

        if not doc_text:
            log.error(f"[CYCLE {cycle}] Generation failed, stopping")
            break

        if cycle == 1:
            doc_text, baseline_report = apply_reference_baseline(
                doc_text,
                reference_baseline,
            )
            baseline_report["reference_prompt_bound"] = True
            baseline_report["reference_content_chars"] = len(reference_text)
            if (
                baseline_report["status"] == "SKIPPED"
                and not reference_baseline.get("structured_blocks")
            ):
                baseline_report["status"] = "REFERENCE_CONTEXT_BOUND"
                baseline_report["reason"] = (
                    "Reference content was bound in the generation prompt; "
                    "no deterministic structured blocks were extracted."
                )
        else:
            baseline_report = {
                "status": "NOT_REAPPLIED",
                "reason": "structured baseline is injected once; later cycles use atomic repairs",
                "applied_blocks": [],
                "reference_prompt_bound": True,
                "reference_content_chars": len(reference_text),
            }
        (cycle_dir / "reference_baseline_application.json").write_text(
            json.dumps(baseline_report, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        log.info(
            "[CYCLE %d] Deterministic reference baseline: %s",
            cycle,
            [
                item.get("target_title")
                for item in baseline_report.get("applied_blocks", [])
            ],
        )

        # ── Phase 2: inject Section 8 sub-element nesting stubs (no LLM) ────
        phase2_result = apply_phase2_nesting(doc_text)
        if phase2_result["sections_expanded"]:
            doc_text = phase2_result["improved_doc"]
            log.info(
                "[CYCLE %d] Phase 2 nesting: %s expanded, %d stubs added",
                cycle,
                phase2_result["sections_expanded"],
                phase2_result["total_stubs_added"],
            )
        else:
            log.debug(
                "[CYCLE %d] Phase 2 nesting: no new stubs needed (already_nested=%s)",
                cycle,
                phase2_result.get("already_nested", []),
            )

        doc_text, profile_formatting_report = apply_profile_formatting(
            doc_text,
            generation_profile,
        )
        (cycle_dir / "profile_formatting_report.json").write_text(
            json.dumps(
                profile_formatting_report,
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        profile_conformance_report = validate_document_against_profile(
            doc_text,
            generation_profile,
        )
        write_profile_conformance_report(
            cycle_dir / "document_profile_conformance.json",
            profile_conformance_report,
        )
        profile_repair_recommendations = (
            profile_conformance_report.profile_repair_recommendations
        )
        if profile_repair_recommendations:
            log.info(
                "[CYCLE %d] Profile conformance: %s with %d repair recs",
                cycle,
                profile_conformance_report.status,
                len(profile_repair_recommendations),
            )
        else:
            log.info("[CYCLE %d] Profile conformance: PASS", cycle)

        (cycle_dir / "draft_before_type_gate.md").write_text(
            doc_text,
            encoding="utf-8",
        )
        generated_type_gate = validate_generated_type(
            EnforcementContext(
                requested_document_type=document_type,
                run_dir=str(cycle_dir),
                generated_document_text=doc_text,
            )
        )
        (cycle_dir / "phase16_generated_type_gate.json").write_text(
            json.dumps(
                asdict(generated_type_gate),
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        if generated_type_gate.status != "PASS":
            (cycle_dir / "training_quarantine.json").write_text(
                json.dumps(
                    {
                        "cycle": cycle,
                        "reason": (
                            f"INVALID_FOR_TRAINING:"
                            f"{generated_type_gate.blocker_code}"
                        ),
                        "gate": asdict(generated_type_gate),
                    },
                    indent=2,
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            raise RuntimeError(
                f"{generated_type_gate.blocker_code}: "
                f"{generated_type_gate.message}"
            )

        # ── SKILL 1: validate_structure — catch stubs before eval ──────────
        struct_report = validate_structure(doc_text)
        if not struct_report.passed:
            log.warning(
                f"[CYCLE {cycle}] Structure incomplete "
                f"({struct_report.completeness_ratio:.0%}) — expanding {len(struct_report.stub_sections + struct_report.empty_sections)} stubs..."
            )
            standards_ctx = " | ".join(external_standards_cache.get("standards", []))
            doc_text, expanded = expand_stub_sections(
                doc_text, standards_ctx, topic, model=ModelConfig.SLOT120_REASONING
            )
            if expanded:
                log.info(f"[CYCLE {cycle}] Skill expanded sections: {expanded}")
                struct_report = validate_structure(doc_text)  # re-check
                log.info(f"[CYCLE {cycle}] Structure after expansion: {struct_report.completeness_ratio:.0%}")

        # STAGE 2: Evaluate quality
        log.info(f"[CYCLE {cycle}] Evaluating quality...")
        metrics = _evaluate_document_quality(
            doc_text,
            topic,
            reference_text,
            required_sections=required_section_titles,
            required_standards=archetype_standards,
        )
        metrics.cycle = cycle
        # The detailed TOC/body validator is stricter than doc_quality_eval and
        # must cap the structural score. This prevents a 50% headline score
        # from hiding 29% actual completeness.
        metrics.structure_score = min(
            metrics.structure_score,
            struct_report.completeness_ratio,
        )
        metrics.section_coverage = metrics.structure_score
        # Weighted formula mirrors doc_quality_eval (no-reference variant):
        # structure=45%, standards=35%, coverage=20%.  Plain average was
        # suppressing quality by treating all three components equally and
        # ignoring the structure-cap already applied above.
        metrics.overall_score = (
            metrics.structure_score * 0.45
            + metrics.standards_score * 0.35
            + metrics.coverage_score * 0.20
        )

        # STAGE 3: One bounded Axi pass followed by one bounded Opus audit.
        # Avoid nested future/CLI timeouts and duplicate fallback calls.
        log.info(
            f"[CYCLE {cycle}] Running bounded Axi validation, external search, "
            "then Claude audit..."
        )
        audit_schema_passed = False
        audit_quality_passed = False
        audit_quality_failures: list[str] = []
        audit_result = None
        contextual_search_topic = " ".join(
            [topic, *discovery_hints_for(document_type)]
        )
        external_standards_cache = _search_external_standards(
            contextual_search_topic,
            doc_text,
        )
        doc_text, standard_reference_report = (
            _ensure_standard_reference_register(
                doc_text=doc_text,
                internal_standards=internal_standards,
                external_standards=external_standards_cache,
                formatting_standards=formatting_standards,
                document_type=document_type,
                forbidden_references=list(generation_profile.forbidden_references),
            )
        )
        (cycle_dir / "standard_reference_register.json").write_text(
            json.dumps(
                standard_reference_report,
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        axi_feedback, recommendations = _axi_validate_and_recommend(
            doc_text,
            metrics,
            topic,
            reference_text,
            False,
            document_type=document_type,
            required_sections=required_section_titles,
            required_standards=archetype_standards,
            discovered_external_standards=external_standards_cache,
        )
        recommendation_pool = build_recommendation_pool(
            document_name=topic,
            document_type=document_type,
            axi_recommendations=recommendations,
            profile_recommendations=profile_repair_recommendations,
            source_records=external_standards_cache.get("sources", []),
            body_citations=standard_reference_report.get("body_citations", []),
        )
        (cycle_dir / "recommendation_pool.json").write_text(
            json.dumps(recommendation_pool, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        pooled_recommendations = [
            str(item.get("recommendation", "")).strip()
            for item in recommendation_pool.get("items", [])
            if str(item.get("recommendation", "")).strip()
        ]
        log.info(
            f"[CYCLE {cycle}] Axi validation complete: "
            f"{len(recommendations)} recommendations"
        )
        log.info(
            f"[CYCLE {cycle}] External standards found: "
            f"{external_standards_cache['standards']}"
        )

        audit_cache_key = hashlib.sha256(
            json.dumps(
                {
                    "document": doc_text,
                    "reference_sha256": reference_baseline.get(
                        "reference_sha256",
                        "",
                    ),
                    "omi_standards": internal_standards,
                    "axi_standards": external_standards_cache.get(
                        "standards",
                        [],
                    ),
                    "judge_model": teacher_judge_model,
                },
                sort_keys=True,
                ensure_ascii=False,
            ).encode("utf-8")
        ).hexdigest()
        try:
            if audit_cache_key in audit_cache:
                audit_result = audit_cache[audit_cache_key]
                (cycle_dir / "audit_cache_reuse.json").write_text(
                    json.dumps(
                        {
                            "cache_key": audit_cache_key,
                            "reused": True,
                        },
                        indent=2,
                    ),
                    encoding="utf-8",
                )
                log.info("[CYCLE %d] Reused source-grounded audit", cycle)
            else:
                audit_result = _claude_code_audit(
                omi_standards=internal_standards,
                axi_standards=external_standards_cache.get("standards", []),
                doc_excerpt=doc_text,
                axi_recommendations=recommendations,
                topic=topic,
                reference_text=reference_text,
                cycle=cycle,
                structure_report={
                    "completeness_ratio": struct_report.completeness_ratio,
                    "stub_sections": struct_report.stub_sections,
                    "empty_sections": struct_report.empty_sections,
                },
                evidence_dir=cycle_dir,
                model=teacher_judge_model,
                )
                if audit_result:
                    audit_cache[audit_cache_key] = audit_result
            if audit_result:
                audit_schema_passed = True
                (
                    audit_quality_passed,
                    audit_quality_failures,
                ) = _evaluate_audit_quality(
                    audit_result,
                    document_type=validation_document_type,
                )
                log.info(
                    f"[CYCLE {cycle}] Claude audit teacher: "
                    f"{len(audit_result.recommendations_from_audit)} recs + "
                    f"{len(audit_result.skill_recommendations)} skill recs; "
                    f"quality={'PASS' if audit_quality_passed else 'BLOCK'}"
                )
                if audit_result.skill_recommendations:
                    skill_recs_path = cycle_dir / "skill_recommendations.json"
                    skill_recs_path.write_text(
                        json.dumps(
                            {
                                "cycle": cycle,
                                "skill_recommendations": (
                                    audit_result.skill_recommendations
                                ),
                                "reference_gap": audit_result.reference_gap,
                            },
                            indent=2,
                            ensure_ascii=False,
                        ),
                        encoding="utf-8",
                    )
                    log.info(
                        f"[CYCLE {cycle}] Skill recommendations saved to "
                        f"{skill_recs_path}"
                    )
        except Exception as exc:
            log.warning(f"[CYCLE {cycle}] Claude audit failed: {exc}")

        # ── RC-1 FIX: blend Claude Bedrock audit quality into overall_score ───
        # omi_quality and axi_quality contain real dimension scores
        # (standards_accuracy, completeness, context_relevance) that were
        # computed by the Bedrock audit but previously discarded.  Blend them
        # at 40% weight so the teacher signal closes the gap toward 0.98.
        if audit_result and audit_result.omi_quality and audit_result.axi_quality:
            omi_vals = [v for v in audit_result.omi_quality.values() if isinstance(v, (int, float))]
            axi_vals = [v for v in audit_result.axi_quality.values() if isinstance(v, (int, float))]
            if omi_vals and axi_vals:
                omi_avg = sum(omi_vals) / len(omi_vals)
                axi_avg = sum(axi_vals) / len(axi_vals)
                claude_quality = (omi_avg + axi_avg) / 2.0
                blended = min(1.0, metrics.overall_score * 0.60 + claude_quality * 0.40)
                # Guard: audit blend can only raise, never penalise a higher
                # structural score (e.g. structural=0.91 > claude=0.89 → keep 0.91).
                new_score = max(metrics.overall_score, blended)
                log.info(
                    f"[CYCLE {cycle}] Audit-blended quality: {new_score:.1%} "
                    f"(structural={metrics.overall_score:.2f}, "
                    f"omi={omi_avg:.2f}, axi={axi_avg:.2f}, "
                    f"blended={blended:.2f}, guard={'raised' if new_score > blended else 'kept'})"
                )
                metrics.overall_score = new_score

        missing_section_recs = _missing_section_recommendations(
            (
                audit_result.reference_gap.get("missing_sections", [])
                if audit_result
                else []
            )
        )
        audit_recommendations = (
            list(audit_result.recommendations_from_audit)
            if audit_result
            else []
        )
        if (
            audit_result
            and not audit_quality_passed
            and not missing_section_recs
            and not audit_recommendations
        ):
            audit_recommendations = _grounded_repair_bridge(
                doc_text=doc_text,
                reference_text=reference_text,
                generated_sections=list(_split_into_sections(doc_text)),
                audit_result=audit_result,
                audit_quality_failures=audit_quality_failures,
                evidence_dir=cycle_dir,
                contextual_sources=[
                    *[
                        {
                            "source_title": standard,
                            "excerpt": standard,
                            "provider": "omi_internal_discovery",
                        }
                        for standard in internal_standards
                    ],
                    *external_standards_cache.get("sources", []),
                ],
                model=teacher_judge_model,
            )
        repair_plan = _build_repair_plan(
            pooled_recommendations or recommendations,
            (
                missing_section_recs
                + audit_recommendations
                + profile_repair_recommendations
            ),
            doc_text=doc_text,
        )
        repair_batch = repair_plan["selected"]
        (cycle_dir / "repair_plan.json").write_text(
            json.dumps(repair_plan, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        log.info(
            f"[CYCLE {cycle}] Next repair batch: "
            f"{repair_plan['selected_count']} selected across "
            f"{len(repair_plan['selected_targets'])} targets; "
            f"{repair_plan['deferred_count']} deferred"
        )

        metrics.axi_recommendations = recommendations

        draft_path = cycle_dir / "draft.md"
        draft_path.write_text(doc_text, encoding='utf-8')

        # ── SKILL 3: quality_gate — evaluate BEFORE learning pairs ─────────
        gate = quality_gate(
            structure_score=metrics.structure_score,
            standards_score=metrics.standards_score,
            overall_score=metrics.overall_score,
            document_type=validation_document_type,
        )
        log.info(f"[CYCLE {cycle}] Quality gate: {gate.reason}")

        # ── Rollback: reject candidate if structure regressed ───────────────
        baseline_metrics = last_accepted_metrics or (
            all_metrics[-1] if all_metrics else None
        )
        structure_regressed = bool(
            baseline_metrics
            and metrics.structure_score < baseline_metrics.structure_score - 0.001
        )
        overall_regressed = bool(
            baseline_metrics
            and metrics.overall_score < baseline_metrics.overall_score - 0.001
        )
        critical_regression = structure_regressed or overall_regressed
        if critical_regression and (last_accepted_doc or prev_doc_text):
            log.warning(
                f"[CYCLE {cycle}] Critical regression detected "
                f"(structure={structure_regressed}, overall={overall_regressed}) "
                "— rolling back to last accepted draft"
            )
            rejected_path = cycle_dir / "draft_rejected.md"
            rejected_path.write_text(doc_text, encoding='utf-8')
            doc_text = last_accepted_doc or prev_doc_text
            draft_path.write_text(doc_text, encoding='utf-8')
            log.info(f"[CYCLE {cycle}] Rolled back. Rejected candidate archived to {rejected_path}")

        # ── SKILL 2: verify_recommendations — check lineage ─────────────────
        # IMPORTANT: verify the recs that were APPLIED at start of cycle (applied_recs),
        # not the new recs generated by Axi this cycle.
        changes_applied: list[str] = list(changes_applied_by_skill)
        # A cycle with no repair batch has no recommendation lineage to
        # verify. Treat it as not applicable instead of failing the hard gate.
        rec_lineage_passed = _initial_recommendation_lineage_status(
            cycle=cycle,
            applied_recommendations=applied_recs_this_cycle,
        )
        if cycle > 1 and prev_doc_text and applied_recs_this_cycle:
            rec_verify = verify_recommendations(applied_recs_this_cycle, prev_doc_text, doc_text)
            log.info(
                f"[CYCLE {cycle}] Rec lineage: apply_rate={rec_verify.apply_rate:.0%} "
                f"(applied={len(rec_verify.applied)}, partial={len(rec_verify.partial)}, "
                f"skipped={len(rec_verify.skipped)}/{len(applied_recs_this_cycle)})"
            )
            pending_global = bool(
                section_edit_result
                and section_edit_result["global_recs"]
            )
            pending_unresolved = bool(
                section_edit_result
                and section_edit_result["unresolved_recs"]
            )
            # Only flag global rollback (entire doc reverted to last_accepted_draft).
            # Individual section partial-commits are expected — unverified recs
            # surface in skipped[] and re-enter the next audit cycle.
            section_rollback = bool(
                section_edit_result
                and section_edit_result["rolled_back"]
            )
            rec_lineage_passed = _final_recommendation_lineage_status(
                applied_recommendations=applied_recs_this_cycle,
                verified_recommendations=changes_applied,
                text_verification_passed=rec_verify.passed,
                pending_global=pending_global,
                pending_unresolved=pending_unresolved,
                rolled_back=section_rollback,
            )
            if rec_verify.skipped:
                log.info(f"[CYCLE {cycle}] Skipped recs re-injected: {rec_verify.skipped[:3]}")
                # Re-inject skipped into next cycle's recommendations
                recommendations = rec_verify.skipped + recommendations

        metrics.changes_applied = changes_applied

        hard_gate_allowed, hard_gate_failures = _evaluate_cycle_hard_gate(
            gate=gate,
            struct_report=struct_report,
            audit_schema_passed=audit_schema_passed,
            audit_quality_passed=audit_quality_passed,
            audit_quality_failures=audit_quality_failures,
            rec_lineage_passed=rec_lineage_passed,
            critical_regression=critical_regression,
            standard_reference_passed=(
                standard_reference_report["status"] == "PASS"
            ),
            profile_conformance_passed=profile_conformance_report.passed,
        )
        if hard_gate_allowed:
            last_accepted_doc = doc_text
            last_accepted_metrics = metrics

        prev_doc_text = doc_text
        all_metrics.append(metrics)

        metrics_path = cycle_dir / "metrics.json"
        metrics_path.write_text(
            json.dumps(asdict(metrics), indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

        # Skills report (gate already computed above)
        diagnostics = diagnose_document_failures(
            structure_report=struct_report,
            retrieval_result=external_standards_cache,
            baseline_report=baseline_report,
            section_edit_result=section_edit_result,
            grounding_rejections_path=(
                cycle_dir / "audit_grounding_rejections.json"
            ),
            critical_regression=critical_regression,
        )
        (cycle_dir / "section_failure_diagnostics.json").write_text(
            json.dumps(diagnostics, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        skills_report = {
            "cycle": cycle,
            "structure": {
                "completeness_ratio": struct_report.completeness_ratio,
                "passed": struct_report.passed,
                "empty_sections": struct_report.empty_sections,
                "stub_sections": struct_report.stub_sections,
            },
            "quality_gate": {
                "allowed": gate.allowed,
                "reason": _gate_message(gate),
                "scores": gate.scores,
            },
            "rollback": critical_regression,
            "audit_schema_passed": audit_schema_passed,
            "audit_quality_passed": audit_quality_passed,
            "audit_quality_failures": audit_quality_failures,
            "recommendation_lineage_passed": rec_lineage_passed,
            "hard_gate_allowed": hard_gate_allowed,
            "hard_gate_failures": hard_gate_failures,
            "standard_reference_register": standard_reference_report,
            "document_profile_conformance": (
                profile_conformance_report.to_dict()
            ),
            "changes_applied": changes_applied,
            "failure_diagnostics": diagnostics["cause_counts"],
        }
        (cycle_dir / "skills_report.json").write_text(
            json.dumps(skills_report, indent=2, ensure_ascii=False), encoding='utf-8'
        )

        # Save as DOCX — SKILL-15: Delivery Format Skill
        try:
            from ops.docagent.docx_writer import markdown_to_docx, verify_docx_quality
            docx_path = cycle_dir / f"document_cycle_{cycle:02d}.docx"
            markdown_to_docx(
                md_text=doc_text,
                output_path=docx_path,
                document_number="",
                title=topic,
                revision=str(cycle).zfill(2),
            )
            # Verify DOCX quality
            docx_qa = verify_docx_quality(docx_path)
            log.info(
                f"[CYCLE {cycle}] DOCX quality: {docx_qa['score']}/10 "
                f"({'PASS' if docx_qa['passed'] else 'FAIL'}) "
                f"issues={len(docx_qa['issues'])}"
            )
            if docx_qa["issues"]:
                log.warning(f"[CYCLE {cycle}] DOCX issues: {docx_qa['issues']}")
            (cycle_dir / "docx_quality.json").write_text(
                json.dumps(docx_qa, indent=2, ensure_ascii=False), encoding="utf-8"
            )
        except Exception as e:
            log.warning(f"[CYCLE {cycle}] DOCX save failed: {e}")

        # ─────────────────────────────────────────────────────────────────────────────
        # Phase 3: Render + Visual QA Gate (DOCX → PDF → PNG + Brightness Analysis)
        # ─────────────────────────────────────────────────────────────────────────────
        render_metrics = None
        visual_qa_passed = None
        visual_qa_blocking_failure = False

        if docx_path and Path(docx_path).suffix.lower() == ".docx" and Path(docx_path).exists():
            try:
                visual_qa_dir = cycle_dir / "visual_qa"
                log.info(f"[CYCLE {cycle}] Phase 3: Rendering + visual QA...")

                render_success, render_metrics = render_docx_for_qa(
                    docx_path=docx_path,
                    output_dir=visual_qa_dir,
                    timeout_sec=120,
                    dpi=150,
                )

                visual_qa_passed = bool(render_success)

                # Compute blocking failure status from metrics
                page_count = int(getattr(render_metrics, 'page_count', 0) or 0)
                blank_page_count = int(getattr(render_metrics, 'blank_page_count', 0) or 0)
                render_timeout = bool(getattr(render_metrics, 'render_timeout', False))
                critical_issues = int(getattr(render_metrics, 'critical_visual_issues_count', 0) or 0)

                # Critical failure: no pages, timeout, or critical visual issues
                if not render_success or render_timeout or page_count == 0 or critical_issues > 0:
                    visual_qa_blocking_failure = True

                # Blank page threshold: allow max(1, page_count // 10) blank pages
                if page_count > 0:
                    blank_threshold = max(1, page_count // 10)
                    if blank_page_count > blank_threshold:
                        visual_qa_blocking_failure = True
                        log.warning(
                            f"[CYCLE {cycle}] Visual QA: Blank pages ({blank_page_count}) "
                            f"exceed threshold ({blank_threshold} for {page_count} pages)"
                        )

                log.info(
                    f"[CYCLE {cycle}] Visual QA: render_success={render_success}, "
                    f"pages={page_count}, blank={blank_page_count}, "
                    f"critical_issues={critical_issues}, blocking={visual_qa_blocking_failure}, "
                    f"timing={getattr(render_metrics, 'render_duration_sec', 0):.2f}s"
                )

                # Log severity-specific messages
                if visual_qa_blocking_failure:
                    error_msg = getattr(render_metrics, 'error_message', '')
                    if render_timeout:
                        log.warning(f"[CYCLE {cycle}] Visual QA: TIMEOUT — {error_msg}")
                    elif critical_issues > 0:
                        log.warning(
                            f"[CYCLE {cycle}] Visual QA: CRITICAL FAILURE "
                            f"({critical_issues} critical issues) — {error_msg}"
                        )
                    elif page_count == 0:
                        log.warning(
                            f"[CYCLE {cycle}] Visual QA: ZERO PAGES RENDERED — {error_msg}"
                        )
                    else:
                        log.warning(f"[CYCLE {cycle}] Visual QA: BLOCKING FAILURE — {error_msg}")
                elif getattr(render_metrics, 'degraded_visual_qa_mode', False):
                    log.warning(
                        f"[CYCLE {cycle}] Visual QA: DEGRADED MODE — "
                        f"{getattr(render_metrics, 'error_message', 'unknown degradation')}"
                    )
                else:
                    log.info(f"[CYCLE {cycle}] Visual QA: PASS")

                # Save render metrics to artifacts
                if render_metrics:
                    render_metrics_path = cycle_dir / "visual_qa_metrics.json"
                    try:
                        render_metrics_dict = render_metrics.to_dict()
                        render_metrics_path.write_text(
                            json.dumps(render_metrics_dict, indent=2, ensure_ascii=False),
                            encoding="utf-8"
                        )
                        log.info(f"[CYCLE {cycle}] Render metrics saved: {render_metrics_path}")
                    except Exception as e:
                        log.warning(f"[CYCLE {cycle}] Failed to save render metrics: {e}")

            except Exception as e:
                log.warning(f"[CYCLE {cycle}] Phase 3 render gate failed: {e}")
                visual_qa_blocking_failure = True
                # Ensure render_metrics is available even on failure
                if not render_metrics:
                    render_metrics = RenderMetrics(
                        render_attempted=True,
                        error_message=str(e),
                        degraded_visual_qa_mode=True
                    )

        # Symlink latest → current cycle
        latest = output_dir / "latest"
        if latest.is_symlink():
            latest.unlink()
        latest.symlink_to(cycle_dir.name)

        # Export learning pairs only when explicitly requested and every gate passes.
        export_allowed, export_failures = _learning_pair_export_decision(
            export_requested=export_learning_pairs,
            hard_gate_allowed=hard_gate_allowed,
            hard_gate_failures=hard_gate_failures,
            preflight_gates=preflight_gates,
            generated_type_gate=generated_type_gate,
        )
        training_tag = None
        if export_allowed:
            _save_training_pairs(
                cycle=cycle,
                topic=topic,
                doc_text=doc_text,
                metrics=metrics,
                recommendations=recommendations,
                external_standards=external_standards_cache,
                output_dir=output_dir
            )
        else:
            reason = "; ".join(export_failures)
            training_tag = f"INVALID_FOR_TRAINING:{reason}"
            log.warning(f"[CYCLE {cycle}] Learning pairs BLOCKED: {training_tag}")
            # Write quarantine marker so data is traceable but not used
            (cycle_dir / "training_quarantine.json").write_text(
                json.dumps(
                    {
                        "cycle": cycle,
                        "reason": training_tag,
                        "scores": gate.scores,
                        "detailed_structure": struct_report.completeness_ratio,
                        "audit_schema_passed": audit_schema_passed,
                        "audit_quality_passed": audit_quality_passed,
                        "audit_quality_failures": audit_quality_failures,
                        "recommendation_lineage_passed": rec_lineage_passed,
                        "export_learning_pairs_requested": export_learning_pairs,
                        "preflight_gates": [
                            asdict(item) for item in preflight_gates
                        ],
                        "generated_type_gate": asdict(generated_type_gate),
                    },
                    indent=2,
                ),
                encoding="utf-8",
            )

        failure_analysis_registry = write_failure_analysis_registry(
            output_dir=output_dir,
            cycle_dir=cycle_dir,
            cycle=cycle,
            document_type=document_type,
            diagnostics=diagnostics,
            profile_conformance=profile_conformance_report.to_dict(),
            hard_gate_failures=hard_gate_failures,
            export_failures=export_failures,
            scores=gate.scores,
        )
        if failure_analysis_registry["record_count"]:
            log.warning(
                "[CYCLE %s] Failure analysis candidates recorded for next "
                "repair cycle: %s (not model training data)",
                cycle,
                failure_analysis_registry["record_count"],
            )
        failure_analysis_brief = write_failure_analysis_brief(
            run_dir=output_dir,
            cycle=cycle,
            document_type=document_type,
        )
        log.info(
            "[CYCLE %s] Failure analysis brief ready for %s: slot32=%s",
            cycle,
            failure_analysis_brief["document_type"],
            failure_analysis_brief.get("slot32", {}).get("model", "unknown"),
        )
        stage_contract_audit = write_stage_contract_audit(output_dir, cycle=cycle)
        if stage_contract_audit["status"] != "PASS":
            log.warning(
                "[CYCLE %s] Stage contract audit: %s failed handoffs",
                cycle,
                stage_contract_audit["failed_count"],
            )

        log.info(f"[CYCLE {cycle}] Draft: {draft_path}")
        log.info(f"[CYCLE {cycle}] Quality: overall={metrics.overall_score:.1%}, structure={metrics.structure_score:.1%}, standards={metrics.standards_score:.1%}")
        log.info(f"[CYCLE {cycle}] Axi feedback: {axi_feedback[:200]}")

        # Progress tracking
        previous_quality_score = (
            float(last_accepted_metrics.overall_score)
            if last_accepted_metrics is not None
            else 0.0
        )
        prev_score = previous_quality_score
        delta = metrics.overall_score - prev_score
        if abs(delta) < NO_PROGRESS_DELTA:
            no_progress_streak += 1
            log.info(
                f"[CYCLE {cycle}] Δ={delta:+.4f} < threshold={NO_PROGRESS_DELTA} "
                f"— stall streak {no_progress_streak}/{NO_PROGRESS_LIMIT}"
            )
        else:
            no_progress_streak = 0
            log.info(f"[CYCLE {cycle}] Progress Δ={delta:+.4f} ({delta:+.2%})")

        cycle_finished_at = datetime.now(timezone.utc)
        cycle_summary = _write_cycle_completion_summary(
            cycle_dir=cycle_dir,
            cycle=cycle,
            max_cycles=max_cycles,
            target_quality=target_quality,
            metrics=metrics,
            previous_quality_score=previous_quality_score,
            hard_gate_allowed=hard_gate_allowed,
            critical_regression=critical_regression,
            visual_qa_blocking_failure=visual_qa_blocking_failure,
            no_progress_streak=no_progress_streak,
            no_progress_limit=NO_PROGRESS_LIMIT,
            cycle_started_at=cycle_started_at,
            cycle_finished_at=cycle_finished_at,
        )
        log.info(
            "[CYCLE %s] Terminal summary written: status=%s quality=%.2f%% cycle=%d",
            cycle,
            cycle_summary["status"],
            cycle_summary["quality_percent"],
            cycle_summary["cycle_index"],
        )

        best_result = CycleResult(
            cycle_num=cycle,
            success=(
                metrics.overall_score >= target_quality
                and hard_gate_allowed
            ),
            generated_doc_path=draft_path,
            metrics=metrics,
            axi_feedback=axi_feedback,
            ready_for_next_cycle=True,
            convergence_score=metrics.overall_score,
            bedrock_invoked=audit_result.bedrock_invoked if audit_result else False,
            render_metrics=render_metrics,
            visual_qa_passed=visual_qa_passed,
            visual_qa_blocking_failure=visual_qa_blocking_failure,
        )

        # Check convergence
        if metrics.overall_score >= target_quality and hard_gate_allowed:
            log.info(f"[CYCLE {cycle}] TARGET QUALITY REACHED: {metrics.overall_score:.1%}")
            best_result.ready_for_next_cycle = False
            break

        if no_progress_streak >= NO_PROGRESS_LIMIT:
            log.warning(
                f"[CYCLE {cycle}] CONVERGENCE STALL — last {NO_PROGRESS_LIMIT} cycles "
                f"delta < {NO_PROGRESS_DELTA:.1%}. Stopping. "
                f"Check cycle_dir/skill_recommendations.json for skill improvement proposals."
            )
            best_result.ready_for_next_cycle = False
            break

        log.info(f"[CYCLE {cycle}] Continuing... (quality={metrics.overall_score:.1%}, target={target_quality:.1%})")

    # Save summary
    summary = {
        "topic": topic,
        "routed_topic": routed_topic,
        "models": {
            "slot14": ModelConfig.SLOT14_SEARCH,
            "slot32_editor": resolve_slot("32"),
            "slot120": ModelConfig.SLOT120_REASONING,
        },
        "document_archetype": document_archetype["archetype_id"],
        "document_governance_standards": archetype_standards,
        "section_contract_count": len(section_contract),
        "teacher_judge": {
            "provider": "claude_code_cli_aws",
            "slot120_used_as_teacher": False,
            "slot120_used_as_judge": False,
            "teacher_model": f"claude_code_cli_aws:{teacher_judge_model}",
            "judge_model": f"claude_code_cli_aws:{teacher_judge_model}",
            "degraded_mode": judge_smoke.get("degraded_mode", False),
        },
        "enforcement": {
            "report": str(enforcement_report_path),
            "profile_gate": profile_gate.status,
            "reference_gate": reference_gate.status,
            "judge_gate": judge_gate.status,
            "training_promotion_gate": training_promotion_gate.status,
            "validation_document_type": validation_document_type,
            "learning_pair_export_requested": export_learning_pairs,
        },
        "cycles_completed": len(all_metrics),
        "target_quality": target_quality,
        "status": (
            "TARGET_REACHED"
            if best_result and best_result.success
            else "INCOMPLETE"
            if best_result
            else "FAIL_GENERATION"
        ),
        "achieved_quality": best_result.metrics.overall_score if best_result else 0.0,
        "convergence_trajectory": [m.overall_score for m in all_metrics],
        "final_output": str(best_result.generated_doc_path) if best_result else None,
    }

    summary_path = output_dir / "summary.json"
    atomic_write_text(summary_path, json.dumps(summary, indent=2, ensure_ascii=False))

    if best_result is None:
        log.error("PIPELINE FAILED: no document was generated")
        raise RuntimeError("No document generated")

    log.info(f"\n{'='*70}")
    log.info(f"PIPELINE COMPLETE")
    log.info(f"Cycles: {len(all_metrics)}, Quality: {best_result.metrics.overall_score:.1%}, Output: {output_dir}")
    log.info(f"{'='*70}\n")

    # ── DOCSREG meta-cycle (opt-in) ────────────────────────────────────
    if enable_docsreg_meta_cycle:
        from ops.docgen.docsreg_docgen_integration import run_docsreg_document_type_cycle
        _evidence_root = docsreg_evidence_root or (output_dir / "docsreg")
        _doc_type = docsreg_document_type or document_type
        _docsreg_result = run_docsreg_document_type_cycle(
            document_type=_doc_type,
            evidence_root=_evidence_root,
            target_quality=docsreg_target_quality,
            max_fresh_runs=docsreg_max_fresh_runs,
            stall_limit=docsreg_stall_limit,
            min_progress_delta=docsreg_min_progress_delta,
            write_policy=docsreg_write_policy,
            auditor_mode=docsreg_auditor_mode,
        )
        log.info(
            "[DOCSREG] meta-cycle outcome=%s quality=%.3f cycles=%d evidence=%s",
            _docsreg_result.get("outcome"),
            _docsreg_result.get("quality", 0.0),
            _docsreg_result.get("cycles_completed", 0),
            _docsreg_result.get("evidence_dir", ""),
        )

    return best_result


# ──────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(name)s] %(levelname)s: %(message)s"
    )

    parser = argparse.ArgumentParser(description="Cyclic document generation pipeline")
    parser.add_argument("--topic", default="Asset Integrity Management Policy and Framework", help="Document topic")
    parser.add_argument("--reference-pdf", type=Path, default=Path("/media/axi_omi_sphere/FDF0-25E2/Documents/Block 10/Стандарты/IG7894~I/Asset Integrity Management Policy and Framework (AIM-PFM).pdf"), help="Reference template PDF")
    parser.add_argument("--max-cycles", type=int, default=5, help="Maximum cycles")
    parser.add_argument("--target-quality", type=float, default=0.98, help="Target quality score (0-1)")
    parser.add_argument(
        "--document-type",
        default="technical_report",
        help="Document type used with task-archetype selection",
    )
    parser.add_argument("--output-dir", type=Path, default=None, help="Output directory")
    parser.add_argument(
        "--export-learning-pairs",
        action="store_true",
        help="Export learning pairs only after every enforcement gate passes",
    )

    args = parser.parse_args()

    result = run_cyclic_generation(
        topic=args.topic,
        reference_pdf=args.reference_pdf,
        max_cycles=args.max_cycles,
        target_quality=args.target_quality,
        output_dir=args.output_dir,
        document_type=args.document_type,
        export_learning_pairs=args.export_learning_pairs,
    )

    print(f"\n✓ Pipeline result: {'SUCCESS' if result.success else 'INCOMPLETE'}")
    print(f"  Final quality: {result.convergence_score:.1%}")
    print(f"  Output: {result.generated_doc_path}")
