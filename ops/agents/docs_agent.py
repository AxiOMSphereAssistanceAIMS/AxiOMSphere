#!/usr/bin/env python3
"""DocsAgent — advisory documentation review gate (never blocks).

FastAPI wrapper around docs_review() from review_gates.py.
Port 8015 | POST /review | GET /health | POST /intake/case

Advisory only: approved is always True regardless of concerns.
Folder inbox scanner: polls DOCS_AGENT_INBOX_DIR every 10 s, routes
case packages by mode. Never registers raw uploads directly.
"""
from __future__ import annotations

import asyncio
import json
import os
import shutil
import uuid
from contextlib import asynccontextmanager
from pathlib import Path

from fastapi import FastAPI, HTTPException, Request
from pydantic import BaseModel

import sys
import pathlib
sys.path.insert(0, str(pathlib.Path(__file__).parent.parent))

from self_healing.review_gates import docs_review
from agents.agent_skill_loader import build_doc_skill_runner


AGENT_NAME = "docs_agent"
PORT = 8015

_doc_runner = build_doc_skill_runner("docs-agent")

_INBOX_DIR      = Path(os.environ.get("DOCS_AGENT_INBOX_DIR",    "/data/docs_agent/inbox"))
_PROCESSED_DIR  = Path(os.environ.get("DOCS_AGENT_PROCESSED_DIR", "/data/docs_agent/processed"))
_FAILED_DIR     = Path(os.environ.get("DOCS_AGENT_FAILED_DIR",    "/data/docs_agent/failed"))
_SCAN_INTERVAL  = int(os.environ.get("DOCS_AGENT_SCAN_INTERVAL_SEC", "10"))
_RAW_REG_FORBIDDEN = os.environ.get("DOCS_AGENT_RAW_UPLOAD_REGISTRATION_FORBIDDEN", "1") == "1"


async def _inbox_scanner() -> None:
    """Poll the inbox folder and route case packages by mode."""
    _INBOX_DIR.mkdir(parents=True, exist_ok=True)
    _PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
    _FAILED_DIR.mkdir(parents=True, exist_ok=True)

    while True:
        await asyncio.sleep(_SCAN_INTERVAL)
        try:
            for case_dir in sorted(_INBOX_DIR.iterdir()):
                if not case_dir.is_dir():
                    continue
                manifest_path = case_dir / "manifest.json"
                if not manifest_path.exists():
                    continue
                try:
                    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
                    if _RAW_REG_FORBIDDEN and manifest.get("raw_upload"):
                        mode = manifest.get("mode", "unknown")
                        result = await _route_case(case_dir, manifest, mode)
                        _write_result(case_dir, result)
                        dst = _PROCESSED_DIR / case_dir.name
                        shutil.move(str(case_dir), str(dst))
                except Exception as exc:
                    err_result = {"status": "error", "error": str(exc)}
                    _write_result(case_dir, err_result)
                    dst = _FAILED_DIR / case_dir.name
                    shutil.move(str(case_dir), str(dst))
        except Exception:
            pass


def _write_result(case_dir: Path, result: dict) -> None:
    try:
        (case_dir / "result.json").write_text(
            json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception:
        pass


async def _route_case(case_dir: Path, manifest: dict, mode: str) -> dict:
    """Route a case package to the appropriate stub handler."""
    if mode in ("register", "multi_document_register"):
        return await _handle_register(case_dir, manifest)
    elif mode in ("compare_consolidate", "multi_document_consolidate"):
        return await _handle_consolidate(case_dir, manifest)
    elif mode in ("evaluate_against_standards", "single_document_docreview"):
        return await _handle_evaluate(case_dir, manifest)
    elif mode == "tuning_pair":
        return {"status": "skipped", "reason": "tuning_pair handled by axi-bot training pipeline"}
    else:
        return {"status": "pending_review", "mode": mode}


async def _handle_register(case_dir: Path, manifest: dict) -> dict:
    """Stub: queue documents for registration approval."""
    return {
        "status": "pending_approval",
        "case_id": manifest.get("case_id"),
        "document_count": manifest.get("document_count", 0),
        "selected_pipeline": manifest.get("selected_pipeline"),
        "registration_forbidden": manifest.get("registration_forbidden", True),
        "note": "Awaiting manual approval before registration.",
    }


async def _handle_consolidate(case_dir: Path, manifest: dict) -> dict:
    """Generate a real consolidated output artifact from uploaded documents."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _consolidate_sync, case_dir, manifest)


def _consolidate_sync(case_dir: Path, manifest: dict) -> dict:
    case_id = manifest.get("case_id", case_dir.name)
    docs_dir = case_dir / "documents"
    consolidated_dir = case_dir / "outputs" / "consolidated"
    reports_dir = case_dir / "outputs" / "reports"
    consolidated_dir.mkdir(parents=True, exist_ok=True)
    reports_dir.mkdir(parents=True, exist_ok=True)

    doc_files = sorted(f for f in docs_dir.iterdir() if f.is_file()) if docs_dir.exists() else []
    exts = {f.suffix.lower() for f in doc_files}

    if exts <= {".xlsx"} and exts:
        primary_output = _consolidate_xlsx(doc_files, consolidated_dir, case_id, reports_dir)
        output_format = "xlsx"
    elif exts <= {".docx"} and exts:
        primary_output = _consolidate_docx(doc_files, consolidated_dir, case_id, reports_dir)
        output_format = "docx"
    else:
        primary_output = _consolidate_mixed(doc_files, consolidated_dir, case_id, reports_dir)
        output_format = "md"

    return {
        "status": "completed_pending_user_approval",
        "case_id": case_id,
        "document_count": manifest.get("document_count", 0),
        "primary_output": str(primary_output.relative_to(case_dir)),
        "output_format": output_format,
        "comparison_report": str((reports_dir / "comparison_report.json").relative_to(case_dir)),
        "recommendations": str((reports_dir / "recommendations.json").relative_to(case_dir)),
        "requires_user_approval": True,
    }



def _safe_xlsx_sheet_title(title: str, used: set | None = None) -> str:
    """Return an Excel-safe worksheet title: no invalid chars, max 31 chars, unique."""
    import re

    used = used if used is not None else set()
    raw = str(title or "Sheet").strip()
    safe = re.sub(r'[\\/*?:\[\]]+', "_", raw)
    safe = re.sub(r"\s+", " ", safe).strip(" '\"_")
    if not safe:
        safe = "Sheet"

    # Excel worksheet title limit is 31 chars.
    base = safe[:31]
    candidate = base
    i = 2
    while candidate in used:
        suffix = f"_{i}"
        candidate = base[: 31 - len(suffix)] + suffix
        i += 1

    used.add(candidate)
    return candidate




def _normalize_form_label(label: object) -> str:
    import re
    text = str(label or "").strip()
    text = re.sub(r"\s+", " ", text)
    text = text.rstrip(":：").strip()
    return text.lower()


def _is_form_like_xlsx(doc_files: list) -> bool:
    """Detect field/value Excel forms where col A is label and B/C/D is value."""
    import openpyxl

    keywords = {
        "presenter", "company", "abstract title", "objective", "introduction",
        "methodology", "results", "key findings", "conclusion", "category",
        "sub-category", "subcategory", "ai", "approval", "biography", "bio",
        "email", "phone", "author", "speaker"
    }

    for doc_path in doc_files:
        try:
            wb = openpyxl.load_workbook(doc_path, data_only=True)
        except Exception:
            continue

        for ws in wb.worksheets[:3]:
            label_value_rows = 0
            keyword_hits = 0
            inspected = 0

            for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 120), max_col=min(ws.max_column, 4), values_only=True):
                inspected += 1
                label = row[0] if row else None
                if label is None:
                    continue

                label_text = str(label).strip()
                if not label_text or len(label_text) > 120:
                    continue

                value = None
                for cell in row[1:4]:
                    if cell not in (None, ""):
                        value = cell
                        break

                if value not in (None, ""):
                    label_value_rows += 1

                norm = _normalize_form_label(label_text)
                if any(k in norm for k in keywords):
                    keyword_hits += 1

            if label_value_rows >= 5 and keyword_hits >= 2:
                return True
            if label_value_rows >= 10:
                return True

    return False


def _extract_form_fields_xlsx(doc_path: Path) -> dict:
    """Extract normalized field -> value from form-like XLSX."""
    import openpyxl

    wb = openpyxl.load_workbook(doc_path, data_only=True)
    fields: dict = {}

    for ws in wb.worksheets:
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=min(ws.max_column, 4), values_only=True):
            if not row or row[0] in (None, ""):
                continue

            label = str(row[0]).strip()
            if not label or len(label) > 160:
                continue

            value = ""
            for cell in row[1:4]:
                if cell not in (None, ""):
                    value = str(cell).strip()
                    break

            key = _normalize_form_label(label)
            if not key:
                continue

            # Keep the most complete value for duplicate labels inside one workbook.
            prev = fields.get(key)
            if prev is None or len(value) > len(str(prev.get("value", ""))):
                fields[key] = {
                    "label": label.rstrip(":：").strip(),
                    "value": value,
                    "source": doc_path.name,
                    "sheet": ws.title,
                }

    return fields


def _choose_final_form_value(field: str, candidates: list[dict]) -> tuple[str, str, str]:
    """Return final_value, rationale, action_required."""
    non_empty = [c for c in candidates if str(c.get("value", "")).strip()]
    if not non_empty:
        return "", "No completed value found in any version.", "Complete before submission"

    norm = _normalize_form_label(field)
    narrative_terms = (
        "abstract", "objective", "introduction", "methodology", "result",
        "finding", "conclusion", "summary", "description", "biography", "bio"
    )
    identity_terms = (
        "name", "presenter", "company", "email", "phone", "category",
        "sub-category", "subcategory", "country", "job title", "position"
    )

    if any(t in norm for t in narrative_terms):
        chosen = max(non_empty, key=lambda c: len(str(c.get("value", ""))))
        rationale = f"Selected the most complete narrative value from {chosen.get('source')}."
    elif any(t in norm for t in identity_terms):
        chosen = non_empty[0]
        values = {str(c.get("value", "")).strip() for c in non_empty}
        if len(values) > 1:
            rationale = f"Conflicting short-field values found; selected first complete value from {chosen.get('source')}."
            return str(chosen.get("value", "")).strip(), rationale, "Review conflict manually"
        rationale = f"Selected completed value from {chosen.get('source')}."
    else:
        chosen = max(non_empty, key=lambda c: len(str(c.get("value", ""))))
        rationale = f"Selected the most complete available value from {chosen.get('source')}."

    action_required = ""
    if any(t in norm for t in ("approval", "approved", "ai", "declaration", "consent")):
        action_required = "Confirm manually before submission"

    return str(chosen.get("value", "")).strip(), rationale, action_required


def _consolidate_xlsx_form_documents(doc_files: list, consolidated_dir: Path, case_id: str, reports_dir: Path) -> Path:
    import json
    import openpyxl
    from openpyxl.styles import Font, Alignment
    from datetime import datetime, timezone

    now_iso = datetime.now(timezone.utc).isoformat()
    out_path = consolidated_dir / f"consolidated_{case_id}.xlsx"

    versions = []
    all_keys = []
    seen = set()

    for idx, doc_path in enumerate(doc_files, start=1):
        fields = _extract_form_fields_xlsx(doc_path)
        versions.append({
            "version": f"Version {idx}",
            "filename": doc_path.name,
            "fields": fields,
        })
        for key in fields:
            if key not in seen:
                seen.add(key)
                all_keys.append(key)

    wb = openpyxl.Workbook()
    ws_final = wb.active
    ws_final.title = _safe_xlsx_sheet_title("Final Consolidated Abstract")
    ws_cmp = wb.create_sheet(_safe_xlsx_sheet_title("Comparison Matrix", {ws_final.title}))
    ws_rec = wb.create_sheet(_safe_xlsx_sheet_title("Recommendations", {ws_final.title, ws_cmp.title}))

    # Final Consolidated Abstract
    ws_final.append(["Field", "Final Value", "Source / Rationale"])
    for cell in ws_final[1]:
        cell.font = Font(bold=True)

    comparison_rows = []
    recommendation_rows = []

    for key in all_keys:
        display_label = None
        candidates = []
        version_values = []

        for v in versions:
            item = v["fields"].get(key)
            if item:
                display_label = display_label or item.get("label") or key
                candidates.append({
                    "source": v["filename"],
                    "version": v["version"],
                    "value": item.get("value", ""),
                })
                version_values.append(item.get("value", ""))
            else:
                version_values.append("")

        display_label = display_label or key
        final_value, rationale, action_required = _choose_final_form_value(display_label, candidates)

        ws_final.append([display_label, final_value, rationale])
        comparison_rows.append([
            display_label,
            version_values[0] if len(version_values) > 0 else "",
            version_values[1] if len(version_values) > 1 else "",
            final_value,
            rationale,
            action_required,
        ])

        norm = _normalize_form_label(display_label)
        if not final_value:
            recommendation_rows.append([
                "High",
                display_label,
                "Complete this field before submission.",
                "No completed value was found in any uploaded version.",
            ])
        elif action_required:
            recommendation_rows.append([
                "High" if any(t in norm for t in ("approval", "ai", "declaration")) else "Medium",
                display_label,
                action_required,
                "This field requires manual confirmation or contains possible conflict.",
            ])
        elif any(t in norm for t in ("objective", "methodology", "result", "finding", "conclusion")) and len(final_value) < 120:
            recommendation_rows.append([
                "Medium",
                display_label,
                "Expand technical detail before final submission.",
                "The selected value appears short for a narrative abstract field.",
            ])

    # Comparison Matrix
    ws_cmp.append(["Field", "Version 1", "Version 2", "Selected Final", "Rationale", "Action Required"])
    for cell in ws_cmp[1]:
        cell.font = Font(bold=True)
    for row in comparison_rows:
        ws_cmp.append(row)

    # Recommendations
    ws_rec.append(["Priority", "Field", "Recommendation", "Reason"])
    for cell in ws_rec[1]:
        cell.font = Font(bold=True)

    if not recommendation_rows:
        recommendation_rows.append([
            "Low",
            "Overall",
            "Review final consolidated abstract before submission.",
            "No critical missing fields were automatically detected.",
        ])

    for row in recommendation_rows:
        ws_rec.append(row)

    for ws in [ws_final, ws_cmp, ws_rec]:
        for col in ws.columns:
            letter = col[0].column_letter
            ws.column_dimensions[letter].width = min(max(16, max(len(str(c.value or "")) for c in col[:50]) + 2), 60)
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(out_path)

    comparison_report = {
        "case_id": case_id,
        "generated_at": now_iso,
        "mode": "form_based_xlsx_consolidation",
        "versions": [{"version": v["version"], "filename": v["filename"], "field_count": len(v["fields"])} for v in versions],
        "fields_compared": len(all_keys),
        "comparison_rows": [
            {
                "field": r[0],
                "version_1": r[1],
                "version_2": r[2],
                "selected_final": r[3],
                "rationale": r[4],
                "action_required": r[5],
            }
            for r in comparison_rows
        ],
    }
    reports_dir.mkdir(parents=True, exist_ok=True)
    (reports_dir / "comparison_report.json").write_text(
        json.dumps(comparison_report, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    recommendations_json = {
        "case_id": case_id,
        "generated_at": now_iso,
        "recommendations": [
            {"priority": r[0], "field": r[1], "recommendation": r[2], "reason": r[3]}
            for r in recommendation_rows
        ],
    }
    (reports_dir / "recommendations.json").write_text(
        json.dumps(recommendations_json, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    return out_path


def _consolidate_xlsx(doc_files: list, consolidated_dir: Path, case_id: str, reports_dir: Path) -> Path:
    if _is_form_like_xlsx(doc_files):
        return _consolidate_xlsx_form_documents(doc_files, consolidated_dir, case_id, reports_dir)

    import openpyxl
    from openpyxl.styles import Font
    from datetime import datetime, timezone

    now_iso = datetime.now(timezone.utc).isoformat()
    out_path = consolidated_dir / f"consolidated_{case_id}.xlsx"
    out_wb = openpyxl.Workbook()
    out_wb.remove(out_wb.active)

    source_wbs: list[tuple[str, object]] = []
    doc_meta: list[dict] = []

    for doc_path in doc_files:
        try:
            wb = openpyxl.load_workbook(doc_path, data_only=True)
            source_wbs.append((doc_path.name, wb))
            meta: dict = {"filename": doc_path.name, "sheets": []}
            for sname in wb.sheetnames:
                ws = wb[sname]
                first_row = next(ws.iter_rows(max_row=1, values_only=True), ())
                headers = [str(h) for h in first_row if h is not None]
                meta["sheets"].append({
                    "name": sname,
                    "row_count": ws.max_row,
                    "col_count": ws.max_column,
                    "headers": headers,
                })
            doc_meta.append(meta)
        except Exception as exc:
            doc_meta.append({"filename": doc_path.name, "error": str(exc), "sheets": []})

    # Copy source sheets
    used_sheet_titles: set = set()
    for src_idx, (fname, wb) in enumerate(source_wbs):
        short = f"Src{src_idx + 1}"
        for sname in wb.sheetnames:
            label = _safe_xlsx_sheet_title(f"{short} - {sname}", used_sheet_titles)
            ws_dst = out_wb.create_sheet(label)
            for row in wb[sname].iter_rows():
                for cell in row:
                    ws_dst.cell(row=cell.row, column=cell.column, value=cell.value)

    # Build union of all column headers
    seen_headers: set = set()
    all_headers: list = []
    for meta in doc_meta:
        for sheet in meta.get("sheets", []):
            for h in sheet.get("headers", []):
                if h not in seen_headers:
                    all_headers.append(h)
                    seen_headers.add(h)

    # Consolidated sheet — merged rows with Source column
    ws_con = out_wb.create_sheet(_safe_xlsx_sheet_title("Consolidated", used_sheet_titles))
    ws_con.append(["Source"] + all_headers)
    for cell in ws_con[1]:
        cell.font = Font(bold=True)
    for fname, wb in source_wbs:
        for sname in wb.sheetnames:
            ws_src = wb[sname]
            row_iter = ws_src.iter_rows(values_only=True)
            src_headers = [str(h) for h in (next(row_iter, None) or []) if h is not None]
            for row_vals in row_iter:
                row_dict = {src_headers[i]: row_vals[i] for i in range(min(len(src_headers), len(row_vals)))}
                ws_con.append([fname] + [row_dict.get(h) for h in all_headers])

    # Comparison Summary sheet
    ws_cmp = out_wb.create_sheet(_safe_xlsx_sheet_title("Comparison Summary", used_sheet_titles))
    ws_cmp.append(["File", "Sheet", "Rows", "Columns", "Column Headers (first 10)"])
    for cell in ws_cmp[1]:
        cell.font = Font(bold=True)
    for meta in doc_meta:
        for sheet in meta.get("sheets", []):
            ws_cmp.append([
                meta["filename"],
                sheet["name"],
                sheet["row_count"],
                sheet["col_count"],
                ", ".join(sheet["headers"][:10]),
            ])

    # Recommendations sheet
    recs = _build_recommendations(doc_meta)
    ws_rec = out_wb.create_sheet(_safe_xlsx_sheet_title("Recommendations", used_sheet_titles))
    ws_rec.append(["Priority", "Recommendation"])
    for cell in ws_rec[1]:
        cell.font = Font(bold=True)
    for rec in recs:
        ws_rec.append([rec["priority"], rec["text"]])

    out_wb.save(out_path)

    # comparison_report.json
    hdrs_by_file = {m["filename"]: set(h for s in m.get("sheets", []) for h in s.get("headers", [])) for m in doc_meta}
    common = set.intersection(*hdrs_by_file.values()) if len(hdrs_by_file) > 1 else set()
    comparison_report = {
        "case_id": case_id,
        "generated_at": now_iso,
        "documents": doc_meta,
        "common_columns": sorted(common),
        "unique_to_each": {fn: sorted(hdrs - common) for fn, hdrs in hdrs_by_file.items()},
    }
    (reports_dir / "comparison_report.json").write_text(
        json.dumps(comparison_report, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # recommendations.json
    (reports_dir / "recommendations.json").write_text(
        json.dumps({"case_id": case_id, "generated_at": now_iso, "recommendations": recs},
                   ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return out_path


def _consolidate_docx(doc_files: list, consolidated_dir: Path, case_id: str, reports_dir: Path) -> Path:
    from datetime import datetime, timezone
    now_iso = datetime.now(timezone.utc).isoformat()
    try:
        from docx import Document as DocxDocument

        out_path = consolidated_dir / f"consolidated_{case_id}.docx"
        out_doc = DocxDocument()
        out_doc.add_heading("Consolidated Document", 0)
        out_doc.add_paragraph(f"Generated: {now_iso}  |  Case: {case_id}")
        out_doc.add_paragraph(f"Sources: {', '.join(f.name for f in doc_files)}")

        for doc_path in doc_files:
            out_doc.add_heading(f"Source: {doc_path.name}", level=1)
            try:
                src = DocxDocument(str(doc_path))
                for para in src.paragraphs:
                    if para.text.strip():
                        out_doc.add_paragraph(para.text)
                for tbl in src.tables:
                    out_doc.add_paragraph(f"[Table: {len(tbl.rows)} rows × {len(tbl.columns)} cols]")
            except Exception as exc:
                out_doc.add_paragraph(f"[Could not read {doc_path.name}: {exc}]")

        out_doc.add_heading("Recommendations", level=1)
        for item in ["Review consolidated content for completeness.", "Approve before registering as master document."]:
            out_doc.add_paragraph(item, style="List Bullet")
        out_doc.save(str(out_path))

        (reports_dir / "comparison_report.json").write_text(
            json.dumps({"case_id": case_id, "generated_at": now_iso, "format": "docx",
                        "documents": [f.name for f in doc_files]}, indent=2),
            encoding="utf-8",
        )
        (reports_dir / "recommendations.json").write_text(
            json.dumps({"case_id": case_id, "generated_at": now_iso, "recommendations": [
                {"priority": "low", "text": "Review consolidated content for completeness."},
                {"priority": "low", "text": "Approve before registering as master document."},
            ]}, indent=2),
            encoding="utf-8",
        )
        return out_path
    except ImportError:
        return _consolidate_mixed(doc_files, consolidated_dir, case_id, reports_dir,
                                  note="python-docx not available in container; falling back to Markdown")


def _consolidate_mixed(doc_files: list, consolidated_dir: Path, case_id: str, reports_dir: Path,
                       note: str = "") -> Path:
    from datetime import datetime, timezone
    now_iso = datetime.now(timezone.utc).isoformat()
    exts = {f.suffix.lower() for f in doc_files}
    out_path = consolidated_dir / f"consolidated_{case_id}.md"

    lines = [
        "# Consolidated Report", "",
        f"**Case ID:** {case_id}",
        f"**Generated:** {now_iso}", "",
    ]
    if note:
        lines += [f"**Note:** {note}", ""]
    if len(exts) > 1:
        lines += [
            f"**Warning:** Mixed source formats: {', '.join(sorted(exts))}.",
            "A format-unified consolidation is not possible. This is a textual summary only.", "",
        ]
    lines += ["## Source Documents", ""]
    for f in doc_files:
        lines.append(f"- `{f.name}` ({f.stat().st_size} bytes)")
    lines += [
        "", "## Recommendations", "",
        "- Review each source document for completeness.",
        "- Standardize formats before re-uploading for a unified consolidated output.",
        "- Approve before registering as master document.", "",
    ]
    out_path.write_text("\n".join(lines), encoding="utf-8")

    (reports_dir / "comparison_report.json").write_text(
        json.dumps({"case_id": case_id, "generated_at": now_iso, "format": "mixed",
                    "documents": [f.name for f in doc_files]}, indent=2),
        encoding="utf-8",
    )
    (reports_dir / "recommendations.json").write_text(
        json.dumps({"case_id": case_id, "generated_at": now_iso, "recommendations": [
            {"priority": "high", "text": "Standardize document formats before consolidation."},
            {"priority": "low", "text": "Approve before registering as master document."},
        ]}, indent=2),
        encoding="utf-8",
    )
    return out_path


def _build_recommendations(doc_meta: list) -> list:
    recs: list = []
    if len(doc_meta) < 2:
        recs.append({"priority": "info", "text": "Only one document provided; consolidation is a direct copy."})
        return recs

    total_rows = [(m["filename"], sum(s.get("row_count", 0) for s in m.get("sheets", []))) for m in doc_meta]
    if len(total_rows) >= 2:
        mx_name, mx_rows = max(total_rows, key=lambda x: x[1])
        mn_name, mn_rows = min(total_rows, key=lambda x: x[1])
        if mn_rows > 0 and mx_rows > mn_rows * 1.2:
            recs.append({
                "priority": "high",
                "text": (f"'{mx_name}' has significantly more rows ({mx_rows}) than "
                         f"'{mn_name}' ({mn_rows}). Check for missing data in the smaller file."),
            })

    hdrs_by_file = {m["filename"]: set(h for s in m.get("sheets", []) for h in s.get("headers", [])) for m in doc_meta}
    fnames = list(hdrs_by_file.keys())
    for i in range(len(fnames)):
        for j in range(i + 1, len(fnames)):
            fi, fj = fnames[i], fnames[j]
            only_i = hdrs_by_file[fi] - hdrs_by_file[fj]
            only_j = hdrs_by_file[fj] - hdrs_by_file[fi]
            if only_i:
                recs.append({"priority": "medium",
                              "text": f"Columns only in '{fi}': {', '.join(sorted(only_i)[:5])}. "
                                      f"Verify whether '{fj}' should include them."})
            if only_j:
                recs.append({"priority": "medium",
                              "text": f"Columns only in '{fj}': {', '.join(sorted(only_j)[:5])}. "
                                      f"Verify whether '{fi}' should include them."})

    recs.append({"priority": "low",
                 "text": "Review the 'Consolidated' sheet for duplicates or conflicts before registering."})
    recs.append({"priority": "low",
                 "text": "Approve this document via the review workflow before use as a reference."})
    return recs


async def _handle_evaluate(case_dir: Path, manifest: dict) -> dict:
    """Stub: queue document for standards evaluation."""
    return {
        "status": "pending_review",
        "case_id": manifest.get("case_id"),
        "document_count": manifest.get("document_count", 0),
        "selected_pipeline": manifest.get("selected_pipeline"),
        "note": "Standards evaluation queued.",
    }


@asynccontextmanager
async def lifespan(app: FastAPI):
    scanner_task = asyncio.create_task(_inbox_scanner())
    print(f"{AGENT_NAME} ready on port {PORT} (advisory — never blocks)")
    print(f"  inbox scanner: {_INBOX_DIR} every {_SCAN_INTERVAL}s")
    yield
    scanner_task.cancel()
    try:
        await scanner_task
    except asyncio.CancelledError:
        pass


app = FastAPI(title="DocsAgent", version="1.0.0", lifespan=lifespan)


class ReviewRequest(BaseModel):
    proposal_id: str = ""
    files_changed: list[str] = []
    patch_diff: str = ""
    risk_level: str = "low"
    root_cause: str = ""
    restorative: bool = True
    concept_preserved: bool = True
    test_result: str = "not_run"
    rollback_notes: str = ""


class DocSkillRequest(BaseModel):
    skill: str
    params: dict = {}


class ReviewResponse(BaseModel):
    approved: bool
    concerns: list[str]
    score: float
    reviewer: str
    review_id: str


class IntakeCaseResponse(BaseModel):
    status: str
    case_id: str
    message: str


@app.post("/review", response_model=ReviewResponse)
def review(req: ReviewRequest, request: Request) -> ReviewResponse:
    proposal = req.model_dump()
    result = docs_review(proposal)
    return ReviewResponse(
        approved=result.approved,  # always True
        concerns=result.concerns,
        score=result.score,
        reviewer=result.reviewer,
        review_id=str(uuid.uuid4()),
    )


@app.post("/intake/case", response_model=IntakeCaseResponse)
async def intake_case(request: Request) -> IntakeCaseResponse:
    """Accept a case package path from axi-bot and place it in the inbox."""
    body = await request.json()
    case_path_str = body.get("case_path")
    if not case_path_str:
        raise HTTPException(status_code=400, detail="case_path required")

    case_path = Path(case_path_str)
    if not case_path.is_dir():
        raise HTTPException(status_code=404, detail=f"case_path not found: {case_path}")

    manifest_path = case_path / "manifest.json"
    if not manifest_path.exists():
        raise HTTPException(status_code=400, detail="manifest.json missing from case package")

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    case_id = manifest.get("case_id", case_path.name)

    dst = _INBOX_DIR / case_id
    if not dst.exists():
        shutil.copytree(str(case_path), str(dst))

    return IntakeCaseResponse(
        status="queued",
        case_id=case_id,
        message=f"Case {case_id} placed in inbox for processing.",
    )


@app.post("/doc-skill")
def doc_skill(req: DocSkillRequest):
    """Invoke a document skill scoped to docs-agent's allowed set."""
    if _doc_runner is None:
        raise HTTPException(status_code=503, detail="doc_runner not available")
    try:
        result = _doc_runner.invoke(req.skill, **req.params)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if result.get("status") == "forbidden":
        raise HTTPException(status_code=403, detail=result.get("notes", "forbidden"))
    return result


@app.get("/health")
def health() -> dict:
    return {
        "status": "ok",
        "service": AGENT_NAME,
        "port": PORT,
        "advisory": True,
        "inbox": str(_INBOX_DIR),
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=PORT)
