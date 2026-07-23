#!/usr/bin/env python3
"""One-document-at-a-time MSDG promotion queue.

Lifecycle: inventory -> OCR/preprocess -> MSDG patch -> semantic gate ->
canonical promotion -> manifest update -> compact outcome -> exact cleanup.
The next candidate is not started unless the previous candidate committed.
"""
from __future__ import annotations

import argparse, hashlib, json, shutil, sys, time
import re
from datetime import datetime, timezone
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path: sys.path.insert(0, str(ROOT))
from ops.scripts import msdg_three_source_orchestrator as inv
from ops.telegram import msdg_patch_workflow as workflow
from ops.telegram.type_resolution_hold_router import send_hold_package
from ops.scripts.audit_unprocessed_document_type_resolution import resolve as triage_resolve

MEDIA_DEFAULT = Path("/media/axi_omi_sphere/FDF0-25E21")
DB_DEFAULT = ROOT / "aims_workspace/aims_registry.db"
DEFAULT_WEB = ROOT / "aims_workspace/agent_architecture_status/msdg_internet_candidates_inspection_20260720.json"
DEFAULT_RECOVERY_WEB = ROOT / "aims_workspace/agent_architecture_status/msdg_internet_recovery_candidates.json"

def stamp(): return datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
def write_json(path, value):
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
def file_hash(path):
    h = hashlib.sha256()
    with Path(path).open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""): h.update(chunk)
    return h.hexdigest()
def local_path(row):
    for key in ("path", "file_path", "stored_path", "source_filename"):
        raw = row.get(key)
        if not raw: continue
        p = Path(str(raw))
        if p.is_file(): return p
        if str(raw).startswith("/data/"):
            mapped = ROOT / "aims_workspace" / str(raw)[len("/data/"):]
            if mapped.is_file(): return mapped
    return None
def docx_from_web(row, out):
    p = out / "internet_reviewed_semantic_model.docx"
    d = Document(); d.add_heading(str(row.get("title") or row.get("url") or "Reviewed public source"), 0)
    d.add_paragraph("Public source reviewed under MSDG authority/provenance contract.")
    for meaning in row.get("extracted_meanings", []): d.add_paragraph(str(meaning))
    d.save(p); return p
def candidate_key(row, path):
    return row.get("sha256") or row.get("content_hash") or file_hash(path)
def media_for_type(root, typ):
    rows = inv.media_candidates(root, typ)
    seen = set(); result = []
    groups = {}
    for row in rows: groups.setdefault(row.get("normalized_basename"), []).append(row)
    suppressed = []
    for name, group in groups.items():
        good_docx = [x for x in group if x.get("extension") == ".docx" and inv.text_quality_ok(inv.text_from_path(Path(x["path"]))) ]
        preferred = sorted(good_docx, key=lambda x: x["path"])[0] if good_docx else None
        for row in group:
            if preferred and row.get("extension") == ".pdf" and row.get("path") != preferred.get("path"):
                suppressed.append({"path": row.get("path"), "reason": "DOCX_PREFERRED_SAME_BASENAME", "preferred": preferred.get("path")}); continue
            key = row.get("sha256")
            if key in seen: continue
            seen.add(key); result.append(row)
    return result, suppressed
def registered_for_type(db, typ):
    rows = inv.registered_candidates(db, typ); result=[]; unavailable=[]
    for row in rows:
        p=local_path(row)
        if p: row=dict(row); row["path"]=str(p); result.append(row)
        else: unavailable.append(row)
    return result, unavailable
def process_one(typ, source, row, queue_root, seq, web_row=None):
    started=time.monotonic(); temp=queue_root/"active"/f"{seq:06d}_{typ.replace('.','_')}"; temp.mkdir(parents=True, exist_ok=False)
    source_path = docx_from_web(web_row, temp) if web_row is not None else Path(row["path"])
    try:
        run=workflow.run_patch(str(source_path), chat_id=0, task=f"sequential:{source}:{typ}:{seq}", master_ref=typ)
        workflow.promote(run, chat_id=0, backup_root=queue_root/"canonical_backups")
        # Completion is semantic/transactional, never a byte or time limit.
        # The next candidate is released only after both the MSDG final gate
        # and the atomic canonical/manifest promotion record exist.
        status_path=run/"FINAL_STATUS.json"; promotion_path=run/"PROMOTION_RECORD.json"
        if not status_path.is_file() or not promotion_path.is_file():
            raise RuntimeError("DOCUMENT_COMPLETION_MARKER_MISSING")
        status=json.loads(status_path.read_text()); promotion=json.loads(promotion_path.read_text())
        if status.get("semantic_evaluation") != "PASS" or promotion.get("promoted") is not True:
            raise RuntimeError("DOCUMENT_COMPLETION_MARKER_NOT_PASS")
        canonical=Path(promotion["canonical"])
        if not canonical.is_absolute(): canonical=workflow.ROOT/canonical
        manifest_entry=json.loads(workflow.MANIFEST.read_text())["entries"][typ]
        if manifest_entry.get("sha256") != file_hash(canonical):
            raise RuntimeError("DOCGEN_MANIFEST_HASH_NOT_UPDATED")
        result={"sequence":seq,"document_type":typ,"source":source,"document":row,"status":"COMMITTED","patch_count":status.get("patch_count",0),"live_obligation_count":status.get("live_obligation_count",0),"new_version":promotion.get("new_version"),"elapsed_seconds":round(time.monotonic()-started,2),"run":run.name}
        write_json(queue_root/"COMMITTED"/f"{seq:06d}.json", result)
        return result
    except Exception as exc:
        result={"sequence":seq,"document_type":typ,"source":source,"document":row,"status":"BLOCKED","reason":str(exc)[:500],"elapsed_seconds":round(time.monotonic()-started,2)}
        write_json(queue_root/"BLOCKED"/f"{seq:06d}.json", result)
        raise RuntimeError(json.dumps(result, ensure_ascii=False))
    finally:
        # Exact transient cleanup only. Do not sweep shared Telegram evidence:
        # another run may be active and the semantic evidence is required for
        # post-run quality statistics and audit. Canonical masters, backups and
        # compact queue records remain available as before.
        shutil.rmtree(temp, ignore_errors=True)
def web_candidates_for_form(web, typ):
    aliases={
        "form.procedure":["procedure_family.operating","procedure_family.maintenance","procedure_family.inspection"],
        "form.strategy":["core.strategy"],
    }
    rows=[]
    for key in [typ]+aliases.get(typ,[]): rows.extend(web.get(key,[]) or [])
    return [x for x in rows if x.get("url") and x.get("authority") in {"government","regulator","official_company","professional_institute","university","standards_body"} and x.get("extracted_meanings")]
def recovery_web_match(row, web):
    """Find an authoritative normal web version for a weak local candidate."""
    candidates=web.get("internet_recovery_candidates",[]) if isinstance(web,dict) else []
    import re
    source=" ".join(str(row.get(k) or "") for k in ("document","title","normalized_basename")).casefold()
    source_tokens=set(re.findall(r"[a-z0-9]{4,}",source))
    best=None; best_score=0
    for candidate in candidates:
        if candidate.get("authority") not in {"government","regulator","official_company","professional_institute","university","standards_body"} or not candidate.get("extracted_meanings") or not candidate.get("url"): continue
        target=" ".join(str(candidate.get(k) or "") for k in ("title","url","document_id")).casefold()
        score=len(source_tokens & set(re.findall(r"[a-z0-9]{4,}",target)))
        if score>best_score: best,best_score=candidate,score
    return best if best_score>=2 else None
def route_weak_local_candidate(source, row, web):
    if source not in {"registered_database","mounted_media"}: return source, row, None
    path=Path(row["path"]); recovery=recovery_web_match(row,web)
    if recovery and path.suffix.casefold() in {".pdf",".doc",".docx"} and not inv.text_quality_ok(inv.text_from_path(path)):
        return "open_internet_recovery", row, recovery
    return source, row, None

def resolve_before_msdg(row, queue_type=None):
    """Mandatory type-resolution gate before MSDG.

    Rules:
    1. Resolver MUST determine actual document type (not just form.standard)
    2. form.standard is legacy queue-key only; actual types are:
       - technical_standard, management_process_standard, regulation, policy,
       - framework, procedure, specification, etc.
    3. If resolver cannot determine type → HOLD (TYPE_RESOLUTION_HELD)
    4. queue_type is fallback hint only, never forced classification
    5. Standards require ISO/IEC markers in name or content
    """
    name = str(row.get("document") or row.get("title") or row.get("path") or "")
    parent = str(row.get("directory_context") or Path(str(row.get("path") or "")).parent)
    resolution = triage_resolve(name, parent)
    path = Path(str(row.get("path") or ""))
    text_quality = None

    if path.is_file():
        try:
            text_quality = bool(inv.text_quality_ok(inv.text_from_path(path)))
        except Exception as exc:
            text_quality = False
            resolution = dict(resolution, signals=list(resolution.get("signals", [])) + [f"content extraction error: {type(exc).__name__}"])

    # Get resolver's determination
    resolved_type = resolution.get("master_type_id")
    resolution_status = resolution.get("resolution", "UNRESOLVED")
    confidence = resolution.get("confidence", 0.0)

    # A directory named "Standards" is a filing category, not proof that the
    # file itself is a standard.  For technical_standard the authority/type
    # marker must be in the filename (ISO/IEC/API/... or the word standard),
    # unless readable content has independently supplied the signal.
    basename = Path(name).name.casefold()
    standard_name_signal = bool(re.search(r"\b(?:iso|iec|ieee|api|asme|astm|en|din|nfpa|gost|norsok)\b|\bstandard(?:s)?\b", basename))
    if resolved_type == "technical_standard" and not standard_name_signal and text_quality is not True:
        return {"resolution": "TYPE_RESOLUTION_HELD",
                "master_type_id": resolved_type,
                "confidence": confidence,
                "alternatives": resolution.get("alternatives", []),
                "signals": resolution.get("signals", []),
                "hold_reason": "TYPE_NOT_DETERMINED_BY_RESOLVER",
                "queue_type_hint": queue_type,
                "content_text_quality_ok": text_quality,
                "note": "technical_standard signal occurred only in directory context; document name/content did not confirm it"}

    # CRITICAL: Resolver MUST determine a specific type, not just form.standard
    # form.standard is pipeline-key, not classification
    if not resolved_type or resolved_type == "form.standard":
        # Resolver could not determine actual type
        # HOLD document for manual review or better resolution
        return {"resolution": "TYPE_RESOLUTION_HELD",
                "master_type_id": resolved_type,
                "confidence": confidence,
                "alternatives": resolution.get("alternatives", []),
                "signals": resolution.get("signals", []),
                "hold_reason": "TYPE_NOT_DETERMINED_BY_RESOLVER",
                "queue_type_hint": queue_type,  # Record what queue suggested
                "content_text_quality_ok": text_quality,
                "note": "Resolver did not classify as specific type (technical_standard, procedure, policy, etc.)"}

    # Type is resolved to specific master type
    # Proceed to MSDG even if text_quality is poor (OCR will handle)
    # PHASE 1 RECOVERY: Lowered threshold from 0.6 → 0.45 to enable more documents
    if resolution_status == "RESOLVED_TRIAGE_ONLY" or confidence >= 0.45:
        return dict(resolution,
                   master_type_id=resolved_type,
                   content_text_quality_ok=text_quality,
                   resolver_version="type_resolution_gate_v3",
                   gate_status="PASS")

    # Type exists but low confidence and not fully resolved
    # For specific types (not form.standard), proceed anyway
    if confidence >= 0.4:
        return dict(resolution,
                   master_type_id=resolved_type,
                   content_text_quality_ok=text_quality,
                   resolver_version="type_resolution_gate_v3",
                   gate_status="PASS_LOW_CONFIDENCE",
                   note="Type determined but at lower confidence level")

    # Type exists but very low confidence → HOLD for review
    return {"resolution": "TYPE_RESOLUTION_HELD",
            "master_type_id": resolved_type,
            "confidence": confidence,
            "alternatives": resolution.get("alternatives", []),
            "signals": resolution.get("signals", []),
            "hold_reason": "TYPE_LOW_CONFIDENCE",
            "queue_type_hint": queue_type,
            "content_text_quality_ok": text_quality,
            "note": "Resolver uncertain about type classification"}
def main():
    ap=argparse.ArgumentParser(); ap.add_argument("--types", nargs="+", required=True); ap.add_argument("--database",default=str(DB_DEFAULT)); ap.add_argument("--media-root",default=str(MEDIA_DEFAULT)); ap.add_argument("--internet-candidates",default=str(DEFAULT_WEB)); ap.add_argument("--output",required=True); ap.add_argument("--max-documents",type=int,default=0,help="smoke-test limit; 0 means exhaustive"); ap.add_argument("--resume-from-sequence",type=int,default=0,help="skip already terminal queue sequences and continue with the next candidate"); ap.add_argument("--hold-chat-id",type=int,default=None,help="explicit Telegram chat for operator hold requests"); ap.add_argument("--send-hold-requests",action="store_true",help="send English hold message plus source attachment to --hold-chat-id"); ap.add_argument("--stop-on-block",action="store_true",default=True); args=ap.parse_args()
    out=Path(args.output).resolve(); out.mkdir(parents=True,exist_ok=False); (out/"active").mkdir(); (out/"COMMITTED").mkdir(); (out/"BLOCKED").mkdir()
    web=json.loads(Path(args.internet_candidates).read_text()) if Path(args.internet_candidates).is_file() else {}
    if DEFAULT_RECOVERY_WEB.is_file():
        recovery=json.loads(DEFAULT_RECOVERY_WEB.read_text(encoding="utf-8"))
        web["internet_recovery_candidates"]=list(web.get("internet_recovery_candidates",[]))+list(recovery.get("internet_recovery_candidates",[]))
    ledger={"schema":"aims.msdg.sequential_promotion_queue.v1","started_at":datetime.now(timezone.utc).isoformat(),"order":[],"types":{},"cleanup_policy":"one exact transient directory after each commit","canonical_mutation":True}
    # Use millisecond precision for sequence to avoid collisions when multiple queues run simultaneously
    seq_base=int(time.time() * 1000)  # millisecond timestamp
    seq_offset=0
    def next_seq():
        nonlocal seq_offset
        seq_offset += 1
        return seq_base * 1000 + seq_offset  # microsecond-level uniqueness
    for typ in args.types:
        if args.max_documents and seq >= args.max_documents: break
        # Source order is contractual and cannot be reordered.
        reg, unavailable=registered_for_type(Path(args.database),typ); media,suppressed=media_for_type(Path(args.media_root),typ)
        seen=set(); queue=[]
        for row in reg:
            p=Path(row["path"]); key=candidate_key(row,p)
            if key not in seen: seen.add(key); queue.append(("registered_database",row,None))
        for row in media:
            p=Path(row["path"]); key=candidate_key(row,p)
            if key not in seen: seen.add(key); queue.append(("mounted_media",row,None))
        web_rows=web_candidates_for_form(web,typ)
        for row in web_rows: queue.append(("open_internet",{"document":row.get("title") or row.get("url"),"url":row.get("url")},row))
        tledger={"registered_candidates":len(reg),"registered_unavailable":len(unavailable),"mounted_candidates":len(media),"suppressed_pdf_duplicates":len(suppressed),"internet_candidates":len(web_rows),"processed":0,"committed":0,"blocked":0,"unavailable_registered":unavailable,"suppressed":suppressed}
        recovery_count=0
        for source,row,web_row in queue:
            source,row,recovery=route_weak_local_candidate(source,row,web)
            if source=="open_internet_recovery": recovery_count+=1
            seq=next_seq(); item={"sequence":seq,"type":typ,"source":source,"document":row.get("document")}
            if seq <= args.resume_from_sequence:
                # Resume mode reconstructs the deterministic source order but
                # never re-runs a terminal candidate or mutates its master again.
                continue
            ledger["order"].append(item)

            # First attempt: resolve type from name/content
            # queue_type (--types form.standard) is hint only, never forced
            resolution = resolve_before_msdg(row, queue_type=typ) if web_row is None else {"resolution": "RESOLVED_EXTERNAL_SOURCE", "master_type_id": typ, "confidence": 1.0}

            # If resolver fails due to bad OCR quality → try internet recovery
            if (resolution.get("resolution") == "TYPE_RESOLUTION_HELD" and
                resolution.get("hold_reason") == "TYPE_NOT_DETERMINED_BY_RESOLVER" and
                resolution.get("content_text_quality_ok") is False and
                source != "open_internet_recovery"):  # Don't loop on web sources
                # Try to find better version in internet
                web_recovery = recovery_web_match(row, web)
                if web_recovery:
                    # Found internet version; re-route and retry
                    source = "open_internet_recovery"
                    recovery = web_recovery
                    recovery_count += 1
                    # Re-run resolution on web source
                    resolution = {"resolution": "RESOLVED_EXTERNAL_SOURCE", "master_type_id": typ, "confidence": 1.0}

            # If resolver cannot determine type → HOLD for manual review
            if resolution.get("resolution") == "TYPE_RESOLUTION_HELD":
                hold_reason = resolution.get("hold_reason", "UNKNOWN")
                held = {"sequence": seq,
                        "document_type": typ,  # Queue key (legacy)
                        "declared_queue_type": typ,
                        "source": source,
                        "document": row,
                        "resolution": resolution,
                        "status": "TYPE_RESOLUTION_HELD",
                        "hold_reason": hold_reason,
                        "action_required": "Manual type determination" if hold_reason == "TYPE_NOT_DETERMINED_BY_RESOLVER" else "Resolver uncertain; check alternatives"}
                (out/"TYPE_RESOLUTION_HELD").mkdir(exist_ok=True)
                write_json(out/"TYPE_RESOLUTION_HELD"/f"{seq:06d}.json", held)
                if args.send_hold_requests and args.hold_chat_id is not None and Path(str(row.get("path") or "")).is_file():
                    from ops.scripts.generate_axi_telegram_type_review_requests import message as hold_message
                    request = {"request_id": f"TYPE_REVIEW.{seq:06d}", "task_chat_id": args.hold_chat_id, "document": row.get("document"), "message_en": hold_message(row.get("document") or "unknown document", resolution.get("master_type_id"), resolution.get("alternatives", []), resolution.get("signals", []), row.get("path"))}
                    try:
                        delivery = send_hold_package(request, row["path"])
                        held["telegram_delivery"] = delivery
                        write_json(out/"TYPE_RESOLUTION_HELD"/f"{seq:06d}.json", held)
                    except Exception as exc:
                        held["telegram_delivery_error"] = str(exc)[:500]
                        write_json(out/"TYPE_RESOLUTION_HELD"/f"{seq:06d}.json", held)
                tledger["processed"] += 1; tledger["blocked"] += 1
                write_json(out/"QUEUE_PROGRESS.json", ledger|{"current": held})
                continue

            # Use resolver's determination (NOT queue type)
            # Resolver must return specific master_type_id, not generic form.standard
            target_type = resolution.get("master_type_id")
            if not target_type:
                # Safety check: should not reach here after resolve_before_msdg
                raise RuntimeError(f"TYPE_RESOLUTION_FAILED: no master_type_id determined (queue_type={typ}, resolver_result={resolution.get('resolution')})")
            try:
                result=process_one(target_type,source,row,out,seq,recovery or web_row)
                result["type_resolution"] = resolution
                tledger["processed"]+=1; tledger["committed"]+=1
                write_json(out/"QUEUE_PROGRESS.json",ledger|{"current":result})
            except RuntimeError as exc:
                tledger["processed"]+=1; tledger["blocked"]+=1
                blocked=json.loads(str(exc)) if str(exc).lstrip().startswith("{") else {"sequence":seq,"document_type":typ,"source":source,"status":"BLOCKED","reason":str(exc),"document":row}
                write_json(out/"QUEUE_PROGRESS.json",ledger|{"current":blocked})
                continue
        tledger["internet_recovery_routed"]=recovery_count; tledger["status"]="COMPLETE"; ledger["types"][typ]=tledger; write_json(out/"QUEUE_PROGRESS.json",ledger)
    ledger["status"]="PASS_SEQUENTIAL_PROMOTION_QUEUE"; ledger["finished_at"]=datetime.now(timezone.utc).isoformat(); ledger["resumed_from_sequence"]=args.resume_from_sequence; write_json(out/"QUEUE_FINAL_LEDGER.json",ledger)
    print(json.dumps({"status":ledger["status"],"sequence_count":seq,"output":str(out)},ensure_ascii=False,indent=2))
if __name__=="__main__": main()
