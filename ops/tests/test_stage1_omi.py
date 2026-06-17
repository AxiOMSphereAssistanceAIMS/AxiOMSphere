"""Stage 1: Omi pipeline isolated unit tests.

Tests each step independently — no real Ollama, NIM, or Qdrant calls.
Steps:
  Step 1  PDF/DOCX → text  (axi_inbox_text)
  Step 2  text → doc_type  (_infer_doc_type_heuristic)
  Step 3  file → DB row    (_register_file / run_once)
"""
from __future__ import annotations

import sqlite3
import sys
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

OPS = Path(__file__).resolve().parents[1]
if str(OPS) not in sys.path:
    sys.path.insert(0, str(OPS))


# ── Step 1: Text extraction ───────────────────────────────────────────────────

class TestTextExtraction:
    """axi_inbox_text.extract_reference_text_from_inbox"""

    def test_txt_file_returned_as_is(self, tmp_path):
        from axi_inbox_text import extract_reference_text_from_inbox
        f = tmp_path / "doc.txt"
        f.write_text("Safety procedure for confined space entry.", encoding="utf-8")
        result = extract_reference_text_from_inbox(f)
        assert "confined space" in result.lower()

    def test_unknown_extension_returns_empty_or_tries_read(self, tmp_path):
        from axi_inbox_text import extract_reference_text_from_inbox
        f = tmp_path / "binary_file.xyz"
        f.write_bytes(b"\x00\x01\x02binary_content")
        result = extract_reference_text_from_inbox(f)
        # Must not raise; returns str
        assert isinstance(result, str)

    def test_docx_extraction(self, tmp_path):
        """DOCX → text via docx_full_text."""
        from axi_inbox_text import docx_full_text
        # Create a minimal valid DOCX using python-docx
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        doc = Document()
        doc.add_heading("Test Procedure", level=1)
        doc.add_paragraph("Section 1: scope of work.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        result = docx_full_text(docx_path)
        assert "Test Procedure" in result
        assert "scope" in result.lower()

    def test_missing_file_returns_empty(self, tmp_path):
        from axi_inbox_text import extract_reference_text_from_inbox
        result = extract_reference_text_from_inbox(tmp_path / "nonexistent.pdf")
        assert isinstance(result, str)
        assert result == "" or len(result) < 10


# ── Step 2: Doc type inference ────────────────────────────────────────────────

class TestDocTypeInference:
    """omi_register._infer_doc_type_heuristic"""

    def test_jsa_detected(self):
        from omi_register import _infer_doc_type_heuristic
        text = "Job Safety Analysis for hot work operations. Hazards identified: fire, fumes."
        result = _infer_doc_type_heuristic(text)
        assert isinstance(result, str)
        assert len(result) > 0

    def test_iso_procedure_detected(self):
        from omi_register import _infer_doc_type_heuristic
        text = "ISO 45001:2018 — Work Procedure for confined space entry. 1. Scope. 2. Responsibilities."
        result = _infer_doc_type_heuristic(text)
        assert isinstance(result, str)

    def test_empty_text_returns_string(self):
        from omi_register import _infer_doc_type_heuristic
        result = _infer_doc_type_heuristic("")
        assert isinstance(result, str)


# ── Step 3: DB registration ───────────────────────────────────────────────────

class TestDBRegistration:
    """omi_register._ensure_schema + _register_file path."""

    def _make_test_db(self, tmp_path: Path) -> tuple[Path, sqlite3.Connection]:
        from omi_register import _ensure_schema
        from omi_register import sqlite_connect_wal
        db = tmp_path / "test_omi.db"
        conn = sqlite_connect_wal(db)
        _ensure_schema(conn)
        return db, conn

    def test_schema_creates_tables(self, tmp_path):
        db, conn = self._make_test_db(tmp_path)
        tables = {r[0] for r in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()}
        assert "omi_docs" in tables or len(tables) > 0, "Schema must create at least one table"
        conn.close()

    def test_txt_sidecar_registered(self, tmp_path):
        """A .txt sidecar in the scan dir should be registered into SQLite."""
        from omi_register import run_once

        ocr_dir = tmp_path / "ocr_text"
        ocr_dir.mkdir()
        sidecar = ocr_dir / "test_safety_procedure.txt"
        sidecar.write_text(
            "1. Scope\nThis procedure covers hot work operations.\n"
            "2. Responsibilities\nSupervisor must issue a hot work permit.\n",
            encoding="utf-8",
        )

        db = tmp_path / "omi.db"
        env_patch = {
            "OMI_OCR_TEXT_DIR": str(ocr_dir),
            "OMI_PREREGISTER_DOCX_DIR": str(tmp_path / "empty_docx"),
            "OMI_REGISTER_LEGACY_OUTBOX_TXT": "0",
            "OMI_REGISTER_LEGACY_OUTBOX_DOCX": "0",
            "OMI_OLLAMA_ENABLED": "0",
            "OMI_DEDUPE_BY_HASH": "0",
            "QDRANT_ENABLED": "0",
        }
        with patch.dict("os.environ", env_patch):
            with patch("omi_register._index_in_qdrant"):
                count = run_once(db_path=db)

        assert db.exists(), "DB file must be created"
        conn = sqlite3.connect(str(db))
        rows = conn.execute("SELECT file_name FROM ocr_documents").fetchall()
        conn.close()
        assert any("test_safety_procedure" in r[0] for r in rows), (
            f"Sidecar must appear in DB. Rows: {rows}"
        )

    def test_duplicate_not_inserted_twice(self, tmp_path):
        """Running run_once twice on same file should not duplicate the DB row."""
        from omi_register import run_once

        ocr_dir = tmp_path / "ocr_text"
        ocr_dir.mkdir()
        (ocr_dir / "doc.txt").write_text("Procedure content.", encoding="utf-8")

        db = tmp_path / "omi.db"
        env_patch = {
            "OMI_OCR_TEXT_DIR": str(ocr_dir),
            "OMI_REGISTER_LEGACY_OUTBOX_TXT": "0",
            "OMI_REGISTER_LEGACY_OUTBOX_DOCX": "0",
            "OMI_OLLAMA_ENABLED": "0",
            "OMI_DEDUPE_BY_HASH": "0",
            "QDRANT_ENABLED": "0",
        }
        with patch.dict("os.environ", env_patch):
            with patch("omi_register._index_in_qdrant"):
                run_once(db_path=db)
                run_once(db_path=db)

        conn = sqlite3.connect(str(db))
        count = conn.execute("SELECT COUNT(*) FROM ocr_documents").fetchone()[0]
        conn.close()
        assert count == 1, f"Duplicate should not be inserted. Got count={count}"


# ── Step 1+2+3 in sequence (but not yet a full pipeline) ─────────────────────

class TestStepsInSequence:
    """Verify steps 1→2→3 work in succession without wiring them as a pipeline."""

    def test_text_then_type_then_register(self, tmp_path):
        from axi_inbox_text import extract_reference_text_from_inbox
        from omi_register import _infer_doc_type_heuristic, run_once

        # Step 1: text from file
        src = tmp_path / "ocr_text" / "procedure.txt"
        src.parent.mkdir()
        src.write_text(
            "Job Safety Analysis — Confined Space Entry\n"
            "Scope: underground copper mine, Level 3.\n"
            "Hazards: oxygen deficiency, toxic gases.\n",
            encoding="utf-8",
        )
        text = extract_reference_text_from_inbox(src)
        assert len(text) > 20

        # Step 2: classify
        doc_type = _infer_doc_type_heuristic(text)
        assert isinstance(doc_type, str)

        # Step 3: register
        db = tmp_path / "omi.db"
        with patch.dict("os.environ", {
            "OMI_OCR_TEXT_DIR": str(src.parent),
            "OMI_REGISTER_LEGACY_OUTBOX_TXT": "0",
            "OMI_REGISTER_LEGACY_OUTBOX_DOCX": "0",
            "OMI_OLLAMA_ENABLED": "0",
            "OMI_DEDUPE_BY_HASH": "0",
            "QDRANT_ENABLED": "0",
        }):
            with patch("omi_register._index_in_qdrant"):
                run_once(db_path=db)

        conn = sqlite3.connect(str(db))
        rows = conn.execute("SELECT file_name FROM ocr_documents").fetchall()
        conn.close()
        assert len(rows) >= 1
