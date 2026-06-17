from __future__ import annotations

import sqlite3
from pathlib import Path

from docx import Document

from ops.aims_report import build_registry_docx, is_registry_request
from ops.omi_telegram.omi_storage import StorageManager


def _make_docx(path: Path, text: str = "Test document") -> Path:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))
    return path


def test_registry_request_detector_recognizes_report_prompt() -> None:
    assert is_registry_request("Покажи AIMS registry report")


def test_build_registry_docx_includes_master_layout_summary(tmp_path) -> None:
    workspace = tmp_path / "workspace"
    master = workspace / "master"
    master_dir = master / "P01_Asset_Strategy"
    master_dir.mkdir(parents=True)
    _make_docx(master_dir / "example.docx")
    (master / "P02_Risk_Assessment").mkdir(parents=True)

    db_path = workspace / "aims_registry.db"
    conn = sqlite3.connect(db_path)
    conn.execute(
        """
        CREATE TABLE documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            file_name TEXT,
            date_added TEXT,
            aims_process TEXT,
            aims_element TEXT
        )
        """
    )
    conn.executemany(
        "INSERT INTO documents (title, file_name, date_added, aims_process, aims_element) VALUES (?,?,?,?,?)",
        [
            ("Policy", "policy.docx", "2026-06-16T08:00:00", "P01", "E01"),
            ("Risk", "risk.docx", "2026-06-16T08:10:00", "P02", "E09"),
        ],
    )
    conn.commit()
    conn.close()

    out = build_registry_docx(db_path=db_path, master_root=master)
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs if p.text)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "Master Layout Summary" in text
    assert "P01_Asset_Strategy" in table_text
    assert "Files on Disk" in table_text
    assert "Registered Docs" in table_text
    assert "2" in table_text


def test_storage_manager_register_generated_bundle_records_process(tmp_path) -> None:
    workspace = tmp_path / "workspace"
    generated = workspace / "generated"
    generated.mkdir(parents=True)
    db_path = workspace / "aims_registry.db"
    docx_path = _make_docx(generated / "bundle.docx")

    storage = StorageManager(db_path=db_path, workspace=workspace)
    result = storage.register_generated_bundle(docx_path, aims_process="P01")

    assert "registered" in result.lower() or "зарегистр" in result.lower()
    conn = sqlite3.connect(db_path)
    row = conn.execute(
        "SELECT file_name, aims_process, source FROM documents WHERE file_path=?",
        (str(docx_path.resolve()),),
    ).fetchone()
    conn.close()

    assert row[0] == "bundle.docx"
    assert row[1] == "P01"
    assert row[2] == "docgen_bundle"
