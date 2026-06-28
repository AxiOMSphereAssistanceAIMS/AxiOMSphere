#!/usr/bin/env python3
"""
omi_bot.py
──────────
Omi — AIMS Database Administrator
Telegram Bot с живым чатом + текстовыми командами.

Запуск:
    python omi_bot.py

Переменные окружения: файл **.env** в каталоге `ops/omi_telegram`:
    TELEGRAM_BOT_TOKEN или OMI_TELEGRAM_TOKEN — токен @BotFather
    TELEGRAM_GROUP_ID    — ID группы (необязательно)
    OLLAMA_BASE_URL      — http://localhost:11434
    OMI_DB_PATH          — aims_registry.db (знаниевая БД Omi)
    AIMS_WORKSPACE       — /aims_workspace
    OMI_DIALOG_LOG_MAX   — число последних сообщений в чате (реплик user+assistant), передаваемых в LLM (по умолчанию 10)
    OMI_DIALOG_STORE_CHARS — макс. длина одной реплики при сохранении (по умолчанию 4000; JSON/длинные ответы — поднимайте)
"""

import os
import sys
import json
import asyncio
import logging
import secrets
import sqlite3
import time
import re
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

from omi_env import load_omi_dotenv
from aims_paths import workspace_root
from document_dialogue_gap_gate import (
    should_run_document_dialogue_gap_check,
    read_document_dialogue_gap_gate_state,
)

load_omi_dotenv()

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.error import Conflict as TelegramConflict, NetworkError as TelegramNetworkError


def _parse_allowed_chat_ids() -> frozenset[int]:
    raw = (
        os.environ.get("OMI_ALLOWED_CHAT_IDS", "").strip()
        or os.environ.get("OMI_ALLOWED_USER_IDS", "").strip()
    )
    if not raw:
        return frozenset()
    out: set[int] = set()
    for part in raw.split(","):
        part = part.strip()
        if not part:
            continue
        try:
            out.add(int(part))
        except ValueError:
            pass
    return frozenset(out)


ALLOWED_CHATS = _parse_allowed_chat_ids()
# Personal chat IDs that should always be able to talk to Omi directly,
# even if ALLOWED_CHATS is restricted.
OMI_DIRECT_CHAT_IDS: frozenset[int] = frozenset({
    8647744602,
})

# В группах по умолчанию отвечаем только если есть @username бота или текст начинается с «omi».
# OMI_GROUP_REQUIRE_MENTION=0 — отвечать на любое текстовое сообщение в разрешённых чатах (шумнее).
OMI_GROUP_REQUIRE_MENTION = os.environ.get(
    "OMI_GROUP_REQUIRE_MENTION", "1"
).strip().lower() in ("1", "true", "yes", "on")

# Как у Axi по политике «сообщение → фоновый прогрев Ollama»: в группе при каждом тексте
# планируется warm для OMI_MODEL (GET /api/ps + ollama_schedule_background_warm), даже если
# OMI_GROUP_REQUIRE_MENTION=1 и ответа бота не будет. 0 — не вызывать прогрев из omi-bot.
OMI_GROUP_WARM_ON_MESSAGE = os.environ.get(
    "OMI_GROUP_WARM_ON_MESSAGE", "1"
).strip().lower() in ("1", "true", "yes", "on")

# If disabled, Omi does not emit upload/duplicate/status messages in groups.
OMI_GROUP_UPLOAD_NOTIFICATIONS = os.environ.get(
    "OMI_GROUP_UPLOAD_NOTIFICATIONS", "0"
).strip().lower() in ("1", "true", "yes", "on")

# If enabled, Omi ignores all user-originated interactions in group chats.
# Backend jobs (e.g., registry polling / cross-handoff) still run.
OMI_GROUP_SILENT_MODE = os.environ.get(
    "OMI_GROUP_SILENT_MODE", "1"
).strip().lower() in ("1", "true", "yes", "on")

# QWEN_PC_ASSIST_STACK: legacy warm-up toggle; do not enable unless explicitly needed.
QWEN_PC_ASSIST_STACK = os.environ.get("QWEN_PC_ASSIST_STACK", "0").strip().lower() in (
    "1",
    "true",
    "yes",
    "on",
)
QWEN_PC_ASSIST_WARM_ON_TELEGRAM = os.environ.get(
    "QWEN_PC_ASSIST_WARM_ON_TELEGRAM", "1"
).strip().lower() in ("1", "true", "yes", "on")

# TODO(AIMS_CONTEXT_VIRTUALIZATION_LAYER_V1): virtualize oversized Telegram
# text/uploads through ops.context.payload_broker before chat/model handoff.
# Omi has several document-ingest routes; patch each route with a context
# handle before enabling large user-upload pass-through.


def _parse_owner_chat_ids() -> frozenset[int]:
    """Чаты владельца: могут добавлять правила поведения бота (OMI_OWNER_CHAT_IDS)."""
    raw = os.environ.get("OMI_OWNER_CHAT_IDS", "").strip()
    if not raw:
        return frozenset()
    out: set[int] = set()
    for part in raw.split(","):
        part = part.strip()
        if not part:
            continue
        try:
            out.add(int(part))
        except ValueError:
            pass
    return frozenset(out)


OWNER_CHATS = _parse_owner_chat_ids()
DOCGEN_UPGRADE_PRIVATE_CHAT_IDS: frozenset[int] = frozenset({8507628575}) | OWNER_CHATS


def _omi_dialog_int(name: str, default: int, lo: int, hi: int) -> int:
    raw = os.environ.get(name, "").strip()
    if not raw:
        return default
    try:
        v = int(raw)
    except ValueError:
        return default
    return max(lo, min(v, hi))


# Rolling dialog in Application.bot_data (пока жив процесс): последние N сообщений → контекст LLM.
# N = OMI_DIALOG_LOG_MAX (по умолчанию 10 реплик: чередуются user/assistant).
# При длинных темах поднимайте OMI_DIALOG_STORE_CHARS (обрезка одной реплики).
OMI_DIALOG_LOG_KEY = "_omi_recent_dialog"
OMI_DIALOG_LOG_MAX = _omi_dialog_int("OMI_DIALOG_LOG_MAX", 10, 2, 40)
OMI_DIALOG_STORE_CHARS = _omi_dialog_int("OMI_DIALOG_STORE_CHARS", 4000, 400, 50000)


def _omi_dialog_messages_for_llm(ctx, chat_id: int | None) -> list[dict]:
    if chat_id is None:
        return []
    store = ctx.application.bot_data.setdefault(OMI_DIALOG_LOG_KEY, {})
    raw = list(store.get(chat_id) or [])
    out: list[dict] = []
    for row in raw:
        if not isinstance(row, dict):
            continue
        role = row.get("role")
        content = row.get("content")
        if role in ("user", "assistant") and isinstance(content, str) and content.strip():
            out.append({"role": role, "content": content.strip()})
    # На случай устаревшего/битого store — не больше OMI_DIALOG_LOG_MAX последних сообщений
    if len(out) > OMI_DIALOG_LOG_MAX:
        out = out[-OMI_DIALOG_LOG_MAX:]
    return out


def _omi_dialog_append(ctx, chat_id: int | None, role: str, content: str) -> None:
    if chat_id is None:
        return
    text = (content or "").strip()
    if not text:
        return
    if len(text) > OMI_DIALOG_STORE_CHARS:
        text = text[: OMI_DIALOG_STORE_CHARS - 1] + "…"
    label = "user" if role == "user" else "assistant"
    store = ctx.application.bot_data.setdefault(OMI_DIALOG_LOG_KEY, {})
    seq = list(store.get(chat_id) or [])
    seq.append({"role": label, "content": text})
    store[chat_id] = seq[-OMI_DIALOG_LOG_MAX:]


def _is_owner(update: Update) -> bool:
    if not OWNER_CHATS:
        return False
    if update.effective_chat is None:
        return False
    return update.effective_chat.id in OWNER_CHATS


def _is_docgen_upgrade_origin(update: Update) -> bool:
    if update.effective_chat is None:
        return False
    if update.effective_chat.type != "private":
        return False
    return update.effective_chat.id in DOCGEN_UPGRADE_PRIVATE_CHAT_IDS


def _format_docgen_upgrade_reply(
    *,
    action_id: str,
    action_path: str,
    execution: dict[str, object],
    run_summary: dict[str, object],
    doc_types: list[str],
) -> str:
    loop_state = str(run_summary.get("loop_state") or execution.get("status") or "unknown")
    shelf_state = str(run_summary.get("shelf_state") or "unknown")
    next_step = str(run_summary.get("next_step") or "unknown")
    target_quality = float(run_summary.get("target_quality") or 0.98)
    best_quality = float(run_summary.get("best_quality") or 0.0)
    best_cycles = int(run_summary.get("best_cycles") or 0)
    avg_quality = float(run_summary.get("avg_quality") or 0.0)
    trace = str(run_summary.get("trajectory_trace") or "").strip() or "n/a"
    launcher_status = str(execution.get("status") or "unknown")
    exit_code = str(execution.get("exit_code") or "unknown")
    evidence_dir = str(execution.get("evidence_dir") or "n/a")
    training_review = bool(run_summary.get("training_review_required"))
    promotion_ready = bool(run_summary.get("promotion_ready"))
    blocked_reasons = [
        str(item)
        for item in (run_summary.get("blocked_reasons") or [])
        if str(item).strip()
    ]
    loop_completed = best_cycles > 0 or best_quality > 0.0
    branch_markers = [
        f"{item.get('document_type')}: {float(item.get('achieved_quality') or 0.0):.1%} / {int(item.get('cycles_completed') or 0)}c"
        for item in run_summary.get("batch_summaries") or []
        if isinstance(item, dict)
    ]
    branch_text = " | ".join(branch_markers) if branch_markers else "n/a"
    doc_types_text = ", ".join(doc_types)
    if loop_state in {"BLOCKED", "REVIEW"} and best_cycles <= 0 and best_quality <= 0.0:
        heading = "❌ DOCGEN runtime preflight failed."
    elif loop_state in {"COMPLETE", "TARGET_REACHED"} or best_quality >= target_quality:
        heading = "✅ DOCGEN upgrade completed."
    elif best_cycles > 0:
        heading = "🔵 DOCGEN runtime started."
    else:
        heading = "🟡 DOCGEN upgrade job accepted."
    blocker_text = ", ".join(blocked_reasons) if blocked_reasons else "n/a"
    return (
        f"{heading}\n"
        f"Action: `{action_id}`\n"
        f"Path: `{action_path}`\n"
        f"Types: `{doc_types_text}`\n"
        f"Loop state: `{loop_state}`\n"
        f"Shelf: `{shelf_state}`\n"
        f"Best quality: `{best_quality:.1%}` / target `{target_quality:.1%}`\n"
        f"Avg quality: `{avg_quality:.1%}`\n"
        f"Cycles: `{best_cycles}`\n"
        f"Trace: `{trace}`\n"
        f"Branches: `{branch_text}`\n"
        f"Training review: `{str(training_review).lower()}`\n"
        f"Promotion ready: `{str(promotion_ready).lower()}`\n"
        f"Next: `{next_step}`\n"
        f"Blocker: `{blocker_text}`\n"
        f"Launcher status: `{launcher_status}`\n"
        f"Exit code: `{exit_code}`\n"
        f"Evidence: `{evidence_dir}`"
    )


def _chat_allowed(update: Update) -> bool:
    # Keep Omi available in direct/private chats.
    if update.effective_chat is not None and update.effective_chat.type == "private":
        return True
    if OMI_GROUP_SILENT_MODE and update.effective_chat is not None:
        if update.effective_chat.type in ("group", "supergroup"):
            return False
    if not ALLOWED_CHATS:
        return True
    if update.effective_chat is None:
        return False
    return update.effective_chat.id in ALLOWED_CHATS


def _reply_lang(text: str) -> str:
    """Язык ответа по тексту пользователя: кириллица → ru, иначе en."""
    if not (text or "").strip():
        return "en"
    return "ru" if any("\u0400" <= ch <= "\u04ff" for ch in text) else "en"



def _registry_without_time_window(s: str) -> bool:
    """Запрос реестра без явного окна (часы/дни/неделя) и без дат в тексте."""
    low = (s or "").lower()
    if any(k in low for k in ("last", "past", "последн", "today", "сегодня", "yesterday", "вчера")):
        return False
    if any(k in low for k in ("week", "weeks", "недел")):
        return False
    if any(k in low for k in ("month", "months", "месяц")):
        return False
    if any(unit in low for unit in (" hour", " hours", " day", " days", " hr", " hrs", "час", "дн", "день", "дней")):
        return False
    if "-" in (s or "") and any(ch.isdigit() for ch in (s or "")):
        return False
    return True


def _normalize_tokens(text: str) -> list[str]:
    t = (text or "").lower()
    for ch in ",.:;!?()[]{}\"'`/\t\r\n":
        t = t.replace(ch, " ")
    return [x for x in t.split(" ") if x]


def _strip_omi_prefix(text: str) -> str:
    clean = (text or "").strip()
    low = clean.lower()
    for pref in ("omi ", "omi,", "omi:", "оми ", "оми,", "оми:"):
        if low.startswith(pref):
            return clean[len(pref):].strip()
    return clean


def _looks_like_register_file_short_request(text: str) -> bool:
    """
    User means the verb: register *a file* into AIMS — not the noun 'registry'.
    Short messages only (avoid matching long strategy texts).
    """
    t = (text or "").strip().lower()
    if not t or len(t) > 160:
        return False
    phrases = (
        "register file",
        "register the file",
        "register this file",
        "register uploaded",
        "зарегистрировать файл",
        "регистрация файла",
        "зарегистрируй файл",
        "зарегистрируйте файл",
        "добавь файл в реестр",
        "добавить файл в реестр",
        "добавить файл в базу",
        "добавь в реестр файл",
        "add file to registry",
        "add file to the registry",
    )
    if not any(p in t for p in phrases):
        return False
    return len(t.split()) <= 16


def _is_command_for_other_bot(clean: str, bot_username: str | None) -> bool:
    if not clean.startswith("/") or "@" not in clean:
        return False
    cmd_head = clean.split(" ", 1)[0]
    if "@" not in cmd_head:
        return False
    target = cmd_head.split("@", 1)[1].strip().lower()
    own = (bot_username or "").strip().lower()
    return bool(target and own and target != own)


def _mentions_axi(text: str) -> bool:
    tokens = _normalize_tokens(text)
    for tok in tokens:
        if tok == "axi" or tok == "акси" or tok.startswith("@axi") or tok.startswith("axi_"):
            return True
    return False


def _mentions_omi(text: str) -> bool:
    tokens = _normalize_tokens(text)
    return "omi" in tokens or "оми" in tokens


def _looks_like_task_then_file_flow(text: str) -> bool:
    low = (text or "").lower()
    analysis_cues = (
        "analy", "review", "check", "assess", "gap",
        "анализ", "проверь", "провер", "оцен", "разбер", "ревью",
    )
    file_cues = (
        "file", "document", "attachment", "attached", "upload",
        "файл", "документ", "влож", "прикреп", "загруж",
    )
    has_analysis = any(k in low for k in analysis_cues)
    has_file_ref = any(k in low for k in file_cues)
    return has_analysis and has_file_ref


def _looks_like_doc_synthesis_request(text: str) -> bool:
    low = (text or "").lower().strip()
    if not low or len(low) > 220:
        return False
    if any(marker in low for marker in ("remove ai labels", "удали метки", "clean up labels")):
        return False
    doc_cues = (
        "policy",
        "framework",
        "procedure",
        "instruction",
        "report",
        "specification",
        "manual",
        "document",
        "doc",
        "политик",
        "рамк",
        "процедур",
        "инструкц",
        "отч",
    )
    create_cues = (
        "create",
        "generate",
        "draft",
        "prepare",
        "write",
        "develop",
        "make",
        "сделай",
        "создай",
        "подготовь",
        "разработай",
        "составь",
    )
    return any(cue in low for cue in create_cues) and any(cue in low for cue in doc_cues)


# ── Pending-confirmation helpers ────────────────────────────────────────────

_PENDING_CONFIRM_TTL = 300  # seconds — window to accept a bare affirmative

_AFFIRMATIVES: frozenset[str] = frozenset({
    "yes", "yeah", "yep", "yup", "sure", "ok", "okay", "ок", "oke",
    "да", "ладно", "конечно", "верно", "правильно",
    "go", "proceed", "apply", "confirm", "confirmed",
    "подтверждаю", "подтвердить", "подтверди", "подтверждено",
    "давай", "давайте", "продолжай", "продолжить", "выполнить", "выполняй",
    "rename", "all", "все", "всё",
    "+", "👍", "✅",
})


def _is_short_affirmative(text: str) -> bool:
    """
    Return True if the message is a brief reply intended as a response —
    affirmatives, short choices (digits, "1", "2 all"), or very short phrases.
    Used to bypass the group mention filter when Omi is waiting for input.
    """
    t = text.strip().lower().rstrip(".,!?")
    if t in _AFFIRMATIVES:
        return True
    # Single digit or "1 all", "2" style option selection
    if t.isdigit():
        return True
    tokens = [tok.strip(".,!?") for tok in t.split() if tok.strip(".,!?")]
    if not tokens:
        return False
    # Short multi-token: up to 4 tokens where at least one is a digit or affirmative
    if len(tokens) <= 4:
        has_digit = any(tok.isdigit() for tok in tokens)
        has_aff = any(tok in _AFFIRMATIVES for tok in tokens)
        if has_digit or has_aff:
            return True
    return False


_CONFIRM_REQUEST_MARKERS = (
    "_confirm_", "shall i", "should i", "want me to",
    "подтверди", "confirm?", "confirm:", "proceed?", "продолжить?",
    "переименовать?", "rename all", "rename them",
    "сделать это", "выполнить?", "применить?", "обработать?",
    "confirmed the list", "before processing",
    "yes` /", "да` /",  # Confirm: `yes` / `no` pattern
)


def _is_confirmation_request(text: str) -> bool:
    """
    Return True when the agent response is waiting for user input —
    either a confirmation OR a clarifying question.
    Any '?' at the end of a line or clarification markers qualify.
    """
    low = (text or "").lower()
    if any(m in low for m in _CONFIRM_REQUEST_MARKERS):
        return True
    # Detect clarify-style responses: ends with a question mark
    stripped = low.rstrip()
    if stripped.endswith("?"):
        return True
    # Detect explicit clarification block markers from LLM
    if "_clarify_" in low or "clarify:" in low or "need to clarify" in low or "уточн" in low:
        return True
    return False


def _set_pending_confirm(ctx, chat_id: int) -> None:
    ctx.application.bot_data.setdefault("_omi_pending_confirms", {})[chat_id] = time.time()


def _clear_pending_confirm(ctx, chat_id: int) -> None:
    ctx.application.bot_data.get("_omi_pending_confirms", {}).pop(chat_id, None)


def _has_pending_confirm(ctx, chat_id: int) -> bool:
    ts = ctx.application.bot_data.get("_omi_pending_confirms", {}).get(chat_id)
    return ts is not None and (time.time() - ts) < _PENDING_CONFIRM_TTL


# ── Execution heartbeat ─────────────────────────────────────────────────────

def _next_task_seq(ctx) -> int:
    """Return next global task sequence number (resets on restart)."""
    n = ctx.application.bot_data.get("_task_seq", 0) + 1
    ctx.application.bot_data["_task_seq"] = n
    return n


async def _exec_heartbeat(msg, task_ref: str, lang: str = "en", interval: float = 6.0) -> None:
    """
    Periodically edit *msg* to show the bot is still executing.
    Cancelled by caller when the agent returns.
    """
    _icons = ("⚙️", "🔄", "⚙️", "🔄")
    step = 0
    elapsed = 0
    while True:
        await asyncio.sleep(interval)
        elapsed += int(interval)
        step += 1
        icon = _icons[step % len(_icons)]
        label = (
            f"{icon} Omi: обрабатываю {task_ref}… ({elapsed}с)"
            if lang == "ru"
            else f"{icon} Omi: processing {task_ref}… ({elapsed}s)"
        )
        try:
            await msg.edit_text(label)
        except Exception:
            return  # message may have been deleted or rate-limited


def _split_multi_values(payload: str, delims: str = ",;\n\t ") -> list[str]:
    buff = payload or ""
    for d in delims:
        buff = buff.replace(d, " ")
    return [x.strip() for x in buff.split(" ") if x.strip()]


def _extract_night_estimate(task: str) -> tuple[str, int | None]:
    text = (task or "").strip()
    low = text.lower()
    if "~" not in low:
        return text, None
    pos = low.rfind("~")
    tail = low[pos + 1 :].strip()
    digits = ""
    for ch in tail:
        if ch.isdigit():
            digits += ch
        else:
            break
    if not digits:
        return text, None
    unit_tail = tail[len(digits):].lstrip()
    unit = ""
    for ch in unit_tail:
        if ch.isalpha() or ch in ("ч",):
            unit += ch
        else:
            break
    val = int(digits)
    unit = unit.lower()
    estimate = val * 60 if unit in ("h", "hr", "hour", "hours", "ч") else val
    # Remove trailing "~30m" / "~2h" from task text.
    stripped = text[:pos].strip()
    return stripped or text, estimate


def _parse_notify_chat_ids() -> tuple[int, ...]:
    raw = os.environ.get("OMI_NOTIFY_CHAT_IDS", "").strip()
    if not raw:
        if ALLOWED_CHATS:
            return tuple(sorted(ALLOWED_CHATS))
        raw = os.environ.get("TELEGRAM_GROUP_ID", "").strip()
    out: list[int] = []
    for part in raw.split(","):
        p = part.strip()
        if not p:
            continue
        try:
            out.append(int(p))
        except ValueError:
            pass
    return tuple(dict.fromkeys(out))


from telegram import BotCommand
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    CallbackQueryHandler, ContextTypes, filters,
)

from omi_agent import ChatResult, OmiAgent, _format_search_rows, _schedule_ollama_background_warm
from omi_docgen import generate_bundle_from_storage
from omi_recovery import attempt_soft_recovery
from omi_selftest import run_selftest_suite
from omi_storage import StorageManager
from ops.knowledge import (
    audit_knowledge_base,
    detect_sqlite_duplicates,
    knowledge_base_ready,
    repair_index_issues,
    search_local_documents,
)

from skill_convert_legacy import convert_inplace_async as _convert_inplace_async, is_legacy as _is_legacy
from skill_rename_ai_qwen import suggest_filename as _suggest_filename, has_meaningful_name as _has_meaningful_name

import test_mode_shared

logging.basicConfig(
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    level=logging.INFO,
)
log = logging.getLogger("omi")

try:
    from registry_audit import (
        cmd_registry_audit as _registry_audit_cmd_handler,
        maybe_handle_registry_audit_message as _maybe_handle_registry_audit_message,
    )
except Exception as _registry_audit_import_err:  # pragma: no cover - optional audit helper
    _registry_audit_cmd_handler = None  # type: ignore[assignment]
    _maybe_handle_registry_audit_message = None  # type: ignore[assignment]
    log.warning("registry audit helpers unavailable: %s", _registry_audit_import_err)

try:
    from docsreg_launch import (
        cmd_docsreg as _docsreg_cmd_handler,
        maybe_handle_docsreg_message as _maybe_handle_docsreg_message,
    )
except Exception as _docsreg_import_err:  # pragma: no cover - optional launch helper
    _docsreg_cmd_handler = None  # type: ignore[assignment]
    _maybe_handle_docsreg_message = None  # type: ignore[assignment]
    log.warning("docsreg launch helpers unavailable: %s", _docsreg_import_err)

# ── Task Registry client (качество обработки запросов) ─────────────────────────
# Импорт отдельным try — если API недоступен, бот продолжает работу без регистрации.
try:
    sys.path.insert(0, str(Path(__file__).parent.parent))
    from models.model_registry import resolve_slot as _resolve_slot
except Exception:
    _resolve_slot = None  # type: ignore[assignment]
try:
    from task_registry_api import TaskRegistryClient as _TRClient
    _tr_client: "_TRClient | None" = _TRClient()
    log.info("Task Registry client initialised: %s", _tr_client.base_url)
except Exception as _tr_err:
    _tr_client = None
    log.warning("Task Registry client unavailable: %s", _tr_err)


def _tr_register(desc: str, chat_id: str, source: str = "omi") -> str:
    """Register a task in Task Registry; returns task_id or '' on failure."""
    if _tr_client is None:
        return ""
    try:
        return _tr_client.register(desc, source=source, chat_id=chat_id) or ""
    except Exception as _e:
        log.debug("_tr_register failed: %s", _e)
        return ""


def _tr_start(task_id: str, assigned_to: str = "omi") -> None:
    if not task_id or _tr_client is None:
        return
    try:
        _tr_client.start(task_id, assigned_to=assigned_to)
    except Exception as _e:
        log.debug("_tr_start failed: %s", _e)


def _tr_done(task_id: str, summary: str = "") -> None:
    if not task_id or _tr_client is None:
        return
    try:
        _tr_client.done(task_id, result_summary=summary[:200])
    except Exception as _e:
        log.debug("_tr_done failed: %s", _e)


def _tr_stuck(task_id: str, error: str = "") -> None:
    if not task_id or _tr_client is None:
        return
    try:
        _tr_client.stuck(task_id, error=error[:200])
    except Exception as _e:
        log.debug("_tr_stuck failed: %s", _e)


_TASK_CLOSE_PHRASES = (
    "выполнено",
    "готово",
    "закрыть задачу",
    "закрой задачу",
    "close task",
    "task done",
    "completed",
    "resolved",
)
_TASK_ID_RE = re.compile(r"\b(task_[0-9a-f]{12,})\b", re.IGNORECASE)


def _extract_task_id(text: str) -> str:
    m = _TASK_ID_RE.search(text or "")
    return (m.group(1) if m else "").lower()


def _looks_like_task_close_intent(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return False
    if _extract_task_id(t):
        return True
    return any(p in t for p in _TASK_CLOSE_PHRASES)


def _close_omi_task_from_chat(chat_id: str, text: str) -> tuple[bool, str]:
    """Close hanging Omi task in this chat instead of creating a new one."""
    if _tr_client is None:
        return False, "Task Registry недоступен."
    wanted_id = _extract_task_id(text)
    try:
        if wanted_id:
            task = _tr_client.get(wanted_id) or {}
            if str(task.get("chat_id", "")) != str(chat_id):
                return True, f"Задача `{wanted_id}` не найдена в этом чате."
            status = str(task.get("status", "")).lower()
            if status in {"done", "failed"}:
                return True, f"Задача `{wanted_id}` уже закрыта (`{status}`)."
            _tr_done(wanted_id, summary="manual_close_by_user")
            return True, f"✅ Закрыл задачу `{wanted_id}` по вашему подтверждению."

        stuck = _tr_client.find_stuck(older_than_minutes=0) or []
        candidates = [
            t for t in stuck
            if str(t.get("chat_id", "")) == str(chat_id)
            and str(t.get("assigned_to", "")).lower() == "omi"
            and str(t.get("status", "")).lower() in {"in_progress", "stuck", "pending"}
        ]
        if not candidates:
            return True, "В этом чате нет зависших задач Omi для закрытия."

        task = sorted(candidates, key=lambda x: (x.get("created_at") or "", x.get("id") or 0))[0]
        tid = str(task.get("task_id", "")).strip()
        if not tid:
            return True, "Нашёл зависшую задачу, но без task_id — не могу закрыть автоматически."
        _tr_done(tid, summary="manual_close_by_user")
        return True, f"✅ Закрыл зависшую задачу `{tid}` по вашему подтверждению."
    except Exception as e:
        log.warning("manual task close failed chat=%s: %s", chat_id, e)
        return True, f"Не смог закрыть задачу автоматически: {type(e).__name__}."

WORKSPACE = workspace_root()
TOKEN = (
    os.environ.get("TELEGRAM_BOT_TOKEN")
    or os.environ.get("OMI_TELEGRAM_TOKEN")
    or os.environ.get("OMI_TELEGRAM_BOT_TOKEN")
    or ""
).strip()
if not TOKEN:
    raise SystemExit(
        "Нет токена: TELEGRAM_BOT_TOKEN / OMI_TELEGRAM_TOKEN / OMI_TELEGRAM_BOT_TOKEN"
    )
DB_PATH = Path(
    os.environ.get("OMI_DB_PATH", str(WORKSPACE / "aims_registry.db"))
)
OCR_REGISTRY_DB = Path(
    os.environ.get("OCR_REGISTRY_DB", str(WORKSPACE / "omi_registry.db"))
)
NOTIFY_CHAT_IDS = _parse_notify_chat_ids()
OMI_NOTIFY_REGISTRATIONS_FROM_DB = os.environ.get("OMI_NOTIFY_REGISTRATIONS_FROM_DB", "0").strip().lower() in (
    "1",
    "true",
    "yes",
    "on",
)
OMI_UPLOAD_AI_RENAME_ENABLED = os.environ.get("OMI_UPLOAD_AI_RENAME_ENABLED", "1").strip().lower() in (
    "1", "true", "yes", "on"
)
OMI_UPLOAD_AI_RENAME_USE_AXI = os.environ.get("OMI_UPLOAD_AI_RENAME_USE_AXI", "0").strip().lower() in (
    "1", "true", "yes", "on"
)

# ── Axi task execution config ─────────────────────────────────────────────────
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "").strip()
ANTHROPIC_MODEL   = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6").strip()
OMI_RESULTS_DIR   = Path(os.environ.get("AXI_RESULTS_DIR", str(WORKSPACE / "result")))

def _name_already_in_aims(file_name: str) -> bool:
    """Check if a document with this filename is already registered in aims_registry."""
    try:
        import sqlite3 as _sq
        c = _sq.connect(str(DB_PATH))
        r = c.execute(
            "SELECT 1 FROM documents WHERE file_name=? OR canonical_file_name=?",
            (file_name, file_name),
        ).fetchone()
        c.close()
        return r is not None
    except Exception:
        return False


agent   = OmiAgent(db_path=DB_PATH, workspace=WORKSPACE)
storage = StorageManager(db_path=DB_PATH, workspace=WORKSPACE)

# Path 2: Telegram file upload — daily stats {chat_id: (date_str, count)}, no hard limit
_DAILY_UPLOADS: dict[int, tuple[str, int]] = {}
_PENDING_ANALYZE: dict[int, dict[str, object]] = {}
BATCH_INBOX = Path(os.environ.get("BATCH_INBOX_DIR", str(WORKSPACE / "batch_inbox")))
# См. ops/omi_batch_ingest.SUPPORTED_EXT — без авто-распаковки (skill 24: .zip — через axi_archive_extract).
_BATCH_UNSUPPORTED_EXTENSIONS = frozenset({".rar", ".7z", ".tar", ".gz", ".tgz", ".bz2"})
_PROCESS_MENU_META: dict[str, str] = {
    "P00": "AIMS General",
    "P01": "Purpose & Context",
    "P02": "Leadership & Governance",
    "P03": "Organization & People",
    "P04": "Strategy & Planning",
    "P05": "Asset Management Decision-Making",
    "P06": "Life Cycle Delivery",
    "P07": "Information Management",
    "P08": "Risk",
    "P09": "Review & Continual Improvement",
    "P10": "Value & Outcomes",
    "P11": "CV_PCV",
}

# ══════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════

def main_menu_keyboard():
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton("📁 Переместить файлы",  callback_data="menu_move"),
            InlineKeyboardButton("🗄 Архивировать",        callback_data="menu_archive"),
        ],
        [
            InlineKeyboardButton("➕ Новый раздел",        callback_data="menu_new_process"),
            InlineKeyboardButton("🚀 Миграция на DGX",    callback_data="menu_migrate"),
        ],
        [
            InlineKeyboardButton("📊 Статус базы",         callback_data="menu_status"),
            InlineKeyboardButton("🔍 Найти документ",      callback_data="menu_search"),
        ],
        [
            InlineKeyboardButton("📂 Разделы (реестр)", callback_data="menu_processes"),
        ],
        [
            InlineKeyboardButton("📝 Переименование файла", callback_data="menu_rename_file"),
        ],
        [
            InlineKeyboardButton("🧪 DOCSREG", callback_data="menu_docsreg"),
        ],
        [
            InlineKeyboardButton("🧪 Test mode (/test_mode)", callback_data="menu_test_mode"),
        ],
    ])


def _truncate_btn_label(text: str, max_len: int = 96) -> str:
    """Telegram inline buttons are safest kept short (client limits vary)."""
    t = (text or "").strip()
    if len(t) <= max_len:
        return t
    return t[: max(0, max_len - 1)] + "…"


def _process_button_title(
    code: str,
    count: int | None = None,
    *,
    registry_name: str | None = None,
) -> str:
    code_u = (code or "").upper()
    rn = (registry_name or "").strip()
    name = rn or _PROCESS_MENU_META.get(code_u)
    if not name:
        return code_u
    if count is None:
        label = f"{code_u} — {name}"
    else:
        # Put count right after code so it's visible even on narrow buttons.
        label = f"{code_u} ({max(0, int(count))}) — {name}"
    return _truncate_btn_label(label)


def process_keyboard(*, back_callback: str = "menu_back") -> InlineKeyboardMarkup:
    buttons = []
    processes = storage.list_processes()
    for p in processes:
        title = _process_button_title(
            p["code"],
            p.get("count"),
            registry_name=str(p.get("name") or "") or None,
        )
        buttons.append([InlineKeyboardButton(title, callback_data=f"proc_{p['code']}")])
    buttons.append([InlineKeyboardButton("« Назад", callback_data=back_callback)])
    return InlineKeyboardMarkup(buttons)

# ══════════════════════════════════════════════════════════════
#  КОМАНДЫ
# ══════════════════════════════════════════════════════════════

async def cmd_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    await update.message.reply_text(
        "👋 Привет, я *Omi* — администратор базы знаний AIMS.\n\n"
        "Пиши на любом языке — отвечу на языке твоего сообщения.\n"
        "*В группе* обязательно упомяни меня `@"
        + (ctx.bot.username or "бот")
        + "` или начни сообщение с `Omi` (или задай `OMI_GROUP_REQUIRE_MENTION=0` на сервере).\n\n"
        "Живой чат: опиши задачу словами или выбери действие ниже.\n"
        "Команда `/test_mode` (или опечатка `/tect_mode`) — следующий запрос с отчётом intent+время (skill 21).\n"
        "Если чего-то не хватает для команды — спрошу уточнение.",
        parse_mode="Markdown",
        reply_markup=main_menu_keyboard(),
    )

async def cmd_help(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    text = (
        "🤖 *Omi — команды*\n\n"
        "*В группе:* напиши `@"
        + (ctx.bot.username or "OmiBot")
        + "` в сообщении *или* начни текст с `Omi` — иначе бот не отвечает "
        "(можно отключить фильтр: `OMI_GROUP_REQUIRE_MENTION=0` в `.env`).\n\n"
        "*Поведение:* ответ на языке запроса; при неполных данных бот задаст вопросы.\n"
        + (
            "*Владелец:* можно добавлять правила поведения фразой «добавь правило: …» "
            "(нужен `OMI_OWNER_CHAT_IDS`).\n\n"
            if OWNER_CHATS
            else ""
        )
        + "*Чат (просто напиши):*\n"
        "• «Найди документы по RBI»\n"
        "• «Переместь отчёт inspection в P04»\n"
        "• «Заархивируй всё старше 2023»\n"
        "• «Создай новый процесс P11 Corrosion Management»\n"
        "• «Omi, create AIMS policy framework» — новый документ через one-pass pipeline\n"
        "• «Смени путь к базе на /new/path»\n\n"
        "*Текстовые команды:*\n"
        "/status — состояние базы и папок\n"
        "/move `<файл>` `<P0X>` — переместить файл\n"
        "/archive `<P0X>` `[год]` — архивировать раздел\n"
        "/newprocess `<код>` `<название>` — новый раздел\n"
        "/migrate `<новый_путь>` — сменить путь к БД\n"
        "/search `<запрос>` — найти документы\n"
        "/omi_kb_status — состояние локальной knowledge base\n"
        "/omi_kb_audit — полный audit knowledge base\n"
        "/omi_kb_search `<query>` — поиск по knowledge base\n"
        "/omi_kb_duplicates — дубликаты knowledge base\n"
        "/omi_kb_reindex — безопасный reindex knowledge base (owner)\n"
        "/omi_kb_exemplar_check `<document_type>` — проверка эталона\n"
        "/registry_sync_status — sync OCR→AIMS: inserted/updated/skipped\n"
        "/rename_by_context `[dry-run|apply]` `<f1,f2,...>` — контекстное имя (preview/apply)\n"
        "/rename_file `[dry-run|apply]` `<f1,f2,...>` — переименование файла по запросу\n"
        "/rename_file_`<new_filename>` — последний/реплай-файл → новое имя; "
        "`/rename_file_<id>_<new>` — по номеру из списка реестра (#348 → `348_…`)\n"
        "/tasks `[pending|done|all]` — очередь задач для структуры БД\n"
        "/close_task `[task_id]` — закрыть зависшую задачу в Task Registry\n"
        "/skills — список skills для Omi-LLM\n"
        "/docgen — one-pass generation only: собрать .docx из реестра (`/docgen` без аргументов — справка)\n"
        "/docgen_upgrade — private-only DOCGEN upgrade batch (no group self-improvement; private chat only)\n"
        "/docsreg — запустить DOCSREG по draft-файлу или папке на DGX\n"
        "/docsreg_start_media — batch alias for folder-based DOCSREG\n"
        "/selftest `[quick|full]` — единый тестовый прогон + чеклист (владелец)\n"
        "/test_mode — следующий запрос: пошаговый отчёт + *Complete successful*, затем test mode OFF (skill `21_test_mode`)\n"
        "/tect_mode — то же, что `/test_mode` (частая опечатка)\n"
        "/menu — главное меню\n"
        + (
            "\n*Бэкапы БД (только владелец, `OMI_OWNER_CHAT_IDS`):*\n"
            "/backup_now — копировать `aims_registry` / `omi_registry` сейчас\n"
            "/backup_list — список файлов `.bak`\n"
            "/backup_restore `<номер_из_списка|имя_файла>` `<aims|omi>` — восстановить\n"
            "/backup_schedule — расписание автокопирования + подсказка для cron\n"
            "/backup_schedule_set `<час>` `<мин>` `[on|off]` — задать ежедневное время\n"
            if OWNER_CHATS
            else ""
        )
    )
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_close_task(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Close hanging Omi task in this chat or by explicit task_id."""
    if not _chat_allowed(update):
        return
    chat_id = str(update.effective_chat.id if update.effective_chat else "")
    task_id = (ctx.args[0].strip() if getattr(ctx, "args", None) else "")
    text = task_id or "закрыть задачу"
    handled, msg = _close_omi_task_from_chat(chat_id, text)
    if handled:
        await update.message.reply_text(msg, parse_mode="Markdown")
        return
    await update.message.reply_text("⚠️ Не удалось обработать /close_task.")


def _set_pending_analyze(chat_id: int, prompt: str) -> None:
    now = time.time()
    _PENDING_ANALYZE[chat_id] = {
        "prompt": prompt,
        "status": "active",
        "created_at": now,
        "updated_at": now,
        "count": 0,
        "notified": False,
    }


def _get_pending_analyze(chat_id: int) -> str | None:
    item = _PENDING_ANALYZE.get(chat_id)
    if not item:
        return None
    if str(item.get("status", "active")).lower() != "active":
        return None
    return str(item.get("prompt", "")).strip() or None


async def cmd_analyze(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if update.effective_chat is None:
        return
    prompt = " ".join(ctx.args).strip() if getattr(ctx, "args", None) else ""
    if not prompt:
        prompt = "Analyze uploaded file(s) and provide concise findings."
    _set_pending_analyze(update.effective_chat.id, prompt)
    lang = _reply_lang(update.message.text or "")
    if lang == "ru":
        await update.message.reply_text(
            "Режим `/analyze` включён. Загрузите файл(ы), я приму их как материалы для анализа.",
            parse_mode="Markdown",
        )
    else:
        await update.message.reply_text(
            "`/analyze` mode enabled. Upload file(s) and I will treat them as analysis input.",
            parse_mode="Markdown",
        )


async def cmd_analyze_done(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if update.effective_chat is None:
        return
    chat_id = update.effective_chat.id
    state = _PENDING_ANALYZE.pop(chat_id, None)
    if not state:
        await update.message.reply_text("Нет активного режима /analyze.")
        return
    count = int(state.get("count", 0))
    lang = _reply_lang(update.message.text or "")
    if lang == "ru":
        await update.message.reply_text(
            f"Режим /analyze завершён вручную (файлов в пакете: {count}). "
            "Это команда завершения загрузки файлов, а не остановка OCR-пайплайна."
        )
    else:
        await update.message.reply_text(
            f"/analyze capture closed manually (files in batch: {count}). "
            "This command finalizes file upload for analysis; it does not stop OCR pipeline."
        )


async def cmd_analyze_end(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Alias for /analyze_done to avoid command mismatch in chats."""
    await cmd_analyze_done(update, ctx)


async def cmd_test_mode(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update) or not update.message:
        return
    cid = update.effective_chat.id if update.effective_chat else None
    if cid is not None:
        test_mode_shared.omi_register_test_mode_pending(cid)
    await update.message.reply_text(
        "Omi: *test mode* только для *следующего* сообщения — отчёт `OmiAgent.chat`, "
        "маршрут skills (master + JSON), затем *Complete successful*; "
        "test mode выключается автоматически после успеха или после отчёта о сбое.\n"
        "Skill: `aims_skills_shared/21_test_mode.md`",
        parse_mode="Markdown",
    )


async def cmd_status(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    status = storage.get_status()
    lines = [
        "📊 *Статус Omi*\n",
        f"🗄 БД: `{status['db_path']}`",
        f"📦 Документов: *{status['total_docs']}*",
        f"⭐ Мастер-файлов: *{status['master_docs']}*",
        f"🔒 Анонимизировано: *{status['anon_docs']}*\n",
        "📁 *Разделы:*",
    ]
    for p in status["processes"]:
        lines.append(f"  `{p['code']}` {p['name']} — {p['count']} файлов")
    lines.append(f"\n💾 Workspace: `{status['workspace']}`")
    lines.append(f"🖥 Платформа: *{status['platform']}*")

    try:
        _ops_root = Path(__file__).resolve().parent.parent
        if str(_ops_root) not in sys.path:
            sys.path.insert(0, str(_ops_root))
        from ollama_resolve import ollama_pc_andrey_small_backup_report

        rep = ollama_pc_andrey_small_backup_report()
        ic = {"ok": "✅", "warn": "⚠️", "fail": "❌", "skip": "ℹ️"}.get(rep.get("level"), "ℹ️")
        lines.append(f"\n{ic} *ПК Андрея* — резерв малой модели")
        lines.append(rep.get("summary") or "—")
    except Exception as e:
        lines.append(f"\n⚠️ ПК Андрея: диагностика недоступна ({e})")

    await update.message.reply_text(
        "\n".join(lines),
        parse_mode="Markdown",
        reply_markup=main_menu_keyboard(),
    )

async def cmd_move(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    args = ctx.args
    if len(args) < 2:
        await update.message.reply_text(
            "Использование: `/move <имя_файла_или_ID> <P0X>`\n"
            "Пример: `/move inspection_report_2024.pdf P04`",
            parse_mode="Markdown",
        )
        return
    query, target = " ".join(args[:-1]), args[-1].upper()
    msg = await update.message.reply_text(f"⏳ Перемещаю `{query}` → `{target}`...", parse_mode="Markdown")
    result = storage.move_documents(query, target)
    await msg.edit_text(result, parse_mode="Markdown")

async def cmd_archive(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    args = ctx.args
    process = args[0].upper() if args else None
    year    = int(args[1]) if len(args) > 1 else None
    msg = await update.message.reply_text("⏳ Архивирую...", parse_mode="Markdown")
    result = storage.archive_documents(process=process, before_year=year)
    await msg.edit_text(result, parse_mode="Markdown")

async def cmd_newprocess(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    args = ctx.args
    if len(args) < 2:
        await update.message.reply_text(
            "Использование: `/newprocess <код> <название>`\n"
            "Пример: `/newprocess P11 Corrosion_Management`",
            parse_mode="Markdown",
        )
        return
    code = args[0].upper()
    name = " ".join(args[1:]).replace("_", " ")
    result = storage.create_process(code, name)
    await update.message.reply_text(result, parse_mode="Markdown")

async def cmd_migrate(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    args = ctx.args
    if not args:
        await update.message.reply_text(
            "Использование: `/migrate <новый_путь_к_БД>`\n"
            "Пример: `/migrate /home/axiomsphere/aims-data/omi_registry.db`\n\n"
            "Для миграции на DGX Spark:\n"
            "`/migrate dgx` — покажет инструкцию",
            parse_mode="Markdown",
        )
        return
    new_path = args[0]
    if new_path.lower() == "dgx":
        text = storage.get_dgx_migration_plan()
    else:
        text = storage.migrate_db(new_path)
    await update.message.reply_text(text, parse_mode="Markdown")

async def cmd_search(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not ctx.args:
        await update.message.reply_text("Использование: `/search <запрос>`", parse_mode="Markdown")
        return
    query = " ".join(ctx.args)
    try:
        storage.sync_from_ocr_registry(limit=4000)
    except Exception:
        pass
    lang = _reply_lang(query)
    rows = storage.hybrid_search_documents(query, limit=15)
    text = _format_search_rows(rows, query=query, lang=lang)
    if not rows:
        text = agent._with_registry_empty_assist(text, query=query, lang=lang, plan=None)
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_docs_today(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    try:
        storage.sync_from_ocr_registry(limit=4000)
    except Exception:
        pass
    text = storage.list_documents_today(lang="en")
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_registry_sync_status(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    ins, upd, sk = storage.sync_from_ocr_registry(limit=10000)
    await update.message.reply_text(
        (
            "📡 *Registry sync status*\n"
            f"- inserted: *{ins}*\n"
            f"- updated: *{upd}*\n"
            f"- skipped: *{sk}*"
        ),
        parse_mode="Markdown",
    )


def _format_kb_summary(audit: dict[str, object], ready: bool, failures: list[str]) -> str:
    sources = audit.get("sources") or []
    sqlite_audits = audit.get("sqlite_audits") or []
    duplicate_clusters = audit.get("duplicate_clusters") or {}
    index_detail = audit.get("index_integrity_detail") or {}
    known = audit.get("known_exemplar") or {}
    return (
        "🧠 Omi KB status\n"
        f"Ready: {str(ready).lower()}\n"
        f"Failures: {', '.join(failures) if failures else 'none'}\n"
        f"Sources: {len(sources)}\n"
        f"SQLite DBs: {len(sqlite_audits)}\n"
        f"Duplicate clusters: {sum(len(v) for v in duplicate_clusters.values()) if isinstance(duplicate_clusters, dict) else 0}\n"
        f"Index status: {index_detail.get('status') if isinstance(index_detail, dict) else 'n/a'}\n"
        f"Known exemplar: {known.get('title') or known.get('file_name') or 'n/a'}"
    )


async def cmd_omi_kb_status(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text("⚠️ Только владелец бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown")
        return
    audit = audit_knowledge_base(storage, expected_document_type="policy_framework")
    ready, failures = knowledge_base_ready(audit)
    await update.message.reply_text(_format_kb_summary(audit, ready, failures), parse_mode="Markdown")


async def cmd_omi_kb_audit(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text("⚠️ Только владелец бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown")
        return
    audit = audit_knowledge_base(storage, expected_document_type="policy_framework")
    ready, failures = knowledge_base_ready(audit)
    dup_count = sum(len(v) for v in (audit.get("duplicate_clusters") or {}).values()) if isinstance(audit.get("duplicate_clusters"), dict) else 0
    msg = (
        "🔎 Omi KB audit\n"
        f"Ready: {str(ready).lower()}\n"
        f"Failures: {', '.join(failures) if failures else 'none'}\n"
        f"Index: {(audit.get('index_integrity_detail') or {}).get('status', 'n/a')}\n"
        f"Known exemplar: {(audit.get('known_exemplar') or {}).get('title') or 'n/a'}\n"
        f"Duplicates: {dup_count}"
    )
    await update.message.reply_text(msg, parse_mode="Markdown")


async def cmd_omi_kb_search(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text("⚠️ Только владелец бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown")
        return
    query = " ".join(ctx.args or []).strip()
    if not query:
        await update.message.reply_text("Использование: `/omi_kb_search <query>`", parse_mode="Markdown")
        return
    result = search_local_documents(storage, query=query, document_type="", task_context=query, limit=10)
    candidates = list(result.get("retrieved_documents") or [])
    top = candidates[0] if candidates else {}
    text = (
        "🔎 Omi knowledge search\n"
        f"Query: {query}\n"
        f"Sources searched: {len((result.get('retrieval_diagnostics') or {}).get('sources_searched') or [])}\n"
        f"Candidates: {len(candidates)}\n"
        f"Top result: {top.get('title') or top.get('file_name') or 'n/a'}\n"
        f"Type: {top.get('document_type') or 'n/a'}\n"
        f"Score: {float(top.get('final_score') or 0.0):.2f}\n"
        f"Path: {top.get('path') or top.get('file_path') or 'n/a'}\n"
        f"Selection: {', '.join(top.get('selection_reasons') or []) or 'n/a'}"
    )
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_omi_kb_duplicates(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text("⚠️ Только владелец бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown")
        return
    clusters = detect_sqlite_duplicates(DB_PATH)
    if not clusters:
        await update.message.reply_text("🧩 Duplicates: none detected")
        return
    lines = ["🧩 Omi duplicates"]
    for cluster in clusters[:10]:
        lines.append(
            f"• {cluster.get('duplicate_type')}: canonical={cluster.get('canonical_record_id')} members={len(cluster.get('members') or [])}"
        )
    await update.message.reply_text("\n".join(lines), parse_mode="Markdown")


async def cmd_omi_kb_reindex(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text("⚠️ Только владелец бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown")
        return
    evidence = repair_index_issues(storage)
    await update.message.reply_text(
        "🔧 Omi KB reindex\n"
        f"FTS rebuilt: {evidence.get('rebuild_fts')}\n"
        f"Embeddings rebuilt: {evidence.get('rebuild_embeddings')}\n"
        f"Synced from OCR: {evidence.get('synced_from_ocr')}\n"
        f"After status: {(evidence.get('issues_after') or {}).get('status', 'n/a')}",
        parse_mode="Markdown",
    )


async def cmd_omi_kb_exemplar_check(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text("⚠️ Только владелец бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown")
        return
    doc_type = " ".join(ctx.args or []).strip() or "policy_framework"
    audit = audit_knowledge_base(storage, expected_document_type=doc_type)
    ready, failures = knowledge_base_ready(audit)
    known = audit.get("known_exemplar") or {}
    text = (
        "📘 Omi exemplar check\n"
        f"Type: {doc_type}\n"
        f"Ready: {str(ready).lower()}\n"
        f"Failures: {', '.join(failures) if failures else 'none'}\n"
        f"Found: {known.get('title') or known.get('file_name') or 'n/a'}\n"
        f"Path: {known.get('path') or known.get('file_path') or 'n/a'}"
    )
    await update.message.reply_text(text, parse_mode="Markdown")


def _parse_rename_args(args: list[str]) -> tuple[bool, bool, bool, list[str]]:
    """
    /rename_by_context [dry-run|apply] [ai|classic] [axi] file1.docx,file2.docx
    По умолчанию: dry-run.
    """
    if not args:
        return False, True, False, []
    mode = (args[0] or "").strip().lower()
    apply_mode = mode in ("apply", "run", "do", "commit")
    rest = args[1:] if mode in ("dry-run", "dry", "preview", "apply", "run", "do", "commit") else args
    ai_mode = True
    use_external_llm = False
    payload_parts: list[str] = []
    for a in rest:
        al = (a or "").strip().lower()
        if al in ("ai", "smart", "llm"):
            ai_mode = True
            continue
        if al in ("classic", "basic"):
            ai_mode = False
            continue
        if al in ("axi", "external"):
            use_external_llm = True
            continue
        payload_parts.append(a)
    payload = " ".join(payload_parts)
    names = [
        Path(x.strip()).name
        for x in _split_multi_values(payload, delims=",;\n")
        if x.strip()
    ]
    return apply_mode, ai_mode, use_external_llm, names


async def cmd_rename_by_context(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    apply_mode, ai_mode, use_external_llm, names = _parse_rename_args(ctx.args or [])
    if not names:
        await update.message.reply_text(
            "Usage:\n"
            "`/rename_by_context dry-run file1.docx,file2.docx`\n"
            "`/rename_by_context apply ai file1.docx,file2.docx`\n"
            "`/rename_by_context dry-run ai axi file1.docx,file2.docx`",
            parse_mode="Markdown",
        )
        return
    try:
        storage.sync_from_ocr_registry(limit=10000)
    except Exception:
        pass
    text = storage.rename_documents_by_context(
        names,
        lang=_reply_lang(update.message.text if update.message else ""),
        apply_changes=apply_mode,
        ai_mode=ai_mode,
        use_external_llm=use_external_llm,
    )
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_rename_file(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Alias for /rename_by_context: file rename by request/context."""
    await cmd_rename_by_context(update, ctx)


async def cmd_rename_file_dynamic(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """
    Dynamic shortcut:
      /rename_file_<new_filename> — новое имя; исходный файл из реплая или last_file_name.
      /rename_file_<id>_<new_filename> — явный id строки в `documents` (как в списке #348).
    """
    if not _chat_allowed(update):
        return
    if update.message is None or not update.message.text:
        return
    text = update.message.text.strip()
    m = re.match(r"^/rename_file_(.+?)(?:@\w+)?$", text, re.IGNORECASE)
    if not m:
        await update.message.reply_text(
            "Usage: `/rename_file_<new_filename>` or `/rename_file_<id>_<new_filename>`\n"
            "Reply to file, or use registry id from the list (e.g. `/rename_file_348_MyPlan.xlsx`).",
            parse_mode="Markdown",
        )
        return
    requested_raw = unicodedata.normalize("NFKC", m.group(1).strip())
    for _z in ("\ufeff", "\u200b", "\u200e", "\u200f"):
        requested_raw = requested_raw.replace(_z, "")
    if not requested_raw:
        await update.message.reply_text(
            "Usage: `/rename_file_<new_filename>`",
            parse_mode="Markdown",
        )
        return

    source_document_id = None
    requested = requested_raw
    for _pat in (
        r"^#?(\d{1,8})_(.+)$",
        r"^#?(\d{1,8})\s+(.+)$",
    ):
        m_id = re.match(_pat, requested_raw, re.DOTALL)
        if m_id and m_id.group(2).strip():
            source_document_id = int(m_id.group(1))
            requested = m_id.group(2).strip()
            break

    source_name = ""
    msg = update.message
    _cid = update.effective_chat.id if update.effective_chat else None
    fallback_doc_id = None
    if msg.reply_to_message and msg.reply_to_message.document:
        source_name = str(msg.reply_to_message.document.file_name or "").strip()
    if _cid is not None:
        try:
            _ctx = storage.get_chat_context(_cid)
            if not source_name:
                source_name = str(_ctx.get("last_file_name") or "").strip()
            lid = _ctx.get("last_doc_id")
            try:
                if lid is not None:
                    fallback_doc_id = int(lid)
            except (TypeError, ValueError):
                fallback_doc_id = None
        except Exception:
            pass
    if not source_name and not source_document_id:
        await update.message.reply_text(
            "⚠️ Source file is unknown. Reply to a file, or use:\n"
            "`/rename_file_<id>_<new_filename>` (registry id + underscore + new name).",
            parse_mode="Markdown",
        )
        return

    out = storage.rename_documents_by_request(
        source_name,
        requested,
        lang=_reply_lang(update.message.text if update.message else ""),
        fallback_doc_id=fallback_doc_id,
        source_document_id=source_document_id,
    )
    await update.message.reply_text(out, parse_mode="Markdown")


async def cmd_fixnoisynames(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """
    Batch auto-fix filenames containing long numeric codes.
    Usage: /fixnoisynames           — dry-run (preview only)
           /fixnoisynames apply     — apply renames in DB and on disk
    """
    if not _chat_allowed(update):
        return
    args = ctx.args or []
    apply_mode = bool(args) and args[0].strip().lower() in ("apply", "run", "do", "commit")
    lang = _reply_lang(update.message.text if update.message else "")
    text = storage.fix_noisy_names_batch(apply_changes=apply_mode, lang=lang)
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_fixnoisyname(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """
    Rename a single document by its registry ID.
    Usage: /fixnoisyname <id> <new_filename>
    Example: /fixnoisyname 42 Политика_ОТиПБ_2023.pdf
    """
    if not _chat_allowed(update):
        return
    args = ctx.args or []
    if len(args) < 2:
        await update.message.reply_text(
            "Usage: `/fixnoisyname <id> <new_filename>`\n"
            "Example: `/fixnoisyname 42 Политика_ОТиПБ.pdf`",
            parse_mode="Markdown",
        )
        return
    try:
        doc_id = int(args[0])
    except ValueError:
        await update.message.reply_text("⚠️ First argument must be a numeric document ID.", parse_mode="Markdown")
        return
    new_name = " ".join(args[1:]).strip()
    lang = _reply_lang(update.message.text if update.message else "")
    out = storage.rename_documents_by_request(
        source_name="",
        requested_name=new_name,
        lang=lang,
        source_document_id=doc_id,
    )
    await update.message.reply_text(out, parse_mode="Markdown")


async def cmd_menu(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    await update.message.reply_text(
        "📋 *Главное меню Omi*",
        parse_mode="Markdown",
        reply_markup=main_menu_keyboard(),
    )


async def cmd_tasks(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Очередь задач для структуры БД (в т.ч. поручения от Axi)."""
    if not _chat_allowed(update):
        return
    st = (ctx.args[0] if ctx.args else "pending").strip().lower()
    if st not in ("pending", "done", "all"):
        st = "pending"
    text = storage.list_tasks(status=st, limit=20)
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_skills(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Список зарегистрированных skills (тело — только во владельческом режиме через агента)."""
    if not _chat_allowed(update):
        return
    text = storage.format_skills_list()
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_docgen(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Сборка .docx из реестра: процесс P0X, список id или последние записи."""
    if not _chat_allowed(update):
        return
    args = ctx.args or []
    if not args:
        await update.message.reply_text(
            "📄 *Генерация пакета документации*\n\n"
            "• `/docgen P04` — документы процесса P04\n"
            "• `/docgen --ids 1,2,3` — по id в `documents`\n"
            "• `/docgen recent` — последние 120 записей (по дате)\n",
            parse_mode="Markdown",
        )
        return
    proc: str | None = None
    doc_ids: list[int] | None = None
    if args[0].strip() == "--ids":
        rest = " ".join(args[1:])
        doc_ids = [int(x) for x in _split_multi_values(rest, delims=",; \n\t") if x.strip().isdigit()]
        if not doc_ids:
            await update.message.reply_text("❌ Укажи числа: `/docgen --ids 1,2,3`", parse_mode="Markdown")
            return
    elif args[0].lower() in ("recent", "last", "последние"):
        proc = None
        doc_ids = None
    else:
        proc = args[0].strip().upper()

    msg = await update.message.reply_text("⏳ Собираю .docx…", parse_mode="Markdown")
    try:
        path, n = generate_bundle_from_storage(
            storage,
            process_code=proc,
            doc_ids=doc_ids,
            limit=120,
        )
    except ValueError as e:
        if str(e) == "no_documents":
            await msg.edit_text("❌ Нет документов по условиям.")
            return
        await msg.edit_text(f"❌ Ошибка: {e}")
        raise
    try:
        await msg.delete()
        await _send_docgen_for_review(
            update,
            ctx,
            path,
            n,
            f"✅ Built a package from *{n}* records (`/docgen`). Please review it with the buttons below.",
            approval_title=f"DOCGEN package ({proc or 'recent'})",
            approval_summary=f"Generated package from {n} registry records.",
            approval_meta={
                "registration_mode": "bundle",
                "aims_process": (proc or "DOCGEN").upper(),
            },
        )
    except Exception as e:
        log.error("docgen send: %s", e)
        await update.message.reply_text(
            f"❌ Файл создан, но не отправлен: `{path}`\n{e}"
        )


async def cmd_docgen_upgrade(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Queue a DOCGEN self-improvement upgrade batch. Private Axi only."""
    if not _chat_allowed(update):
        return
    if not _is_docgen_upgrade_origin(update):
        await update.message.reply_text(
            "⚠️ `/DOCGEN_UPGRADE` is private-only for Axi. "
            "Group chats stay generation-only and do not start the self-improvement loop.\n"
            "Use `/docgen` for generation-only runs.",
            parse_mode="Markdown",
        )
        return

    try:
        from ops.logi.docgen_skill_scope_policy import DOCGEN_DOCUMENT_TYPES  # noqa: PLC0415
        from ops.docgen.universal_overlay.nightly_improvement_orchestrator import (  # noqa: PLC0415
            build_docgen_upgrade_batch_action,
            make_docgen_upgrade_run_tag,
            split_docgen_upgrade_request,
            run_docgen_upgrade_batch_now,
            stream_docgen_upgrade_progress,
        )
    except Exception:
        try:
            from logi.docgen_skill_scope_policy import DOCGEN_DOCUMENT_TYPES  # type: ignore  # noqa: PLC0415
            from docgen.universal_overlay.nightly_improvement_orchestrator import (  # type: ignore  # noqa: PLC0415
                build_docgen_upgrade_batch_action,
                make_docgen_upgrade_run_tag,
                split_docgen_upgrade_request,
                run_docgen_upgrade_batch_now,
                stream_docgen_upgrade_progress,
            )
        except Exception as exc:  # pragma: no cover - import depends on runtime layout
            await update.message.reply_text(
                f"❌ DOCGEN upgrade launcher unavailable: {type(exc).__name__}",
                parse_mode="Markdown",
            )
            return

    topic_prefix, requested_types = split_docgen_upgrade_request(
        ctx.args or [],
        DOCGEN_DOCUMENT_TYPES,
    )
    doc_types = requested_types or list(DOCGEN_DOCUMENT_TYPES)
    run_tag = make_docgen_upgrade_run_tag()
    action_preview = build_docgen_upgrade_batch_action(
        document_types=doc_types,
        created_by="axi",
        self_improvement_enabled=True,
        max_cycles=0,
        topic_prefix=" ".join(topic_prefix).strip() or None,
        run_tag=run_tag,
    )
    progress_msg = await update.message.reply_text("⏳ DOCGEN upgrade started...", parse_mode="Markdown")
    chat_id = update.effective_chat.id
    stop_event = asyncio.Event()

    async def _run_upgrade_and_report() -> None:
        try:
            log.info("DOCGEN upgrade: starting batch run for %s", ",".join(doc_types))
            monitor_task = asyncio.create_task(
                stream_docgen_upgrade_progress(
                    action=action_preview,
                    document_types=doc_types,
                    chat_id=chat_id,
                    bot=ctx.bot,
                    stop_event=stop_event,
                )
            )
            action = await asyncio.to_thread(
                run_docgen_upgrade_batch_now,
                document_types=doc_types,
                created_by="axi",
                self_improvement_enabled=True,
                max_cycles=0,
                topic_prefix=" ".join(topic_prefix).strip() or None,
                run_tag=run_tag,
            )
            action_payload = action.get("action") or {}
            action_id = str(action_payload.get("action_id") or "")
            action_path = str(action.get("action_path") or "")
            execution = action.get("execution_result") or {}
            run_summary = action.get("run_summary") or {}
            summary = run_summary if isinstance(run_summary, dict) else {}
            final_msg = _format_docgen_upgrade_reply(
                action_id=action_id,
                action_path=action_path,
                execution=execution,
                run_summary=summary,
                doc_types=doc_types,
            )
            completion_msg = _format_docgen_upgrade_completion(
                action_id=action_id,
                run_summary=summary,
            )
            try:
                await progress_msg.edit_text(final_msg, parse_mode="Markdown")
            except Exception:
                pass
            await ctx.bot.send_message(chat_id=chat_id, text=completion_msg, parse_mode="Markdown")
            log.info("DOCGEN upgrade: completion message sent for %s", action_id or "unknown")
            stop_event.set()
            try:
                await monitor_task
            except Exception:
                pass
        except Exception as exc:
            stop_event.set()
            err = f"❌ DOCGEN upgrade failed: {type(exc).__name__}: {exc}"
            try:
                await progress_msg.edit_text(err, parse_mode="Markdown")
            except Exception:
                await ctx.bot.send_message(chat_id=chat_id, text=err, parse_mode="Markdown")

    await _run_upgrade_and_report()
    return


async def cmd_docsreg(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if _docsreg_cmd_handler is None:
        await update.message.reply_text(
            "❌ DOCSREG launcher недоступен в этом окружении.",
            parse_mode="Markdown",
        )
        return
    await _docsreg_cmd_handler(update, ctx)


async def _handle_group_doc_synthesis(update: Update, ctx: ContextTypes.DEFAULT_TYPE, clean: str, lang: str) -> bool:
    chat_id = update.effective_chat.id if update.effective_chat else None
    if chat_id is None:
        return False
    run_tag = f"group_docgen_{chat_id}_{int(time.time())}"
    run_root = None

    thinking = await update.message.reply_text(
        "⚙️ Omi: генерирую документ…" if lang == "ru" else "⚙️ Omi: generating document…"
    )
    task_id = _tr_register(clean[:120] or "(document synthesis)", chat_id=str(chat_id), source="group")
    _tr_start(task_id, assigned_to="omi")

    def _infer_document_type(text: str) -> str:
        low = (text or "").lower()
        if "procedure" in low or "instruction" in low or "процед" in low or "инструк" in low:
            return "procedure"
        if "report" in low or "отч" in low:
            return "technical_report"
        if "specification" in low or "requirements" in low:
            return "requirements_specification"
        return "policy_framework"

    try:
        from docgen.document_architecture import (  # noqa: PLC0415
            build_requirement_graph,
            retrieve_local_knowledge,
            validate_architecture_gate,
        )
        from doc_agent_api import DocAgentClient  # noqa: PLC0415
    except Exception as exc:
        await thinking.edit_text(
            f"⛔ DOCGEN one-pass blocked.\n\nBlocker: {exc}",
            parse_mode="Markdown",
        )
        return True

    doc_type = _infer_document_type(clean)
    try:
        retrieval = retrieve_local_knowledge(
            storage,
            query=clean,
            document_type=doc_type,
            task_context=clean,
            limit=8,
            require_exemplar=True,
        )
        graph = build_requirement_graph(
            request_text=clean,
            document_type=doc_type,
            retrieved_documents=list(retrieval["retrieved_documents"] or []),
            selected_exemplar=retrieval["selected_exemplar"],
            retrieval_diagnostics=retrieval["retrieval_diagnostics"],
        )
        gate = validate_architecture_gate(
            retrieval_passed=bool(retrieval["retrieval_passed"]),
            expected_exemplar_found=bool(retrieval["selected_exemplar"]),
            requirement_count=len(graph.nodes),
            final_section_map=graph.final_section_map,
        )
        if not gate.passed:
            blocker = ", ".join(gate.blockers) or "EXPECTED_EXEMPLAR_NOT_RETRIEVED"
            await thinking.edit_text(
                "⛔ DOCGEN one-pass blocked.\n\n"
                f"Blocker: {blocker}\n"
                "Next: fix retrieval / architecture gate",
                parse_mode="Markdown",
            )
            return True
    except Exception as exc:
        await thinking.edit_text(
            f"⛔ DOCGEN one-pass blocked.\n\nBlocker: {exc}",
            parse_mode="Markdown",
        )
        return True

    source_docs = list(retrieval["retrieved_documents"] or [])
    selected = retrieval["selected_exemplar"] or {}
    if selected:
        source_docs = [selected, *source_docs]
    source_docs = source_docs[:8]
    source_docs.insert(
        0,
        {
            "title": "Document architecture",
            "summary": json.dumps(
                {
                    "document_type": doc_type,
                    "final_section_map": list(graph.final_section_map),
                    "requirement_graph": graph.to_dict(),
                    "retrieval_diagnostics": retrieval["retrieval_diagnostics"],
                },
                indent=2,
                ensure_ascii=False,
            ),
            "aims_process": "DOCGEN",
            "iso_clause": "architecture",
            "is_master": True,
        },
    )

    try:
        _api_url = os.environ.get("DOC_AGENT_API_URL", "http://doc-agent:8767")
        _gen_dir = getattr(storage, "workspace", None)
        if _gen_dir:
            run_root = Path(_gen_dir) / "generated" / "docgen_one_pass" / run_tag
            run_root.mkdir(parents=True, exist_ok=True)
        _out_dir = str(run_root) if run_root else None
        result = DocAgentClient(_api_url).generate_dual(
            clean,
            title=f"{doc_type.replace('_', ' ').title()}",
            source_docs=source_docs,
            out_dir=_out_dir,
            architecture_context={
                "generation_mode": "FULL_DOCUMENT_FROM_FINAL_ARCHITECTURE",
                "document_type": doc_type,
                "base_archetype": list(graph.base_archetype),
                "final_section_map": list(graph.final_section_map),
                "requirement_graph": graph.to_dict(),
                "selected_exemplar": selected,
                "retrieval_diagnostics": retrieval["retrieval_diagnostics"],
                "architecture_gate": {
                    "passed": True,
                    "checks": gate.checks,
                    "blockers": [],
                },
            },
        )
        if not result.get("ok"):
            raise RuntimeError(result.get("error", "document synthesis failed"))
        path = Path(result["path"])
        preview = result.get("preview", "")
        await thinking.delete()
        await _send_docgen_for_review(
            update,
            ctx,
            path,
            len(source_docs),
            (
                "✅ DOCGEN one-pass accepted.\n\n"
                f"Project: AIMS\nType: {doc_type}\n"
                f"Quality: {result.get('compliance_pct', 'n/a')}%"
                if lang != "ru"
                else "✅ DOCGEN one-pass accepted.\n\n"
                f"Проект: AIMS\nТип: {doc_type}\n"
                f"Качество: {result.get('compliance_pct', 'n/a')}%"
            ),
            approval_title=path.stem,
            approval_summary=preview or "Generated document ready for review.",
            approval_meta={
                "registration_mode": "document",
                "aims_process": "DOCGEN",
                "source_text": clean,
                "architecture": graph.to_dict(),
                "cleanup_dir": str(run_root or path.parent),
                "review_mode": "customer_approval",
                "run_tag": run_tag,
            },
            approval_buttons=("approve", "rework"),
        )
        _tr_done(task_id, summary=f"docgen:{path.name}")
        return True
    except Exception as exc:
        _tr_stuck(task_id, error=str(exc)[:200])
        await thinking.edit_text(
            f"⛔ DOCGEN one-pass failed.\n\nBlocker: {type(exc).__name__}: {exc}",
            parse_mode="Markdown",
        )
        return True


async def cmd_nightplan(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Ночной план: /nightplan — статус очереди, /nightplan add <задача> — добавить."""
    if not _chat_allowed(update):
        return
    args = ctx.args or []
    verb = args[0].lower() if args else ""

    if verb in ("add", "добавить", "+"):
        task = " ".join(args[1:]).strip()
        if not task:
            await update.message.reply_text(
                "Использование: `/nightplan add <задача> [~30m | ~2h]`",
                parse_mode="Markdown",
            )
            return
        task, est = _extract_night_estimate(task)
        result = storage.enqueue_night_task(task, estimate_min=est, source="cmd_nightplan")
        await update.message.reply_text(result, parse_mode="Markdown")
        return

    await update.message.reply_text(storage.format_night_plan(), parse_mode="Markdown")


async def cmd_backup_now(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text(
            "⚠️ Только *владелец* бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown"
        )
        return
    msg = await update.message.reply_text("⏳ Делаю копию…")
    loop = asyncio.get_event_loop()

    def _run():
        return agent.backup.run_backup_now()

    ok, text = await loop.run_in_executor(None, _run)
    await msg.edit_text(text, parse_mode="Markdown")


async def cmd_backup_list(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text(
            "⚠️ Только *владелец* бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown"
        )
        return
    text = agent.backup.format_backup_list()
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_backup_restore(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text(
            "⚠️ Только *владелец* бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown"
        )
        return
    args = ctx.args
    if len(args) < 2:
        await update.message.reply_text(
            "Использование: `/backup_restore <номер|имя_файла> <aims|omi>`\n"
            "Пример: `/backup_restore 1 aims`",
            parse_mode="Markdown",
        )
        return
    target = args[-1]
    index_or_name = " ".join(args[:-1])
    msg = await update.message.reply_text("⏳ Восстанавливаю…")
    loop = asyncio.get_event_loop()
    text = await loop.run_in_executor(
        None, lambda: agent.backup.restore(index_or_name, target)
    )
    await msg.edit_text(text, parse_mode="Markdown")


async def cmd_backup_schedule(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text(
            "⚠️ Только *владелец* бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown"
        )
        return
    text = agent.backup.format_schedule()
    await update.message.reply_text(text, parse_mode="Markdown")


async def cmd_backup_schedule_set(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text(
            "⚠️ Только *владелец* бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown"
        )
        return
    args = ctx.args
    if len(args) < 2:
        await update.message.reply_text(
            "Использование: `/backup_schedule_set <час_0_23> <мин_0_59> [on|off]`\n"
            "Пример: `/backup_schedule_set 2 30 on` — каждый день в 02:30",
            parse_mode="Markdown",
        )
        return
    try:
        hour = int(args[0])
        minute = int(args[1])
    except ValueError:
        await update.message.reply_text("❌ Час и минуты должны быть числами.")
        return
    enabled = True
    if len(args) >= 3:
        enabled = args[2].strip().lower() not in ("off", "false", "0", "выкл", "нет")
    text = agent.backup.set_schedule(enabled, hour, minute, "every_2_days")
    await update.message.reply_text(text, parse_mode="Markdown")


# ══════════════════════════════════════════════════════════════
#  TEMPLATE COMMANDS  /save_template  /templates  /use_template
# ══════════════════════════════════════════════════════════════

async def cmd_save_template(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Сохранить шаблон из приложенного документа (ответ на файл + /save_template [name])."""
    if not _chat_allowed(update):
        return
    args = ctx.args or []
    name = " ".join(args).strip() if args else ""

    # Determine source: reply-to-document OR direct document in same message
    doc_msg = None
    if update.message.reply_to_message and update.message.reply_to_message.document:
        doc_msg = update.message.reply_to_message
    elif update.message.document:
        doc_msg = update.message

    if not doc_msg:
        await update.message.reply_text(
            "📋 *Сохранение шаблона*\n\n"
            "Пришли документ (.docx, .xlsx, .pptx) и ответь на него:\n"
            "`/save_template Название шаблона`\n\n"
            "Или пришли документ с подписью `/save_template Название`.",
            parse_mode="Markdown",
        )
        return

    ext = Path(doc_msg.document.file_name or "").suffix.lower()
    if ext not in (".docx", ".xlsx", ".pptx"):
        await update.message.reply_text(
            f"❌ Формат `{ext}` не поддерживается для шаблонов.\n"
            "Поддерживаются: `.docx`, `.xlsx`, `.pptx`",
            parse_mode="Markdown",
        )
        return

    if not name:
        # Use filename stem as default name
        name = Path(doc_msg.document.file_name or "template").stem

    msg = await update.message.reply_text("⏳ Извлекаю шаблон…")

    import tempfile as _tmpmod
    tmp_dl = Path(_tmpmod.mktemp(suffix=ext, prefix="_tpl_dl_"))
    try:
        tg_file = await doc_msg.document.get_file()
        await tg_file.download_to_drive(str(tmp_dl))

        from skill_doc_template import save_template, ensure_templates_table
        ensure_templates_table(DB_PATH)

        info = save_template(
            src=tmp_dl,
            name=name,
            db_path=DB_PATH,
            workspace=WORKSPACE,
            description=f"Из: {doc_msg.document.file_name}",
        )
        await msg.edit_text(
            f"✅ Шаблон сохранён!\n\n"
            f"*{info.name}* (`.{info.format}`)\n"
            f"Slug: `{info.slug}`\n"
            f"Сохранён: `{info.created_at[:10]}`\n\n"
            f"Используй `/use_template {info.slug}` чтобы получить шаблон.",
            parse_mode="Markdown",
        )
    except ValueError as e:
        await msg.edit_text(f"❌ {e}")
    except Exception as e:
        log.exception("cmd_save_template error")
        await msg.edit_text(f"❌ Ошибка: {e}")
    finally:
        tmp_dl.unlink(missing_ok=True)


async def cmd_templates(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Список сохранённых шаблонов."""
    if not _chat_allowed(update):
        return
    from skill_doc_template import list_templates, ensure_templates_table
    ensure_templates_table(DB_PATH)
    items = list_templates(DB_PATH)
    if not items:
        await update.message.reply_text(
            "📋 Библиотека шаблонов пуста.\n\n"
            "Сохрани шаблон: пришли .docx/.xlsx/.pptx и ответь\n"
            "`/save_template Название`",
            parse_mode="Markdown",
        )
        return
    lines = ["📋 *Библиотека шаблонов:*\n"]
    for it in items:
        src = f" ← {it.source_name}" if it.source_name else ""
        lines.append(f"• {it.label()}{src}")
    lines.append(f"\nПолучить шаблон: `/use_template <slug>`")
    await update.message.reply_text("\n".join(lines), parse_mode="Markdown")


async def cmd_use_template(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Отправить шаблон из библиотеки."""
    if not _chat_allowed(update):
        return
    args = ctx.args or []
    name_or_slug = " ".join(args).strip()
    if not name_or_slug:
        await update.message.reply_text(
            "Использование: `/use_template <название или slug>`\n\n"
            "Список шаблонов: /templates",
            parse_mode="Markdown",
        )
        return
    from skill_doc_template import get_template, ensure_templates_table
    ensure_templates_table(DB_PATH)
    info = get_template(DB_PATH, name_or_slug)
    if not info:
        await update.message.reply_text(
            f"❌ Шаблон `{name_or_slug}` не найден.\n\nСписок: /templates",
            parse_mode="Markdown",
        )
        return
    tpl_path = Path(info.file_path)
    if not tpl_path.is_file():
        await update.message.reply_text(f"❌ Файл шаблона не найден: `{tpl_path.name}`", parse_mode="Markdown")
        return
    caption = f"📋 Шаблон: *{info.name}* (`.{info.format}`)"
    if info.description:
        caption += f"\n_{info.description}_"
    await update.message.reply_document(
        document=str(tpl_path),
        filename=f"template_{info.slug}.{info.format}",
        caption=caption,
        parse_mode="Markdown",
    )


async def cmd_delete_template(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Удалить шаблон из библиотеки."""
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text("⚠️ Только *владелец* может удалять шаблоны.", parse_mode="Markdown")
        return
    args = ctx.args or []
    name_or_slug = " ".join(args).strip()
    if not name_or_slug:
        await update.message.reply_text("Использование: `/delete_template <slug>`", parse_mode="Markdown")
        return
    from skill_doc_template import delete_template, ensure_templates_table
    ensure_templates_table(DB_PATH)
    ok = delete_template(DB_PATH, name_or_slug)
    if ok:
        await update.message.reply_text(f"🗑 Шаблон `{name_or_slug}` удалён.", parse_mode="Markdown")
    else:
        await update.message.reply_text(f"❌ Шаблон `{name_or_slug}` не найден.", parse_mode="Markdown")


async def cmd_task_quality(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Показать отчёт качества Task Registry за последние N часов."""
    if not _chat_allowed(update):
        return
    hours = 24
    if ctx.args:
        try:
            hours = int(ctx.args[0])
        except ValueError:
            pass
    if _tr_client is None:
        await update.message.reply_text(
            "⚠️ Task Registry недоступен (TASK_REGISTRY_URL не настроен или сервис не запущен).",
            parse_mode="Markdown",
        )
        return
    try:
        report = _tr_client.quality_report(hours=hours)
        if not report:
            report = f"Нет данных за последние {hours}ч."
        await update.message.reply_text(report, parse_mode="Markdown")
    except Exception as e:
        await update.message.reply_text(f"⚠️ Task Registry ошибка: {e}")


async def cmd_selftest(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    if not _is_owner(update):
        await update.message.reply_text(
            "⚠️ Только *владелец* бота (`OMI_OWNER_CHAT_IDS`).", parse_mode="Markdown"
        )
        return
    mode = (ctx.args[0].strip().lower() if ctx.args else "quick")
    if mode not in ("quick", "full"):
        mode = "quick"
    msg = await update.message.reply_text(f"🧪 Запускаю selftest ({mode})…")
    loop = asyncio.get_event_loop()
    report = await loop.run_in_executor(
        None,
        lambda: run_selftest_suite(storage, agent, mode=mode),
    )
    checks = report.get("checks", [])
    passed = sum(1 for c in checks if c.get("status") == "PASS")
    failed = sum(1 for c in checks if c.get("status") == "FAIL")
    skipped = sum(1 for c in checks if c.get("status") == "SKIP")
    recovery_note = ""
    if failed > 0:
        try:
            rec = await loop.run_in_executor(None, lambda: attempt_soft_recovery(force_restart=False))
            if rec.get("ok"):
                recovery_note = "\n- recovery: *applied* ✅"
            else:
                recovery_note = "\n- recovery: *attempted* ⚠️"
            storage._log("selftest_recovery", detail=json.dumps(rec, ensure_ascii=False)[:500])
        except Exception as e:
            recovery_note = "\n- recovery: *error*"
            storage._log("selftest_recovery", detail=f"error={e}")
    summary = (
        f"✅ *Selftest complete* ({mode})\n"
        f"- total: *{len(checks)}*\n"
        f"- pass/fail/skip: *{passed}/{failed}/{skipped}*\n"
        f"- report: `{report.get('report_md','')}`"
        f"{recovery_note}"
    )
    await msg.edit_text(summary, parse_mode="Markdown")
    try:
        md = Path(report.get("report_md", ""))
        if md.is_file():
            with md.open("rb") as fh:
                await update.message.reply_document(
                    document=fh,
                    filename=md.name,
                    caption="📋 Omi selftest checklist report",
                )
    except Exception as e:
        log.warning("selftest: send report failed: %s", e)


# ══════════════════════════════════════════════════════════════
#  NIGHT PLAN — черновик плана → подтверждение → batch enqueue
# ══════════════════════════════════════════════════════════════

_TYPE_EMOJI = {
    "ocr": "📄", "ai_summary": "🤖", "aims_summary": "📋",
    "backup": "💾", "integrity": "🔍", "sync": "🔄",
    "optimize": "⚡", "cleanup": "🗑", "rename_by_content": "✏️",
    "dedup_delete": "🗂", "custom": "📌",
}


def _night_plan_store_pending(ctx, tasks: list[dict], chat_id: int) -> str:
    token = secrets.token_hex(4)
    ctx.application.bot_data.setdefault("_night_plan_pending", {})[token] = {
        "tasks": tasks,
        "chat_id": chat_id,
    }
    return token


def _format_night_plan_draft(tasks: list[dict]) -> str:
    lines = ["🌙 *Ночной план* — старт 01:00 Dubai\n", "```"]
    lines.append(f"{'#':<3} {'Задача':<35} {'Тип':<12} {'Оценка'}")
    lines.append("─" * 60)
    total = 0
    for i, t in enumerate(tasks, 1):
        emoji = _TYPE_EMOJI.get(t["type"], "📌")
        est = t["estimate_min"]
        total += est
        h, m = divmod(est, 60)
        est_str = f"{h}ч{m:02d}м" if h else f"{m}м"
        lines.append(f"{i:<3} {t['title'][:35]:<35} {emoji}{t['type']:<11} ~{est_str}")
    lines.append("─" * 60)
    th, tm = divmod(total, 60)
    lines.append(f"{'ИТОГО':<50} ~{th}ч{tm:02d}м")
    lines.append("```")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════
#  DOCGEN — отправка файла на проверку + регистрация в реестре
# ══════════════════════════════════════════════════════════════


def _docgen_store_pending(ctx: ContextTypes.DEFAULT_TYPE, path: Path, chat_id: int) -> str:
    token = secrets.token_hex(4)
    ctx.application.bot_data.setdefault("_docgen_pending", {})[token] = {
        "path": str(path),
        "chat_id": chat_id,
    }
    return token


async def _send_docgen_for_review(
    update: Update,
    ctx: ContextTypes.DEFAULT_TYPE,
    path: Path,
    bundle_count: int,
    text_reply: str,
    *,
    approval_title: str | None = None,
    approval_summary: str | None = None,
    approval_meta: dict | None = None,
) -> None:
    """Send the generated file, then a review dialog with Approve/Revise/Skip."""
    chat_id = update.effective_chat.id if update.effective_chat else 0
    cap = f"Ready\nStatus: DOCX package submitted for review ({bundle_count})."
    with path.open("rb") as fh:
        await update.message.reply_document(
            document=fh,
            filename=path.name,
            caption=cap,
            parse_mode="Markdown",
        )
    try:
        import doc_approval as _da  # noqa: PLC0415
    except Exception:
        _da = None  # type: ignore[assignment]
    if _da is None:
        if text_reply:
            await update.message.reply_text(text_reply, parse_mode="Markdown")
        return
    await _da.send_approval_request(
        ctx.bot,
        ctx.application.bot_data,
        chat_id=chat_id,
        file_path=path,
        title=approval_title or path.stem,
        summary=approval_summary or text_reply or f"Generated document {path.name}.",
        meta=approval_meta or {},
    )


_LOCAL_PREVIEW_LINK_RE = re.compile(
    r"!\[[^\]]*\]\((?:https?://)?(?:127\.0\.0\.1|localhost)[^)]+\)|"
    r"!\((?:https?://)?(?:127\.0\.0\.1|localhost)[^)]+\)|"
    r"(?:https?://)?(?:127\.0\.0\.1|localhost)\S*",
    re.IGNORECASE,
)
_LOCAL_DOC_NAME_RE = re.compile(
    r"\b([A-Za-z0-9][A-Za-z0-9 _\-.]{2,}\.(?:docx|xlsx|pptx|pdf|txt|md))\b",
    re.IGNORECASE,
)


def _strip_local_preview_links(text: str) -> tuple[str, bool]:
    cleaned = _LOCAL_PREVIEW_LINK_RE.sub("", text or "")
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned, cleaned != (text or "").strip()


def _resolve_generated_file_from_text(text: str) -> Path | None:
    names = [m.group(1).strip() for m in _LOCAL_DOC_NAME_RE.finditer(text or "")]
    if not names:
        return None
    ws = workspace_root()
    search_dirs = [
        ws / "generated",
        ws / "result",
        ws / "anonymized",
        ws / "report",
        ws / "outbox",
        ws / "inbox" / "income",
        ws / "inbox",
        ws / "master",
    ]
    for name in names:
        for d in search_dirs:
            if not d.exists():
                continue
            direct = d / name
            if direct.is_file():
                return direct
            try:
                for p in d.rglob(name):
                    if p.is_file():
                        return p
            except Exception:
                continue
    return None


# ══════════════════════════════════════════════════════════════
#  ЖИВОЙ ЧАТ — обработка свободного текста через агента
# ══════════════════════════════════════════════════════════════

async def handle_chat(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not update.message or not update.message.text:
        return
    _chat_id = update.effective_chat.id if update.effective_chat else None
    log.info("handle_chat: chat=%s text=%r", _chat_id, (update.message.text or "")[:80])
    text = update.message.text.strip()

    if not _chat_allowed(update):
        log.info("handle_chat: chat %s not allowed (ALLOWED_CHATS=%s)", _chat_id, sorted(ALLOWED_CHATS))
        return

    chat_type = update.effective_chat.type if update.effective_chat else getattr(getattr(update.message, "chat", None), "type", "")
    bot_username = ctx.bot.username
    # В группе: по умолчанию только @бот или префикс «omi…» — иначе молчим (см. OMI_GROUP_REQUIRE_MENTION).
    # Исключение: короткое подтверждение ("yes", "да", "👍" …) пропускается если есть ожидающее подтверждение.
    if chat_type in ("group", "supergroup") and OMI_GROUP_REQUIRE_MENTION:
        named_in_text = _mentions_omi(text)
        log.info(
            "handle_chat: group mention check named=%s @mentioned=%s text=%r",
            named_in_text,
            f"@{bot_username}" in text,
            text[:60],
        )
        if f"@{bot_username}" not in text and not named_in_text:
            # Allow bare affirmatives through when Omi is waiting for a confirmation
            if _chat_id and _is_short_affirmative(text) and _has_pending_confirm(ctx, _chat_id):
                log.info("handle_chat: bare affirmative %r passed — pending confirm for chat %s", text, _chat_id)
            else:
                return

    if _maybe_handle_registry_audit_message is not None:
        try:
            if await _maybe_handle_registry_audit_message(update, ctx, update.message.text or ""):
                return
        except Exception as e:
            log.warning("registry audit routing failed: %s", e)
    if _maybe_handle_docsreg_message is not None:
        try:
            if await _maybe_handle_docsreg_message(update, ctx, update.message.text or ""):
                return
        except Exception as e:
            log.warning("docsreg launch handling failed: %s", e)

    # Та же линия, что AXI_CHAT_NIM_UNTIL_OLLAMA_READY + фоновый warm: подгрузка OMI_MODEL в Ollama
    # на каждое сообщение (без обязательного ответа Omi), чтобы при @Omi уже шёл qwen/локалка.
    try:
        if QWEN_PC_ASSIST_STACK and QWEN_PC_ASSIST_WARM_ON_TELEGRAM:
            _schedule_ollama_background_warm()
        elif chat_type in ("group", "supergroup") and OMI_GROUP_WARM_ON_MESSAGE:
            _schedule_ollama_background_warm()
    except Exception:
        log.debug("handle_chat: Ollama warm schedule failed", exc_info=True)

    # Убрать упоминание бота из текста (@ и словесный префикс "Omi,")
    clean = text.replace(f"@{bot_username}", "").strip()
    clean = _strip_omi_prefix(clean)

    # Stage 4E Option A: live handler path exists behind gate.
    # Gate defaults to disabled now because registry runtime_activation is false.
    if should_run_document_dialogue_gap_check():
        try:
            from document_dialogue_gap_route_wrapper import route_document_dialogue_gap_check  # noqa: PLC0415

            _payload = {
                "source": "omi_live_handler",
                "chat_id": str(_chat_id or ""),
                "user_text": clean,
                "available_documents": [],
                "context": {"workflow": "document_dialogue_gap_check", "test_only": False},
            }
            _r = route_document_dialogue_gap_check(_payload, runtime_activation=True)
            _rs = str(_r.get("route_status") or "")
            _reason = str(_r.get("reason") or "")
            _handoff = str(_r.get("recommended_handoff") or "")

            if _rs == "ROUTE_BLOCKED":
                _msg = (
                    f"🛑 Request blocked by document dialogue safety gate.\n{_reason}"
                    if lang != "ru"
                    else f"🛑 Запрос заблокирован защитным шлюзом диалога по документам.\n{_reason}"
                )
                await update.message.reply_text(_msg)
                return
            if _rs == "ROUTE_NEEDS_CONTEXT":
                _missing = _r.get("missing_inputs") or []
                _miss = ", ".join(str(x) for x in _missing) if isinstance(_missing, list) else ""
                _msg = (
                    f"ℹ️ Need more context before routing.\n{_reason}\nMissing: {_miss}".strip()
                    if lang != "ru"
                    else f"ℹ️ Нужен дополнительный контекст перед маршрутизацией.\n{_reason}\nНе хватает: {_miss}".strip()
                )
                await update.message.reply_text(_msg)
                return
            if _rs == "ROUTE_READY_FOR_HANDOFF":
                _msg = (
                    f"✅ Request validated for safe handoff ({_handoff})."
                    if lang != "ru"
                    else f"✅ Запрос проверен для безопасной передачи ({_handoff})."
                )
                await update.message.reply_text(_msg)
                return
        except Exception as _dg_err:
            _gate_state = read_document_dialogue_gap_gate_state()
            log.warning(
                "document-dialogue-gap-check gated route error: %s | gate=%s",
                _dg_err,
                _gate_state,
            )

    # Не перехватываем команды, адресованные другим ботам: /tasks@AxiOMsphere_bot
    if _is_command_for_other_bot(clean, bot_username):
        return
    # И не отвечаем на явные обращения к Axi в группе.
    if _mentions_axi(clean):
        return

    # ── Doc approval: intercept rework description ─────────────
    if _chat_id:
        import doc_approval as _da
        if _da.get_rework_token(ctx.application.bot_data, _chat_id):
            consumed = await _da.handle_rework_text(
                _chat_id,
                clean,
                ctx.application.bot_data,
                lambda t, **kw: update.message.reply_text(t, **kw),
            )
            if consumed:
                return

    # ── Intake pipeline: intercept path input when awaiting source ─────────────
    if _chat_id:
        try:
            from omi_intake_handler import handle_intake_path_input
            if await handle_intake_path_input(update, _chat_id, clean):
                return
        except Exception as e:
            log.warning("handle_intake_path_input failed: %s", e)

    # Auto-capture "task first, file later" flow without requiring /analyze.
    if _chat_id and not _get_pending_analyze(_chat_id) and _looks_like_task_then_file_flow(clean):
        _set_pending_analyze(_chat_id, clean[:1200])
        log.info("handle_chat: auto-enabled pending analyze for chat=%s", _chat_id)

    if _looks_like_task_close_intent(clean):
        handled, msg = _close_omi_task_from_chat(str(_chat_id or ""), clean)
        if handled:
            await update.message.reply_text(msg, parse_mode="Markdown")
            return

    if _looks_like_doc_synthesis_request(clean):
        if await _handle_group_doc_synthesis(update, ctx, clean, _reply_lang(clean)):
            return

    lang = _reply_lang(clean)

    # ── Short "register file" after upload: explain pipeline; do not ask for URL (LLM clarify noise).
    if _chat_id and _looks_like_register_file_short_request(clean):
        last_fn = ""
        try:
            _ctx = storage.get_chat_context(_chat_id)
            last_fn = str(_ctx.get("last_file_name") or "").strip()
        except Exception:
            pass
        if last_fn:
            if lang == "ru":
                _reg_help = (
                    f"Файл **`{last_fn}`** уже принят и стоит в очереди **batch-pipeline**: "
                    "OCR → анонимизация → запись в реестр **AIMS**.\n\n"
                    "Отдельно ничего присылать не нужно — запись в базе появится после прохождения очереди.\n"
                    "Чтобы позже указать раздел **P01–P11** (когда строка уже в реестре): например "
                    f"`Переместь {last_fn} в P06`.\n\n"
                    "Статус: `/status`."
                )
            else:
                _reg_help = (
                    f"**`{last_fn}`** is already queued: **batch-pipeline** "
                    "(OCR → anonymize → register in AIMS).\n\n"
                    "You do **not** need to send a link again — the row appears when processing finishes.\n"
                    "To set **process P01–P11** after the document is in the registry: e.g. "
                    f"`move {last_fn} to P06`.\n\n"
                    "Status: `/status`."
                )
        else:
            if lang == "ru":
                _reg_help = (
                    "Чтобы **зарегистрировать файл**, сначала **пришлите его вложением** в этот чат — "
                    "тогда он попадёт в очередь OCR → анонимизация → реестр.\n\n"
                    "Если файл уже загружали, напишите **точное имя файла** (как в сообщении с документом)."
                )
            else:
                _reg_help = (
                    "To **register a file**, **upload it here** as an attachment first — it will join the "
                    "OCR → anonymize → registry pipeline.\n\n"
                    "If you already uploaded it, send the **exact filename** as shown in Telegram."
                )
        await update.message.reply_text(_reg_help, parse_mode="Markdown")
        _omi_dialog_append(ctx, _chat_id, "user", clean)
        _omi_dialog_append(ctx, _chat_id, "assistant", _reg_help.replace("*", "").replace("`", ""))
        return

    # ── Fast-path: AIMS registry report → DOCX (no AI needed) ────────────
    try:
        from aims_report import is_registry_request, build_registry_docx
        if is_registry_request(clean):
            _thinking = await update.message.reply_text(
                "Omi: generating AIMS Registry report..." if lang != "ru"
                else "Omi: генерирую отчёт по реестру AIMS..."
            )
            docx_path = build_registry_docx()
            with docx_path.open("rb") as _fh:
                await update.message.reply_document(
                    document=_fh,
                    filename=docx_path.name,
                    caption="Omi: AIMS Document Registry — full hierarchy report",
                )
            docx_path.unlink(missing_ok=True)
            await _thinking.delete()
            return
    except Exception as _reg_err:
        log.warning("aims_report fast-path failed: %s", _reg_err)

    # ── Contextual document-work routing: choose DOCSREG vs DOCGEN from intent ─
    try:
        from chat_intent_router import DOCUMENT_WORK_CMDS, classify  # noqa: PLC0415

        _doc_dialog = _omi_dialog_messages_for_llm(ctx, _chat_id)
        _doc_routed = await asyncio.to_thread(
            classify,
            clean,
            DOCUMENT_WORK_CMDS,
            dialog_messages=_doc_dialog,
        )
        if _doc_routed:
            _doc_cmd = _doc_routed[0]
            if _doc_cmd == "docsreg" and _maybe_handle_docsreg_message is not None:
                try:
                    if await _maybe_handle_docsreg_message(update, ctx, clean):
                        return
                except Exception as e:
                    log.warning("docsreg launch handling failed: %s", e)
            elif _doc_cmd == "docgen":
                if await _handle_group_doc_synthesis(update, ctx, clean, _reply_lang(clean)):
                    return
    except Exception as e:
        log.debug("document work routing failed: %s", e, exc_info=True)

    # NLP intent routing — deterministic pipeline via registry
    try:
        from pipeline_registry import OMI_REGISTRY, params_check, args_to_params  # noqa: PLC0415
        from chat_intent_router import classify, OMI_CMDS  # noqa: PLC0415

        _clarify_store = ctx.application.bot_data.setdefault("_omi_clarify_pending", {})

        # ── Resume pending clarification if user replied to a question ────────
        _pending = _clarify_store.get(_chat_id) if _chat_id else None
        if _pending:
            _p_intent  = _pending["intent"]
            _p_params  = _pending["params"]
            _p_missing = _pending["missing_param"]
            _tmpl_p = OMI_REGISTRY.get(_p_intent)
            if _tmpl_p:
                _p_params[_p_missing] = clean.strip()
                _still_missing = params_check(_tmpl_p, _p_params)
                if _still_missing:
                    # Another param needed — ask next question
                    _np = _still_missing[0]
                    _clarify_store[_chat_id] = {
                        "intent": _p_intent, "params": _p_params,
                        "missing_param": _np.name,
                    }
                    _q = _np.clarify_ru if lang == "ru" else _np.clarify_en
                    await update.message.reply_text(_q)
                    return
                else:
                    _clarify_store.pop(_chat_id, None)
                    await _tmpl_p.execute(update, ctx, _p_params)
                    return

        # ── Fresh intent classification ────────────────────────────────────────
        _omi_routed = await asyncio.to_thread(classify, clean, OMI_CMDS)
        if _omi_routed:
            _omi_cmd, _omi_args = _omi_routed
            _tmpl = OMI_REGISTRY.get(_omi_cmd)
            if _tmpl:
                _params = args_to_params(_tmpl, _omi_args)
                _missing = params_check(_tmpl, _params)
                if _missing:
                    _mp = _missing[0]
                    if _chat_id:
                        _clarify_store[_chat_id] = {
                            "intent": _omi_cmd, "params": _params,
                            "missing_param": _mp.name,
                        }
                    _q = _mp.clarify_ru if lang == "ru" else _mp.clarify_en
                    await update.message.reply_text(_q)
                    return
                await _tmpl.execute(update, ctx, _params)
                return
    except Exception:
        pass

    try:
        from bot.axi_scope_guard import ScopeDecision, build_safe_internal_replacement_plan, classify_omi_scope
        from bot.axi_response_templates import REFUSAL_UNRELATED
        from bot.axi_public_answers import try_public_answer

        guard = classify_omi_scope(clean)
        if not guard.allowed:
            reply = guard.response or REFUSAL_UNRELATED
            if guard.decision == ScopeDecision.SENSITIVE_INTERNAL_REQUEST:
                low = clean.lower()
                if any(term in low for term in ("code", "python", "implement", "build")):
                    reply = (
                        f"{reply}\n\n"
                        f"{build_safe_internal_replacement_plan('AIMS admin workflow integration')}"
                    )
            await update.message.reply_text(reply)
            return

        public_answer = try_public_answer(clean)
        if public_answer:
            await update.message.reply_text(public_answer)
            return
    except ImportError:
        log.warning("bot scope guard unavailable for Omi; continuing without guard")

    task_seq = _next_task_seq(ctx)
    # Не «task #N» — у Axi это номер пайплайн-задачи; здесь только порядковый номер запроса к Omi.
    task_ref = f"запрос №{task_seq}" if lang == "ru" else f"request #{task_seq}"
    _tr_task_id = _tr_register(
        clean[:120] or "(пустое сообщение)",
        chat_id=str(_chat_id or ""),
        source="group" if chat_type in ("group", "supergroup") else "omi",
    )
    _tr_start(_tr_task_id, assigned_to="omi")
    thinking = await update.message.reply_text(
        f"⚙️ Omi: обрабатываю {task_ref}…" if lang == "ru" else f"⚙️ Omi: processing {task_ref}…"
    )
    _heartbeat = asyncio.create_task(_exec_heartbeat(thinking, task_ref, lang=lang))

    _tm_omi = False
    try:
        chat_id = update.effective_chat.id if update.effective_chat else None
        _tm_omi = bool(chat_id and test_mode_shared.omi_pop_pending_and_enable(chat_id))
        owner = _is_owner(update)
        _dialogue_prior = _omi_dialog_messages_for_llm(ctx, chat_id)

        def _run_omi_chat():
            return agent.chat(
                clean,
                chat_id=chat_id,
                is_owner=owner,
                dialogue_context=_dialogue_prior,
            )

        _agent_t0 = time.perf_counter()
        try:
            result = await asyncio.get_event_loop().run_in_executor(None, _run_omi_chat)
        finally:
            _heartbeat.cancel()
            try:
                await _heartbeat
            except asyncio.CancelledError:
                pass

        _agent_dt = time.perf_counter() - _agent_t0
        _pre_tm, _foot_tm = ("", "")
        if _tm_omi and chat_id is not None:
            _pre_tm, _foot_tm = test_mode_shared.omi_tm_prefix_and_footer(
                chat_id, _agent_dt, lang == "ru"
            )

        _tr_action = ""
        if isinstance(result, ChatResult):
            _tr_action = getattr(result, "action", "") or ""
        _tr_done(_tr_task_id, summary=_tr_action or f"dt={_agent_dt:.1f}s")

        _omi_dialog_append(ctx, chat_id, "user", clean)
        if isinstance(result, ChatResult) and getattr(result, "silent_handoff", False):
            await thinking.delete()
            return

        if isinstance(result, ChatResult):
            _asst = (result.text or "").strip()
            if not _asst and result.document_path:
                _asst = f"[файл: {result.document_path.name}]"
            if not _asst and result.night_plan_tasks:
                _asst = "[ночной план: черновик]"
            if _asst:
                _omi_dialog_append(ctx, chat_id, "assistant", _asst)
        else:
            _omi_dialog_append(ctx, chat_id, "assistant", str(result))

        if isinstance(result, ChatResult) and result.night_plan_tasks:
            await thinking.delete()
            tasks = result.night_plan_tasks
            draft_text = _format_night_plan_draft(tasks)
            token = _night_plan_store_pending(ctx, tasks, update.effective_chat.id)
            markup = InlineKeyboardMarkup([[
                InlineKeyboardButton("✅ Запланировать", callback_data=f"np_yes:{token}"),
                InlineKeyboardButton("✗ Отмена",         callback_data=f"np_no:{token}"),
            ]])
            await update.message.reply_text(
                _pre_tm + draft_text + "\n\nПодтвердить план?" + _foot_tm,
                parse_mode="Markdown",
                reply_markup=markup,
            )
        elif isinstance(result, ChatResult) and result.document_path and result.direct_send:
            await thinking.delete()
            path = result.document_path
            with path.open("rb") as fh:
                await update.message.reply_document(
                    document=fh,
                    filename=path.name,
                    caption=f"📎 {path.name}",
                )
            _short_status = (
                f"Готово\nСтатус: Активная задача — отправка `{path.name}` в чат."
                if lang == "ru"
                else f"Done\nStatus: Active task — sent `{path.name}` to chat."
            )
            try:
                await update.message.reply_text(_short_status, parse_mode="Markdown")
            except Exception:
                await update.message.reply_text(_short_status)
            if _foot_tm.strip():
                await update.message.reply_text(_foot_tm.strip(), parse_mode="Markdown")
        elif isinstance(result, ChatResult) and result.document_path:
            await thinking.delete()
            await _send_docgen_for_review(
                update,
                ctx,
                result.document_path,
                result.bundle_count or 0,
                (result.text or "") + _foot_tm,
                approval_title=result.document_path.stem,
                approval_summary=(result.text or "Generated document ready for review."),
                approval_meta={
                    "registration_mode": "document",
                    "aims_process": "DOCGEN",
                    "source_text": clean,
                },
            )
        else:
            text = result.text if isinstance(result, ChatResult) else str(result)
            text, _had_local_links = _strip_local_preview_links(text)
            if _had_local_links:
                _resolved = _resolve_generated_file_from_text(result.text if isinstance(result, ChatResult) else str(result))
                if _resolved and _resolved.is_file():
                    if _chat_id:
                        if _is_confirmation_request(text):
                            _set_pending_confirm(ctx, _chat_id)
                        else:
                            _clear_pending_confirm(ctx, _chat_id)
                    await thinking.delete()
                    with _resolved.open("rb") as _fh:
                        await update.message.reply_document(
                            document=_fh,
                            filename=_resolved.name,
                            caption=f"📎 {_resolved.name}",
                        )
                    _note = text or (
                        "Файл отправлен напрямую в чат для проверки."
                        if lang == "ru"
                        else "File sent to chat directly for review."
                    )
                    try:
                        await update.message.reply_text(_pre_tm + _note + _foot_tm, parse_mode="Markdown")
                    except Exception:
                        await update.message.reply_text(_pre_tm + _note + _foot_tm)
                    return
            low_text = text.lower()
            if (
                ("ollama failed" in low_text)
                or ("connection refused" in low_text)
                or ("errno 111" in low_text)
                or ("модель недоступна" in low_text)
            ):
                text = (
                    "Omi: I am online, but AI analysis is temporarily unavailable. "
                    "Use live registry requests and /help commands."
                    if lang != "ru"
                    else "Omi: я онлайн, но AI-анализ временно недоступен. "
                    "Используйте запросы к реестру и команду /help."
                )
            # Track whether agent is waiting for user confirmation
            if _chat_id:
                if _is_confirmation_request(text):
                    _set_pending_confirm(ctx, _chat_id)
                else:
                    _clear_pending_confirm(ctx, _chat_id)
            _TG_LIMIT = 4000
            _chunks: list[str] = []
            if len(text) <= _TG_LIMIT:
                _chunks = [text]
            else:
                # Split on newlines to avoid cutting mid-line
                _buf = ""
                for _line in text.splitlines(keepends=True):
                    if len(_buf) + len(_line) > _TG_LIMIT:
                        if _buf:
                            _chunks.append(_buf.rstrip())
                        _buf = _line
                    else:
                        _buf += _line
                if _buf.strip():
                    _chunks.append(_buf.rstrip())

            for _ci, _chunk in enumerate(_chunks):
                _piece = _chunk
                if _ci == 0:
                    _piece = _pre_tm + _piece
                if _ci == len(_chunks) - 1:
                    _piece = _piece + _foot_tm
                try:
                    if _ci == 0:
                        await thinking.edit_text(_piece, parse_mode="Markdown")
                    else:
                        await update.message.reply_text(_piece, parse_mode="Markdown")
                except Exception as _md_err:
                    log.warning("handle_chat markdown failed (part %d): %s — retrying plain", _ci + 1, _md_err)
                    try:
                        if _ci == 0:
                            await thinking.edit_text(_piece)
                        else:
                            await update.message.reply_text(_piece)
                    except Exception as _plain_err:
                        log.error("handle_chat send failed (part %d): %s", _ci + 1, _plain_err)
                        break
    except Exception as e:
        _fail_tm = ""
        if _tm_omi and _chat_id:
            _fail_tm = "\n\n" + test_mode_shared.test_mode_failure_report(
                str(e or ""), lang_ru=(lang == "ru")
            )
        if _chat_id:
            test_mode_shared.omi_tm_abort(_chat_id)
        _tr_stuck(_tr_task_id, error=str(e or "")[:200])
        log.error(f"Agent error: {e}")
        err_text = str(e or "")
        # Do not leak transport/runtime details to users; keep chat human-friendly.
        if ("ollama" in err_text.lower()) or ("connection refused" in err_text.lower()) or ("errno 111" in err_text.lower()):
            # If this looked like a registry request, provide deterministic DB output even with LLM down.
            if any(k in clean.lower() for k in ("register", "registry", "реестр", "documents", "files", "документ", "файл")):
                rl = _reply_lang(clean)
                if _registry_without_time_window(clean):
                    _snap = storage.list_registry_snapshot(75, lang=rl)
                    try:
                        await thinking.edit_text(_snap, parse_mode="Markdown")
                    except Exception:
                        await thinking.edit_text(_snap)
                else:
                    _docs = storage.list_documents_last_hours(24, lang=rl)
                    try:
                        await thinking.edit_text(_docs, parse_mode="Markdown")
                    except Exception:
                        await thinking.edit_text(_docs)
                if _fail_tm.strip():
                    try:
                        await update.message.reply_text(_fail_tm.strip(), parse_mode="Markdown")
                    except Exception:
                        await update.message.reply_text(_fail_tm.strip())
            else:
                _ollama_msg = (
                    "Omi: I am online, but AI analysis is temporarily unavailable. "
                    "You can still use live registry requests and /help commands."
                )
                try:
                    await thinking.edit_text(_ollama_msg + _fail_tm, parse_mode="Markdown")
                except Exception:
                    await thinking.edit_text(_ollama_msg + _fail_tm)
            return
        try:
            await thinking.edit_text(
                "Omi: temporary processing error. Please try again or use /help." + _fail_tm,
                parse_mode="Markdown",
            )
        except Exception:
            await thinking.edit_text(
                "Omi: temporary processing error. Please try again or use /help." + _fail_tm
            )


async def _notify_new_registrations_job(ctx: ContextTypes.DEFAULT_TYPE):
    """
    Omi self-voice: announce newly registered OCR documents in chat.
    Keeps last seen id in bot_data to avoid duplicates.
    """
    if not NOTIFY_CHAT_IDS or not OCR_REGISTRY_DB.exists():
        return

    key = "_omi_last_notified_ocr_id"
    last_id = int(ctx.application.bot_data.get(key, 0) or 0)
    try:
        con = sqlite3.connect(str(OCR_REGISTRY_DB))
        cur = con.cursor()
        if last_id <= 0:
            row = cur.execute("SELECT COALESCE(MAX(id), 0) FROM ocr_documents").fetchone()
            ctx.application.bot_data[key] = int(row[0] or 0)
            con.close()
            return
        rows = cur.execute(
            """
            SELECT id, file_name
            FROM ocr_documents
            WHERE id > ?
            ORDER BY id ASC
            LIMIT 50
            """,
            (last_id,),
        ).fetchall()
        con.close()
    except Exception as e:
        log.warning("notify_new_registrations: read failed: %s", e)
        return

    if not rows:
        return

    names = [str(r[1] or "").strip() for r in rows if str(r[1] or "").strip()]
    if names:
        if len(names) == 1:
            text = f"Omi: document registered successfully: `{names[0]}`"
        else:
            lines = ["Omi: documents registered successfully:"]
            lines.extend([f"- `{n}`" for n in names[:30]])
            if len(names) > 30:
                lines.append(f"- ... and {len(names) - 30} more")
            text = "\n".join(lines)
        for chat_id in NOTIFY_CHAT_IDS:
            try:
                await ctx.bot.send_message(chat_id=chat_id, text=text, parse_mode="Markdown")
            except Exception as e:
                log.warning("notify_new_registrations: send failed chat=%s err=%s", chat_id, e)

    ctx.application.bot_data[key] = int(rows[-1][0])

# ══════════════════════════════════════════════════════════════
#  INLINE KEYBOARD CALLBACKS
# ══════════════════════════════════════════════════════════════

async def handle_callback(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not _chat_allowed(update):
        return
    query = update.callback_query
    await query.answer()
    data = query.data or ""

    # ── Document approval (da_yes / da_rework / da_skip) ───────
    if data.startswith(("da_yes:", "da_rework:", "da_skip:")):
        import doc_approval as _da
        if data.startswith("da_yes:"):
            await _da.handle_da_yes(query, ctx.application.bot_data)
        elif data.startswith("da_rework:"):
            await _da.handle_da_rework(query, ctx.application.bot_data)
        else:
            await _da.handle_da_skip(query, ctx.application.bot_data)
        return

    # ── Night plan confirmation ─────────────────────────────────
    if data.startswith("np_yes:"):
        token = data.split(":", 1)[1]
        pending = ctx.application.bot_data.setdefault("_night_plan_pending", {}).pop(token, None)
        if not pending or pending.get("chat_id") != query.message.chat_id:
            await query.answer("Запрос устарел.", show_alert=True)
            return
        tasks = pending["tasks"]
        results = []
        for t in tasks:
            r = storage.enqueue_night_task(
                t["title"], estimate_min=t["estimate_min"], source="night_plan"
            )
            results.append(r)
        summary = f"✅ *{len(tasks)} задач* добавлено в ночной план.\nСтарт: *01:00 Dubai*"
        await query.edit_message_text(
            summary + "\n\n" + storage.format_night_plan(),
            parse_mode="Markdown",
            reply_markup=None,
        )
        return

    if data.startswith("np_no:"):
        token = data.split(":", 1)[1]
        ctx.application.bot_data.setdefault("_night_plan_pending", {}).pop(token, None)
        await query.edit_message_text("❌ Ночной план отменён.", reply_markup=None)
        return

    if data.startswith("dg_yes:"):
        token = data.split(":", 1)[1]
        pending = ctx.application.bot_data.setdefault("_docgen_pending", {}).pop(
            token, None
        )
        if (
            not pending
            or pending.get("chat_id") != query.message.chat_id
        ):
            await query.answer("Запрос устарел или из другого чата.", show_alert=True)
            return
        path = Path(pending["path"])
        reg = storage.register_generated_bundle(path)
        await query.answer("Записано в реестр.")
        try:
            await query.edit_message_caption(
                caption=f"✅ *{reg}*\n\n`{path.name}`",
                parse_mode="Markdown",
                reply_markup=None,
            )
        except Exception:
            await query.message.reply_text(reg, parse_mode="Markdown")
        return

    if data.startswith("dg_no:"):
        token = data.split(":", 1)[1]
        ctx.application.bot_data.setdefault("_docgen_pending", {}).pop(token, None)
        await query.answer("Без регистрации.")
        try:
            await query.edit_message_caption(
                caption=(
                    "❌ *Регистрация отменена.* Файл остаётся в каталоге `generated/` "
                    "на сервере.\n\n_Повторить: `/docgen`_"
                ),
                parse_mode="Markdown",
                reply_markup=None,
            )
        except Exception:
            await query.message.reply_text("Регистрация отменена.")
        return

    if data == "menu_status":
        status = storage.get_status()
        lines = [f"📊 *Статус Omi*\n",
                 f"🗄 БД: `{status['db_path']}`",
                 f"📦 Документов: *{status['total_docs']}*\n",
                 "📁 *Разделы:*"]
        for p in status["processes"]:
            lines.append(f"  `{p['code']}` {p['name']} — {p['count']} файлов")
        await query.edit_message_text(
            "\n".join(lines), parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [
                    InlineKeyboardButton("🔘 Разделы (кнопки)", callback_data="menu_processes"),
                ],
                [
                    InlineKeyboardButton("« Назад", callback_data="menu_back"),
                ],
            ]),
        )

    elif data == "menu_move":
        await query.edit_message_text(
            "📁 *Переместить файлы*\n\n"
            "Выбери целевой раздел кнопкой ниже или напиши текстом:\n"
            "• `Переместь <файл или ключевое слово> в P04`\n\n"
            "Команда:\n"
            "`/move <имя_файла> <P0X>`",
            parse_mode="Markdown",
            reply_markup=process_keyboard(),
        )

    elif data == "menu_processes":
        await query.edit_message_text(
            "📂 *Разделы AIMS* — названия и число документов из реестра.\n\n"
            "Нажми раздел для подсказки команд или «Назад» в главное меню.",
            parse_mode="Markdown",
            reply_markup=process_keyboard(),
        )

    elif data.startswith("proc_"):
        code = (data.replace("proc_", "") or "").strip().upper()
        if not code:
            await query.edit_message_text(
                "⚠️ Некорректный код раздела.",
                reply_markup=process_keyboard(),
            )
            return
        cnt = 0
        reg_name = ""
        for p in storage.list_processes():
            if str(p.get("code") or "").upper() == code:
                cnt = int(p.get("count") or 0)
                reg_name = str(p.get("name") or "").strip()
                break
        disp = _process_button_title(code, cnt, registry_name=reg_name or None)
        await query.edit_message_text(
            f"📂 *Раздел:* {disp}\n\n"
            f"Примеры:\n"
            f"• `Переместь <файл> в {code}`\n"
            f"• `/move <имя_файла> {code}`\n"
            f"• `Найди документы в {code}`",
            parse_mode="Markdown",
            reply_markup=process_keyboard(),
        )

    elif data == "menu_archive":
        buttons = [[
            InlineKeyboardButton(
                _truncate_btn_label(
                    "🗄 "
                    + _process_button_title(
                        p["code"],
                        p.get("count"),
                        registry_name=str(p.get("name") or "") or None,
                    )
                ),
                callback_data=f"archive_{p['code']}",
            )
        ] for p in storage.list_processes()]
        buttons.append([InlineKeyboardButton("Архив всё до 2023", callback_data="archive_all_2023")])
        buttons.append([InlineKeyboardButton("« Назад", callback_data="menu_back")])
        await query.edit_message_text(
            "🗄 *Архивировать документы*\nВыбери раздел:",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(buttons),
        )

    elif data.startswith("archive_"):
        target = data.replace("archive_", "")
        if target == "all_2023":
            result = storage.archive_documents(before_year=2023)
        else:
            result = storage.archive_documents(process=target)
        await query.edit_message_text(result, parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("« Назад", callback_data="menu_back")
            ]]))

    elif data == "menu_new_process":
        await query.edit_message_text(
            "➕ *Создать новый раздел*\n\n"
            "Напиши в чате:\n"
            "`Создай новый процесс P11 Corrosion Management`\n\n"
            "Или команда:\n"
            "`/newprocess P11 Corrosion_Management`",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("« Назад", callback_data="menu_back")
            ]])
        )

    elif data == "menu_migrate":
        plan = storage.get_dgx_migration_plan()
        await query.edit_message_text(
            plan, parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("« Назад", callback_data="menu_back")
            ]])
        )

    elif data == "menu_search":
        await query.edit_message_text(
            "🔍 *Найти документы*\n\n"
            "Напиши в чате что ищешь, например:\n"
            "• `Найди все документы по RBI`\n"
            "• `Покажи мастер-файлы P02`\n\n"
            "Или команда: `/search <запрос>`",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("« Назад", callback_data="menu_back")
            ]])
        )

    elif data == "menu_rename_file":
        await query.edit_message_text(
            "📝 *Переименование файла по запросу*\n\n"
            "• `/rename_file_<новое_имя>` — reply на файл или последний файл в контексте.\n"
            "• `/rename_file_<id>_<новое_имя>` — по номеру из списка реестра (как `#348` → `348_…`).\n\n"
            "Для пакетного/AI rename: `/rename_by_context`.",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("« Назад", callback_data="menu_back")
            ]])
        )

    elif data == "menu_docsreg":
        await query.edit_message_text(
            "📋 *DOCSREG — регистрация документов*\n\n"
            "Проверяет и регистрирует стандарты из папки или файла.\n"
            "Поддерживаемые форматы: `.md` `.txt` `.rst` `.docx` `.pdf` `.pptx` `.xlsx` `.xls` `.csv` `.html`\n\n"
            "*Запуск:*\n"
            "`/docsreg` — обработать папку Standards по умолчанию\n"
            "`/docsreg /media/.../Standards` — указать папку явно\n"
            "`/docsreg /path/to/file.docx` — один файл\n\n"
            "*Остановить во время работы:*\n"
            "`stop` или `стоп` — завершить после текущего файла",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("« Назад", callback_data="menu_back")
            ]])
        )

    elif data == "menu_test_mode":
        await query.edit_message_text(
            "🧪 *Test mode*\n\n"
            "Отправь команду в чат (в группе — с `@"
            + (ctx.bot.username or "бот")
            + "` или префиксом `Omi`):\n"
            "`/test_mode`\n\n"
            "Следующее сообщение к боту получит пошаговый отчёт (intent, skills, время) и "
            "*Complete successful*, затем режим выключится.\n\n"
            "Опечатка `/tect_mode` тоже работает.",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("« Назад", callback_data="menu_back")
            ]])
        )

    elif data == "menu_back":
        await query.edit_message_text(
            "📋 *Главное меню Omi*",
            parse_mode="Markdown",
            reply_markup=main_menu_keyboard(),
        )

# ══════════════════════════════════════════════════════════════
#  PATH 2: TELEGRAM FILE UPLOAD (max 5/day per chat)
# ══════════════════════════════════════════════════════════════

async def handle_file_upload(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    """Receive a document or photo via Telegram → drop into batch_inbox/ for pipeline processing."""
    if not _chat_allowed(update):
        return
    if update.message is None:
        return

    chat_id = update.effective_chat.id
    is_group_chat = bool(update.effective_chat and update.effective_chat.type in ("group", "supergroup"))
    send_upload_notice = (not is_group_chat) or OMI_GROUP_UPLOAD_NOTIFICATIONS
    analyze_prompt = _get_pending_analyze(chat_id)
    # Upload pipeline notices are intentionally fixed in English.
    lang = "en"
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # Count uploads for logging (no hard limit for Omi Telegram)
    _date, _count = _DAILY_UPLOADS.get(chat_id, ("", 0))
    if _date != today:
        _date, _count = today, 0

    # Get the file object
    tg_file = None
    original_name = "document"
    if update.message.document:
        doc = update.message.document
        original_name = doc.file_name or f"upload_{doc.file_id[-8:]}"
        tg_file = await ctx.bot.get_file(doc.file_id)
    elif update.message.photo:
        photo = update.message.photo[-1]  # largest size
        original_name = f"photo_{photo.file_id[-8:]}.jpg"
        tg_file = await ctx.bot.get_file(photo.file_id)
    else:
        return

    # Prepare batch_inbox
    BATCH_INBOX.mkdir(parents=True, exist_ok=True)
    # Download to temp first, then place with SHA256 dedup logic
    _tmp = BATCH_INBOX / f".tmp_{original_name}"
    try:
        await tg_file.download_to_drive(str(_tmp))
    except Exception as exc:
        log.error("File download failed: %s", exc)
        await update.message.reply_text("Download error — please try again.")
        return

    # SHA256 dedup: identical content → skip; different content → _v2, _v3 ...
    import hashlib as _hashlib
    _incoming_sha = _hashlib.sha256(_tmp.read_bytes()).hexdigest()
    dest = BATCH_INBOX / original_name
    if dest.exists():
        if _hashlib.sha256(dest.read_bytes()).hexdigest() == _incoming_sha:
            _tmp.unlink(missing_ok=True)
            _dup_msg = (
                f"⚠️ `{original_name}` уже в очереди — файл идентичен (SHA256). Загрузка пропущена."
                if lang == "ru" else
                f"⚠️ `{original_name}` is already queued — file is identical (SHA256). Upload skipped."
            )
            if send_upload_notice:
                await update.message.reply_text(_dup_msg, parse_mode="Markdown")
            return
        # Different content — use versioned name _v2, _v3 ...
        stem, suffix = Path(original_name).stem, Path(original_name).suffix
        for _n in range(2, 100):
            _candidate = BATCH_INBOX / f"{stem}_v{_n}{suffix}"
            if not _candidate.exists():
                dest = _candidate
                break
    _tmp.rename(dest)
    original_name = dest.name

    # --- Silent legacy format conversion (.doc→.docx, .xls→.xlsx, .ppt→.pptx) ---
    if _is_legacy(dest):
        try:
            dest, _converted = await _convert_inplace_async(dest)
            if _converted:
                original_name = dest.name
        except Exception as _ce:
            log.warning("legacy_convert_failed: %s", _ce)

    ext_up = dest.suffix.lower()
    try:
        from axi_archive_extract import maybe_fix_extension_from_file_magic

        dest, ext_up, original_name = maybe_fix_extension_from_file_magic(
            dest, ext_up, original_name or ""
        )
    except Exception as _mf:
        log.warning("maybe_fix_extension_from_file_magic: %s", _mf)

    # AI rename on upload (before registration pipeline) based on file content.
    # Only rename files with coded/meaningless names (not human-readable titles).
    rename_note = ""
    if OMI_UPLOAD_AI_RENAME_ENABLED and not _has_meaningful_name(dest.name):
        try:
            new_name, reason = _suggest_filename(
                dest,
                dest.name,
                use_external=OMI_UPLOAD_AI_RENAME_USE_AXI,
            )
            if new_name and new_name != dest.name:
                target = dest.with_name(new_name)
                if target.exists():
                    stem, suffix = target.stem, target.suffix
                    target = target.with_name(f"{stem}_{int(datetime.now(timezone.utc).timestamp())}{suffix}")
                dest.rename(target)
                rename_note = f"\nAI rename: `{dest.name}` → `{target.name}` ({reason})."
                dest = target
            else:
                rename_note = f"\nAI rename: kept `{dest.name}` ({reason})."
        except Exception as e:
            log.warning("upload ai rename skipped: %s", e)
            rename_note = "\nAI rename: skipped (temporary error)."

    ext_up = dest.suffix.lower()

    if ext_up == ".zip":
        import shutil as _shutil

        from axi_archive_extract import extract_zip_members_flat

        staged, err = extract_zip_members_flat(dest)
        try:
            dest.unlink(missing_ok=True)
        except OSError:
            pass
        if err:
            if lang == "ru":
                await update.message.reply_text(f"ZIP: {err}")
            else:
                await update.message.reply_text(f"ZIP: {err}")
            log.info("upload zip extract failed: %s", err)
            return
        parent = staged[0].parent
        try:
            for p in staged:
                tgt = BATCH_INBOX / p.name
                if tgt.exists():
                    tgt = BATCH_INBOX / f"{p.stem}_{int(datetime.now(timezone.utc).timestamp())}{p.suffix}"
                _shutil.move(str(p), str(tgt))
        finally:
            _shutil.rmtree(parent, ignore_errors=True)
        _DAILY_UPLOADS[chat_id] = (today, _count + 1)
        if lang == "ru":
            await update.message.reply_text(
                f"ZIP распакован: **{len(staged)}** файл(ов) в batch_inbox → очередь batch-pipeline.\n"
                f"Загрузок сегодня: {_count + 1}.",
                parse_mode="Markdown",
            )
        else:
            await update.message.reply_text(
                f"ZIP unpacked: **{len(staged)}** file(s) into batch_inbox → batch pipeline queue.\n"
                f"Uploads today: {_count + 1}.",
                parse_mode="Markdown",
            )
        log.info("Path2 zip unpacked %d files (chat=%s)", len(staged), chat_id)
        return

    if dest.suffix.lower() in _BATCH_UNSUPPORTED_EXTENSIONS:
        try:
            dest.unlink(missing_ok=True)
        except Exception:
            pass
        if lang == "ru":
            await update.message.reply_text(
                f"Архив `{dest.name}` **не обрабатывается** batch-pipeline (в очереди только PDF/Office/изображения). "
                "Файл удалён из batch_inbox. Распакуйте и загрузите PDF или документ.\n\n"
                "_(Дашборд Ollama/VRAM к этому не относится — OCR batch не грузит большую модель.)_",
                parse_mode="Markdown",
            )
        else:
            await update.message.reply_text(
                f"Archive `{dest.name}` is **not** processed by batch-pipeline (only PDF/office/images). "
                "Removed from batch_inbox. Unpack and upload PDF or a document.\n\n"
                "_(Ollama VRAM dashboard is unrelated — batch OCR does not load the big LLM.)_",
                parse_mode="Markdown",
            )
        log.info("upload rejected (archive): %s", dest.name)
        return

    # Dedup check by filename: if already in registry, skip pipeline
    if _name_already_in_aims(dest.name):
        try:
            dest.unlink(missing_ok=True)
        except Exception:
            pass
        msg = (
            f"⚠️ Файл `{dest.name}` уже есть в реестре — регистрация пропущена."
            if lang == "ru" else
            f"⚠️ File `{dest.name}` is already in the registry — skipped."
        )
        if send_upload_notice:
            await update.message.reply_text(msg, parse_mode="Markdown")
        log.info("upload skipped (name dup): %s (chat=%s)", dest.name, chat_id)
        return

    # Update daily counter (tracking only, no limit)
    _DAILY_UPLOADS[chat_id] = (today, _count + 1)
    # Persist uploaded file as chat entity memory so the next text message
    # can resolve "attached/that file" even before OCR/registration is done.
    try:
        storage.set_chat_context(
            chat_id,
            last_file_name=dest.name,
            last_file_path=str(dest),
        )
    except Exception:
        pass
    _omi_dialog_append(
        ctx,
        chat_id,
        "user",
        f"[uploaded file queued for OCR: {dest.name}]",
    )

    if analyze_prompt:
        st = _PENDING_ANALYZE.get(chat_id, {})
        st["status"] = "active"
        st["count"] = int(st.get("count", 0)) + 1
        st["updated_at"] = time.time()
        _PENDING_ANALYZE[chat_id] = st
        if not st.get("notified"):
            st["notified"] = True
            _PENDING_ANALYZE[chat_id] = st
            await update.message.reply_text(
                "Режим /analyze: принял пакет файлов, дождусь завершения загрузки и обработаю их вместе."
                if lang == "ru" else
                "/analyze mode: file batch accepted, I will wait for uploads to finish and process them together."
            )
        # Не шлём подробный ответ на каждый файл в analyze-режиме.
        return

    if send_upload_notice:
        if lang == "ru":
            await update.message.reply_text(
                f"Файл принят: `{dest.name}`\n"
                f"Поставлен в очередь batch-pipeline (OCR → anonymize → register).\n"
                f"Загрузок сегодня: {_count + 1}."
                f"{rename_note}",
                parse_mode="Markdown",
            )
        else:
            await update.message.reply_text(
                f"File received: `{dest.name}`\n"
                f"Queued for batch pipeline (OCR → anonymize → register).\n"
                f"Uploads today: {_count + 1}."
                f"{rename_note}",
                parse_mode="Markdown",
            )
    log.info("Path2 upload queued: %s (chat=%s)", dest.name, chat_id)

    # Task Registry: зарегистрировать загрузку файла как завершённую задачу
    _tr_fu_id = _tr_register(
        f"upload:{dest.name}",
        chat_id=str(chat_id),
        source="omi",
    )
    _tr_start(_tr_fu_id, assigned_to="omi")
    _tr_done(_tr_fu_id, summary=f"queued batch-pipeline: {dest.name}")


# ══════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════

async def _job_omi_deliver_cross_handoffs(context: ContextTypes.DEFAULT_TYPE) -> None:
    """Доставка очереди Axi→Omi: одно сообщение «принял в работу» от Omi."""
    try:
        from cross_bot_handoff import (
            claim_pending_for_target,
            handoff_delivery_enabled,
            infer_lang_ru,
            mark_delivered,
            mark_failed,
        )
    except ImportError:
        return
    if not handoff_delivery_enabled():
        return
    omi_name = os.environ.get("OMI_NAME", "Omi")
    rows = claim_pending_for_target("omi", limit=6)
    bot = context.application.bot
    for r in rows:
        hid = int(r["id"])
        payload = (r.get("payload") or "").strip()
        reason = (r.get("reason") or "").strip()
        chat_id = int(r["chat_id"])
        try:
            ru = infer_lang_ru(payload)
            if ru:
                body = f"{omi_name}: Принял в работу"
                if reason:
                    body += f" ({reason})"
                body += ":\n" + payload
            else:
                body = f"{omi_name}: On it"
                if reason:
                    body += f" ({reason})"
                body += ":\n" + payload
            if len(body) > 4096:
                body = body[:4090] + "…"
            await bot.send_message(chat_id=chat_id, text=body)
            mark_delivered(hid)
        except Exception as e:
            log.warning("cross handoff deliver omi id=%s: %s", hid, e)
            mark_failed(hid, repr(e))


# ── Axi-task helpers (internal DB strategy generation) ────────────────────────

def _omi_extract_keywords(text: str) -> list[str]:
    import re as _re
    # Very tight stop list — keep ALL domain terms (preservation, equipment, procedure, etc.)
    stop = {
        "the", "a", "an", "and", "or", "for", "of", "in", "on", "at", "to", "is",
        "based", "from", "database", "related", "this", "that", "are", "with",
        "по", "из", "на", "для", "и", "в", "с", "от", "как",
        "axi", "omi", "aims", "separate", "each", "their", "also", "have",
        "develop", "compare", "best", "practice", "using", "internal",
    }
    words = _re.findall(r"[a-zA-Zа-яёА-ЯЁ]{3,}", text)
    seen: dict[str, int] = {}
    for w in words:
        lw = w.lower()
        if lw not in stop:
            seen[lw] = seen.get(lw, 0) + 1
    # Sort by frequency, then alphabetically for stability; return up to 20
    return [k for k, _ in sorted(seen.items(), key=lambda x: -x[1])][:20]


def _omi_resolve_doc_path(raw_path: str | None) -> Path | None:
    if not raw_path:
        return None
    p = Path(raw_path)
    parts = p.parts
    try:
        idx = next(i for i, part in enumerate(parts) if part == "aims_workspace")
        candidate = WORKSPACE / Path(*parts[idx + 1:])
        if candidate.exists():
            return candidate
    except StopIteration:
        pass
    if p.exists():
        return p
    candidate2 = WORKSPACE / p.name
    if candidate2.exists():
        return candidate2
    return None


def _omi_fetch_db_docs(keywords: list[str], max_docs: int = 20) -> list[dict]:
    """
    Fetch relevant documents from AIMS DB using hybrid search (FTS5 + Qdrant via StorageManager).
    Falls back to raw FTS5 if StorageManager unavailable.
    Reads text content from .md/.txt/.docx OCR outputs, falls back to summary.
    """
    if not DB_PATH.exists():
        log.warning("_omi_fetch_db_docs: DB not found: %s", DB_PATH)
        return []
    results: list[dict] = []

    # Build query string from keywords
    query = " ".join(keywords) if keywords else ""

    # Try hybrid search via StorageManager (FTS5 + Qdrant RRF)
    try:
        rows_hybrid = storage.hybrid_search_documents(query, limit=max_docs) if query else []
        if rows_hybrid:
            seen_ids: set[int] = set()
            for row in rows_hybrid:
                did = row.get("id")
                if did in seen_ids:
                    continue
                seen_ids.add(did)
                content = _omi_read_doc_content(row)
                results.append({
                    "id": did,
                    "title": row.get("title") or row.get("file_name") or f"doc#{did}",
                    "process": row.get("aims_process") or "—",
                    "content": content,
                    "keywords": row.get("keywords") or "",
                })
            log.info("_omi_fetch_db_docs: hybrid search returned %d docs", len(results))
            return results
    except Exception as e:
        log.warning("_omi_fetch_db_docs hybrid error: %s — falling back to FTS5", e)

    # Fallback: raw FTS5 SQL
    try:
        conn = sqlite3.connect(str(DB_PATH))
        conn.row_factory = sqlite3.Row
        if keywords:
            fts_query = " OR ".join(keywords)
            rows = conn.execute(
                "SELECT d.id, d.title, d.aims_process, d.stored_path, "
                "d.anonymized_result_path, d.summary, d.keywords "
                "FROM documents d "
                "JOIN documents_fts fts ON fts.rowid = d.id "
                "WHERE documents_fts MATCH ? "
                "AND d.file_name NOT LIKE '%I_m_ready%' "
                "AND d.file_name NOT LIKE '%skill%' "
                "ORDER BY d.id DESC LIMIT ?",
                (fts_query, max_docs),
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT id, title, aims_process, stored_path, anonymized_result_path, "
                "summary, keywords FROM documents ORDER BY id DESC LIMIT ?",
                (max_docs,),
            ).fetchall()
        conn.close()
        for row in rows:
            content = _omi_read_doc_content(dict(row))
            if not content and row["summary"]:
                content = row["summary"]
            results.append({
                "id": row["id"],
                "title": row["title"] or row["stored_path"] or f"doc#{row['id']}",
                "process": row["aims_process"] or "—",
                "content": content,
                "keywords": row["keywords"] or "",
            })
    except Exception as e:
        log.warning("_omi_fetch_db_docs FTS5 error: %s", e)
    return results


def _omi_read_doc_content(row: dict) -> str:
    """
    Read document text content from stored files.
    Priority: anonymized_result_path (.txt/.md) → stored_path (.txt/.md/.docx)
    → summary field.
    Returns up to 20000 chars.
    """
    for path_key in ("anonymized_result_path", "stored_path", "file_path"):
        raw_path = row.get(path_key)
        if not raw_path:
            continue
        doc_path = _omi_resolve_doc_path(str(raw_path))
        if not doc_path:
            continue
        try:
            if doc_path.suffix.lower() in (".md", ".txt"):
                return doc_path.read_text(errors="replace")[:20000]
            if doc_path.suffix.lower() == ".docx":
                from docx import Document as _DocxDoc
                d = _DocxDoc(str(doc_path))
                text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
                return text[:20000]
        except Exception:
            pass
    # Last resort: summary stored in DB
    return (row.get("summary") or "")[:2000]


def _nim_url() -> str:
    """Get OmniRoute URL from environment."""
    return os.environ.get("OMNIROUTE_BASE_URL", "http://127.0.0.1:20129/v1").rstrip("/")


def _nim_key() -> str:
    """Get OmniRoute API key from environment."""
    return os.environ.get("OMNI_TOKEN", "").strip()


async def _omi_local_draft(prompt: str) -> str:
    """
    Step 1 — generate document draft using local Ollama synthesis model.
    Raises RuntimeError if Ollama unavailable or returns empty.
    """
    import httpx

    synthesis_model = os.environ.get("OMI_SYNTHESIS_MODEL") or (_resolve_slot("32") if _resolve_slot else "")
    synthesis_timeout = int(os.environ.get("OMI_SYNTHESIS_TIMEOUT", "300"))
    num_ctx = int(os.environ.get("OMI_SYNTHESIS_NUM_CTX", "16000"))
    num_predict = int(os.environ.get("OMI_SYNTHESIS_NUM_PREDICT", "6000"))

    try:
        from ollama_resolve import effective_ollama_base_url
        ollama_url = effective_ollama_base_url()
    except Exception:
        ollama_url = os.environ.get("DGX_OLLAMA_URL") or os.environ.get(
            "OLLAMA_BASE_URL", "http://host.docker.internal:11434"
        )

    log.info("_omi_local_draft: model=%s url=%s", synthesis_model, ollama_url)
    async with httpx.AsyncClient(timeout=synthesis_timeout) as client:
        resp = await client.post(
            f"{ollama_url}/api/generate",
            json={
                "model": synthesis_model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.3,
                    "num_ctx": num_ctx,
                    "num_predict": num_predict,
                },
            },
        )
        resp.raise_for_status()
        text = resp.json().get("response", "").strip()
        if not text:
            raise RuntimeError(f"Ollama returned empty response (model={synthesis_model})")
        log.info("_omi_local_draft: OK chars=%d", len(text))
        return text


async def _axi_review_and_correct(
    draft: str,
    original_task: str,
    source_titles: list[str],
) -> str:
    """
    Step 2 — Axi reviews the local draft and returns the final corrected document.

    Backend order:
      1. OmniRoute: gemini-free-fallback (primary)
      2. Anthropic: Claude (fallback via ANTHROPIC_API_KEY)

    Raises RuntimeError if every backend fails.
    """
    import httpx

    sources_note = (
        "\n".join(f"  - {t}" for t in source_titles[:20])
        if source_titles else "  (no source documents)"
    )
    review_prompt = (
        "You are Axi, a senior technical writer and HSE document specialist.\n\n"
        "A local AI model has generated a draft document from an internal database. "
        "Review this draft and produce the FINAL corrected version.\n\n"
        f"ORIGINAL TASK:\n{original_task}\n\n"
        f"SOURCE DOCUMENTS USED:\n{sources_note}\n\n"
        f"LOCAL MODEL DRAFT:\n{draft}\n\n"
        "REVIEW INSTRUCTIONS:\n"
        "1. Verify the document fully satisfies the original task (all required sections present).\n"
        "2. Correct factual errors, missing steps, or incomplete tables.\n"
        "3. Ensure Markdown tables cover all required equipment types / plant units / timeframes.\n"
        "4. Improve technical precision and formal language where needed.\n"
        "5. Do NOT add irrelevant content or change structure unless required by the task.\n\n"
        "OUTPUT: Final corrected document only. No preamble, no review commentary."
    )

    nim_url = _nim_url()
    nim_key = _nim_key()
    last_err: Exception | None = None

    # ── 1. OmniRoute — primary backend ───────────────────────────────────────
    if nim_url:
        try:
            log.info("_axi_review: using OmniRoute")
            if not nim_key:
                raise RuntimeError("OMNI_TOKEN missing for OmniRoute project1.")
            omniroute_hdrs: dict = {"Content-Type": "application/json", "Authorization": f"Bearer {nim_key}"}
            async with httpx.AsyncClient(timeout=300) as client:
                resp = await client.post(
                    f"{nim_url}/chat/completions",
                    json={
                        "model": os.environ.get("OMNIROUTE_MODEL", os.environ.get("OMNIROUTE_TEACHER_MODEL", "project1")),
                        "messages": [
                            {"role": "system", "content": "You are Axi, a senior technical writer."},
                            {"role": "user", "content": review_prompt},
                        ],
                        "temperature": 0.3,
                        "max_tokens": 8192,
                    },
                    headers=omniroute_hdrs,
                )
                if resp.status_code == 200:
                    choices = resp.json().get("choices") or []
                    if choices:
                        msg = choices[0].get("message") or {}
                        text = msg.get("content", "").strip()
                        if text:
                            log.info("_axi_review: OmniRoute OK chars=%d", len(text))
                            return text
                else:
                    log.warning("_axi_review: OmniRoute status=%d", resp.status_code)
                    last_err = RuntimeError(f"OmniRoute HTTP {resp.status_code}")
        except (httpx.TimeoutException, httpx.NetworkError) as e:
            log.warning("_axi_review: OmniRoute network error: %s", e)
            last_err = e
        except Exception as e:
            log.warning("_axi_review: OmniRoute error: %s", e)
            last_err = e

    # ── 2. Anthropic — fallback ──────────────────────────────────────────────
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if anthropic_key:
        try:
            log.info("_axi_review: falling back to Anthropic model=%s", ANTHROPIC_MODEL)
            async with httpx.AsyncClient(timeout=300) as client:
                resp = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    json={
                        "model": ANTHROPIC_MODEL,
                        "max_tokens": 8192,
                        "messages": [{"role": "user", "content": review_prompt}],
                    },
                    headers={
                        "x-api-key": anthropic_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                )
                resp.raise_for_status()
                text = resp.json()["content"][0]["text"].strip()
                if text:
                    log.info("_axi_review: Anthropic OK chars=%d", len(text))
                    return text
        except Exception as e:
            log.error("_axi_review: Anthropic also failed: %s", e)
            last_err = e

    raise RuntimeError(
        f"All backends failed (OmniRoute + Anthropic). "
        f"Last: {last_err}"
    )


async def _omi_llm_generate(prompt: str) -> str:
    """Legacy single-call generate for non-pipeline LLM calls (summaries etc.).
    Order: OmniRoute → Anthropic last.
    """
    import httpx

    nim_url = _nim_url()
    nim_key = _nim_key()
    last_err: Exception | None = None

    # ── 1. OmniRoute — primary backend ───────────────────────────────────────
    if nim_url:
        try:
            if not nim_key:
                raise RuntimeError("OMNI_TOKEN missing for OmniRoute project1.")
            gen_hdrs: dict = {"Content-Type": "application/json", "Authorization": f"Bearer {nim_key}"}
            async with httpx.AsyncClient(timeout=120) as client:
                resp = await client.post(
                    f"{nim_url}/chat/completions",
                    json={
                        "model": os.environ.get("OMNIROUTE_MODEL", os.environ.get("OMNIROUTE_TEACHER_MODEL", "project1")),
                        "messages": [{"role": "user", "content": prompt}],
                        "temperature": 0.3,
                        "max_tokens": 8192,
                    },
                    headers=gen_hdrs,
                )
                if resp.status_code == 200:
                    choices = resp.json().get("choices") or []
                    if choices:
                        msg = choices[0].get("message") or {}
                        text = msg.get("content", "").strip()
                        if text:
                            return text
                else:
                    last_err = RuntimeError(f"OmniRoute HTTP {resp.status_code}")
        except Exception as e:
            last_err = e

    # ── 2. Anthropic — fallback ──────────────────────────────────────────────
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if anthropic_key:
        try:
            async with httpx.AsyncClient(timeout=120) as client:
                resp = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    json={
                        "model": ANTHROPIC_MODEL,
                        "max_tokens": 4096,
                        "messages": [{"role": "user", "content": prompt}],
                    },
                    headers={
                        "x-api-key": anthropic_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                )
                resp.raise_for_status()
                return resp.json()["content"][0]["text"]
        except Exception as e:
            last_err = e

    raise RuntimeError(f"All LLM backends failed (NIM + Anthropic): {last_err}")


def _omi_build_docx(content: str, filename_stem: str, title: str | None = None) -> Path:
    import re as _re
    from docx import Document as _Document
    OMI_RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe = _re.sub(r"[^\w\-]", "_", filename_stem)[:60]
    path = OMI_RESULTS_DIR / f"{safe}_{stamp}.docx"
    doc = _Document()
    if title:
        doc.add_heading(title, level=0)
    bold_re = _re.compile(r"\*\*(.+?)\*\*")

    def _add_para(paragraph, line: str) -> None:
        for i, part in enumerate(bold_re.split(line)):
            run = paragraph.add_run(part)
            if i % 2 == 1:
                run.bold = True

    for line in content.splitlines():
        s = line.rstrip()
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith(("- ", "* ")):
            _add_para(doc.add_paragraph(style="List Bullet"), s[2:])
        elif s == "":
            doc.add_paragraph("")
        else:
            _add_para(doc.add_paragraph(), s)
    doc.save(str(path))
    return path


async def _omi_execute_axi_task(task_id: str, description: str, chat_id: int, bot) -> None:
    """
    Full internal-resource workflow delegated by Axi.

    Pipeline:
      1. Search AIMS DB — FTS5 + Qdrant hybrid, up to 20 docs
      2. Local Ollama draft  — internal model generates raw document
      3. Axi Anthropic review — corrects, completes, polishes final document
      4. Send .docx to chat
    """
    _tr_start(task_id, assigned_to="omi")
    try:
        await bot.send_message(chat_id=chat_id, text="Omi: 🔍 Ищу документы в базе AIMS…")

        keywords = _omi_extract_keywords(description)
        loop = asyncio.get_event_loop()
        docs = await loop.run_in_executor(None, _omi_fetch_db_docs, keywords)

        if not docs:
            await bot.send_message(
                chat_id=chat_id,
                text="Omi: ⚠️ Документы по теме не найдены в базе AIMS. "
                     "Уточните запрос или сначала зарегистрируйте документы.",
            )
            _tr_stuck(task_id, "no db docs found")
            return

        docs_with_content = [d for d in docs if d["content"]]
        found_list = "\n".join(f"• [{d['process']}] {d['title']}" for d in docs)
        await bot.send_message(
            chat_id=chat_id,
            text=(
                f"Omi: 📂 Найдено {len(docs)} документов "
                f"({len(docs_with_content)} с содержимым):\n{found_list}\n\n"
                "⚙️ Шаг 1/2 — генерирую черновик (локальная модель)…"
            ),
        )

        # Build document context
        parts = []
        for d in docs:
            header = f"### [{d['process']}] {d['title']}"
            body = d["content"][:12000] if d["content"] else "[file pending OCR — use title as reference]"
            if d.get("keywords"):
                body = f"Keywords: {d['keywords']}\n{body}"
            parts.append(f"{header}\n{body}")
        doc_context = "\n\n---\n\n".join(parts)

        content_note = (
            "Use the document content above as primary source. "
            "Reproduce specific requirements, tables, and checklists faithfully. "
            "Fill gaps with applicable industry best practice."
            if docs_with_content else
            "Document text is pending OCR — use titles and metadata as structural anchors. "
            "Generate based on oil & gas / heavy industrial best practice."
        )

        draft_prompt = (
            f"{description}\n\n"
            f"[AIMS DATABASE — {len(docs)} DOCUMENTS, {len(docs_with_content)} WITH FULL TEXT]:\n\n"
            f"{doc_context}\n\n"
            f"[INSTRUCTION]: {content_note}\n\n"
            "RULES:\n"
            "- Formal technical English. Document content only — no preamble.\n"
            "- Use # Title, ## Section, ### Subsection headers.\n"
            "- For procedures: scope, applicability, responsibilities, step-by-step instructions, "
            "acceptance criteria, references.\n"
            "- For tables (by equipment type, by plant, by term): use Markdown tables.\n"
            "- Cite source documents in [square brackets] where data is drawn from them.\n"
            "- Reference standards from source documents; supplement with API 686, ISO 55001/55002, "
            "NFPA, ASME, EN 13460 as applicable.\n"
        )

        # ── Step 1: Local Ollama draft ──────────────────────────────────────────
        try:
            draft = await _omi_local_draft(draft_prompt)
        except Exception as draft_err:
            log.warning("_omi_execute_axi_task: local draft failed: %s — skipping to review", draft_err)
            draft = (
                f"[LOCAL MODEL UNAVAILABLE — skipping draft step]\n\n"
                f"Task: {description}\n\n"
                f"Sources used:\n" + "\n".join(f"- {d['title']}" for d in docs)
            )

        # ── Step 2: Axi Anthropic review + correction ──────────────────────────
        await bot.send_message(
            chat_id=chat_id,
            text="Omi: ✏️ Шаг 2/2 — Axi проверяет и корректирует черновик…",
        )

        source_titles = [d["title"] for d in docs]
        try:
            final = await _axi_review_and_correct(draft, description, source_titles)
        except RuntimeError as review_err:
            # Anthropic unavailable — use local draft as final
            log.warning("_omi_execute_axi_task: Axi review failed: %s — using local draft", review_err)
            await bot.send_message(
                chat_id=chat_id,
                text=f"Omi: ⚠️ Axi недоступна ({review_err}). "
                     "Отправляю черновик локальной модели.",
            )
            final = draft

        first_line = next(
            (ln.strip().lstrip("#").strip() for ln in final.splitlines() if ln.strip()),
            "Document",
        )
        stem = "aims_doc"

        docx_path = await loop.run_in_executor(
            None, _omi_build_docx, final, stem, first_line
        )

        with docx_path.open("rb") as fh:
            await bot.send_document(
                chat_id=chat_id,
                document=fh,
                filename=docx_path.name,
                caption=(
                    f"📄 {first_line}\n"
                    f"Источники: {len(docs_with_content)} из {len(docs)} документов AIMS\n"
                    f"Модели: local draft → Axi review"
                ),
            )

        _tr_done(task_id, summary=f"omi_docgen:{docx_path.name}:{len(docs)}docs")
        log.info("_omi_execute_axi_task done: %s → %s", task_id, docx_path.name)

    except Exception as e:
        log.error("_omi_execute_axi_task %s: %s", task_id, e, exc_info=True)
        _tr_stuck(task_id, error=str(e)[:200])
        try:
            await bot.send_message(
                chat_id=chat_id,
                text=f"Omi: ⚠️ Ошибка при генерации документа: {type(e).__name__}",
            )
        except Exception:
            pass


async def _job_poll_task_registry(context: ContextTypes.DEFAULT_TYPE) -> None:
    """
    Poll Task Registry every 15 s for pending tasks from Axi (source='axi').
    Spawns _omi_execute_axi_task as a background asyncio task per new task.
    """
    if _tr_client is None:
        return
    key = "_omi_processed_task_ids"
    seen: set = context.application.bot_data.setdefault(key, set())
    try:
        tasks = _tr_client.recent(status="pending", limit=30)
    except Exception as e:
        log.debug("poll_task_registry fetch error: %s", e)
        return

    for t in tasks:
        task_id = t.get("task_id", "")
        source = (t.get("source") or "").lower()
        description = (t.get("description") or "").strip()
        chat_id_raw = t.get("chat_id", "")

        # Only tasks delegated by Axi
        if not task_id or task_id in seen:
            continue
        if source != "axi":
            continue
        if not description or not chat_id_raw:
            continue

        try:
            chat_id_int = int(str(chat_id_raw).strip())
        except (ValueError, TypeError):
            log.warning("poll_task_registry: bad chat_id %r for %s", chat_id_raw, task_id)
            continue

        seen.add(task_id)
        log.info("poll_task_registry: new Axi task %s chat=%s: %s", task_id, chat_id_int, description[:80])

        # Spawn execution as background task — don't block the scheduler
        asyncio.create_task(
            _omi_execute_axi_task(task_id, description, chat_id_int, context.bot)
        )


async def _post_init_omi_commands(app: Application) -> None:
    """Команды в меню Telegram (кнопка / в поле ввода) — личка и группы."""
    cmds = [
        BotCommand("start", "Главное меню и подсказки"),
        BotCommand("help", "Полный список команд"),
        BotCommand("analyze", "Ожидать загрузку файлов для анализа"),
        BotCommand("analyze_done", "Завершить пакет /analyze сейчас"),
        BotCommand("analyze_end", "Алиас завершения загрузки"),
        BotCommand("intake", "Массовая загрузка документов из папки"),
        BotCommand("intake_preflight", "Проверка готовности intake pipeline"),
        BotCommand("menu", "Кнопочное меню"),
        BotCommand("test_mode", "Следующий запрос: отчёт intent+время (skill 21)"),
        BotCommand("tect_mode", "То же что /test_mode (опечатка)"),
        BotCommand("status", "Состояние БД и разделов"),
        BotCommand("search", "Поиск в реестре"),
        BotCommand("omi_kb_status", "Статус Omi knowledge base"),
        BotCommand("omi_kb_audit", "Полный audit knowledge base"),
        BotCommand("omi_kb_search", "Поиск по knowledge base"),
        BotCommand("omi_kb_duplicates", "Дубликаты knowledge base"),
        BotCommand("omi_kb_reindex", "Безопасный reindex knowledge base"),
        BotCommand("omi_kb_exemplar_check", "Проверка эталона по типу документа"),
        BotCommand("tasks", "Очередь задач структуры БД"),
        BotCommand("close_task", "Закрыть зависшую задачу"),
        BotCommand("skills", "Skills Omi-LLM"),
        BotCommand("docgen", "Сборка .docx из реестра"),
        BotCommand("docgen_upgrade", "Запуск DOCGEN upgrade batch (private only)"),
        BotCommand("docsreg", "Запуск DOCSREG по файлу на DGX"),
        BotCommand("registry_audit", "Read-only registry consistency audit"),
        BotCommand("docsreg_start_media", "DOCSREG batch launch from folder on DGX"),
        BotCommand("gocsreg_start_media", "Alias: DOCSREG batch launch from folder on DGX"),
        BotCommand("rename_by_context", "Переименовать файлы по контексту"),
        BotCommand("rename_file", "Переименование файла по запросу"),
        BotCommand("fixnoisynames", "Авто-очистка числовых кодов в именах файлов"),
        BotCommand("fixnoisyname", "Переименовать файл по ID реестра"),
        BotCommand("registry_sync_status", "Статус sync OCR→AIMS"),
    ]
    try:
        await app.bot.set_my_commands(cmds)
    except Exception as e:
        log.warning("set_my_commands failed: %s", e)
    try:
        from ollama_resolve import ollama_schedule_telegram_stack_warm
        ollama_schedule_telegram_stack_warm()
        log.info("omi post_init: model warm scheduled (heavy=SLOT32 + small=SLOT14)")
    except Exception as e:
        log.warning("omi post_init: model warm failed: %s", e)
    if app.job_queue is not None:
        try:
            from cross_bot_handoff import handoff_delivery_enabled, handoff_poll_interval_sec

            if handoff_delivery_enabled():
                app.job_queue.run_repeating(
                    _job_omi_deliver_cross_handoffs,
                    interval=handoff_poll_interval_sec(),
                    first=3,
                    name="omi_cross_handoff",
                )
        except Exception as e:
            log.warning("omi cross-handoff job not started: %s", e)


def main():
    app = Application.builder().token(TOKEN).post_init(_post_init_omi_commands).build()

    async def _error_handler(update: object, context) -> None:
        err = context.error
        if isinstance(err, TelegramConflict):
            log.debug("telegram conflict (duplicate polling session) — resolves automatically")
            return
        if isinstance(err, TelegramNetworkError):
            log.warning("telegram network error: %s", err)
            return
        log.exception("unhandled telegram error: %s", err, exc_info=err)

    app.add_error_handler(_error_handler)

    # Команды
    app.add_handler(CommandHandler("start",      cmd_start))
    app.add_handler(CommandHandler("help",       cmd_help))
    app.add_handler(CommandHandler("analyze",    cmd_analyze))
    app.add_handler(CommandHandler("analyze_done", cmd_analyze_done))
    app.add_handler(CommandHandler("analyze_end", cmd_analyze_end))
    app.add_handler(CommandHandler("test_mode",  cmd_test_mode))
    app.add_handler(CommandHandler("tect_mode",  cmd_test_mode))
    app.add_handler(CommandHandler("status",     cmd_status))
    app.add_handler(CommandHandler("move",       cmd_move))
    app.add_handler(CommandHandler("archive",    cmd_archive))
    app.add_handler(CommandHandler("newprocess", cmd_newprocess))
    app.add_handler(CommandHandler("migrate",    cmd_migrate))
    app.add_handler(CommandHandler("search",     cmd_search))
    app.add_handler(CommandHandler("omi_kb_status", cmd_omi_kb_status))
    app.add_handler(CommandHandler("omi_kb_audit", cmd_omi_kb_audit))
    app.add_handler(CommandHandler("omi_kb_search", cmd_omi_kb_search))
    app.add_handler(CommandHandler("omi_kb_duplicates", cmd_omi_kb_duplicates))
    app.add_handler(CommandHandler("omi_kb_reindex", cmd_omi_kb_reindex))
    app.add_handler(CommandHandler("omi_kb_exemplar_check", cmd_omi_kb_exemplar_check))
    app.add_handler(CommandHandler("docs_today", cmd_docs_today))
    app.add_handler(CommandHandler("registry_sync_status", cmd_registry_sync_status))
    app.add_handler(CommandHandler("close_task", cmd_close_task))
    app.add_handler(CommandHandler("rename_by_context", cmd_rename_by_context))
    app.add_handler(CommandHandler("rename_file",       cmd_rename_file))
    app.add_handler(CommandHandler("fixnoisynames",     cmd_fixnoisynames))
    app.add_handler(CommandHandler("fixnoisyname",      cmd_fixnoisyname))
    app.add_handler(CommandHandler("menu",       cmd_menu))
    app.add_handler(CommandHandler("tasks",      cmd_tasks))
    app.add_handler(CommandHandler("skills",     cmd_skills))
    app.add_handler(CommandHandler("docgen",     cmd_docgen))
    app.add_handler(CommandHandler("docgen_upgrade", cmd_docgen_upgrade))
    app.add_handler(CommandHandler("docsreg",    cmd_docsreg))
    app.add_handler(CommandHandler("docsreg_start_media", cmd_docsreg))
    app.add_handler(CommandHandler("gocsreg_start_media", cmd_docsreg))
    if _registry_audit_cmd_handler is not None:
        app.add_handler(CommandHandler("registry_audit", _registry_audit_cmd_handler))
    app.add_handler(CommandHandler("nightplan",  cmd_nightplan))
    app.add_handler(CommandHandler("backup_now", cmd_backup_now))
    app.add_handler(CommandHandler("backup_list", cmd_backup_list))
    app.add_handler(CommandHandler("backup_restore", cmd_backup_restore))
    app.add_handler(CommandHandler("backup_schedule", cmd_backup_schedule))
    app.add_handler(CommandHandler("backup_schedule_set", cmd_backup_schedule_set))
    app.add_handler(CommandHandler("selftest",     cmd_selftest))
    app.add_handler(CommandHandler("task_quality", cmd_task_quality))

    # Intake pipeline commands
    try:
        from omi_intake_handler import cmd_intake, cmd_intake_preflight
        app.add_handler(CommandHandler("intake", cmd_intake))
        app.add_handler(CommandHandler("intake_preflight", cmd_intake_preflight))
    except Exception as e:
        log.warning("intake commands not loaded: %s", e)
    app.add_handler(CommandHandler("save_template",   cmd_save_template))
    app.add_handler(CommandHandler("templates",       cmd_templates))
    app.add_handler(CommandHandler("use_template",    cmd_use_template))
    app.add_handler(CommandHandler("delete_template", cmd_delete_template))

    # Кнопки
    app.add_handler(CallbackQueryHandler(handle_callback))

    # Path 2: Telegram file uploads (max OMI_UPLOAD_DAILY_MAX/day per chat)
    app.add_handler(MessageHandler(filters.Document.ALL | filters.PHOTO, handle_file_upload))
    app.add_handler(MessageHandler(filters.Regex(r"^/rename_file_.+"), cmd_rename_file_dynamic))
    app.add_handler(MessageHandler(filters.Regex(r"^/DOCGEN_UPGRADE(?:\s|$)"), cmd_docgen_upgrade))

    # Живой чат
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_chat))
    # run_repeating для cross-handoff подключён в post_init (_post_init_omi_commands)
    if OMI_NOTIFY_REGISTRATIONS_FROM_DB and app.job_queue is not None:
        app.job_queue.run_repeating(_notify_new_registrations_job, interval=30, first=10)

    # Task Registry polling — react to tasks from Axi
    if app.job_queue is not None and _tr_client is not None:
        app.job_queue.run_repeating(_job_poll_task_registry, interval=15, first=5)
        log.info("Task Registry polling job started (interval=15s)")
    else:
        log.warning("Task Registry polling job NOT started (job_queue=%s, _tr_client=%s)",
                    app.job_queue, _tr_client)

    # Periodic heartbeat — refresh /tmp/bot_alive so Docker healthcheck
    # detects a truly stuck bot (not just that the process started once).
    async def _heartbeat_refresh(_ctx):
        try:
            open("/tmp/bot_alive", "w").close()
        except OSError:
            pass

    if app.job_queue is not None:
        app.job_queue.run_repeating(_heartbeat_refresh, interval=20, first=0)

    log.info("Omi bot starting...")
    open("/tmp/bot_alive", "w").close()
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
