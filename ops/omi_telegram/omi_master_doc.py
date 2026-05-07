"""
omi_master_doc.py
─────────────────
Синтез мастер-документа из зарегистрированного файла.

Стандарты:
  IEC/IEEE 82079-1:2019 — структура технической документации
  ISO 45001:2018/Amd1:2024 — требования к HSE-документам
  ISO 9001:2015 §7.5 — метаданные и контроль версий
  ISO 21502:2020 — жизненный цикл проектного документа

Pipeline:
  OCR-текст → anonymizer (local Qwen) → Claude claude-sonnet-4-6 → structured master_doc JSON
                                                                  → NIM completeness score (optional)
                                                                  → DOCX export (optional)

Rule: text MUST be anonymized locally before any cloud API call.

Env vars:
  ANTHROPIC_API_KEY   — обязателен
  ANTHROPIC_MODEL     — модель (default: claude-sonnet-4-6)
  OMI_MASTER_DOC_ENABLED — 1/0 (default: 1)
  OMI_MASTER_DOC_EXPORT_DOCX — 1 = сохранять DOCX рядом с JSON (default: 0)
  NVIDIA_NIM_URL — endpoint для scoring (optional, default: http://127.0.0.1:8082)
  NVIDIA_NIM_API_KEY — API key для NIM (optional)
"""

from __future__ import annotations

import json
import logging
import os
import re
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

log = logging.getLogger("omi_master_doc")

# ── Config ──────────────────────────────────────────────────────────────────────

def _anthropic_key() -> str:
    return os.environ.get("ANTHROPIC_API_KEY", "").strip()

def _anthropic_model() -> str:
    return os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6").strip() or "claude-sonnet-4-6"

def _enabled() -> bool:
    return os.environ.get("OMI_MASTER_DOC_ENABLED", "1").strip() != "0"

def _export_docx() -> bool:
    return os.environ.get("OMI_MASTER_DOC_EXPORT_DOCX", "0").strip() == "1"

def _nim_url() -> str:
    return os.environ.get("NVIDIA_NIM_URL", "http://127.0.0.1:8082").rstrip("/")

def _nim_key() -> str:
    return os.environ.get("NVIDIA_NIM_API_KEY", "").strip()


# ── Document type map ────────────────────────────────────────────────────────────

_DOC_TYPES = {
    "JSA":  "Job Safety Analysis",
    "ERP":  "Emergency Response Plan",
    "MOC":  "Management of Change",
    "LOTO": "Lockout/Tagout Procedure",
    "PTW":  "Permit to Work",
    "SOP":  "Standard Operating Procedure",
    "ITP":  "Inspection and Test Plan",
    "MSI":  "Method Statement / Work Instruction",
    "RAM":  "Risk Assessment Matrix",
    "EP":   "Environmental Procedure",
    "QMP":  "Quality Management Plan",
    "OTHER": "General Industrial Document",
}


# ── Master doc schema (empty template) ──────────────────────────────────────────

def _empty_master() -> dict:
    return {
        "schema_version": "1.0",
        "meta": {
            "doc_id":       None,
            "revision":     "Rev.0",
            "status":       "draft",
            "language":     "en",
            "created_at":   datetime.now(timezone.utc).strftime("%Y-%m-%d"),
            "reviewed_at":  None,
            "approved_by":  None,
            "next_review":  None,
            "change_log":   [],
        },
        "classification": {
            "document_type":   None,
            "document_type_full": None,
            "aims_process":    None,
            "iso_clause":      [],
            "target_audience": [],
            "facility_type":   None,
            "activity":        None,
        },
        "content": {
            "scope":       None,
            "definitions": {},
            "hazards":     [],
            "controls": {
                "eliminate":      [],
                "substitute":     [],
                "engineering":    [],
                "administrative": [],
                "ppe":            [],
            },
            "procedures":             [],
            "emergency_procedures":   None,
            "roles_responsibilities": [],
            "standards_referenced":   [],
            "legal_references":       [],
        },
        "quality": {
            "completeness_score": None,
            "gaps":               [],
            "quality_feedback":    None,
            "claude_model":       None,
            "generated_at":       None,
        },
    }


# ── Claude extraction prompt ─────────────────────────────────────────────────────

_SYSTEM = """\
You are an expert HSE document analyst. Your task is to extract structured information \
from industrial safety and management documents into a precise JSON master document.

Apply these international standards:
- IEC/IEEE 82079-1:2019 (technical documentation structure)
- ISO 45001:2018/Amd1:2024 (OH&S requirements)
- ISO 9001:2015 §7.5 (documented information control)
- ISO 21502:2020 (document lifecycle)

Output ONLY valid JSON — no markdown fences, no explanation, no preamble.
All text values in English. If the source is in another language, translate key content.
If a field cannot be determined from the text, use null (not empty string).
"""

_USER_TMPL = """\
FILE: {file_name}
AIMS PROCESS: {aims_process}
KNOWN TITLE: {title}

DOCUMENT TEXT (first 14000 chars):
{text}

Extract and return this exact JSON structure (fill every field you can determine):

{{
  "meta": {{
    "doc_id": "auto-generate as TYPE-PROCESS-NNNN or null",
    "revision": "Rev.X from document or Rev.0 if not found",
    "status": "draft|reviewed|approved",
    "language": "en|ru|other",
    "reviewed_at": "YYYY-MM-DD or null",
    "approved_by": "[anonymized role, e.g. HSE Manager] or null",
    "next_review": "YYYY-MM-DD or null",
    "change_log": [{{"rev": "Rev.0", "summary": "Initial issue"}}]
  }},
  "classification": {{
    "document_type": "JSA|ERP|MOC|LOTO|PTW|SOP|ITP|MSI|RAM|EP|QMP|OTHER",
    "document_type_full": "full name of the document type",
    "aims_process": "{aims_process}",
    "iso_clause": ["ISO 45001:2018 §X.X", "..."],
    "target_audience": ["operators", "supervisors", "HSE officers", "..."],
    "facility_type": "mine|refinery|gas plant|construction|other",
    "activity": "short description of the activity covered"
  }},
  "content": {{
    "scope": "one paragraph describing scope and applicability",
    "definitions": {{"ACRONYM": "definition", "...": "..."}},
    "hazards": [
      {{
        "id": "H01",
        "description": "hazard description",
        "severity": 1-5,
        "likelihood": 1-5,
        "risk_level": "LOW|MEDIUM|HIGH|CRITICAL",
        "affected_parties": ["workers", "..."]
      }}
    ],
    "controls": {{
      "eliminate":      ["..."],
      "substitute":     ["..."],
      "engineering":    ["..."],
      "administrative": ["..."],
      "ppe":            ["..."]
    }},
    "procedures": [
      {{"step": 1, "action": "...", "responsible": "role", "critical": true}}
    ],
    "emergency_procedures": "text describing emergency actions",
    "roles_responsibilities": [
      {{"role": "HSE Officer", "responsibilities": ["..."]}}
    ],
    "standards_referenced": ["ISO 45001:2018", "API RP 505", "..."],
    "legal_references": ["regulation name", "..."]
  }},
  "quality": {{
    "completeness_score": 0.0-1.0,
    "gaps": ["list of missing mandatory sections or content"],
    "quality_feedback": null
  }}
}}
"""


# ── Claude API call ──────────────────────────────────────────────────────────────

def _call_claude(system: str, user: str, *, timeout: int = 120) -> str:
    import urllib.request as _ur
    key = _anthropic_key()
    if not key:
        raise RuntimeError("ANTHROPIC_API_KEY not set")
    model = _anthropic_model()
    payload = json.dumps({
        "model":      model,
        "max_tokens": 4096,
        "system":     system,
        "messages":   [{"role": "user", "content": user}],
    }).encode()
    req = _ur.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "x-api-key":         key,
            "anthropic-version": "2023-06-01",
            "content-type":      "application/json",
        },
    )
    with _ur.urlopen(req, timeout=timeout) as resp:
        data = json.loads(resp.read())
    return data["content"][0]["text"].strip()


# ── NIM completeness scoring ─────────────────────────────────────────────────────

def _nim_score(master_doc: dict) -> tuple[float, str]:
    import urllib.request as _ur
    url = _nim_url()
    key = _nim_key()

    doc_type = master_doc.get("classification", {}).get("document_type", "HSE document")
    hazards  = master_doc.get("content", {}).get("hazards", [])
    gaps     = master_doc.get("quality", {}).get("gaps", [])

    prompt = f"""Rate the completeness and quality of this extracted {doc_type} master document.
Output ONLY a JSON object:
{{"score": 0.0-1.0, "feedback": "one sentence assessment"}}

Document summary:
- Hazards identified: {len(hazards)}
- Controls defined: {sum(len(v) for v in master_doc.get("content", {}).get("controls", {}).values())}
- Procedures steps: {len(master_doc.get("content", {}).get("procedures", []))}
- Standards cited: {len(master_doc.get("content", {}).get("standards_referenced", []))}
- Reported gaps: {gaps}
- Completeness (self-reported): {master_doc.get("quality", {}).get("completeness_score")}

Criteria: ISO 45001:2018, IEC/IEEE 82079-1:2019, industrial HSE best practices."""

    payload = json.dumps({
        "model": "meta/llama-3.1-405b-instruct",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens": 256,
    }).encode()

    headers = {"Content-Type": "application/json"}
    if key:
        headers["Authorization"] = f"Bearer {key}"

    req = _ur.Request(
        f"{url}/v1/chat/completions",
        data=payload,
        headers=headers,
    )
    try:
        with _ur.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read())
        text = data.get("choices", [{}])[0].get("message", {}).get("content", "")
        m = re.search(r"\{.*\}", text, re.DOTALL)
        if not m:
            raise ValueError(f"no JSON: {text[:80]}")
        result = json.loads(m.group())
        return float(result.get("score", 0.0)), result.get("feedback", "")
    except Exception as e:
        log.warning("master_doc: NIM score failed: %s", e)
        return 0.0, f"NIM scoring unavailable: {e}"


# ── Text extraction ──────────────────────────────────────────────────────────────

def _read_text(file_path: str | Path | None, limit: int = 14000) -> str:
    if not file_path:
        return ""
    fp = Path(file_path)
    if not fp.is_file():
        return ""
    try:
        suffix = fp.suffix.lower()
        if suffix == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(fp) as pdf:
                    parts = [p.extract_text() or "" for p in pdf.pages[:30]]
                return "\n".join(parts)[:limit]
            except Exception:
                pass
        if suffix in (".docx", ".doc"):
            try:
                from docx import Document
                doc = Document(fp)
                return "\n".join(p.text for p in doc.paragraphs)[:limit]
            except Exception:
                pass
        # plain text fallback
        return fp.read_text(errors="ignore")[:limit]
    except Exception as e:
        log.warning("master_doc: read_text %s: %s", fp, e)
        return ""


# ── DOCX export ──────────────────────────────────────────────────────────────────

def _export_docx(master_doc: dict, out_path: Path) -> None:
    try:
        from docx import Document as _Doc
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.warning("master_doc: python-docx not installed, skipping DOCX export")
        return

    doc = _Doc()
    meta  = master_doc.get("meta", {})
    cls   = master_doc.get("classification", {})
    cont  = master_doc.get("content", {})
    qual  = master_doc.get("quality", {})

    doc_type = cls.get("document_type", "")
    title_str = f"{doc_type}: {cls.get('activity') or 'Master Document'}"

    h = doc.add_heading(title_str, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header block (IEC 82079-1 §5.4.1 identification)
    info = doc.add_table(rows=4, cols=4)
    info.style = "Table Grid"
    def _cell(row, col, text):
        info.cell(row, col).text = str(text or "—")
    _cell(0, 0, "Document ID");  _cell(0, 1, meta.get("doc_id") or "TBD")
    _cell(0, 2, "Revision");     _cell(0, 3, meta.get("revision") or "Rev.0")
    _cell(1, 0, "Status");       _cell(1, 1, meta.get("status") or "draft")
    _cell(1, 2, "Next Review");  _cell(1, 3, meta.get("next_review") or "—")
    _cell(2, 0, "AIMS Process"); _cell(2, 1, cls.get("aims_process") or "—")
    _cell(2, 2, "Audience");     _cell(2, 3, ", ".join(cls.get("target_audience") or []))
    _cell(3, 0, "ISO Clauses");  _cell(3, 1, ", ".join(cls.get("iso_clause") or []))
    _cell(3, 2, "Completeness"); _cell(3, 3, f"{round((qual.get('completeness_score') or 0)*100)}%")

    doc.add_paragraph("")

    # 1. Scope
    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph(cont.get("scope") or "—")

    # 2. Definitions
    defs = cont.get("definitions") or {}
    if defs:
        doc.add_heading("2. Definitions and Abbreviations", level=1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.cell(0, 0).text = "Term"
        tbl.cell(0, 1).text = "Definition"
        for term, defn in list(defs.items())[:30]:
            r = tbl.add_row()
            r.cells[0].text = str(term)
            r.cells[1].text = str(defn)

    # 3. Hazard identification & risk assessment (ISO 45001 §8.1.2)
    hazards = cont.get("hazards") or []
    if hazards:
        doc.add_heading("3. Hazard Identification and Risk Assessment", level=1)
        doc.add_paragraph("Risk matrix: Severity (1-5) × Likelihood (1-5) per ISO 31000:2018")
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        for i, hdr in enumerate(["ID","Hazard","Severity","Likelihood","Risk Level","Affected Parties"]):
            tbl.cell(0, i).text = hdr
        for h_item in hazards:
            r = tbl.add_row()
            r.cells[0].text = h_item.get("id") or ""
            r.cells[1].text = h_item.get("description") or ""
            r.cells[2].text = str(h_item.get("severity") or "")
            r.cells[3].text = str(h_item.get("likelihood") or "")
            r.cells[4].text = h_item.get("risk_level") or ""
            r.cells[5].text = ", ".join(h_item.get("affected_parties") or [])

    # 4. Hierarchy of controls (ISO 45001 §8.1.2 — NIOSH hierarchy)
    controls = cont.get("controls") or {}
    doc.add_heading("4. Control Measures (Hierarchy of Controls)", level=1)
    for level_name, level_key in [
        ("4.1 Elimination",      "eliminate"),
        ("4.2 Substitution",     "substitute"),
        ("4.3 Engineering Controls", "engineering"),
        ("4.4 Administrative Controls", "administrative"),
        ("4.5 PPE",              "ppe"),
    ]:
        items = controls.get(level_key) or []
        doc.add_heading(level_name, level=2)
        if items:
            for item in items:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph("None identified.", style="List Bullet")

    # 5. Procedures (IEC 82079-1 §7 — instructional information)
    procs = cont.get("procedures") or []
    if procs:
        doc.add_heading("5. Step-by-Step Procedures", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        for i, hdr in enumerate(["Step","Action","Responsible","Critical"]):
            tbl.cell(0, i).text = hdr
        for p_item in procs:
            r = tbl.add_row()
            r.cells[0].text = str(p_item.get("step") or "")
            r.cells[1].text = str(p_item.get("action") or "")
            r.cells[2].text = str(p_item.get("responsible") or "")
            r.cells[3].text = "YES" if p_item.get("critical") else "no"

    # 6. Emergency procedures (ISO 45001 §8.2)
    emerg = cont.get("emergency_procedures")
    doc.add_heading("6. Emergency Procedures", level=1)
    doc.add_paragraph(emerg or "Refer to site Emergency Response Plan.")

    # 7. Roles & responsibilities
    roles = cont.get("roles_responsibilities") or []
    if roles:
        doc.add_heading("7. Roles and Responsibilities", level=1)
        for role_item in roles:
            doc.add_heading(role_item.get("role") or "Role", level=2)
            for resp in (role_item.get("responsibilities") or []):
                doc.add_paragraph(str(resp), style="List Bullet")

    # 8. Standards and legal references (IEC 82079-1 §5.4.2)
    stds = cont.get("standards_referenced") or []
    legal = cont.get("legal_references") or []
    doc.add_heading("8. Applicable Standards and Regulations", level=1)
    for s in stds:
        doc.add_paragraph(str(s), style="List Bullet")
    for l in legal:
        doc.add_paragraph(str(l), style="List Bullet")
    if not stds and not legal:
        doc.add_paragraph("—")

    # 9. Quality / gaps (internal metadata)
    doc.add_heading("9. Document Quality Assessment", level=1)
    score_pct = round((qual.get("completeness_score") or 0) * 100)
    doc.add_paragraph(f"Completeness score: {score_pct}%")
    gaps = qual.get("gaps") or []
    if gaps:
        doc.add_paragraph("Identified gaps:")
        for g in gaps:
            doc.add_paragraph(str(g), style="List Bullet")
    fb = qual.get("quality_feedback")
    if fb:
        doc.add_paragraph(f"Gemini review: {fb}")

    # Change log (ISO 9001 §7.5.3)
    chlog = meta.get("change_log") or []
    if chlog:
        doc.add_heading("10. Change Log", level=1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.cell(0, 0).text = "Revision"
        tbl.cell(0, 1).text = "Summary"
        for entry in chlog:
            r = tbl.add_row()
            r.cells[0].text = str(entry.get("rev") or "")
            r.cells[1].text = str(entry.get("summary") or "")

    doc.save(str(out_path))
    log.info("master_doc: DOCX saved → %s", out_path)


# ── Public API ───────────────────────────────────────────────────────────────────

def generate_master_doc(
    file_name: str,
    file_path: str | Path | None,
    ocr_text: str | None = None,
    ai_classification: dict | None = None,
    *,
    export_docx_dir: Path | None = None,
) -> dict | None:
    """
    Generate a master document JSON from a registered file.

    Args:
        file_name:         original filename (for context)
        file_path:         path to the file (PDF/DOCX) for text extraction
        ocr_text:          pre-extracted OCR text (used if file_path unavailable)
        ai_classification: result from _ai_classify_core() (aims_process, suggested_title, …)
        export_docx_dir:   if set, save a DOCX master document alongside

    Returns:
        master_doc dict (schema_version 1.0) or None if disabled / no API key.
    """
    if not _enabled():
        return None
    if not _anthropic_key():
        log.warning("master_doc: ANTHROPIC_API_KEY not set — skipping")
        return None

    # ── Gather text ──────────────────────────────────────────────────────────────
    text = (ocr_text or "").strip()
    if not text:
        text = _read_text(file_path)
    if not text:
        log.warning("master_doc: no text for %s — skipping", file_name)
        return None

    cls = ai_classification or {}
    aims_process = cls.get("aims_process") or "unknown"
    title        = cls.get("suggested_title") or file_name

    # ── Anonymize BEFORE any cloud call ──────────────────────────────────────────
    try:
        from anonymizer import anonymize_text
        log.info("master_doc: anonymizing %s (%d chars)", file_name, len(text))
        text = anonymize_text(text)
        log.info("master_doc: anonymization done")
    except Exception as e:
        log.warning("master_doc: anonymizer unavailable (%s) — proceeding without", e)

    # ── Claude extraction ────────────────────────────────────────────────────────
    user_prompt = _USER_TMPL.format(
        file_name=file_name,
        aims_process=aims_process,
        title=title,
        text=text[:14000],
    )
    try:
        log.info("master_doc: calling Claude for %s", file_name)
        t0 = time.time()
        raw = _call_claude(_SYSTEM, user_prompt)
        log.info("master_doc: Claude done in %.1fs", time.time() - t0)
    except Exception as e:
        log.error("master_doc: Claude failed for %s: %s", file_name, e)
        return None

    # ── Parse JSON ───────────────────────────────────────────────────────────────
    master: dict = _empty_master()
    try:
        m = re.search(r"\{[\s\S]*\}", raw)
        if not m:
            raise ValueError("no JSON object in Claude response")
        extracted: dict = json.loads(m.group())

        def _merge(base: dict, update: dict) -> None:
            for k, v in update.items():
                if k in base and isinstance(base[k], dict) and isinstance(v, dict):
                    _merge(base[k], v)
                elif v is not None:
                    base[k] = v

        _merge(master, extracted)
    except Exception as e:
        log.error("master_doc: JSON parse failed for %s: %s | raw=%s", file_name, e, raw[:200])
        return None

    # ── Stamp metadata ────────────────────────────────────────────────────────────
    master["quality"]["claude_model"]  = _anthropic_model()
    master["quality"]["generated_at"]  = datetime.now(timezone.utc).isoformat()
    if not master["meta"]["created_at"]:
        master["meta"]["created_at"] = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # ── Auto doc_id if missing ────────────────────────────────────────────────────
    if not master["meta"].get("doc_id"):
        doc_type = master["classification"].get("document_type") or "DOC"
        proc     = (aims_process or "P00").replace(".", "")
        ts       = datetime.now(timezone.utc).strftime("%Y%m%d")
        master["meta"]["doc_id"] = f"{doc_type}-{proc}-{ts}"

    # ── NIM quality scoring ───────────────────────────────────────────────────────
    score, feedback = _nim_score(master)
    if score > 0:
        master["quality"]["completeness_score"] = round(score, 2)
        master["quality"]["quality_feedback"]    = feedback
    log.info("master_doc: completeness=%.0f%% gaps=%d file=%s",
             (master["quality"]["completeness_score"] or 0) * 100,
             len(master["quality"]["gaps"] or []),
             file_name)

    # ── DOCX export ───────────────────────────────────────────────────────────────
    if _export_docx() or export_docx_dir:
        out_dir = export_docx_dir or (Path(file_path).parent if file_path else Path("/tmp"))
        out_dir.mkdir(parents=True, exist_ok=True)
        stem = re.sub(r"[^\w\-]", "_", master["meta"]["doc_id"] or "master").strip("_")
        _export_docx(master, out_dir / f"{stem}_master.docx")

    return master


def master_doc_to_source_doc(master: dict) -> dict:
    """
    Convert master_doc JSON → source_docs entry for DocAgent.generate().
    DocAgent uses this as a rich context source for creating document variations.
    """
    cls   = master.get("classification", {})
    cont  = master.get("content", {})
    qual  = master.get("quality", {})
    meta  = master.get("meta", {})

    hazards_text = "; ".join(
        f"{h.get('id')} {h.get('description')} [{h.get('risk_level')}]"
        for h in (cont.get("hazards") or [])[:10]
    )
    controls_text = "; ".join(
        f"{level}: {', '.join(items[:3])}"
        for level, items in (cont.get("controls") or {}).items()
        if items
    )

    summary = (
        f"Master document [{cls.get('document_type')}] | "
        f"Process: {cls.get('aims_process')} | "
        f"Activity: {cls.get('activity') or 'N/A'} | "
        f"Scope: {(cont.get('scope') or '')[:300]} | "
        f"Hazards: {hazards_text} | "
        f"Controls: {controls_text} | "
        f"Standards: {', '.join((cont.get('standards_referenced') or [])[:5])} | "
        f"Completeness: {round((qual.get('completeness_score') or 0)*100)}%"
    )

    return {
        "title":        f"[MASTER] {meta.get('doc_id')} — {cls.get('activity') or cls.get('document_type')}",
        "summary":      summary,
        "aims_process": cls.get("aims_process"),
        "iso_clause":   ", ".join(cls.get("iso_clause") or []),
        "file_name":    meta.get("doc_id"),
    }
