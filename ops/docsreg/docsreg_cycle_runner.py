"""
DOCSREG certification cycle runner.

Provides a direct entry point for running the DOCSREG document-type
certification loop with either a local deterministic auditor or the
Claude Code CLI auditor as an advisory teacher layer.

The runner is intentionally thin:
  - Redis-backed orchestrator for the DOCSREG worker pipeline
  - deterministic structure + reference-governance checks
  - optional teacher audit via ``docsreg_auditor_cli``
  - CLI entry point for operators
"""
from __future__ import annotations

import argparse
import json
import logging
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal, Optional

from ops.agents.skills.docsreg_reference_governance import (
    ReferenceGateDecision,
    run_reference_governance_gate,
)
from ops.cyclic_skills import validate_structure
from ops.docsreg.docsreg_auditor_cli import (
    AUDITOR_VERDICT_BLOCKED,
    AUDITOR_VERDICT_PASS,
    AuditorRequest,
    make_claude_code_auditor,
)
from ops.docsreg.docsreg_document_type_cycle import (
    AUDIT_STATUS_COMPONENT_FAIL_REPAIRABLE,
    AUDIT_STATUS_COMPONENT_PASS,
    DocumentTypeCertificationLoop,
)
from ops.docsreg.docsreg_orchestrator import DocsregOrchestrator
from ops.docsreg.docsreg_retry_policy import STANDARD_RETRY
from ops.docsreg.docsreg_run_manifest import DocsregRunManifest

log = logging.getLogger("docsreg_cycle_runner")

TeacherMode = Literal["noop", "claude_code"]


@dataclass(frozen=True)
class DocsregCycleRunResult:
    """Structured result for a DOCSREG certification cycle run."""

    document_type: str
    draft_path: str
    evidence_root: str
    outcome: str
    passed: bool
    cycles_run: int
    best_quality: float
    notes: str
    teacher_mode: str
    run_id: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "document_type": self.document_type,
            "draft_path": self.draft_path,
            "evidence_root": self.evidence_root,
            "outcome": self.outcome,
            "passed": self.passed,
            "cycles_run": self.cycles_run,
            "best_quality": round(self.best_quality, 4),
            "notes": self.notes,
            "teacher_mode": self.teacher_mode,
            "run_id": self.run_id,
        }


def _load_redis_client(redis_url: str) -> Any:
    """Load a Redis client lazily. Raises RuntimeError on missing dependency."""
    try:
        import redis  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise RuntimeError(
            "redis package is not installed; cannot run DOCSREG with Redis backend"
        ) from exc
    return redis.Redis.from_url(redis_url)


def _read_document_text(draft_path: str | Path) -> str:
    path = Path(draft_path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        try:
            from docx import Document  # type: ignore[import-not-found]
        except Exception as exc:  # pragma: no cover - environment dependent
            raise RuntimeError("python-docx is required to read DOCX drafts") from exc
        doc = Document(str(path))
        chunks: list[str] = []
        for para in doc.paragraphs:
            txt = (para.text or "").strip()
            if txt:
                chunks.append(txt)
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join((cell.text or "").strip() for cell in row.cells).strip()
                if row_text:
                    chunks.append(row_text)
        return "\n".join(chunks)
    if suffix == ".pdf":
        pdf_text_parts: list[str] = []
        try:
            import pdfplumber  # type: ignore[import-not-found]

            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    page_text = (page.extract_text() or "").strip()
                    if page_text:
                        pdf_text_parts.append(page_text)
        except Exception:
            try:
                from pypdf import PdfReader  # type: ignore[import-not-found]

                reader = PdfReader(str(path))
                for page in reader.pages:
                    page_text = (page.extract_text() or "").strip()
                    if page_text:
                        pdf_text_parts.append(page_text)
            except Exception as exc:  # pragma: no cover - environment dependent
                raise RuntimeError(f"PDF text extraction failed for {path.name}: {exc}") from exc
        text = "\n".join(pdf_text_parts).strip()
        if not text:
            raise RuntimeError(f"PDF text extraction returned empty text for {path.name}")
        return text
    return path.read_text(encoding="utf-8")


def build_docsreg_auditor(
    *,
    document_type: str,
    evidence_root: str | Path,
    teacher_mode: TeacherMode = "noop",
    claude_bin: str = "claude",
    claude_timeout: int = 120,
):
    """Build an auditor function compatible with DocumentTypeCertificationLoop.

    The returned callable expects a DocsregRunManifest and returns a dict with:
    - status
    - quality
    - section_coverage
    - reference_governance
    - recommendations
    """

    evidence_root_path = Path(evidence_root)

    claude_auditor = None
    if teacher_mode == "claude_code":
        claude_auditor = make_claude_code_auditor(
            claude_bin=claude_bin,
            timeout=claude_timeout,
            evidence_root=evidence_root_path,
        )

    def _auditor(manifest: DocsregRunManifest) -> dict[str, Any]:
        draft_path = Path(manifest.draft_path)
        cycle_evidence_dir = evidence_root_path / document_type / manifest.run_id
        cycle_evidence_dir.mkdir(parents=True, exist_ok=True)

        try:
            document_text = _read_document_text(draft_path)
        except Exception as exc:  # noqa: BLE001
            log.error("DOCSREG auditor: cannot read draft %s: %s", draft_path, exc)
            return {
                "status": AUDITOR_VERDICT_BLOCKED,
                "quality": 0.0,
                "section_coverage": 0.0,
                "reference_governance": "FAIL",
                "recommendations": [f"draft unreadable: {exc}"],
                "notes": f"draft read failed: {exc}",
            }

        structure_report = validate_structure(document_text)
        reference_report = run_reference_governance_gate(document_text)

        structure_pass = bool(structure_report.passed)
        reference_pass = reference_report.decision == ReferenceGateDecision.PASS

        quality = 0.0
        if structure_pass and reference_pass:
            quality = 1.0
        else:
            quality = round(
                max(0.0, min(1.0, (structure_report.completeness_ratio * 0.7))),
                4,
            )
            if not reference_pass:
                quality = round(min(quality, 0.59), 4)

        notes = []
        recommendations: list[str] = []

        if not structure_pass:
            notes.append(
                f"structure completeness={structure_report.completeness_ratio:.4f}"
            )
            recommendations.append("increase section coverage and fill stub sections")

        if not reference_pass:
            notes.append(f"reference decision={reference_report.decision.value}")
            recommendations.extend(reference_report.repair_plan)

        status = AUDITOR_VERDICT_PASS if (structure_pass and reference_pass) else AUDITOR_VERDICT_BLOCKED

        if claude_auditor is not None:
            teacher_response = claude_auditor(
                AuditorRequest(
                    document_type=document_type,
                    cycle_id=manifest.run_id,
                    evidence_dir=str(cycle_evidence_dir),
                    quality=quality,
                    notes="; ".join(notes) if notes else "docsreg composite audit",
                )
            )
            if teacher_response.verdict != AUDITOR_VERDICT_BLOCKED:
                quality = max(quality, teacher_response.quality)
                if teacher_response.notes:
                    notes.append(teacher_response.notes)
                if teacher_response.verdict == AUDITOR_VERDICT_PASS and structure_pass and reference_pass:
                    status = AUDITOR_VERDICT_PASS
            else:
                status = AUDITOR_VERDICT_BLOCKED
                notes.append(teacher_response.notes or "teacher auditor blocked")

        return {
            "status": status if status == AUDITOR_VERDICT_PASS else AUDITOR_VERDICT_BLOCKED,
            "quality": quality,
            "section_coverage": structure_report.completeness_ratio,
            "reference_governance": "PASS" if reference_pass else "FAIL",
            "recommendations": recommendations,
            "notes": "; ".join(notes),
        }

    return _auditor


def run_docsreg_cycle(
    *,
    document_type: str,
    draft_path: str | Path,
    evidence_root: str | Path = "aims_workspace/docsreg_evidence",
    redis_url: str = "redis://aims-redis:6379/0",
    target_quality: float = 0.98,
    max_cycles: int = 7,
    stall_window: int = 3,
    min_quality_delta: float = 0.005,
    teacher_mode: TeacherMode = "noop",
    claude_bin: str = "claude",
    claude_timeout: int = 120,
) -> DocsregCycleRunResult:
    """Run a DOCSREG certification cycle and return a structured result."""
    evidence_root_path = Path(evidence_root)
    draft_path = Path(draft_path)
    redis_client = _load_redis_client(redis_url)
    orchestrator = DocsregOrchestrator(redis_client)
    auditor_fn = build_docsreg_auditor(
        document_type=document_type,
        evidence_root=evidence_root_path,
        teacher_mode=teacher_mode,
        claude_bin=claude_bin,
        claude_timeout=claude_timeout,
    )
    loop = DocumentTypeCertificationLoop(
        orchestrator=orchestrator,
        auditor_fn=auditor_fn,
        evidence_root=evidence_root_path,
        max_cycles=max_cycles,
        target_quality=target_quality,
        min_quality_delta=min_quality_delta,
        stall_window=stall_window,
    )

    outcome = loop.run_document_type(
        document_type=document_type,
        draft_path=str(draft_path),
        retry_policy=STANDARD_RETRY,
    )

    return DocsregCycleRunResult(
        document_type=document_type,
        draft_path=str(draft_path),
        evidence_root=str(evidence_root_path),
        outcome=outcome.outcome,
        passed=outcome.passed,
        cycles_run=outcome.cycles_run,
        best_quality=outcome.best_quality,
        notes=outcome.notes,
        teacher_mode=teacher_mode,
        run_id="",
    )


def main(argv: Optional[list[str]] = None) -> int:
    """CLI entry point for the DOCSREG certification cycle runner."""
    parser = argparse.ArgumentParser(description="DOCSREG certification cycle runner")
    parser.add_argument("--document-type", required=True)
    parser.add_argument("--draft-path", required=True)
    parser.add_argument("--evidence-root", default="aims_workspace/docsreg_evidence")
    parser.add_argument("--redis-url", default=os.environ.get("AIMS_REDIS_URL", "redis://aims-redis:6379/0"))
    parser.add_argument("--target-quality", type=float, default=0.98)
    parser.add_argument("--max-cycles", type=int, default=7)
    parser.add_argument("--stall-window", type=int, default=3)
    parser.add_argument("--min-quality-delta", type=float, default=0.005)
    parser.add_argument("--teacher-mode", choices=("noop", "claude_code"), default="noop")
    parser.add_argument("--claude-bin", default="claude")
    parser.add_argument("--claude-timeout", type=int, default=120)
    parser.add_argument("--output-json", default="")

    args = parser.parse_args(argv)

    result = run_docsreg_cycle(
        document_type=args.document_type,
        draft_path=args.draft_path,
        evidence_root=args.evidence_root,
        redis_url=args.redis_url,
        target_quality=args.target_quality,
        max_cycles=args.max_cycles,
        stall_window=args.stall_window,
        min_quality_delta=args.min_quality_delta,
        teacher_mode=args.teacher_mode,
        claude_bin=args.claude_bin,
        claude_timeout=args.claude_timeout,
    )

    payload = result.to_dict()
    print(json.dumps(payload, indent=2, ensure_ascii=False))

    if args.output_json:
        out_path = Path(args.output_json)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")

    return 0 if result.passed else 1


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
