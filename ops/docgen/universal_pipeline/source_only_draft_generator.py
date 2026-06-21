from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _unique(values: Iterable[Any]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = " ".join(str(value or "").split())
        key = text.casefold()
        if text and key not in seen:
            seen.add(key)
            result.append(text)
    return result


def _add_table(
    document: Document,
    headers: tuple[str, ...],
    rows: list[tuple[str, ...]],
) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def _add_field(run: Any, instruction: str, placeholder: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    elements = [begin, code, separate]
    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        elements.append(text)
    elements.append(end)
    run._r.extend(elements)


def _add_lead_paragraph(
    document: Document,
    label: str,
    text: str,
) -> None:
    paragraph = document.add_paragraph()
    lead = paragraph.add_run(label.rstrip(".") + ". ")
    lead.bold = True
    paragraph.add_run(text)


def _apply_layout(document: Document, title: str) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        header = section.header.paragraphs[0]
        header.text = f"{title} | Controlled Document | Revision 0"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(
            "Controlled copy — verify current revision before use | Page "
        )
        _add_field(footer.add_run(), "PAGE")


def _add_front_matter(document: Document, request: dict[str, Any]) -> None:
    document.add_heading("Document Control", level=1)
    _add_table(
        document,
        ("Field", "Value"),
        [
            ("Document title", str(request.get("title") or "")),
            ("Project / organization", str(request.get("project") or "")),
            ("Status", "Generated for controlled review"),
            ("Owner", "To be assigned by the approving organization"),
            ("Generation basis", "Request, approved sources, and document archetype"),
        ],
    )
    document.add_heading("Document Distribution", level=1)
    _add_table(
        document,
        ("Copy", "Recipient / repository", "Control"),
        [
            ("Master", "Approved document repository", "Controlled master"),
            ("Working", "Authorized users", "Verify revision before use"),
        ],
    )
    document.add_heading("Revision History", level=1)
    _add_table(
        document,
        ("Revision", "Date", "Description", "Approval status"),
        [("0", date.today().isoformat(), "Initial generated issue", "For review")],
    )
    document.add_heading("Table of Contents", level=1)
    paragraph = document.add_paragraph()
    _add_field(
        paragraph.add_run(),
        'TOC \\o "1-3" \\h \\z \\u',
        "Update field to refresh the table of contents.",
    )


def _pdca_phase(title: str) -> str:
    text = title.casefold()
    if any(
        word in text
        for word in (
            "improvement",
            "continual",
            "management review",
        )
    ):
        return "Act"
    if any(
        word in text
        for word in (
            "audit",
            "assurance",
            "performance",
            "reporting",
            "investigation",
            "quality control",
        )
    ):
        return "Check"
    if any(
        word in text
        for word in (
            "policy",
            "strategy",
            "risk",
            "scope",
            "governance",
            "planning",
            "lifecycle",
        )
    ):
        return "Plan"
    return "Do"


def _add_dynamic_policy_matrices(
    document: Document,
    section_content_plan: dict[str, Any] | None,
) -> None:
    plans = list((section_content_plan or {}).get("sections") or [])
    if not plans:
        return
    all_topics = [
        str(topic)
        for plan in plans
        for topic in plan.get("required_topics") or []
    ]
    if any("pdca" in topic.casefold() for topic in all_topics):
        document.add_heading(
            "Plan-Do-Check-Act (PDCA) Mapping",
            level=1,
        )
        _add_table(
            document,
            ("PDCA phase", "Policy sections", "Verification basis"),
            [
                (
                    phase,
                    "; ".join(
                        str(plan.get("section_title") or "")
                        for plan in plans
                        if _pdca_phase(
                            str(plan.get("section_title") or "")
                        )
                        == phase
                    ),
                    "Section controls, records, monitoring, and review "
                    "requirements.",
                )
                for phase in ("Plan", "Do", "Check", "Act")
            ],
        )

    owner_rows: list[tuple[str, str, str, str]] = []
    for plan in plans:
        roles = _unique(plan.get("required_roles") or [])
        controls = _unique(plan.get("required_controls") or [])
        records = _unique(plan.get("required_records") or [])
        if not roles and not controls and not records:
            continue
        owner_rows.append(
            (
                str(plan.get("section_title") or ""),
                roles[0] if roles else "To be assigned from approved context",
                ", ".join(controls),
                ", ".join(records),
            )
        )
    if owner_rows:
        document.add_heading(
            "Policy Section Ownership and Evidence Matrix",
            level=1,
        )
        _add_table(
            document,
            ("Policy section", "Lead role", "Key controls", "Key records"),
            owner_rows,
        )

    lifecycle_plans = [
        plan
        for plan in plans
        if "life cycle" in str(plan.get("section_title") or "").casefold()
        or "lifecycle" in str(plan.get("section_title") or "").casefold()
    ]
    stage_terms = (
        "concept",
        "design",
        "procurement",
        "construction",
        "commissioning",
        "operation",
        "maintenance",
        "ageing",
        "life extension",
        "decommissioning",
        "abandonment",
    )
    lifecycle_rows: list[tuple[str, str, str]] = []
    for plan in lifecycle_plans:
        roles = _unique(plan.get("required_roles") or [])
        records = _unique(plan.get("required_records") or [])
        for topic in _unique(plan.get("required_topics") or []):
            if not any(term in topic.casefold() for term in stage_terms):
                continue
            lifecycle_rows.append(
                (
                    topic,
                    ", ".join(roles)
                    or "Roles defined by approved lifecycle context",
                    ", ".join(records)
                    or "Source-linked lifecycle evidence",
                )
            )
    if lifecycle_rows:
        document.add_heading(
            "Lifecycle Responsibility and Evidence Matrix",
            level=1,
        )
        _add_table(
            document,
            ("Lifecycle stage / aspect", "Participating roles", "Evidence"),
            lifecycle_rows,
        )


def _postulate_text(
    postulate: dict[str, Any],
    section_title: str,
    ordinal: int,
) -> str:
    identifiers = _unique(postulate.get("requirement_ids") or [])
    focus = str(postulate.get("focus") or "the source requirement")
    roles = _unique(postulate.get("roles") or [])
    controls = _unique(postulate.get("controls") or [])
    records = _unique(postulate.get("records") or [])
    kpis = _unique(postulate.get("kpis") or [])
    role = roles[ordinal % len(roles)] if roles else "designated accountable owner"
    control = (
        controls[ordinal % len(controls)]
        if controls
        else "an approved implementation and verification control"
    )
    record = (
        records[ordinal % len(records)]
        if records
        else "a traceable implementation record"
    )
    trace = ", ".join(identifiers)
    text = (
        f"Source requirement {trace}. The {role} shall establish and maintain "
        f"{focus} within {section_title}. Implementation shall use "
        f"{control}, define decision and escalation authority, and retain "
        f"{record} as evidence of execution, review, exceptions, and closure."
    )
    if kpis:
        text += (
            f" Effectiveness shall be monitored using "
            f"{kpis[ordinal % len(kpis)]}."
        )
    return text


def _add_section(
    document: Document,
    plan: dict[str, Any],
    request: dict[str, Any],
    adaptive_config: dict[str, Any] | None = None,
    authored_section: dict[str, Any] | None = None,
) -> None:
    purpose = str(plan.get("section_purpose") or "")
    document.add_paragraph(
        f"{purpose} This section applies within the scope defined for "
        f"{request.get('project') or 'the requesting organization'}."
    )
    postulates = plan.get("requirement_postulates") or []
    topics = _unique(plan.get("required_topics") or [])
    source_ids = _unique(plan.get("source_requirements") or [])
    if topics:
        midpoint = max(1, (len(topics) + 1) // 2)
        trace_suffix = (
            " Source requirements: " + ", ".join(source_ids) + "."
            if source_ids
            else ""
        )
        document.add_paragraph(
            "The required policy subjects are "
            + "; ".join(topics[:midpoint])
            + ". They shall be expressed as clear policy intent, bounded by "
            "the approved request context and the cited source obligations."
            + trace_suffix
        )
        if len(topics) > midpoint:
            document.add_paragraph(
                "Implementation shall also address "
                + "; ".join(topics[midpoint:])
                + ". The document shall preserve their source meaning while "
                "stating the required outcome, evidence, control, and review "
                "expectation appropriate to this document type."
                + trace_suffix
            )
    if authored_section:
        for paragraph in authored_section.get("narrative_paragraphs") or []:
            document.add_paragraph(str(paragraph))
        authored_postulates = (
            authored_section.get("normative_postulates") or []
        )
        if authored_postulates:
            _add_lead_paragraph(
                document,
                "Normative policy postulates",
                "The following source-linked statements define the mandatory "
                "policy position for this section.",
            )
            for postulate in authored_postulates:
                document.add_paragraph(
                    str(postulate.get("text") or ""),
                    style="List Bullet",
                )
        summary = str(
            authored_section.get("records_and_controls_summary") or ""
        )
        if summary:
            _add_lead_paragraph(
                document,
                "Controls and records summary",
                summary,
            )
    elif postulates:
        _add_lead_paragraph(
            document,
            "Normative requirements",
            "The following source-linked statements define the mandatory "
            "policy position for this section.",
        )
        for index, postulate in enumerate(postulates):
            document.add_paragraph(
                _postulate_text(
                    postulate,
                    str(plan.get("section_title") or ""),
                    index,
                ),
                style="List Bullet",
            )
    subsections = [
        item
        for item in _unique(plan.get("required_subsections") or [])
        if not item.casefold().startswith("iso ")
    ]
    for subsection in subsections:
        document.add_heading(subsection, level=2)
        related = [
            item
            for item in postulates
            if subsection.casefold()
            in str(item.get("subsection") or "").casefold()
        ]
        if related:
            document.add_paragraph(
                "Source Traceability: "
                + ", ".join(
                    identifier
                    for item in related
                    for identifier in item.get("requirement_ids") or []
                )
                + "."
            )
    roles = _unique(plan.get("required_roles") or [])
    controls = _unique(plan.get("required_controls") or [])
    records = _unique(plan.get("required_records") or [])
    kpis = _unique(plan.get("required_kpis") or [])
    interfaces = _unique(plan.get("interfaces") or [])
    if roles:
        _add_lead_paragraph(
            document,
            "Responsibilities and authorities",
            "Accountability and participation shall be assigned to "
            + ", ".join(roles)
            + "; delegated authority and escalation shall be documented.",
        )
    if controls or records:
        parts: list[str] = []
        if controls:
            parts.append("controls: " + ", ".join(controls))
        if records:
            parts.append("records and evidence: " + ", ".join(records))
        _add_lead_paragraph(
            document,
            "Controls and evidence",
            "; ".join(parts) + ".",
        )
    if kpis:
        _add_lead_paragraph(
            document,
            "Monitoring and review",
            "Effectiveness shall be evaluated using "
            + ", ".join(kpis)
            + "; adverse trends and overdue actions shall be escalated.",
        )
    if interfaces:
        _add_lead_paragraph(
            document,
            "Interfaces",
            "Inputs, outputs, decision rights, handover evidence, and "
            "escalation shall be defined for interfaces with "
            + ", ".join(interfaces)
            + ".",
        )
    _add_adaptive_source_grounded_content(
        document,
        plan,
        adaptive_config or {},
    )


def _add_adaptive_source_grounded_content(
    document: Document,
    plan: dict[str, Any],
    adaptive_config: dict[str, Any],
) -> None:
    rules = adaptive_config.get("generation_rules") or {}
    title = str(plan.get("section_title") or "")
    target_sections = {
        str(item).casefold()
        for item in rules.get("target_sections") or []
        if str(item).strip()
    }
    if target_sections and title.casefold() not in target_sections:
        return
    terminology_level = int(
        rules.get("terminology_enrichment_level") or 0
    )
    depth_level = int(rules.get("section_depth_level") or 0)
    specificity_level = int(
        rules.get("section_specificity_level") or 0
    )
    if max(terminology_level, depth_level, specificity_level) <= 0:
        return

    topics = _unique(plan.get("required_topics") or [])
    controls = _unique(plan.get("required_controls") or [])
    records = _unique(plan.get("required_records") or [])
    roles = _unique(plan.get("required_roles") or [])
    interfaces = _unique(plan.get("interfaces") or [])
    source_ids = _unique(plan.get("source_requirements") or [])
    trace = ", ".join(source_ids)

    if terminology_level > 0 and topics:
        _add_lead_paragraph(
            document,
            "Source-grounded policy subjects",
            "The policy treatment in this section shall preserve the "
            "source meaning of "
            + "; ".join(topics)
            + f". Traceability: {trace}.",
        )

    if specificity_level > 0 and topics:
        _add_lead_paragraph(
            document,
            "Section-specific expectations",
            "The following source-grounded subjects require explicit "
            "implementation treatment.",
        )
        maximum = min(len(topics), 4 + specificity_level * 4)
        for topic in topics[:maximum]:
            document.add_paragraph(
                f"{topic}: the approved implementation shall state the "
                "applicable decision, control, evidence, and review "
                f"expectation supported by {trace}.",
                style="List Bullet",
            )

    if depth_level > 0:
        elements: list[str] = []
        if controls:
            elements.append("controls: " + ", ".join(controls))
        if records:
            elements.append("records: " + ", ".join(records))
        if roles:
            elements.append(
                "participating roles: " + ", ".join(roles)
            )
        if interfaces:
            elements.append("interfaces: " + ", ".join(interfaces))
        if elements:
            _add_lead_paragraph(
                document,
                "Implementation logic",
                "Implementation shall connect "
                + "; ".join(elements)
                + ". Role ownership and approval authority shall be assigned "
                "only where supported by the approved source context. "
                f"Traceability: {trace}.",
            )
        if depth_level > 1 and topics:
            document.add_paragraph(
                "Effectiveness review shall test whether the stated controls "
                "and records demonstrate implementation of "
                + ", ".join(topics[: min(8, len(topics))])
                + "; unresolved evidence gaps shall remain open actions "
                "until verified against the cited source requirements."
            )


def _add_compliance_matrix(
    document: Document,
    compliance_matrix: dict[str, Any] | None,
) -> None:
    rows = [
        (
            str(item.get("standard_reference") or ""),
            str(item.get("requirement_intent") or ""),
            str(item.get("aims_section") or ""),
            str(item.get("aims_element") or ""),
            str(item.get("implementation_evidence") or ""),
            str(item.get("records") or ""),
            str(item.get("verification_method") or ""),
            str(item.get("coverage_status") or ""),
        )
        for item in (compliance_matrix or {}).get("rows") or []
    ]
    _add_table(
        document,
        (
            "Standard / source reference",
            "Requirement intent",
            "Document section",
            "Requirement group",
            "Implementation evidence",
            "Records",
            "Verification method",
            "Coverage",
        ),
        rows,
    )


def _add_references(
    document: Document,
    reference_register: dict[str, Any] | None,
) -> None:
    rows = [
        (
            str(item.get("code") or ""),
            str(item.get("title") or ""),
            str(item.get("publication_year") or ""),
        )
        for item in (reference_register or {}).get("rows") or []
    ]
    _add_table(
        document,
        ("Code", "Title", "Publication year / edition"),
        rows,
    )
    for notice in (reference_register or {}).get("revision_notices") or []:
        document.add_paragraph(
            f"{notice.get('used_code')} is the controlled interpretation "
            f"basis and is under review against "
            f"{notice.get('current_official_code')} "
            f"(official publication date "
            f"{notice.get('current_official_publication_date')}; "
            f"source: {notice.get('official_source_url')})."
        )


def generate_source_only_draft(
    *,
    request: dict[str, Any],
    section_model: dict[str, Any],
    implementation_model: dict[str, Any],
    section_content_plan: dict[str, Any] | None = None,
    compliance_matrix: dict[str, Any] | None = None,
    reference_register: dict[str, Any] | None = None,
    adaptive_config: dict[str, Any] | None = None,
    authored_sections: dict[str, dict[str, Any]] | None = None,
    output_path: str | Path,
) -> Path:
    _ = implementation_model
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    title = str(request.get("title") or "Generated Document")
    document.core_properties.title = title
    document.core_properties.subject = str(
        section_model.get("document_type") or ""
    )
    document.core_properties.author = "AIMS Universal DOCGEN Pipeline"
    _apply_layout(document, title)
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = metadata.add_run(
        f"Project: {request.get('project')} | "
        f"Document type: {section_model.get('document_type')} | "
        f"Date: {date.today().isoformat()}"
    )
    run.font.size = Pt(10)
    _add_front_matter(document, request)
    _add_dynamic_policy_matrices(document, section_content_plan)

    plans = {
        str(item.get("section_title") or ""): item
        for item in (section_content_plan or {}).get("sections") or []
    }
    authored_sections = authored_sections or {}
    for index, section in enumerate(
        section_model.get("sections") or [],
        start=1,
    ):
        title_text = str(section.get("title") or "")
        document.add_heading(f"{index}. {title_text}", level=1)
        plan = plans.get(title_text)
        if plan:
            _add_section(
                document,
                plan,
                request,
                adaptive_config=adaptive_config,
                authored_section=authored_sections.get(title_text),
            )
        if "compliance matrix" in title_text.casefold():
            _add_compliance_matrix(document, compliance_matrix)
        if "reference" in title_text.casefold():
            _add_references(document, reference_register)

    document.save(output_path)
    return output_path
